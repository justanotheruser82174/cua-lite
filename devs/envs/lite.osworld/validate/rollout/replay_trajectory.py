"""Replay a past agent trajectory and run the evaluator for diagnosis.

Loads turn actions from a rollout log, replays them step-by-step in a fresh
container, then runs the task evaluator and prints detailed diagnostics.

Replay calibration: replay env_kwargs and inter-turn pacing MUST match the
original rollout, otherwise the same agent actions reproduce a different end
state. Two failure modes if mismatched:

  (a) Env-kwargs mismatch — rollout configs can pin resolution, max_steps,
      loop detection, or domain_overrides. Replaying with different kwargs can
      move click targets or change the step budget, so the same actions no
      longer reproduce the rollout state.
  (b) Race conditions — without an inter-turn sleep, replay fires turns
      back-to-back; LO's GUI hasn't yet updated (font dropdown still open,
      paragraph hasn't re-rendered) when the next click fires. Rollout has
      natural pacing from agent thinking time (LLM API call latency).

Defaults set both: `--config-path $LITE_OSWORLD_CONFIG` when set, otherwise
`scripts/configs/gpt/default/lite.osworld.yaml`, and `--inter-turn-sleep 1.0`.

Usage:
    # Auto-find latest rollout that has the task. Inherits GPT rollout
    # env_kwargs (currently max_steps=30, loop_detect=0) by default:
    uv run python devs/envs/lite.osworld/validate/rollout/replay_trajectory.py <task_id>

    # Specify a rollout directory explicitly:
    uv run python devs/envs/lite.osworld/validate/rollout/replay_trajectory.py <task_id> \\
        --rollout .logs/rollout/Qwen_Qwen3.5-9B/lite.osworld/20260504_202251

    # Override config — match a different rollout (e.g. Claude or local Qwen):
    uv run python devs/envs/lite.osworld/validate/rollout/replay_trajectory.py <task_id> \\
        --config-path scripts/configs/claude/default/lite.osworld.yaml

    # Override env_kwargs ad-hoc on top of the config:
    uv run python devs/envs/lite.osworld/validate/rollout/replay_trajectory.py <task_id> \\
        --env-kwargs '{"max_steps":30,"loop_detect":0}'

    # Replay against a private whole-config (required for private-build checks):
    LITE_OSWORLD_CONFIG=$PWD/.tmp/osworld.private.yaml \\
    uv run python devs/envs/lite.osworld/validate/rollout/replay_trajectory.py \\
        osworld_multi_apps_716a6079

    # Or override the image ad-hoc on top of the selected config:
    uv run python devs/envs/lite.osworld/validate/rollout/replay_trajectory.py \\
        osworld_multi_apps_716a6079 \\
        --env-kwargs '{"image":"cua-lite/lite.osworld:mine"}'

    # Disable config inheritance entirely (use env defaults — usually wrong
    # for reproducing a real rollout):
    uv run python devs/envs/lite.osworld/validate/rollout/replay_trajectory.py <task_id> \\
        --config-path ''

    # Adjust inter-turn pacing for slow-rendering apps (Impress / multi_apps):
    uv run python devs/envs/lite.osworld/validate/rollout/replay_trajectory.py <task_id> \\
        --inter-turn-sleep 2.0

    # Stop after N turns (don't replay the full trajectory):
    uv run python devs/envs/lite.osworld/validate/rollout/replay_trajectory.py \\
        <task_id> --max-turns 3
"""
from __future__ import annotations

import argparse
import asyncio
import json
import logging
import os
import pathlib
import subprocess

from lite.core.tools.action_space import LITE_ACTION_BATCH_TOOL_NAMES
from lite.core.tools.calls import (
    tool_call_arguments,
    tool_call_name,
    validate_lite_tool_call,
)
from lite.gym.types import LiteEnvStepResult
from lite.infer.debug.log_layout import ACTIONS_JSON, turn_dirs

logging.basicConfig(level=logging.WARNING, format="%(name)s %(levelname)s %(message)s")

DATA_FILES = [
    "lite/gym/envs/lite/osworld/data/train.synth.jsonl",
    "lite/gym/envs/lite/osworld/data/train.perturb.jsonl",
    "lite/gym/envs/lite/osworld/data/eval.jsonl",
]

ROLLOUT_BASE = pathlib.Path(".logs/rollout")
ROLLOUT_SPLITS = ("train", "train.synth", "train.perturb", "eval")
DEFAULT_REPLAY_CONFIG_PATH = "scripts/configs/gpt/default/lite.osworld.yaml"
DEFAULT_REPLAY_IMAGE = "cua-lite/lite.osworld:latest"
REMOTE_REPLAY_ENV_VARS = ("CUA_LITE_ENV_SERVER_URL", "CUA_LITE_ENV_SERVER_TOKEN")
BARE_REPLAY_ENV_KEYS = frozenset({
    "container",
    "computer_config",
    "debug",
    "display_resolution",
    "image",
    "loop_detect",
    "max_steps",
    "post_action_delay",
    "reset_timeout",
    "seed",
    "step_timeout",
    "vnc_port",
})
# Intentionally absent: the retired cursor wrapper key and ``cursor``.
# This replay path constructs a bare SandboxBaseEnv for local diagnostics; the
# Linux sandbox capture path composites its own cursor unconditionally and does
# not accept either public cursor key.


def _force_direct_local_replay() -> bool:
    """Replay uses a process-local temp env and local Docker diagnostics."""
    removed = {name: os.environ.pop(name, None) for name in REMOTE_REPLAY_ENV_VARS}
    if any(value for value in removed.values()):
        print(
            "  WARN: ignoring CUA_LITE_ENV_SERVER_URL/TOKEN for local "
            "rollout replay"
        )
        return True
    return False


def _load_task(task_id: str) -> dict:
    for path in DATA_FILES:
        p = pathlib.Path(path)
        if not p.exists():
            continue
        for line in p.read_text().splitlines():
            line = line.strip()
            if not line:
                continue
            t = json.loads(line)
            if t["task_id"] == task_id:
                return t
    raise SystemExit(f"Task {task_id!r} not found in any data file.")


def _find_rollout_dir(task_id: str) -> pathlib.Path:
    """Return the most-recent (by mtime) sample dir containing this task_id.

    Scans `.logs/rollout/<model>/lite.osworld/<timestamp>/<split>/<task_id>/sample_00/summary.json`
    across every model directory, so the script works regardless of which model rolled
    the trajectory (azure_gpt-5.5, Qwen_*, Claude_*, etc.).
    """
    candidates = []
    if not ROLLOUT_BASE.exists():
        raise SystemExit(f"No rollout base directory at {ROLLOUT_BASE!s}.")
    for model_dir in ROLLOUT_BASE.iterdir():
        env_dir = model_dir / "lite.osworld"
        if not env_dir.is_dir():
            continue
        for ts_dir in env_dir.iterdir():
            if not ts_dir.is_dir():
                continue
            for split in ROLLOUT_SPLITS:
                summary = ts_dir / split / task_id / "sample_00" / "summary.json"
                if summary.exists():
                    candidates.append((summary.stat().st_mtime, summary.parent))
    if not candidates:
        raise SystemExit(
            f"No completed rollout found for {task_id!r} under {ROLLOUT_BASE}/*/lite.osworld/.\n"
            "Use --rollout <path> to specify one explicitly."
        )
    chosen = max(candidates)[1]  # most recent mtime
    rollout_label = (
        f"{chosen.parent.parent.parent.parent.name}/"
        f"{chosen.parent.parent.parent.name}"
    )
    print(f"  Auto-selected rollout: {rollout_label}")
    return chosen


def _load_actions(sample_dir: pathlib.Path, max_turns: int | None) -> list[list[dict]]:
    """Load canonical Lite tool calls from each turn's 03_actions.json."""
    turns = []
    for turn_dir in turn_dirs(sample_dir):
        if max_turns is not None and len(turns) >= max_turns:
            break
        af = turn_dir / ACTIONS_JSON
        if not af.exists():
            continue
        d = json.loads(af.read_text())
        tool_calls = d.get("lite_message", {}).get("tool_calls", [])
        if tool_calls:
            for index, tool_call in enumerate(tool_calls):
                reason = validate_lite_tool_call(
                    tool_call,
                    f"{turn_dir.name}/{ACTIONS_JSON}.lite_message.tool_calls[{index}]",
                    require_id=False,
                )
                if reason is not None:
                    raise ValueError(reason)
            turns.append(tool_calls)
    return turns


def _tool_call_names_for_display(tool_calls: list[dict]) -> list[str | None]:
    """Logical action names for printing, without changing replay payloads."""
    names: list[str | None] = []
    for call in tool_calls:
        if not isinstance(call, dict):
            continue
        name = tool_call_name(call)
        args = tool_call_arguments(call)
        actions = args.get("actions")
        if name in LITE_ACTION_BATCH_TOOL_NAMES and isinstance(actions, list):
            names.extend(
                action.get("action")
                for action in actions
                if isinstance(action, dict)
            )
        else:
            names.append(name)
    return names


def _extract_file_paths(spec: dict) -> tuple[str | None, str | None]:
    """Return (result_path, expected_path) from evaluator config, or (None, None).

    Handles both single-evaluator (result/expected = dict) and multi-evaluator
    (result/expected = list of dicts, e.g. `compare_docx_tables × 2`). For
    multi-evaluator, returns the FIRST file pair (sufficient for inspection).
    """
    ev = spec["metadata"].get("evaluator", {})
    result_cfg = ev.get("result", {})
    expected_cfg = ev.get("expected", {})
    if isinstance(result_cfg, list):
        result_cfg = result_cfg[0] if result_cfg else {}
    if isinstance(expected_cfg, list):
        expected_cfg = expected_cfg[0] if expected_cfg else {}
    if not isinstance(result_cfg, dict):
        result_cfg = {}
    if not isinstance(expected_cfg, dict):
        expected_cfg = {}
    result_path = result_cfg.get("path") or result_cfg.get("dest")
    expected_path = expected_cfg.get("path") or expected_cfg.get("dest")
    return result_path, expected_path


def _load_env_kwargs_from_config(config_path: str | None) -> dict:
    """Load raw env_kwargs from a rollout yaml config.

    Returns empty dict if config_path is None or yaml has no env_kwargs section.
    """
    if not config_path:
        return {}
    try:
        import yaml
    except ImportError:
        print("  WARN: PyYAML not available; --config-path ignored")
        return {}
    with open(config_path) as f:
        cfg = yaml.safe_load(f) or {}
    return dict(cfg.get("env_kwargs", {}) or {})


def _resolve_replay_config_path(
    config_path_arg: str | None,
    env: dict[str, str] | None = None,
) -> str | None:
    """Resolve replay config precedence without hiding explicit empty-string opt out."""
    if config_path_arg == "":
        return None
    if config_path_arg:
        return config_path_arg
    env = os.environ if env is None else env
    return env.get("LITE_OSWORLD_CONFIG") or DEFAULT_REPLAY_CONFIG_PATH


def _apply_domain_overrides(env_kwargs: dict, task_domain: str) -> dict:
    env_kwargs = dict(env_kwargs)
    overrides = env_kwargs.pop("domain_overrides", {}) or {}
    if task_domain in overrides:
        env_kwargs.update(overrides[task_domain] or {})
    return env_kwargs


def _normalize_replay_env_kwargs_for_bare_sandbox(env_kwargs: dict) -> dict:
    """Translate lite.osworld rollout kwargs to bare SandboxBaseEnv kwargs."""
    env_kwargs = dict(env_kwargs)
    computer = env_kwargs.pop("computer", None)
    if isinstance(computer, dict):
        computer_display = computer.get("display_resolution")
        if (
            computer_display is not None
            and "display_resolution" not in env_kwargs
            and "resolution" not in env_kwargs
        ):
            env_kwargs["display_resolution"] = computer_display
        if "image" not in env_kwargs and isinstance(computer.get("image"), str):
            env_kwargs["image"] = computer["image"]

    if "resolution" in env_kwargs and "display_resolution" not in env_kwargs:
        env_kwargs["display_resolution"] = env_kwargs["resolution"]
    env_kwargs.pop("resolution", None)

    if "display_resolution" in env_kwargs:
        env_kwargs["display_resolution"] = tuple(env_kwargs["display_resolution"])

    return {k: v for k, v in env_kwargs.items() if k in BARE_REPLAY_ENV_KEYS}


def _merge_replay_env_kwargs(
    config_path: str | None,
    task_domain: str,
    extra_env_kwargs: dict | None,
) -> dict:
    from lite.gym import finalize_env_kwargs
    from lite.utils.config import deep_merge

    env_kwargs = deep_merge(
        _load_env_kwargs_from_config(config_path),
        extra_env_kwargs or {},
    )
    env_kwargs = finalize_env_kwargs(env_kwargs)
    env_kwargs = _apply_domain_overrides(env_kwargs, task_domain)
    return _normalize_replay_env_kwargs_for_bare_sandbox(env_kwargs)


def _effective_replay_image(env_kwargs: dict | None) -> str:
    if env_kwargs and isinstance(env_kwargs.get("image"), str) and env_kwargs["image"]:
        return env_kwargs["image"]
    computer_config = env_kwargs.get("computer_config") if env_kwargs else None
    if (
        isinstance(computer_config, dict)
        and isinstance(computer_config.get("image"), str)
        and computer_config["image"]
    ):
        return computer_config["image"]
    return DEFAULT_REPLAY_IMAGE


def _replay_computer_config(image: str) -> dict:
    return {
        "image": image,
        "display": "1920x1080",
        # Match lite.osworld's own configs/default.yaml: replaying a GEGL-heavy
        # GIMP trajectory under 4GB reproduces the container OOM-kill rather
        # than the trajectory, which is exactly how that failure got
        # misdiagnosed as an agent capability ceiling.
        "memory": "8GB",
        "cpu": "1",
        "timeout": 180,
    }


async def _run(
    task_id: str, sample_dir: pathlib.Path, spec: dict, max_turns: int | None,
    no_wait: bool = False,
    config_path: str | None = None,
    extra_env_kwargs: dict | None = None,
    inter_turn_sleep: float = 0.0,
) -> None:
    _force_direct_local_replay()

    import lite.gym as gym
    from lite.gym.envs.lite.osworld.src.utils.setup import setup_fn
    from lite.gym.envs.lite.osworld.src.utils.verify import evaluate_final_fn
    from lite.gym.sandbox import SandboxTaskConfig, register_tasks

    actions = _load_actions(sample_dir, max_turns)
    print(f"  Loaded {len(actions)} turns from {sample_dir}")
    for i, turn in enumerate(actions):
        names = _tool_call_names_for_display(turn)
        print(f"    turn_{i:04d}: {names}")

    meta = spec["metadata"]
    task_domain = meta.get("others", {}).get("domain", "")
    env_kwargs = _merge_replay_env_kwargs(config_path, task_domain, extra_env_kwargs)
    replay_image = _effective_replay_image(env_kwargs)
    task = SandboxTaskConfig(
        task_id=task_id,
        instruction=spec["instruction"],
        platform="desktop",
        computer=_replay_computer_config(replay_image),
        max_steps=30,
        metadata=meta,
        setup_fn=setup_fn,
        evaluate_final_fn=evaluate_final_fn,
    )

    temp_env_id = f"_replay_{task_id}"
    register_tasks(temp_env_id, {"train": [task]})

    # Load env_kwargs from the rollout config to mirror rollout behavior
    # (resolution, max_steps, loop_detect, domain_overrides). Without this,
    # replay can run with a different step budget or screen shape than the
    # original sample, so coordinates can land on the wrong widgets.
    print(f"  replay image: {replay_image}")
    if env_kwargs:
        print(f"  env_kwargs: {env_kwargs}")
    env = gym.make(f"{temp_env_id}@{task_id}", **env_kwargs)

    # Cycle 34: save per-turn screenshots so subagent can inspect what
    # happened during replay (the rollout trajectory_logger saves these for
    # the original rollout; replay was previously lossy on this).
    shot_dir = pathlib.Path("/tmp") / f"replay_{task_id}_screenshots"
    shot_dir.mkdir(parents=True, exist_ok=True)

    def _save_shot(png: bytes | None, name: str) -> None:
        if not png:
            return
        try:
            (shot_dir / f"{name}.png").write_bytes(png)  # raw PNG bytes now
        except Exception as ex:
            print(f"    (failed to save shot {name}: {ex})")

    def _first_result_image(result: LiteEnvStepResult) -> bytes | None:
        for tool_result in result.results:
            images = getattr(tool_result, "images", ()) or ()
            if images:
                return images[-1]
        return None

    print("\n  Resetting env (setup)...")
    obs = await env.reset()
    print(f"  Instruction: {(obs.text or '')[:120]}")
    _save_shot(obs.image, "turn_00_initial")

    inner = env
    while hasattr(inner, "env"):
        inner = inner.env
    computer = inner._computer
    container_name = getattr(inner, "_container_name", None)
    print(f"  Container: {container_name}")
    print(f"  Per-turn screenshots → {shot_dir}/")

    try:
        for i, turn_actions in enumerate(actions):
            names = _tool_call_names_for_display(turn_actions)
            print(f"\n  step turn_{i:04d}: {names}")
            result = await env.step(turn_actions)
            print(
                f"    reward={result.reward}  terminated={result.terminated}  "
                f"truncated={result.truncated}"
            )
            _save_shot(_first_result_image(result), f"turn_{i:04d}_after")
            if result.terminated or result.truncated:
                break
            if inter_turn_sleep > 0 and i + 1 < len(actions):
                # Cycle 34: rollout has natural inter-turn delay (agent thinks
                # before next turn — usually multiple seconds for LLM API calls
                # in agent rollout; for local Qwen ~1s; eval rollout 5-15s+
                # depending on model). Replay fires turns back-to-back which
                # can race LO's GUI update (e.g. font picker dropdown still
                # open when next click happens, click lands on dropdown not
                # on intended widget). Insert a small sleep to mimic the
                # inter-turn pause.
                await asyncio.sleep(inter_turn_sleep)

        # Run evaluator
        print("\n  Running evaluate_final_fn...")
        reward = await evaluate_final_fn(task, computer)
        print(f"  Final score: {reward}")

        # Pull result + expected files for inspection
        result_path, expected_path = _extract_file_paths(spec)
        if result_path and container_name:
            _pull_and_inspect(container_name, result_path, expected_path, spec)

        if reward == 1.0:
            print("\n  ✓ PASS")
        else:
            print(f"\n  ✗ FAIL (score={reward})")
            if container_name and not no_wait:
                print("\n  Container still alive — inspect interactively:")
                print(f"    docker exec -it {container_name} bash")
                print("\n  Non-interactive (subagent-safe):")
                print(
                    f"    docker exec {container_name} bash -c "
                    "'cat /home/user/.config/GIMP/2.10/gimprc'"
                )
                print("\n  [diag] sleeping 1800s instead of waiting on stdin")
                await asyncio.sleep(1800)
            elif container_name and no_wait:
                print(
                    f"\n  Container still alive (--no-wait): "
                    f"docker exec {container_name} bash -c '...'"
                )
                print("  Container will be closed when this script exits.")

    finally:
        await env.close()


def _lo_normalize_in_container(container: str, container_path: str, fmt: str) -> None:
    """Mirror eval framework's `_lo_normalize_ooxml` — run soffice headless
    inside the container to normalize the file in place. Without this, the
    debug `_inspect_docx` output shows raw python-docx style names ('Normal')
    while the actual eval-framework `_get_expected` would have normalized them
    to 'LO-Normal' (matching the agent's LO-saved result), so the per-paragraph
    "diff" printed below would falsely flag style mismatches that the real
    metric never sees.
    """
    cmd = (
        f"tmpd=$(mktemp -d) && "
        f"DISPLAY=:1 soffice --headless --norestore --nofirststartwizard "
        f"--convert-to {fmt} --outdir \"$tmpd\" '{container_path}' 2>/dev/null && "
        f"[ -f \"$tmpd/$(basename '{container_path}')\" ] && "
        f"cp \"$tmpd/$(basename '{container_path}')\" '{container_path}'; "
        f"rm -rf \"$tmpd\"; true"
    )
    subprocess.run(
        ["docker", "exec", container, "bash", "-c", cmd],
        capture_output=True,
        timeout=60,
    )


def _pull_and_inspect(
    container: str,
    result_path: str,
    expected_path: str | None,
    spec: dict,
) -> None:
    """Pull result (and expected if available) from container and print diff summary."""
    ev_func = spec["metadata"].get("evaluator", {}).get("func", "")
    result_local = "/tmp/dbg_result"
    expected_local = "/tmp/dbg_expected"
    ext = pathlib.Path(result_path).suffix

    # Cycle 34 Fix A: also normalize result for symmetric comparison. The
    # eval framework's `_get_result` now LO-normalizes too (matching
    # `_get_expected`); mirror the same here so debug diff reflects the
    # actual metric input. Without this, agent's interactive Ctrl+S save can
    # leave style.name as 'normal' (e.g. simple docs like HK_train_record)
    # while metric sees the normalized 'LO-normal' version.
    if result_path.endswith(".docx"):
        _lo_normalize_in_container(container, result_path, "docx")
    elif result_path.endswith(".pptx"):
        _lo_normalize_in_container(container, result_path, "pptx")
    r = subprocess.run(
        ["docker", "cp", f"{container}:{result_path}", result_local + ext],
        capture_output=True,
    )
    result_status = "OK" if r.returncode == 0 else r.stderr.decode().strip()
    print(f"\n  Pulled result ({result_path}, LO-normalized): {result_status}")

    if expected_path:
        # Cycle 34: mirror eval framework's `_get_expected` behavior — it auto-
        # normalizes expected OOXML files before metric eval. Without doing the
        # same here, debug diff shows raw python-docx output (style 'Normal',
        # different run splits) vs LO-saved result, producing FALSE diffs that
        # don't reflect what compare_docx_strict actually sees.
        if expected_path.endswith(".docx"):
            _lo_normalize_in_container(container, expected_path, "docx")
        elif expected_path.endswith(".pptx"):
            _lo_normalize_in_container(container, expected_path, "pptx")
        r2 = subprocess.run(
            ["docker", "cp", f"{container}:{expected_path}", expected_local + ext],
            capture_output=True,
        )
        expected_status = "OK" if r2.returncode == 0 else r2.stderr.decode().strip()
        print(f"  Pulled expected ({expected_path}, LO-normalized): {expected_status}")

    if "docx" in ev_func or ext == ".docx":
        _inspect_docx(result_local + ext, expected_local + ext if expected_path else None, spec)


def _inspect_docx(result_local: str, expected_local: str | None, spec: dict) -> None:
    try:
        from docx import Document

        from lite.gym.envs.lite.osworld.src.eval.metrics import (
            _paragraph_format_signature,
            compare_docx_strict,
        )
    except ImportError as e:
        print(f"  (docx inspect skipped: {e})")
        return

    dr = Document(result_local)
    print(f"\n  RESULT ({result_local}): {len(dr.paragraphs)} paragraphs")
    for i, p in enumerate(dr.paragraphs[:5]):
        bolds = [r.bold for r in p.runs]
        print(f"    para[{i}]: text={p.text[:50]!r}  bold={bolds}  style={p.style.name!r}")

    if not expected_local or not pathlib.Path(expected_local).exists():
        return

    de = Document(expected_local)
    print(f"\n  EXPECTED ({expected_local}): {len(de.paragraphs)} paragraphs")
    for i, p in enumerate(de.paragraphs[:5]):
        bolds = [r.bold for r in p.runs]
        print(f"    para[{i}]: text={p.text[:50]!r}  bold={bolds}  style={p.style.name!r}")

    mismatches = []
    for i, (p1, p2) in enumerate(zip(dr.paragraphs, de.paragraphs)):
        s1 = _paragraph_format_signature(p1)
        s2 = _paragraph_format_signature(p2)
        if s1 != s2:
            mismatches.append((i, s1, s2))
    if mismatches:
        print(f"\n  Para-format mismatches: {len(mismatches)}")
        for i, a, b in mismatches[:5]:
            print(f"    para[{i}]: result={a}  expected={b}")

    ev = spec["metadata"].get("evaluator", {})
    opts = ev.get("options", {})
    score = compare_docx_strict(result_local, expected_local, **{
        k: opts.get(k, False)
        for k in ("examine_font_name", "examine_font_size", "examine_color",
                  "examine_highlight", "examine_images")
    })
    print(f"\n  compare_docx_strict score: {score}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Replay a past agent trajectory for a task.")
    parser.add_argument("task_id", help="Task ID to replay (e.g. synth_writer_bold_text_0001)")
    parser.add_argument(
        "--rollout",
        help="Path to rollout timestamp dir (auto-detected if omitted)",
    )
    parser.add_argument("--max-turns", type=int, default=None, help="Stop after N turns")
    parser.add_argument("--no-wait", action="store_true",
                        help="Skip the interactive 'Press Enter' pause on FAIL (subagent-safe)")
    parser.add_argument(
        "--config-path", default=None,
        help="YAML rollout config to inherit env_kwargs (resolution / max_steps / "
             "domain_overrides) so replay matches the original rollout config. "
             "Default: $LITE_OSWORLD_CONFIG if set, otherwise "
             f"{DEFAULT_REPLAY_CONFIG_PATH}. Pass empty string '' to skip config "
             "loading.",
    )
    parser.add_argument(
        "--env-kwargs", default=None, type=json.loads,
        help='JSON env_kwargs to merge on top of --config-path env_kwargs. '
             'Example: \'{"max_steps":30,"loop_detect":0}\'. '
             'Use \'{"image":"cua-lite/lite.osworld:mine"}\' to replay '
             'against a private image tag.',
    )
    parser.add_argument(
        "--inter-turn-sleep", type=float, default=1.0,
        help="Seconds to sleep between turns (default: 1.0). Mimics the natural "
             "agent thinking pause between turns in rollout — replay fires turns "
             "back-to-back, which can race LO GUI updates (e.g. font dropdown "
             "still open when next click happens). Set 0 to disable.",
    )
    args = parser.parse_args()

    print(f"\n{'='*60}")
    print(f"TASK: {args.task_id}")
    print(f"{'='*60}")

    spec = _load_task(args.task_id)
    print(f"  Instruction: {spec['instruction'][:100]}")

    if args.rollout:
        rollout = pathlib.Path(args.rollout)
        # Accept either <timestamp>/ or <timestamp>/train/ or <timestamp>/eval/
        sample_dir = None
        for candidate in [
            rollout / "sample_00",
            rollout / "train" / args.task_id / "sample_00",
            rollout / "eval" / args.task_id / "sample_00",
            rollout / "train.perturb" / args.task_id / "sample_00",
            rollout / "train.synth" / args.task_id / "sample_00",
        ]:
            if candidate.exists():
                sample_dir = candidate
                break
        if sample_dir is None:
            raise SystemExit(f"sample_00 not found under {rollout}")
    else:
        sample_dir = _find_rollout_dir(args.task_id)

    config_path = _resolve_replay_config_path(args.config_path)
    asyncio.run(_run(
        args.task_id, sample_dir, spec, args.max_turns,
        no_wait=args.no_wait,
        config_path=config_path,
        extra_env_kwargs=args.env_kwargs,
        inter_turn_sleep=args.inter_turn_sleep,
    ))


if __name__ == "__main__":
    main()
