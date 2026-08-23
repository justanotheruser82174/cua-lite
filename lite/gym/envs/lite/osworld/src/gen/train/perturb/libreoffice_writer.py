"""Writer domain per-task perturbation (Track B, TYPE_1 + TYPE_2 + TYPE_3).

8 op types with deterministic oracle (python-docx): bold, font, find_replace,
spacing, uppercase, strikethrough, lowercase, italic. TYPE_1 resamples the
eval task's own op with different params (eligible tasks only, see _TYPE1_FNS).
TYPE_2 draws from an explicit per-task candidate pool (_T2_VARIANTS, ≥4 ops
per task) so max_type2>2 is always meaningful and distribution match holds
for any max_type2 value.

TYPE_3 (P3-4-writer) — per-base archetype rows targeting evaluators that
compare_docx_strict does NOT cover (3 of 19 missing eval-evaluator rows are
real gaps; the other 16 are subsumed by strict). Each archetype declares its
own evaluator + oracle; runs *in addition to* TYPE_1/TYPE_2 on its base. See
_TYPE3_BASES + _emit_{tabstops,subscript,pdf_export}_rows.

Instruction style (D5 — length + multi-step diversity):

Each op carries a 9-12 entry paraphrase pool partitioned into three length
buckets so the emitted dataset spans p25 ≤ 14 words, p50 ≈ 28 words, and
p75 ≥ 40 words simultaneously (matching eval's bimodal short/long shape):
  - 4 short  (~6-12 words)     → covers eval p25 (~10 words)
  - 4 medium (~24-32 words)    → covers eval mean (~27.7 words)
  - 4 long   (~42-65 words)    → covers eval p75/max (~47-85 words)
TYPE_3 archetype pools (tabstops/subscript/pdf_export) follow the same
short/medium/long structure with smaller short tier (2/3-5/3-4) since
they're called less often.

A subset of variants per pool embed multi-step sequence keywords
(then / next / first / once / after) so the multi-step ratio in the
emitted dataset clears the ≥18% target while preserving the V3 mean
within ±5pp of the eval baseline (~27.7 words). Pool means are tuned
to ~28-31 words so realised per-row mean stays near 30-32. Polite share
(~30%) is shaped by ~20% polite-leading templates per pool. The "save"
keyword is intentionally excluded from every variant (V3 save% = 0).

Design doc: devs/envs/lite.osworld/perturb/libreoffice_writer.md

Usage:
    uv run python -m lite.gym.envs.lite.osworld.src.gen.train --track perturb --domain libreoffice_writer
"""

from __future__ import annotations

import json
import random
import textwrap
from pathlib import Path

from lite.gym.envs.lite.osworld.src.gen.common import (
    LO_SAVE_POSTCONFIG,
)
from lite.gym.envs.lite.osworld.src.gen.train.perturb._utils import (
    make_perturb_row,
)

_ANALYSIS_PATH = Path("/tmp/writer_full.json")

PERTURB_TYPE_1 = "type1"
PERTURB_TYPE_2 = "type2"
# TYPE_3: per-base archetype rows that target the **specific eval evaluator**
# instead of compare_docx_strict. Used for the 3 writer eval bases that fall
# outside strict's coverage (P3-4-writer):
#   - check_tabstops               (0a0faba3)
#   - compare_docx_files+compare_subscript_contains  (0b17a146)
#   - compare_pdfs+compare_pdfs+compare_pdfs+compare_pdfs  (4bcb1253)
# See devs/envs/lite.osworld/perturb/libreoffice_writer.md "Archetypes (TYPE_3)".
PERTURB_TYPE_3 = "type3"

# ---------------------------------------------------------------------------
# Core helpers
# ---------------------------------------------------------------------------

def _get_file_path(eval_row: dict) -> str | None:
    for step in eval_row["metadata"]["config"]:
        if step.get("type") == "download":
            files = step["parameters"].get("files", [])
            if files:
                return files[0]["path"]
        if step.get("type") == "open":
            return step["parameters"].get("path")
    return None


def _make_config_step(py_code: str) -> dict:
    return {"type": "execute", "parameters": {
        "command": f"python3 << 'PYEOF'\n{py_code}\nPYEOF", "shell": True,
    }}


def _lo_normalize_cmd(path: str, fmt: str) -> str:
    return (
        f"tmpd=$(mktemp -d) && "
        f"DISPLAY=:1 soffice --headless --norestore --nofirststartwizard "
        f"--convert-to {fmt} --outdir \"$tmpd\" '{path}' 2>/dev/null && "
        f"[ -f \"$tmpd/$(basename '{path}')\" ] && "
        f"cp \"$tmpd/$(basename '{path}')\" '{path}'; "
        f"rm -rf \"$tmpd\"; true"
    )


# ---------------------------------------------------------------------------
# Oracle builder
# ---------------------------------------------------------------------------

def _build_oracle(file_path: str, expected_path: str) -> list[dict]:
    # LO normalization asymmetry fix: normalize expected_path once (step 1),
    # then cp to file_path (step 2) — a byte-identical copy of the normalized
    # expected. The eval result-getter (_lo_normalize_docx in runner.py) applies
    # exactly ONE more normalize pass to BOTH result and expected at eval time,
    # so result = f(f1) and expected = f(f1) — identical on every field.
    # NOTE: LO's --convert-to normalize is NOT idempotent (trailing empty para,
    # line_spacing/indent drift), so the earlier step-3 re-normalize of file_path
    # gave result one MORE pass than expected -> drift -> spurious 0. Removed.
    return [
        {"type": "execute", "parameters": {
            "command": _lo_normalize_cmd(expected_path, "docx"), "shell": True,
        }},
        {"type": "execute", "parameters": {
            "command": f"cp '{expected_path}' '{file_path}'", "shell": True,
        }},
    ]


# ---------------------------------------------------------------------------
# Evaluator builder
# ---------------------------------------------------------------------------

_TEXT_ONLY_OP_INDICES = {2, 4, 6}  # find_replace, uppercase, lowercase


def _build_evaluator(op_idx: int, file_path: str, expected_path: str) -> dict:
    # Text-only ops (find_replace / uppercase / lowercase) don't change
    # formatting — use OSWorld's `compare_docx_files` (paragraph-text
    # match, no format check) so LO Writer's docx round-trip style.name
    # renames don't false-fail an otherwise-correct edit.
    if op_idx in _TEXT_ONLY_OP_INDICES:
        return {
            "func": "compare_docx_files",
            "result": {"type": "vm_file", "path": file_path, "dest": file_path.split("/")[-1]},
            "expected": {"type": "vm_file", "path": expected_path, "dest": "expected_file"},
            "postconfig": LO_SAVE_POSTCONFIG,
        }

    # Format ops (bold / font / spacing / strikethrough / italic) MUST keep
    # `compare_docx_strict` — that's the whole point of these tasks.
    # Per-op options: relax char-format fields LibreOffice's docx round-trip
    # normalizes on untouched runs, BUT keep strict on the field this op
    # targets. _make_font (op_idx=1) targets font_name → keep strict.
    # _make_bold/_make_spacing target bold/line_spacing → relaxed at the
    # char-format layer (paragraph-format-signature still checks line_spacing).
    eval_options = {
        "examine_font_name": op_idx == 1,  # _make_font keeps strict
        "examine_font_size": False,
        "examine_color": False,
        "examine_highlight": False,
        "examine_images": False,
    }
    return {
        "func": "compare_docx_strict",
        "result": {"type": "vm_file", "path": file_path, "dest": file_path.split("/")[-1]},
        "expected": {"type": "vm_file", "path": expected_path, "dest": "expected_file"},
        "options": eval_options,
        "postconfig": LO_SAVE_POSTCONFIG,
    }


# ---------------------------------------------------------------------------
# Row builder
# ---------------------------------------------------------------------------

def _build_perturb_row(
    eval_row: dict,
    op_idx: int,
    instruction: str,
    py_code: str,
    file_path: str,
    expected_path: str,
    perturb_type: str = PERTURB_TYPE_2,
) -> dict:
    oracle = _build_oracle(file_path, expected_path)
    evaluator = _build_evaluator(op_idx, file_path, expected_path)
    # oracle_after_postconfig=True: validate.py runs postconfig (Ctrl+S)
    # first, then kills soffice, then runs oracle. The oracle just needs to
    # cp the expected file over; LO is already dead and _postconfig_done=True
    # prevents the evaluator from sending another Ctrl+S that would revert it.
    return make_perturb_row(
        eval_row=eval_row,
        knob_assignment={"type": perturb_type, "op": op_idx, "variant": instruction[:30]},
        new_instruction=instruction,
        new_oracle=oracle,
        new_evaluator=evaluator,
        perturb_config_step=_make_config_step(py_code),
        oracle_after_postconfig=True,
    )


# ---------------------------------------------------------------------------
# Instruction paraphrase pools — narrative-style with context + motivation.
# NO save instructions (AGENTS.md hard constraint).
#
# Design targets (matching eval distribution, not exceeding):
#   - 5 paraphrases per op
#   - 20-35 words per template (eval avg: ~28 words)
#   - ~20% polite-leading templates (eval polite share: 17%)
#   - Mixed lead-ins: "I'm reviewing...", "While editing...", "For a writeup...",
#     "Quick edit:", "Could you..." — so distribution is varied
# ---------------------------------------------------------------------------

_ORDINALS = {0: "first", 1: "second", 2: "third", 3: "fourth", 4: "fifth"}

# Bases where "the {ord} paragraph" instructions are visually unambiguous.
# Each value is the list of python-docx paragraph indices that match the
# user's natural visual counting of body paragraphs (skipping titles,
# headings, captions, metadata, line-break-merged paragraphs).
#
# When generating an ordinal-paragraph perturb op for a base IN this dict:
#   ord_pos = rng.randint(0, min(4, len(body_idxs)-1))
#   target_idx = body_idxs[ord_pos]                  # python-docx index for oracle
#   ord_word   = _ORDINALS[ord_pos]                  # ordinal in instruction
#
# Bases NOT in this dict have ambiguous structure (interleaved headings,
# small-caps lead paragraphs, banner+image+table layouts, multi-page legal
# covers, etc.) where agent's visual paragraph counting diverges from
# python-docx's <w:p> ordering. Para-idx ops (bold/spacing/upper/strike/
# lower/italic) are dropped from the T2 pool for these bases — only
# doc-wide ops (font, find_replace) remain.
#
# Validated 2026-05-09 against rendered PDFs at /tmp/writer_png/<short>.png
# (validation: systematic ordinal-ambiguity avoidance).
_ORDINAL_SAFE_BASES: dict[str, list[int]] = {
    "0810415c": [0, 2, 4, 6, 8],    # 6+ uniform body, Heading 1 only at para[13]
    "0a0faba3": [0, 1, 2, 3, 4],    # 10 uniform body, no title
    "0e763496": [0, 1, 2, 3, 4],    # 14 uniform body, no title
    "66399b0d": [2, 4, 6, 8, 10],   # title (centered bold) + 5 uniform body
    "6f81754e": [0, 1, 2, 3, 4],    # 100 uniform train-record data lines
    "72b810ef": [0, 2, 4],          # 3 uniform body, no title (n_body=3)
    "8472fece": [1],                # title (large bold) + 1 body (n_body=1)
    "adf5e2c3": [4, 6, 8, 10, 12],  # Heading 1 + topic line + 5 essay body
    "b21acd93": [0, 1, 2],          # 3 uniform body, no title (n_body=3)
}

# Pool composition (D5: instruction length + multi-step diversity):
#   2 short  (~6-12 words)  → covers eval p25 (~10 words)
#   4 medium (~28-32 words) → matches eval p50 (~14, padded toward avg ~28)
#   2 long   (~50-70 words) → covers eval p75/max (~47-85 words)
# A subset of variants per pool embed multi-step sequence keywords
# (then / next / first / once / after) so the multi-step ratio in the
# emitted dataset clears the ≥18% target. Word counts are tuned so the
# pool mean stays near 28 (V3 mean ±5pp), and short/long variants
# alternate in pool order so distinct paraphrase sampling stays diverse.
_BOLD_VARIANTS = [
    # short
    "Bold the {ord} paragraph for me.",
    "Please make the {ord} paragraph bold throughout.",
    "Apply bold formatting to the {ord} paragraph.",
    "Bold every word of the {ord} paragraph.",
    # medium (existing narrative pads)
    "I'm reviewing a draft and want the {ord} paragraph to stand out as a key point — please apply bold formatting to that whole paragraph for me.",
    "While editing this document I noticed the {ord} paragraph carries the main argument; could you make the entire {ord} paragraph bold so readers spot it quickly?",
    "I want to draw attention to the {ord} paragraph because it summarises the section, so first scroll to it and then apply bold formatting to every word in that paragraph.",
    # multi-step medium (then/first sequence keywords)
    "First locate the {ord} paragraph in the document, then apply bold formatting to every word inside it so the emphasis reads consistently across the entire paragraph for my reviewers.",
    # long narrative pad
    "I'm finalising this writeup before tomorrow's stakeholder review and the {ord} paragraph is the executive summary that needs to read as the most important sentence on the page — please apply bold formatting to every word in that paragraph so the printed handout draws the reader's eye there first when the team passes copies around the conference room.",
    # long multi-step (after/then sequence keywords)
    "I'm preparing the report copy for an offsite and want a clean visual hierarchy — first scan the document until you find the {ord} paragraph, then bold every word in that paragraph from start to finish so reviewers immediately see it carries the main argument once they open the file on Monday.",
    # extra long narrative pad (no multi-step keyword)
    "For a writeup I'm finalizing tonight, the {ord} paragraph is the takeaway my reviewers need to see immediately when they open the file — please apply bold formatting across every word of the {ord} paragraph so the emphasis directs their attention straight to that line on the page.",
    # long multi-step variant 2 (then sequence keyword)
    "I'm cleaning up the visual hierarchy in this writeup before tomorrow's print run — first scan to the {ord} paragraph in the document, then bold every word in that paragraph end to end, leaving the rest of the body text untouched.",
]

_FONT_VARIANTS = [
    # short
    'Change the font to "{font}" everywhere.',
    'Set every paragraph and table to "{font}".',
    'Use "{font}" throughout the document.',
    'Apply "{font}" to all body and table text.',
    # medium narrative pads
    'I\'m preparing this document for submission and the style guide requires a single font throughout — could you change the font to "{font}" for every paragraph and table cell?',
    'While polishing this draft I realised the typography is inconsistent. Please change the font of the entire document, including any tables, to "{font}" so it reads uniformly.',
    'I need the typography to look uniform before I print this. First set the font to "{font}" throughout the body paragraphs, then apply the same font to every cell in the document tables.',
    # multi-step medium (first/then sequence keywords)
    'First select all of the document text including any tables, then change the font to "{font}" so every paragraph and cell renders in the same typeface across the whole file.',
    # long narrative pad
    'I\'m getting this manuscript ready for the journal submission portal and the editorial team has a strict house style — the body font must be "{font}" everywhere; please update every paragraph in the document plus any table cells so the entire file uses "{font}" as the single body face for the final PDF.',
    # long multi-step (then/once sequence keywords)
    'For a report I\'m finalising before tomorrow\'s deadline, the team standardised on "{font}" as the body font; first apply "{font}" to every body paragraph, then sweep through the tables and set each cell to "{font}" too, so once the document goes out for sign-off the typography reads consistently.',
    # extra long narrative pad
    'I need the typography in this document to look uniform before printing — please set the font to "{font}" throughout, covering every body paragraph plus the text in the tables, so the printed copies read with the same typeface from cover to back.',
    # long multi-step variant 2 (after sequence keyword)
    'For a polished version of this document, please update the entire body to use "{font}" as the single typeface — first walk through every paragraph and apply "{font}", after that handle every cell in every table the same way so the file reads consistently.',
]

_FIND_REPLACE_VARIANTS = [
    # short
    'Replace the most frequent 4+ letter word with "{new}".',
    'Swap the top long word for "{new}" everywhere.',
    'Substitute the most common 4+ letter word with "{new}".',
    'Find the most-used 4+ letter word and replace it with "{new}".',
    # medium narrative pads
    'I\'m revising this document and the most-used long word (4+ letters) has become repetitive — find that word and replace every occurrence with "{new}" to vary the prose.',
    'While proofreading I noticed one 4+ letter word dominates the document. Could you locate the most frequent long word and swap every instance for "{new}" so it reads less repetitively?',
    'I want to reduce repetition in this writeup. First find the single most frequent word (at least four letters), then substitute "{new}" in every place that word appears across the document.',
    # multi-step medium (first/then sequence keywords)
    'First scan the document to identify the single word of four or more letters that appears most often, then replace every instance of that word with "{new}" across the whole file.',
    # long narrative pad
    'I\'m polishing this writeup before sending it to my editor and her main feedback was that one long word keeps cropping up on every page — please find the most frequent 4+ letter word in the document and replace every occurrence with "{new}" so the prose finally reads with more variety.',
    # long multi-step (first/then/once sequence keywords)
    'For a polish pass on this draft, first identify the most frequently occurring 4+ letter word, then replace it everywhere with "{new}"; once that swap is complete the prose should read with noticeably more variety, which is exactly what my reviewer flagged in their last round.',
    # extra long narrative pad
    'I want to reduce repetition in this writeup before tomorrow\'s editorial deadline — please find the single most frequent word of at least four letters across the document and substitute "{new}" everywhere it appears, so the prose reads with more variety on every printed page.',
    # long multi-step variant 2 (next/then sequence keyword)
    'The reviewers flagged that one long word keeps repeating throughout this draft — first identify the most frequently occurring 4+ letter word across the file, next replace every occurrence of that word with "{new}", after which the document should read with the variety the editor wants.',
]

_SPACING_VARIANTS = [
    # short
    "Set {spacing} line spacing on the {ord} paragraph.",
    "Change the {ord} paragraph to {spacing} line spacing.",
    "Apply {spacing} line spacing to the {ord} paragraph.",
    "Use {spacing} line spacing for the {ord} paragraph.",
    # medium narrative pads
    "I'm formatting this document for review and need extra room between lines in the {ord} paragraph — please set the line spacing of the {ord} paragraph to {spacing}.",
    "While editing I realised the {ord} paragraph is dense and hard to annotate. Change the line spacing of the {ord} paragraph to {spacing} so reviewers can mark it up.",
    "The {ord} paragraph contains a long quotation that I want easier to read; first locate that paragraph, then update its line spacing to {spacing} so the quotation breathes more on the page.",
    # multi-step medium (first/then sequence keywords)
    "First find the {ord} paragraph in the document, then change its line spacing to {spacing} so it has enough vertical room for inline annotations during the next review pass.",
    # long narrative pad
    "I'm putting together a printed handout for tomorrow's offsite and the {ord} paragraph carries a long block quote attendees will mark up by hand during the workshop — please set the line spacing of the {ord} paragraph to {spacing} so there's enough vertical room for annotations on the printed copies.",
    # long multi-step (after/then sequence keywords)
    "For a printed reference sheet I'm preparing, the {ord} paragraph needs much looser spacing than the rest of the body — first locate the {ord} paragraph, then apply {spacing} line spacing to it; after that change the printed text should breathe enough that reviewers can comfortably annotate between lines on the hard copies.",
    # extra long narrative pad
    "I'd like the {ord} paragraph to use {spacing} line spacing because it contains a long block of quoted material I want easier to read on the printed page — please update the {ord} paragraph's line spacing to {spacing} so the handout has breathing room for inline notes.",
    # long multi-step variant 2 (then/once sequence keywords)
    "The {ord} paragraph is dense and hard for reviewers to annotate during a markup pass — first navigate to the {ord} paragraph in the body, then change its line spacing to {spacing}, so once printed reviewers have room to scribble inline notes between lines.",
]

# Spacing variants that assert the target paragraph *is* a long quotation /
# block quote (matched by phrase, not index, so reordering the list cannot
# silently mis-tag them). These are only emitted for bases whose audited body
# paragraphs have enough long-quote candidates for the chosen ordinal. For these
# the oracle targets the ord_pos-th paragraph whose word count clears
# _LONG_QUOTE_MIN_WORDS rather than body_idxs[ord_pos]. Every OTHER spacing
# variant keeps the ordinal-body target.
# (§C #19: 66399b0d "first paragraph contains a long quotation" — body[0] is a
# 9-word line, but the two real quotation blocks are body[1]/body[2], wc 88/86.)
_SPACING_LONG_QUOTE_VARIANTS = frozenset(
    v for v in _SPACING_VARIANTS
    if "long quotation" in v
    or "long block quote" in v
    or "long block of quoted material" in v
)
# 66399b0d body word counts: [9, 88, 86, 25, 23]. The two quotation blocks are
# 88/86 words; the next-longest body line is 25. A 40-word cutoff sits cleanly
# in the (25, 86] gap, separating genuine long quotations from short body lines.
_LONG_QUOTE_MIN_WORDS = 40
_SPACING_LONG_QUOTE_LIMITS = {
    # body word counts: [9, 88, 86, 25, 23]
    "66399b0d": 2,
}

_UPPERCASE_VARIANTS = [
    # short
    "Make the {ord} paragraph all uppercase.",
    "Convert the {ord} paragraph to uppercase letters.",
    "Uppercase every word in the {ord} paragraph.",
    "Change the {ord} paragraph to all caps.",
    # medium narrative pads
    "I'm preparing a notice and the {ord} paragraph is the warning line that absolutely must shout — convert the entire {ord} paragraph to uppercase letters so it reads as a banner.",
    "While editing this draft I realised the {ord} paragraph is a header that should be visually loud. Please change all text in the {ord} paragraph to uppercase across every word.",
    "I want the {ord} paragraph to function as an attention-grabbing label in this writeup; first locate it in the document, then uppercase every letter so it reads as a banner above the body prose.",
    # multi-step medium (first/then sequence keywords)
    "First locate the {ord} paragraph in the document, then convert every character in that paragraph to uppercase so the text reads as a loud all-caps banner across the whole line.",
    # long narrative pad
    "I'm finalising the safety-notice section of this document for the warehouse team and the {ord} paragraph contains the warning line that absolutely cannot be missed when staff scan the printout on the wall — please convert the entire {ord} paragraph to uppercase so it reads as an unmissable banner above the rest of the prose.",
    # long multi-step (first/then/once sequence keywords)
    "For a flyer I'm finalising for an event, the {ord} paragraph carries the headline copy and needs to dominate the page visually; first locate the {ord} paragraph, then uppercase every character in it, so once the file is printed at A3 scale the headline reads as a strong all-caps title across the whole top margin.",
    # extra long narrative pad
    "I want the {ord} paragraph in this writeup to function as an attention-grabbing label for the warehouse team reading from a distance — please uppercase every letter in the {ord} paragraph so the header stays readable from across the room when staff scan the page on the wall.",
    # long multi-step variant 2 (then/after sequence keywords)
    "I decided the {ord} paragraph should function as a loud header line that cannot be skimmed past during the next review — first locate the {ord} paragraph, then uppercase every character in it; after that the line will visually shout above the rest of the body prose.",
]

_STRIKETHROUGH_VARIANTS = [
    # short
    "Strike through the {ord} paragraph.",
    "Apply strikethrough to the {ord} paragraph.",
    "Cross out the {ord} paragraph entirely.",
    "Mark the {ord} paragraph with strikethrough.",
    # medium narrative pads
    "I'm peer-reviewing this draft and the {ord} paragraph is content I want flagged as redundant — apply strikethrough formatting to every word in the {ord} paragraph for me.",
    "While editing I decided the {ord} paragraph should be marked for removal but kept visible for now. Please add strikethrough across the entire {ord} paragraph so it reads as crossed out.",
    "I'm tracking edits manually and the {ord} paragraph needs to look struck through; first locate that paragraph, then apply strikethrough to every word so it reads as crossed out for the next review.",
    # multi-step medium (first/then sequence keywords)
    "First find the {ord} paragraph in the document, then apply strikethrough formatting to every word in that paragraph so it visibly reads as pending deletion during the next review.",
    # long narrative pad
    "I'm peer-reviewing this draft for a colleague and the {ord} paragraph is a stretch of content I want flagged as redundant without actually deleting it yet, because she may want to keep some of the phrasing — please apply strikethrough to every word of the {ord} paragraph so it reads as crossed-out content during our next editing meeting.",
    # long multi-step (first/then/after sequence keywords)
    "For a revision pass on this document, the {ord} paragraph is text I want struck through to mark it as pending deletion; first locate the {ord} paragraph in the file, then apply strikethrough formatting to every word inside it, after which the visual cue should make it obvious to my reviewer the paragraph is a deletion candidate.",
    # extra long narrative pad
    "While editing this draft I decided the {ord} paragraph should be marked for removal but kept visible for now so I can discuss the cut with my co-author later this week — please add strikethrough across every word in the {ord} paragraph so it reads as crossed out for our next call.",
    # long multi-step variant 2 (then/once sequence keywords)
    "I'm tracking edits manually before sending the file to my co-author — the {ord} paragraph needs to look struck through so she sees it as a deletion candidate; first navigate to the {ord} paragraph, then apply strikethrough across every word so she sees the cue when she opens the file.",
]

_LOWERCASE_VARIANTS = [
    # short
    "Lowercase the {ord} paragraph.",
    "Make the {ord} paragraph all lowercase.",
    "Convert every letter in the {ord} paragraph to lowercase.",
    "Change the {ord} paragraph to lowercase.",
    # medium narrative pads
    "I'm cleaning up this document and the {ord} paragraph contains shouting all-caps text that should match the rest of the prose — convert the {ord} paragraph to lowercase letters.",
    "While normalising the formatting in this draft I realised the {ord} paragraph is in upper case and looks out of place. Please change all text in the {ord} paragraph to lowercase.",
    "The surrounding sections are already lowercase so the {ord} paragraph stands out — first locate it, then convert every character in the {ord} paragraph to lowercase so the document reads with uniform capitalisation.",
    # multi-step medium (first/then sequence keywords)
    "First locate the {ord} paragraph in the document, then convert every character in that paragraph to lowercase so it stops shouting and matches the calmer tone of the surrounding sections.",
    # long narrative pad
    "I'm doing a stylistic cleanup pass on this document before sending it to my supervisor and the {ord} paragraph is the only block left in shouting upper-case while the rest of the prose has already been normalised — please convert every character in the {ord} paragraph to lowercase so the entire writeup finally reads with a single consistent capitalisation throughout.",
    # long multi-step (first/then/once sequence keywords)
    "For a stylistic pass on this writeup, the {ord} paragraph reads too loud at its current capitalisation; first find the {ord} paragraph in the document, then lowercase every character inside it, so once the change is in place the paragraph blends smoothly with the calm tone of the surrounding text and the whole section reads with consistent capitalisation.",
    # extra long narrative pad
    "I want the {ord} paragraph to read in lowercase only because the surrounding sections have been normalised that way and the inconsistency is distracting on the printed copy — please change all the text in the {ord} paragraph to lowercase so the document reads with consistent capitalisation.",
    # long multi-step variant 2 (then/after sequence keywords)
    "Quick consistency fix before this draft goes to print: the {ord} paragraph is in mixed or upper case but should be lowercase to match the rest — first scan to the {ord} paragraph, then lowercase every character inside it; after that the paragraph blends with the surrounding tone.",
]

_ITALIC_VARIANTS = [
    # short
    "Italicise the {ord} paragraph.",
    "Make the {ord} paragraph italic throughout.",
    "Apply italic formatting to the {ord} paragraph.",
    "Set the {ord} paragraph to italic.",
    # medium narrative pads
    "I'm finalising this manuscript and the {ord} paragraph is a quotation that should be visually distinct — please italicise the entire {ord} paragraph so readers see it as quoted text.",
    "While reviewing this draft I noticed the {ord} paragraph carries an aside that needs softer emphasis. Apply italic formatting to every word of the {ord} paragraph for me.",
    "The {ord} paragraph is an inline excerpt I want set off from the main body; first locate it, then italicise every word in the {ord} paragraph so the excerpt reads as quoted material.",
    # multi-step medium (first/then sequence keywords)
    "First locate the {ord} paragraph in the document, then italicise every word in that paragraph so the quotation is visually set off from the surrounding prose for readers.",
    # long narrative pad
    "I'm finalising the manuscript for a journal submission and the {ord} paragraph is a verbatim block quotation that the journal's house style requires to be set in italics across every word — please apply italic formatting to the entire {ord} paragraph so the typesetter doesn't kick the file back during the desk-review stage of the submission.",
    # long multi-step (first/then/after sequence keywords)
    "For a report I'm putting together, the {ord} paragraph contains a long inline citation that house style says must be italic; first locate the {ord} paragraph, then apply italic formatting to every word inside it, after that the citation should read as a clearly set-off block of quoted material rather than blending into the regular body prose.",
    # extra long narrative pad
    "I'd like the {ord} paragraph to display in italics because it's a long inline excerpt from a third-party source I want set off cleanly from the main body — please format the {ord} paragraph in italics across every word so the excerpt visually reads as quoted material.",
    # long multi-step variant 2 (then/once sequence keywords)
    "The {ord} paragraph carries an aside that needs softer emphasis than the surrounding body prose — first navigate to the {ord} paragraph, then italicise every word inside it, so once readers reach that paragraph the styling cues them to read it as an aside.",
]

# NOTE: Georgia and Calibri are NOT installed in the LO VM (per synth validation
# audit at writer/find_default_font_v2_0001). Restrict to fonts verified
# present in the VM build so font ops produce visible changes (not silent
# fallback to default → trivial_pass risk).
_FONTS = ["Arial", "Times New Roman", "Courier New", "Verdana"]
_REPLACE_WORDS = ["crucial", "framework", "overview", "strategy", "initiative", "essential"]
_SPACINGS = [("double", 2.0), ("1.5 lines", 1.5)]


def _build_instruction(phrase: str, rng: random.Random) -> str:
    """Pass-through. Each variant pool already encodes its own narrative lead
    (mixing imperative, polite, and contextual openings at the desired ratio),
    so no extra polite-prefix augmentation is applied here. Kept as a no-op
    so callers keep a single integration point — easy to re-enable styling
    layers later without touching the op makers.
    """
    return phrase


# ---------------------------------------------------------------------------
# Op makers — each returns (instruction, py_code)
# ---------------------------------------------------------------------------

def _resolve_idx_and_ord(rng, n_paras, body_idxs):
    """Pick (target_idx, ord_word) for ordinal-paragraph ops.

    When ``body_idxs`` is provided (ordinal-safe base), pick ``ord_pos`` in
    [0, min(4, len(body_idxs)-1)] and use ``body_idxs[ord_pos]`` as the
    python-docx index, with ``_ORDINALS[ord_pos]`` as the agent-facing
    ordinal. This guarantees agent's visual counting (which skips titles
    and headings) matches the oracle's modification target.

    When ``body_idxs`` is None (unsafe base), legacy behavior — but the
    dispatcher should NOT call ordinal ops on unsafe bases, so this branch
    is only reached for back-compat / direct test invocations.
    """
    if body_idxs:
        ord_pos = rng.randint(0, min(4, len(body_idxs) - 1))
        return body_idxs[ord_pos], _ORDINALS[ord_pos]
    idx = rng.randint(0, min(4, n_paras - 1))
    return idx, _ORDINALS.get(idx, f"paragraph {idx+1}")


def _make_bold(rng: random.Random, n_paras: int, body_idxs=None) -> tuple[str, str]:
    idx, ord_word = _resolve_idx_and_ord(rng, n_paras, body_idxs)
    phrase = rng.choice(_BOLD_VARIANTS).format(ord=ord_word)
    instruction = _build_instruction(phrase, rng)
    py_code = textwrap.dedent(f"""\
        import sys
        from docx import Document

        def _style_bold(style):
            while style is not None:
                if style.font.bold is True:
                    return True
                if style.font.bold is False:
                    return False
                style = style.base_style
            return False

        def _effectively_bold(p):
            # Skip paragraphs where every run is already bold — either via
            # explicit font.bold=True OR inherited from the paragraph style.
            # Using is True alone misses style-inherited bold (font.bold=None):
            # LO normalization strips redundant explicit bold, making expected==original.
            # Also skip paragraphs with no runs: can't apply bold via python-docx.
            runs = [r for r in p.runs if r.text]
            if not runs:
                return False
            sb = _style_bold(p.style)
            return all(
                (r.font.bold is True) or (r.font.bold is None and sb)
                for r in runs
            )

        doc = Document('{{file_path}}')
        if doc.paragraphs:
            target = min({idx}, len(doc.paragraphs) - 1)
            while target < len(doc.paragraphs) - 1 and (
                not doc.paragraphs[target].text.strip()
                or not doc.paragraphs[target].runs
                or _effectively_bold(doc.paragraphs[target])
            ):
                target += 1
            p = doc.paragraphs[target]
            if not p.text.strip() or not p.runs or _effectively_bold(p):
                sys.exit(42)
            for r in p.runs:
                r.font.bold = True
        doc.save('{{expected_path}}')
    """)
    return instruction, py_code


def _make_font(
    rng: random.Random, _n_paras: int,
    exclude_fonts: set[str] | None = None,
) -> tuple[str, str]:
    eligible = [f for f in _FONTS if f not in (exclude_fonts or set())]
    if not eligible:
        raise ValueError("no eligible fonts")
    font = rng.choice(eligible)
    phrase = rng.choice(_FONT_VARIANTS).format(font=font)
    instruction = _build_instruction(phrase, rng)
    # Apply the font change to BOTH paragraphs and table cells. Earlier
    # version only iterated `doc.paragraphs`, leaving table-cell runs in
    # the original font. A correct agent (Ctrl+A → change font) covers
    # tables → expected mismatch on table runs, eval=0.
    py_code = textwrap.dedent(f"""\
        import sys
        from docx import Document

        def _style_font(style):
            while style is not None:
                if style.font.name:
                    return style.font.name
                style = style.base_style
            return None

        def _effective_font(run, para, normal_font):
            # run explicit -> run char-style -> paragraph style -> Normal default
            if run.font.name:
                return run.font.name
            cs = getattr(run, 'style', None)
            if cs is not None and _style_font(cs):
                return _style_font(cs)
            if _style_font(para.style):
                return _style_font(para.style)
            return normal_font

        doc = Document('{{file_path}}')
        try:
            _normal = doc.styles['Normal'].font.name
        except Exception:
            _normal = None
        _runs = []
        for _p in doc.paragraphs:
            for _r in _p.runs:
                if _r.text:
                    _runs.append((_r, _p))
        for _tbl in doc.tables:
            for _row in _tbl.rows:
                for _cell in _row.cells:
                    for _p in _cell.paragraphs:
                        for _r in _p.runs:
                            if _r.text:
                                _runs.append((_r, _p))
        # No-op guard (#154): if every run already renders as the target font
        # (effective: run -> char-style -> para-style -> Normal default), the
        # perturb would yield expected==source -> trivial pass (eval=1.0 before
        # the oracle). exit(42) so the framework drops/rerolls the task. General
        # self-correction for any mis-identified exclude_fonts registry entry,
        # across every source and split.
        if _runs and all(_effective_font(_r, _p, _normal) == '{font}' for _r, _p in _runs):
            sys.exit(42)
        def _apply(paragraphs):
            for p in paragraphs:
                for r in p.runs:
                    r.font.name = '{font}'
        _apply(doc.paragraphs)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    _apply(cell.paragraphs)
        doc.save('{{expected_path}}')
    """)
    return instruction, py_code


def _make_find_replace(rng: random.Random, _n_paras: int) -> tuple[str, str]:
    new = rng.choice(_REPLACE_WORDS)
    phrase = rng.choice(_FIND_REPLACE_VARIANTS).format(new=new)
    instruction = _build_instruction(phrase, rng)
    # NOTE: prior version saved expected unchanged when no
    # 4+letter words matched (e.g. HK_train_record.docx with only station
    # codes E201/SHL1/FAL2 — each letter+digit so regex matches none).
    # Eval compare_docx_strict(out, expected) then returned 1 for ANY
    # unchanged file → vacuous pass. Fix: progressively relax the regex
    # (4+ letters → 3+ → 2+ → any alphanumeric token) so SOME word is
    # always found and the resulting expected file actually differs.
    py_code = textwrap.dedent(f"""\
        import re, sys
        from collections import Counter
        from docx import Document
        doc = Document('{{file_path}}')
        all_text = " ".join(p.text for p in doc.paragraphs)
        # Try strictest regex first, fall back progressively until we find one.
        old_word = None
        for pattern in [r'[A-Za-z]{{{{4,}}}}', r'[A-Za-z]{{{{3,}}}}',
                        r'[A-Za-z]{{{{2,}}}}', r'[A-Za-z0-9]+']:
            words = [w.lower() for w in re.findall(pattern, all_text)]
            counts = Counter(words)
            if counts:
                old_word = counts.most_common(1)[0][0]
                break
        if old_word is None:
            # Document is empty or contains only whitespace/punctuation.
            # Mark the perturb un-runnable; setup fails fast rather than
            # vacuously saving an unchanged file.
            sys.exit(42)
        # Per-run substitution (not run-collapse). Earlier code collapsed
        # every modified paragraph into a single run, which made successful
        # GUI Find&Replace fail compare_docx_strict on character formatting.
        for p in doc.paragraphs:
            if old_word not in p.text.lower():
                continue
            for r in p.runs:
                if old_word in r.text.lower():
                    r.text = re.sub(re.escape(old_word), '{new}', r.text, flags=re.IGNORECASE)
        doc.save('{{expected_path}}')
    """)
    return instruction, py_code


def _make_spacing(
    rng: random.Random, n_paras: int,
    exclude_idx: set[int] | None = None,
    exclude_vals: set[float] | None = None,
    body_idxs=None,
    long_quote_limit: int | None = None,
) -> tuple[str, str]:
    """Generate a paragraph-line-spacing perturb instruction + oracle py_code.

    ``body_idxs`` is required (must be non-empty). The dispatcher only routes
    spacing ops to bases in ``_ORDINAL_SAFE_BASES``, so this function is never
    called without one. ``exclude_idx`` is interpreted as ord_pos values to
    exclude (used by T1 to avoid duplicating the eval target — e.g.,
    ``exclude_idx={0,1}`` for 0810415c whose eval task targets body[0,1]).
    ``long_quote_limit`` is the audited count of body paragraphs that can
    honestly satisfy the long-quotation wording.
    """
    if not body_idxs:
        raise ValueError("_make_spacing requires body_idxs (ordinal-safe bases only)")
    eligible_spacings = [(n, v) for n, v in _SPACINGS if v not in (exclude_vals or set())]
    if not eligible_spacings:
        raise ValueError("no eligible spacing params")
    eligible_pos = [i for i in range(min(5, len(body_idxs))) if i not in (exclude_idx or set())]
    if not eligible_pos:
        raise ValueError("no eligible spacing ord_pos")
    ord_pos = rng.choice(eligible_pos)
    idx = body_idxs[ord_pos]
    ord_word = _ORDINALS[ord_pos]
    name, val = rng.choice(eligible_spacings)
    variant = rng.choice(_SPACING_VARIANTS)          # single rng draw, unchanged
    if variant in _SPACING_LONG_QUOTE_VARIANTS and (
        long_quote_limit is None or ord_pos >= long_quote_limit
    ):
        variant = rng.choice([
            v for v in _SPACING_VARIANTS
            if v not in _SPACING_LONG_QUOTE_VARIANTS
        ])
    phrase = variant.format(ord=ord_word, spacing=name)
    instruction = _build_instruction(phrase, rng)

    # Shared oracle preamble: imports + style-chain spacing resolver.
    head = textwrap.dedent("""\
        import sys
        from docx import Document

        def _effective_spacing(paragraph):
            # Resolve line_spacing through the style chain so that paragraphs
            # whose spacing comes from the style (pf.line_spacing=None) are
            # compared correctly against the target value.
            pf = paragraph.paragraph_format
            if pf.line_spacing is not None:
                return pf.line_spacing
            style = paragraph.style
            while style is not None:
                spf = style.paragraph_format
                if spf.line_spacing is not None:
                    return spf.line_spacing
                style = style.base_style
            return None

        doc = Document('{file_path}')
        if not doc.paragraphs:
            doc.save('{expected_path}')
            sys.exit(0)
    """)
    # SCOPED: only the "long quotation" variants resolve the target by content
    # (ord_pos-th paragraph whose word count >= _LONG_QUOTE_MIN_WORDS). If the
    # base has no such paragraph for this ordinal, sys.exit(42) marks the perturb
    # un-runnable (fails cleanly, no vacuous pass). ALL other variants keep the
    # existing ordinal-body target UNCHANGED.
    if variant in _SPACING_LONG_QUOTE_VARIANTS:
        resolve = textwrap.dedent(f"""\
            long_idxs = [i for i, _p in enumerate(doc.paragraphs)
                         if len(_p.text.split()) >= {_LONG_QUOTE_MIN_WORDS}]
            if len(long_idxs) <= {ord_pos}:
                sys.exit(42)  # no long paragraph for this ordinal — un-runnable
            target = long_idxs[{ord_pos}]
        """)
    else:
        resolve = f"target = min({idx}, len(doc.paragraphs) - 1)\n"
    tail = textwrap.dedent(f"""\
        p = doc.paragraphs[target]
        if _effective_spacing(p) == {val}:
            sys.exit(42)  # already at target spacing — perturb would be a no-op
        p.paragraph_format.line_spacing = {val}
        doc.save('{{expected_path}}')
    """)
    py_code = head + resolve + tail
    return instruction, py_code


def _make_uppercase(rng: random.Random, n_paras: int, body_idxs=None) -> tuple[str, str]:
    idx, ord_word = _resolve_idx_and_ord(rng, n_paras, body_idxs)
    phrase = rng.choice(_UPPERCASE_VARIANTS).format(ord=ord_word)
    instruction = _build_instruction(phrase, rng)
    # Pick the first non-empty paragraph at or after idx — earlier code
    # would no-op on empty separator paragraphs (vacuous score=1.0 because
    # uppercase("") == "" and agent did nothing).
    # Safety net: if every paragraph from idx onward is empty/no-runs OR the
    # target text is already uppercase, sys.exit(42) so the variant fails to
    # produce expected_path (oracle then fails cleanly rather than triggering
    # a vacuous trivial_pass). The dispatcher restricts idx to the non-empty
    # paragraph range via n_paras_safe to avoid this in normal cases.
    py_code = textwrap.dedent(f"""\
        import sys
        from docx import Document
        doc = Document('{{file_path}}')
        if doc.paragraphs:
            target = min({idx}, len(doc.paragraphs) - 1)
            # advance past empty / runless / already-uppercase paragraphs
            def _is_skippable_upper(p):
                if not p.text.strip() or not p.runs:
                    return True
                txt = "".join(r.text for r in p.runs)
                return txt == txt.upper()  # already all-upper → no-op
            while target < len(doc.paragraphs) - 1 and _is_skippable_upper(doc.paragraphs[target]):
                target += 1
            p = doc.paragraphs[target]
            if _is_skippable_upper(p):
                sys.exit(42)
            for r in p.runs:
                r.text = r.text.upper()
        else:
            sys.exit(42)
        doc.save('{{expected_path}}')
    """)
    return instruction, py_code


def _make_strikethrough(rng: random.Random, n_paras: int, body_idxs=None) -> tuple[str, str]:
    idx, ord_word = _resolve_idx_and_ord(rng, n_paras, body_idxs)
    phrase = rng.choice(_STRIKETHROUGH_VARIANTS).format(ord=ord_word)
    instruction = _build_instruction(phrase, rng)
    py_code = textwrap.dedent(f"""\
        import sys
        from docx import Document

        def _effectively_struck(p):
            runs = [r for r in p.runs if r.text]
            if not runs:
                return False
            return all(r.font.strike is True for r in runs)

        doc = Document('{{file_path}}')
        if doc.paragraphs:
            target = min({idx}, len(doc.paragraphs) - 1)
            while target < len(doc.paragraphs) - 1 and (
                not doc.paragraphs[target].text.strip()
                or not doc.paragraphs[target].runs
                or _effectively_struck(doc.paragraphs[target])
            ):
                target += 1
            p = doc.paragraphs[target]
            if not p.text.strip() or not p.runs or _effectively_struck(p):
                sys.exit(42)
            for r in p.runs:
                r.font.strike = True
        doc.save('{{expected_path}}')
    """)
    return instruction, py_code


def _make_lowercase(rng: random.Random, n_paras: int, body_idxs=None) -> tuple[str, str]:
    idx, ord_word = _resolve_idx_and_ord(rng, n_paras, body_idxs)
    phrase = rng.choice(_LOWERCASE_VARIANTS).format(ord=ord_word)
    instruction = _build_instruction(phrase, rng)
    # Safety net mirrors _make_uppercase — see comment there.
    py_code = textwrap.dedent(f"""\
        import sys
        from docx import Document
        doc = Document('{{file_path}}')
        if doc.paragraphs:
            target = min({idx}, len(doc.paragraphs) - 1)
            def _is_skippable_lower(p):
                if not p.text.strip() or not p.runs:
                    return True
                txt = "".join(r.text for r in p.runs)
                return txt == txt.lower()  # already all-lower → no-op
            while target < len(doc.paragraphs) - 1 and _is_skippable_lower(doc.paragraphs[target]):
                target += 1
            p = doc.paragraphs[target]
            if _is_skippable_lower(p):
                sys.exit(42)
            for r in p.runs:
                r.text = r.text.lower()
        else:
            sys.exit(42)
        doc.save('{{expected_path}}')
    """)
    return instruction, py_code


def _make_italic(rng: random.Random, n_paras: int, body_idxs=None) -> tuple[str, str]:
    idx, ord_word = _resolve_idx_and_ord(rng, n_paras, body_idxs)
    phrase = rng.choice(_ITALIC_VARIANTS).format(ord=ord_word)
    instruction = _build_instruction(phrase, rng)
    py_code = textwrap.dedent(f"""\
        import sys
        from docx import Document

        def _style_italic(style):
            while style is not None:
                if style.font.italic is True:
                    return True
                if style.font.italic is False:
                    return False
                style = style.base_style
            return False

        def _effectively_italic(p):
            runs = [r for r in p.runs if r.text]
            if not runs:
                return False
            si = _style_italic(p.style)
            return all(
                (r.font.italic is True) or (r.font.italic is None and si)
                for r in runs
            )

        doc = Document('{{file_path}}')
        if doc.paragraphs:
            target = min({idx}, len(doc.paragraphs) - 1)
            while target < len(doc.paragraphs) - 1 and (
                not doc.paragraphs[target].text.strip()
                or not doc.paragraphs[target].runs
                or _effectively_italic(doc.paragraphs[target])
            ):
                target += 1
            p = doc.paragraphs[target]
            if not p.text.strip() or not p.runs or _effectively_italic(p):
                sys.exit(42)
            for r in p.runs:
                r.font.italic = True
        doc.save('{{expected_path}}')
    """)
    return instruction, py_code


_OP_POOL = [
    _make_bold,          # 0
    _make_font,          # 1
    _make_find_replace,  # 2
    _make_spacing,       # 3
    _make_uppercase,     # 4
    _make_strikethrough, # 5
    _make_lowercase,     # 6
    _make_italic,        # 7
]


# ---------------------------------------------------------------------------
# TYPE_3 archetype builders — target evaluators that compare_docx_strict does
# NOT cover. Each archetype returns a fully-formed (instruction, oracle_actions,
# evaluator) triple, not just (instruction, py_code), because:
#   - check_tabstops: result must be a vanilla docx file pair (no LO normalize)
#   - compare_subscript_contains: needs subscript runs preserved through LO
#     round-trip, so the standard _build_oracle's double-normalize is fine but
#     the evaluator is a 2-func compound (compare_docx_files + subscript).
#   - compare_pdfs (×4): result/expected are .pdf at 4 different paths; the
#     oracle script must export the docx to all 4 destinations. No docx
#     compare at all.
# ---------------------------------------------------------------------------

# 1) check_tabstops archetype — eval base 0a0faba3 ----------------------------
# The eval task asks the agent to align "first three words left, rest right"
# using a tabstop on a right margin. We perturb by varying:
#   (a) tabstop alignment   — RIGHT (eval) / CENTER / LEFT-fixed-position
#   (b) split point         — first-N-words → tab → rest, with N ∈ {2, 3, 4}
# The oracle py rewrites every non-empty paragraph to put the split tab + a
# trailing tabstop at the right margin (or a fixed position for LEFT/CENTER).

# #155 #20: the gold guard tabs only paragraphs with MORE than {n} words
# (len(words) >= n+1) — an exactly-{n}-word paragraph has nothing after word {n}.
# The instructions therefore say "more than {n} words", so a correct agent that
# leaves the {n}-word paragraph untabbed matches the gold (was: "every paragraph",
# which mis-implied tabbing the {n}-word one too -> FN).
_TABSTOPS_VARIANTS = [
    # short
    "In each paragraph with more than {n} words, add a tab after word {n} and a right tabstop at the margin.",
    "Tab-split every paragraph that has more than {n} words after word {n}, right-aligned at the margin.",
    # medium narrative pads
    "I'm preparing this handout and want each paragraph with more than {n} words reformatted so the first {n} words sit at the left and the remainder hangs at the right margin via a tab stop — please apply that tabstop layout throughout the document.",
    "While polishing this draft I realised the alignment looks ragged. Could you split every paragraph that has more than {n} words at the {n}-th word with a tab and add a right-aligned tab stop at the page margin so the second half snaps right?",
    "Quick layout pass: take every paragraph with more than {n} words and insert a tab after the first {n} words, then attach a right-aligned tab stop at the right margin so the trailing words flush right.",
    "Please reformat each paragraph that has more than {n} words by separating the leading {n} words from the remainder with a tab and registering a right-aligned tab stop at the right margin so the layout reads as two columns.",
    # long narrative pad
    "I'm finalising a printed reference sheet for the team and the formatting needs to read as two columns on every line — please walk through every paragraph with more than {n} words in the document, keep the first {n} words anchored on the left, and use a right-aligned tab stop at the page margin to push the remainder of each such paragraph flush against the right edge.",
    # long multi-step (first/then/once sequence keywords)
    "For a printed handout I'm preparing, the layout needs the first {n} words of each paragraph with more than {n} words at the left and the rest hanging on the right; first walk through every such paragraph and insert a tab after the {n}-th word, then attach a right-aligned tab stop at the right margin, so once the document prints the body text reads as a clean two-column layout end to end.",
    # extra long narrative pad
    "For a printed reference sheet I'm finalising tonight, I'd like the first {n} words of each paragraph with more than {n} words kept anchored on the left while the remainder is pushed flush to the right via a right-aligned tab stop at the paragraph margin, so the printed copy reads as a clean two-column layout that the team can scan during tomorrow's planning meeting.",
]


def _make_tabstops(rng: random.Random) -> tuple[str, str]:
    """Return (instruction, py_code) for the tabstops archetype.

    Oracle rewrites every non-empty paragraph: collapse runs into one, insert
    a tab after the n-th word, register a right-aligned tab stop at the right
    margin (paragraph_width). check_tabstops compares per-paragraph tab stop
    positions/alignments against the expected file produced here.
    """
    n = rng.choice([2, 3, 4])
    phrase = rng.choice(_TABSTOPS_VARIANTS).format(n=n)
    instruction = _build_instruction(phrase, rng)
    py_code = textwrap.dedent(f"""\
        import sys
        from docx import Document
        from docx.enum.text import WD_TAB_ALIGNMENT

        doc = Document('{{file_path}}')
        section = doc.sections[0]
        pw = section.page_width - section.left_margin - section.right_margin

        applied = 0
        for p in doc.paragraphs:
            if not p.text.strip():
                continue
            words = p.text.split()
            if len(words) < {n} + 1:
                continue
            head = ' '.join(words[:{n}])
            tail = ' '.join(words[{n}:])
            # collapse all runs into the first run
            if not p.runs:
                continue
            for r in p.runs:
                r.text = ''
            p.runs[0].text = head + '\\t' + tail
            p.paragraph_format.tab_stops.add_tab_stop(pw, WD_TAB_ALIGNMENT.RIGHT)
            applied += 1
        if applied == 0:
            sys.exit(42)
        doc.save('{{expected_path}}')
    """)
    return instruction, py_code


# 2) compare_subscript_contains archetype — eval base 0b17a146 ----------------
# The eval task asks the agent to make the "2" in H2O a subscript. We perturb
# by picking a different paragraph and a different short token (digit/letter)
# to subscript. Evaluator is the same compound as eval (compare_docx_files +
# compare_subscript_contains), so any paragraph that has at least one subscript
# run AND the same overall text passes.

_SUBSCRIPT_VARIANTS = [
    # short
    "Subscript the first digit in the {ord} non-empty paragraph.",
    "Make the first digit in the {ord} paragraph a subscript.",
    # medium narrative pads
    "I'm typesetting a chemistry note and the {ord} non-empty paragraph contains a number that should display as a subscript — please convert the first standalone digit in that paragraph to subscript formatting.",
    "While editing this writeup I realised the {ord} non-empty paragraph has a digit that needs to read as a chemical subscript. Could you mark the first digit you find in that paragraph as subscript so it sits below the baseline?",
    "Quick formatting fix: in the {ord} non-empty paragraph, the first lone digit should render as a subscript instead of regular text — turn that digit into a subscript.",
    "Please format the {ord} non-empty paragraph so its first standalone digit renders as a subscript, leaving the rest of the paragraph and the document untouched.",
    # multi-step medium (first/then sequence keywords)
    "First find the {ord} non-empty paragraph in the document, then locate the first standalone digit inside it and convert that digit to subscript formatting so it sits below the baseline.",
    # long narrative pad
    "I'm typesetting the chemistry section of a handout for tomorrow's lab tutorial and the {ord} non-empty paragraph contains a chemical formula whose digit needs to render as a true subscript instead of inline text — please find the first standalone digit in that paragraph and apply subscript formatting to it so the formula reads correctly when the students see the printout in class.",
    # long multi-step (first/then/once sequence keywords)
    "For a science handout I'm preparing, the {ord} non-empty paragraph carries a chemical name whose digit should sit below the baseline; first scan the document until you reach the {ord} non-empty paragraph, then find the first standalone digit there and apply subscript formatting to it, so once the file is printed the digit reads as a proper chemical subscript.",
    # extra long narrative pad
    "I'm putting together study notes for an undergraduate chemistry tutorial and the {ord} non-empty paragraph contains a chemical formula whose first standalone digit needs to render as a subscript so the formula reads correctly to students seeing the printout — please apply subscript formatting to that digit while leaving the rest of the paragraph and the surrounding document text completely untouched.",
]


def _make_subscript(rng: random.Random) -> tuple[str, str]:
    """Return (instruction, py_code) for the subscript archetype.

    Oracle finds the n-th non-empty paragraph, locates the first standalone
    digit run (or splits a longer run around the digit), and sets
    run.font.subscript = True. compare_subscript_contains needs at least one
    aligned (run1, run2) pair in some paragraph where both are subscript;
    compare_docx_files needs the paragraph text identical (which it is — we
    only change formatting, not text).
    """
    nth = rng.choice([0, 1, 2])
    phrase = rng.choice(_SUBSCRIPT_VARIANTS).format(ord=_ORDINALS[nth])
    instruction = _build_instruction(phrase, rng)
    py_code = textwrap.dedent(f"""\
        import sys, re
        import copy
        from docx import Document

        doc = Document('{{file_path}}')
        # Filter to non-empty BODY paragraphs that contain a digit. validation
        # audit: previously included Heading/Title paragraphs (e.g., 0b17a146's
        # "H2O—Soak up the Science" Title at paragraphs[1]), so candidates[0]
        # was the Title — agent visually skipped it and picked the next body
        # paragraph, producing ordinal mismatch. Skipping Heading/Title style
        # aligns oracle's n-th candidate with agent's natural body counting.
        candidates = []
        for p in doc.paragraphs:
            if not p.text.strip():
                continue
            style_name = p.style.name if p.style else ""
            if "Heading" in style_name or "Title" in style_name:
                continue
            if any(re.search(r'\\d', r.text or '') for r in p.runs):
                candidates.append(p)
        if len(candidates) <= {nth}:
            sys.exit(42)
        target = candidates[{nth}]

        # Find first run that contains a digit; split it so the digit lives in
        # its own run, with subscript=True on the digit run. We insert the new
        # runs *immediately after* the source run via XML element-level
        # insertion (NOT add_run, which appends to paragraph end and would
        # break text order in multi-run paragraphs).
        applied = False
        for run in list(target.runs):
            m = re.search(r'\\d', run.text or '')
            if not m:
                continue
            before = run.text[:m.start()]
            digit = run.text[m.start():m.start()+1]
            after = run.text[m.start()+1:]
            # 1) shrink original run to the prefix
            run.text = before
            # 2) clone the run's <w:r> to inherit rPr (bold/italic/font/etc.),
            #    then patch text + subscript on the clone for the digit, and
            #    a second clone for the trailing tail (if any).
            src_xml = run._r
            digit_xml = copy.deepcopy(src_xml)
            # clear text nodes in the clone
            for t in digit_xml.findall('.//{{{{*}}}}t'):
                t.getparent().remove(t)
            # add a fresh text node + vertAlign=subscript to rPr
            from docx.oxml.ns import qn
            from lxml import etree
            t_el = etree.SubElement(digit_xml, qn('w:t'))
            t_el.text = digit
            t_el.set(qn('xml:space'), 'preserve')
            rPr = digit_xml.find(qn('w:rPr'))
            if rPr is None:
                rPr = etree.SubElement(digit_xml, qn('w:rPr'))
                digit_xml.insert(0, rPr)
            # remove any prior vertAlign and add subscript
            for va in rPr.findall(qn('w:vertAlign')):
                rPr.remove(va)
            va = etree.SubElement(rPr, qn('w:vertAlign'))
            va.set(qn('w:val'), 'subscript')
            src_xml.addnext(digit_xml)
            if after:
                tail_xml = copy.deepcopy(src_xml)
                for t in tail_xml.findall('.//{{{{*}}}}t'):
                    t.getparent().remove(t)
                t_el2 = etree.SubElement(tail_xml, qn('w:t'))
                t_el2.text = after
                t_el2.set(qn('xml:space'), 'preserve')
                digit_xml.addnext(tail_xml)
            applied = True
            break
        if not applied:
            sys.exit(42)
        doc.save('{{expected_path}}')
    """)
    return instruction, py_code


# 3) compare_pdfs ×4 archetype — eval base 4bcb1253 ---------------------------
# The eval task asks the agent to export the docx to PDF at 4 standard paths
# (Desktop / Documents / Downloads / home), and the evaluator compares each
# of those 4 PDFs against the gold PDF using fuzzy text match (conj=or, so
# any one PDF matching is enough). We perturb by varying which subset of
# locations the export targets (the evaluator passes if ANY one match exists
# under conj=or; we always export to all 4 to keep oracle robust).
#
# Oracle: shell-based — uses soffice --convert-to pdf to write the PDF to all
# 4 standard locations from the source docx. The evaluator paths are inherited
# unchanged from the eval row's evaluator.

_PDF_EXPORT_VARIANTS = [
    # short
    "Export the document to PDF with the same name.",
    "Convert this document into a PDF.",
    # medium narrative pads
    "I'm finalising this report for distribution and need a PDF copy of the document — please export it to PDF, keeping the same file name, so I can attach it to an email later.",
    "Could you export the current document to PDF using the same base file name? I want to share a non-editable version with my reviewers and need the PDF available on the desktop.",
    "Quick export task: turn this document into a PDF (same filename, .pdf extension) so I can hand it off to a colleague who can't open .docx files.",
    "I'd like a PDF version of this writeup with the same name as the docx — export the document to PDF for me, please, so I have both formats side by side.",
    # multi-step medium (first/then sequence keywords)
    "First open the export menu in LibreOffice Writer, then convert the current document to PDF using the same base file name so the PDF sits alongside the original docx for me to share.",
    # long narrative pad
    "I'm finalising this report for tomorrow's distribution to external partners and several of them have flagged that they can't reliably open editable docx files in their corporate browsers — please export the current document to PDF with the same base name as the docx, so I can send around a non-editable copy that renders identically across every recipient's mail client.",
    # long multi-step (first/then/once sequence keywords)
    "For an upcoming submission, please convert this document into a PDF that preserves the same file name; first export the docx to PDF using LibreOffice's File menu, then make sure the PDF lands in a standard location, so once the export finishes I can pick up the PDF for upload to the submission portal alongside the original docx.",
    # extra long narrative pad
    "For an upcoming submission to a partner organisation, please convert this document into a PDF that preserves the original docx file name so the PDF lives next to the docx in the same folder for me to upload — several partners explicitly ask for PDF renderings rather than editable docx files when they pull material into their internal review systems.",
]


def _make_pdf_export(rng: random.Random, file_path: str) -> tuple[str, str]:
    """Return (instruction, shell_oracle_cmd) for the PDF export archetype.

    The oracle is a shell command (not python-docx), unlike the other writer
    archetypes. It uses soffice --convert-to pdf to render the docx to PDF
    at all 4 standard target locations the eval evaluator inspects:
        /home/user/Desktop/<base>.pdf
        /home/user/Documents/<base>.pdf
        /home/user/Downloads/<base>.pdf
        /home/user/<base>.pdf
    """
    phrase = rng.choice(_PDF_EXPORT_VARIANTS)
    instruction = _build_instruction(phrase, rng)
    base = file_path.rsplit("/", 1)[-1]
    if base.lower().endswith(".docx"):
        pdf_name = base[:-5] + ".pdf"
    else:
        pdf_name = base + ".pdf"
    # Shell snippet: convert with LO once into a tmp dir, then copy to all 4
    # target dirs (mkdir -p so missing dirs are created). Used as the
    # perturb_config_step's command body.
    shell_cmd = (
        f"tmpd=$(mktemp -d) && "
        f"DISPLAY=:1 soffice --headless --norestore --nofirststartwizard "
        f"--convert-to pdf --outdir \"$tmpd\" '{file_path}' 2>/dev/null; "
        f"src=\"$tmpd/{pdf_name}\"; "
        f"if [ -f \"$src\" ]; then "
        f"  for d in /home/user/Desktop /home/user/Documents /home/user/Downloads /home/user; do "
        f"    mkdir -p \"$d\" && cp \"$src\" \"$d/{pdf_name}\"; "
        f"  done; "
        f"fi; "
        f"rm -rf \"$tmpd\"; true"
    )
    return instruction, shell_cmd


# 4) has_page_numbers_in_footers archetype — eval base 0e47de2a --------------
# Eval: "Add page number for every page at the bottom left". Evaluator checks
# footer paragraph contains a digit. Oracle modifies file_path's footer to
# contain a page-number digit; runs after postconfig so it overrides agent state.

_PAGE_NUMBER_VARIANTS = [
    # short
    "Add page numbers in the footer of every page.",
    "Insert a page number field at the bottom of the document.",
    # medium
    "I'm preparing this document for printing and the pages need to be numbered — please add a page number to the footer of every page.",
    "While polishing this draft I noticed the printed handout has no page references — could you add page numbers in the footer so reviewers can cite specific pages?",
    "Quick formatting fix: insert page numbers into the footer of the document so each printed page shows its number at the bottom.",
    "Please add a page-number field to the footer of this document so each page is clearly numbered when the file is printed or shared as PDF.",
    # long
    "I'm finalising the document for an internal review and the editorial team flagged that page numbers are missing — please add a page number to the footer of every page so reviewers can refer to specific pages during the markup pass.",
    "For a printed handout I'm preparing tonight, the body of the document has no page numbers in the footer; first add a page-number field to the footer so each printed page shows its number, then save the file so the next print job picks up the change.",
]


def _make_page_number(rng: random.Random) -> tuple[str, str]:
    """Return (instruction, py_code) for the page-number archetype.

    Oracle adds a digit to footer text in every section. Evaluator
    (has_page_numbers_in_footers) only checks footer contains a digit.
    """
    phrase = rng.choice(_PAGE_NUMBER_VARIANTS)
    instruction = _build_instruction(phrase, rng)
    py_code = textwrap.dedent("""\
        from docx import Document
        doc = Document('{file_path}')
        for sec in doc.sections:
            footer = sec.footer
            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            if not any(c.isdigit() for c in para.text):
                para.text = (para.text or '') + '1'
        doc.save('{file_path}')
    """)
    return instruction, py_code


# 5) contains_page_break archetype — eval base ecc2413d ----------------------
# Eval: "throw in a blank page right after this one" (page_break_count: 5; doc
# already has 4 page breaks). Oracle inserts 1 extra page break, evaluator
# checks final count == 5.

_PAGE_BREAK_VARIANTS = [
    # short
    "Insert a page break after the first paragraph.",
    "Add a blank page right after the first section.",
    # medium
    "I'm reorganising this document for printing and need a page break after the first paragraph so the second section starts on a fresh page — please add one.",
    "Could you insert a page break after the first paragraph? The next reviewer wants the body content starting on its own page.",
    "Quick layout fix: add a blank page after the first paragraph by inserting a page break, so the document flows onto a fresh page after the opening line.",
    # long
    "For a printed copy I'm preparing for tomorrow's review session, the opening paragraph should sit on its own page; first place the cursor at the end of the first paragraph, then insert a page break so the rest of the document starts on a fresh page when printed.",
]


def _make_page_break(rng: random.Random) -> tuple[str, str]:
    """Return (instruction, py_code) for the page-break archetype.

    Oracle inserts a page break at the end of paragraph 0. Evaluator
    (contains_page_break) checks final break count matches the rule. For
    ecc2413d (4 existing breaks), expected final count is 5.
    """
    phrase = rng.choice(_PAGE_BREAK_VARIANTS)
    instruction = _build_instruction(phrase, rng)
    py_code = textwrap.dedent("""\
        from docx import Document
        from docx.enum.text import WD_BREAK
        doc = Document('{file_path}')
        run = doc.paragraphs[0].add_run()
        run.add_break(WD_BREAK.PAGE)
        doc.save('{file_path}')
    """)
    return instruction, py_code


# Per-base archetype registry. Each entry maps a writer eval short_id to a
# list of archetype builder callables; the perturb dispatcher emits up to
# `max_type3` rows per base by sampling distinct paraphrases from the pool.
# (Implemented inline in perturb_writer_per_task since each archetype needs a
# different post-processing path — tabstops/subscript share the docx-LO oracle
# pattern with a custom evaluator override; pdf_export uses a pure-shell oracle
# with no docx normalize; page_number/page_break write back to file_path post-save.)
_TYPE3_BASES = frozenset({"0a0faba3", "0b17a146", "4bcb1253", "0e47de2a", "ecc2413d"})

# Run-format ops (bold=0, strikethrough=5, italic=7) share a feasibility constraint:
# e528b65e has 4 paras with the last being empty/no-run. Exclude these ops for that
# task to prevent rc=42 rows. Handled via _RUN_FORMAT_INFEASIBLE at generation time;
# _T2_VARIANTS for e528b65e already omits ops 0/5/7 so the check is a safety net only.
_RUN_FORMAT_INFEASIBLE = frozenset({
    "osworld_libreoffice_writer_e528b65e",
})
_RUN_FORMAT_OP_INDICES = frozenset({0, 5, 7})  # bold, strikethrough, italic

# Ops that pick a paragraph index. Used by the dispatcher to select which
# n_paras bound to pass: idx ops get the safe bound (last_non_empty_idx + 1)
# so the runtime advance loop can always land on a non-empty paragraph.
# Font (1) and find_replace (2) are doc-wide / text-search ops with no idx.
_PARA_IDX_OPS = frozenset({0, 3, 4, 5, 6, 7})  # bold, spacing, upper, strike, lower, italic

# TYPE_1 eligible tasks: eval op is in the perturb pool; resample same op class.
# Maps short_id → (op_idx, fn(rng, n_paras) → (instruction, py_code))
# Validation oracle/validate fix: per-base source font registry used by
# TYPE_2 font op (op_idx=1) to exclude the source's current font and avoid
# vacuous-pass perturb rows. Detected via oracle/validate trivial_pass on
# 6ada715d (source docx already Verdana → "change font to Verdana" no-op).
# Add entries here whenever a base's source docx has a known font that's
# in `_FONTS`.
_BASE_SOURCE_FONT: dict[str, str] = {
    "6ada715d": "Verdana",
}


_TYPE1_FNS: dict[str, tuple[int, object]] = {
    # eval: spacing of first two paras to double → TYPE_1: para 2+ (any spacing)
    # exclude_idx={0,1} excludes ord_pos 0,1 (i.e., body[0],body[1]) so T1 picks
    # a different paragraph than eval's "first two paragraph" target.
    "0810415c": (3, lambda rng, n, b: _make_spacing(rng, n, exclude_idx={0, 1}, body_idxs=b)),
    # eval: spacing (single para) → TYPE_1: different para or 1.5 lines spacing
    "b21acd93": (3, lambda rng, n, b: _make_spacing(rng, n, exclude_idx={0}, body_idxs=b)),
    # eval: font → "Times New Roman" → TYPE_1: different font (no para idx; b ignored)
    "0e763496": (1, lambda rng, n, b: _make_font(rng, n, exclude_fonts={"Times New Roman"})),
    # eval: source `loa-one-time-submission-sealand.docx` renders as ARIAL (all
    # runs font.name=None; Normal-style default = Arial). The old "Times New Roman"
    # exclude was a mis-identified source font, so _make_font could pick Arial → a
    # no-op perturb → compare_docx_strict's effective-font cascade scores it 1.0
    # BEFORE the oracle (trivial-pass FP). Exclude Arial so the perturb is a real
    # font change. (#154 oracle-validation trivial_pass; verified: source effective
    # fonts = {'Arial'}, no Times-New-Roman runs.)
    "f178a4a9": (1, lambda rng, n, b: _make_font(rng, n, exclude_fonts={"Arial"})),
}

# Explicit per-task TYPE_2 candidate op indices (≥4 per task so max_type2>2 is meaningful).
# Indices: 0=bold, 1=font, 2=find_replace, 3=spacing, 4=uppercase, 5=strike, 6=lower, 7=italic
# Design constraints applied at pool-design time (pool-level invariant, not sampling-time):
#   - find_replace(2) excluded from long docs (n_paras>20) — agent can't scan whole doc in 15 turns
#   - T1 op excluded from T2 pool for T1-eligible tasks — avoids duplicating the eval op
#   - bold(0)/strike(5)/italic(7) excluded for e528b65e — empty last para causes rc=42
# The pool composition ensures all 8 ops appear across all tasks with roughly equal frequency,
# so distribution match holds regardless of which max_type2 value is chosen at runtime.
# validation: bases not in _ORDINAL_SAFE_BASES have ordinal-paragraph ops
# (0=bold, 3=spacing, 4=upper, 5=strike, 6=lower, 7=italic) silently skipped at
# dispatch time. The pool entries below are kept as-is for traceability of the
# original op design — runtime filter applies. Where the unfiltered remainder
# would yield 0 ops, font(1) is added to ensure each base produces at least one
# T2 row (bb8ccc78 still excluded — eval infeasible).
_T2_VARIANTS: dict[str, list[int]] = {
    # Validation note: removed op 2 (find_replace) globally because the
    # `_make_find_replace` template asks the agent to "find the most frequent
    # 4+ letter word and replace it" — requires whole-document word-frequency
    # analysis the agent can't reliably do via GUI. 3+ bases had this template
    # consistently FAIL (0a0faba3 5/5 FAIL = uniform-zero; b21acd93 + 0b17a146
    # INFEASIBLE_CLAIM_TRAIN @ turn_00/early). Pool entries below preserved
    # for traceability; runtime filter on op 2 produces same-shape mix
    # without the unwinnable find_replace variant.
    "0810415c": [0, 1, 4, 7],     # long+T1(spac): bold, font, upper, italic
    "0a0faba3": [1, 5, 6],        # was [1,2,5,6] — drop fr (uniform-zero source bug)
    "0b17a146": [0, 3, 7],        # was [0,2,3,7] — drop fr
    "0e47de2a": [0, 1, 5, 6],     # unsafe long: bold, font, strike, lower → font only
    "0e763496": [0, 3, 5],        # was [0,2,3,5] — drop fr
    "3ef2b351": [1, 4, 7],        # was [1,2,4,7] — drop fr
    "4bcb1253": [0, 1, 3, 5],     # unsafe long: bold, font, spacing, strike → font only
    "66399b0d": [1, 3, 7],        # was [1,2,3,7] — drop fr
    "6ada715d": [0, 1, 3, 4, 6],  # unsafe long: + font(1) so non-empty after filter
    "6f81754e": [1, 3, 5, 6],     # long: font, spacing, strike, lower
    "72b810ef": [0, 4, 7],        # was [0,2,4,7] — drop fr
    "8472fece": [1, 4, 5],        # was [1,2,4,5] — drop fr
    "88fe4b2d": [0, 1, 3, 6, 7],  # unsafe long: + font(1) so non-empty after filter
    "936321ce": [1, 3, 4, 5],     # unsafe long: font, spacing, upper, strike → font only
    "adf5e2c3": [0, 1, 5, 7],     # long: bold, font, strike, italic
    "b21acd93": [0, 4, 6],        # was [0,2,4,6] — drop fr
    # bb8ccc78 removed: eval base is infeasible ("share document in real-time").
    # Dispatcher now skips infeasible eval bases (perturb/__init__.py); this
    # entry is dead but kept commented for traceability.
    "d53ff5ee": [4, 6, 7],        # was [2,4,6,7] — drop fr
    "e246f6d8": [0, 1, 3, 6, 7],  # unsafe long: + font(1) so non-empty after filter
    "e528b65e": [1, 3, 4],        # was [1,2,3,4] — drop fr
    "ecc2413d": [0, 1, 3, 4, 7],  # unsafe long: + font(1) so non-empty after filter
    "f178a4a9": [0, 3, 5, 6],     # unsafe long+T1(font): T1 already provides font;
                                   # T2 would conflict so leave empty (filter→[]).
}

# ---------------------------------------------------------------------------
# Analysis loader
# ---------------------------------------------------------------------------

_ANALYSIS: dict | None = None


def _load_analysis() -> dict:
    global _ANALYSIS
    if _ANALYSIS is None:
        _ANALYSIS = json.loads(_ANALYSIS_PATH.read_text()) if _ANALYSIS_PATH.exists() else {}
    return _ANALYSIS


# ---------------------------------------------------------------------------
# TYPE_3 row builder — uses an archetype-specific evaluator instead of
# compare_docx_strict. Follows the same oracle pattern (LO normalize +
# cp + LO normalize) for tabstops/subscript so docx round-trip parity holds;
# the PDF archetype uses a pure-shell oracle (soffice --convert-to pdf).
# ---------------------------------------------------------------------------

def _build_archetype_row(
    eval_row: dict,
    short: str,
    archetype_name: str,
    instruction: str,
    file_path: str,
    expected_path: str,
    py_code_filled: str | None,
    shell_cmd: str | None,
    evaluator: dict,
    oracle: list[dict],
) -> dict:
    """Build a TYPE_3 perturb row.

    Exactly one of ``py_code_filled`` / ``shell_cmd`` must be provided:
      - ``py_code_filled``: full python-docx snippet that writes ``expected_path``
        (heredoc-wrapped into a python3 execute step)
      - ``shell_cmd``: raw shell command body (no heredoc, run in shell=True)
    """
    if py_code_filled is not None:
        config_step = _make_config_step(py_code_filled)
    else:
        assert shell_cmd is not None
        config_step = {"type": "execute", "parameters": {
            "command": shell_cmd, "shell": True,
        }}
    return make_perturb_row(
        eval_row=eval_row,
        knob_assignment={
            "type": PERTURB_TYPE_3,
            "archetype": archetype_name,
            "variant": instruction,
        },
        new_instruction=instruction,
        new_oracle=oracle,
        new_evaluator=evaluator,
        perturb_config_step=config_step,
        oracle_after_postconfig=True,
    )


def _emit_tabstops_rows(
    eval_row: dict, short: str, file_path: str, max_rows: int, rng: random.Random,
) -> list[dict]:
    """Emit up to max_rows tabstops-archetype rows for base 0a0faba3."""
    expected_path = f"/tmp/perturb_expected_{eval_row['task_id'][-8:]}_tabstops.docx"
    evaluator = {
        "func": "check_tabstops",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "expected": {"type": "vm_file", "path": expected_path,
                     "dest": "expected_file"},
        "postconfig": LO_SAVE_POSTCONFIG,
    }
    oracle = _build_oracle(file_path, expected_path)
    out: list[dict] = []
    seen: set[str] = set()
    for _ in range(max_rows * 5):
        if len(out) >= max_rows:
            break
        instruction, py_code = _make_tabstops(rng)
        if instruction in seen:
            continue
        seen.add(instruction)
        py_code_filled = py_code.format(
            file_path=file_path, expected_path=expected_path,
        )
        out.append(_build_archetype_row(
            eval_row, short, "tabstops", instruction,
            file_path, expected_path, py_code_filled, None,
            evaluator, oracle,
        ))
    return out


def _emit_subscript_rows(
    eval_row: dict, short: str, file_path: str, max_rows: int, rng: random.Random,
) -> list[dict]:
    """Emit up to max_rows subscript-archetype rows for base 0b17a146.

    Evaluator is the eval-shape compound (compare_docx_files ∧
    compare_subscript_contains) with list-form result/expected.
    """
    expected_path = f"/tmp/perturb_expected_{eval_row['task_id'][-8:]}_subscript.docx"
    evaluator = {
        "func": ["compare_docx_files", "compare_subscript_contains"],
        "result": [
            {"type": "vm_file", "path": file_path,
             "dest": file_path.split("/")[-1]},
            {"type": "vm_file", "path": file_path,
             "dest": file_path.split("/")[-1]},
        ],
        "expected": [
            {"type": "vm_file", "path": expected_path, "dest": "expected_file"},
            {"type": "vm_file", "path": expected_path, "dest": "expected_file"},
        ],
        "postconfig": LO_SAVE_POSTCONFIG,
    }
    oracle = _build_oracle(file_path, expected_path)
    out: list[dict] = []
    seen: set[str] = set()
    for _ in range(max_rows * 5):
        if len(out) >= max_rows:
            break
        instruction, py_code = _make_subscript(rng)
        if instruction in seen:
            continue
        seen.add(instruction)
        py_code_filled = py_code.format(
            file_path=file_path, expected_path=expected_path,
        )
        out.append(_build_archetype_row(
            eval_row, short, "subscript", instruction,
            file_path, expected_path, py_code_filled, None,
            evaluator, oracle,
        ))
    return out


def _emit_pdf_export_rows(
    eval_row: dict, short: str, file_path: str, max_rows: int, rng: random.Random,
) -> list[dict]:
    """Emit up to max_rows pdf-export-archetype rows for base 4bcb1253.

    Evaluator inherits the eval row's compare_pdfs ×4 list-form structure but
    with paths rewritten to the standard 4 export targets keyed off the docx
    base name (so it works for any task that follows this archetype, even if
    we later extend to additional bases).
    """
    base = file_path.rsplit("/", 1)[-1]
    pdf_name = (base[:-5] + ".pdf") if base.lower().endswith(".docx") else base + ".pdf"
    target_dirs = [
        "/home/user/Desktop",
        "/home/user/Documents",
        "/home/user/Downloads",
        "/home/user",
    ]
    # The eval evaluator uses cloud_file expected (HuggingFace gold pdf). We
    # don't have a gold pdf — we *generate* expected on the fly by exporting
    # the same docx to a /tmp path and reusing it 4×. The oracle script writes
    # a single expected.pdf and copies it to all 4 result paths via SAVE_
    # POSTCONFIG-then-cp; the evaluator then trivially passes (text identical).
    expected_pdf = f"/tmp/perturb_expected_{eval_row['task_id'][-8:]}.pdf"
    evaluator = {
        "func": ["compare_pdfs", "compare_pdfs", "compare_pdfs", "compare_pdfs"],
        "conj": "or",
        "result": [
            {"type": "vm_file", "path": f"{d}/{pdf_name}", "dest": f"result_{i}.pdf"}
            for i, d in enumerate(target_dirs, 1)
        ],
        "expected": [
            {"type": "vm_file", "path": expected_pdf, "dest": f"expected_{i}.pdf"}
            for i in range(1, 5)
        ],
    }
    # Oracle: convert docx → pdf once, copy to all 4 target dirs AND to
    # expected_pdf. No LO_SAVE_POSTCONFIG: the agent's job is to PDF-export, not
    # save the docx; if oracle ran after Ctrl+S, it would race with LO save.
    # We use oracle_after_postconfig=True (so postconfig kills LO first), and
    # the agent's saved docx state is irrelevant to the evaluator.
    oracle = [
        {"type": "execute", "parameters": {
            "command": (
                f"tmpd=$(mktemp -d) && "
                f"DISPLAY=:1 soffice --headless --norestore --nofirststartwizard "
                f"--convert-to pdf --outdir \"$tmpd\" '{file_path}' 2>/dev/null; "
                f"src=\"$tmpd/{pdf_name}\"; "
                f"if [ -f \"$src\" ]; then "
                f"  cp \"$src\" '{expected_pdf}'; "
                f"  for d in {' '.join(target_dirs)}; do "
                f"    mkdir -p \"$d\" && cp \"$src\" \"$d/{pdf_name}\"; "
                f"  done; "
                f"fi; "
                f"rm -rf \"$tmpd\"; true"
            ),
            "shell": True,
        }},
    ]
    # BUG-3 fix: the gold ``expected_pdf`` MUST be generated at
    # CONFIG time, not only in the oracle. The evaluator compares the agent's
    # exported PDF against ``expected_pdf``; during a real agent rollout the
    # oracle actions DO NOT run, so if the gold is built only there it never
    # exists → ``compare_pdfs`` gets a missing ``pdf2_path`` → every PDF-export
    # task scores 0 (4bcb1253 family, 0/3). Oracle VALIDATION passed because the
    # oracle creates it — masking the bug (same pattern as the thunderbird
    # default-pref seed). synth's PDF-export tasks already do this in pre_config
    # (_soffice_docx_to_pdf_step), which is why synth PDF yields ~87% vs perturb 0.
    # The docx is a pure export here (instruction adds no transform), so a
    # config-time headless convert produces the same text the agent's GUI export
    # will → compare_pdfs (text fuzz_ratio) matches. The oracle still copies to
    # the 4 target dirs so oracle validation keeps working.
    perturb_config_step = {
        "type": "execute",
        "parameters": {
            "command": (
                f"tmpd=$(mktemp -d) && "
                f"DISPLAY=:1 soffice --headless --norestore --nofirststartwizard "
                f"--convert-to pdf --outdir \"$tmpd\" '{file_path}' 2>/dev/null; "
                f"[ -f \"$tmpd/{pdf_name}\" ] && cp \"$tmpd/{pdf_name}\" '{expected_pdf}'; "
                f"rm -rf \"$tmpd\"; true"
            ),
            "shell": True,
        },
    }
    out: list[dict] = []
    seen: set[str] = set()
    for _ in range(max_rows * 5):
        if len(out) >= max_rows:
            break
        instruction, _shell_cmd_unused = _make_pdf_export(rng, file_path)
        if instruction in seen:
            continue
        seen.add(instruction)
        out.append(make_perturb_row(
            eval_row=eval_row,
            knob_assignment={
                "type": PERTURB_TYPE_3,
                "archetype": "pdf_export",
                "variant": instruction,
            },
            new_instruction=instruction,
            new_oracle=oracle,
            new_evaluator=evaluator,
            perturb_config_step=perturb_config_step,
            oracle_after_postconfig=True,
        ))
    return out


def _emit_page_number_rows(
    eval_row: dict, short: str, file_path: str, max_rows: int, rng: random.Random,
) -> list[dict]:
    """Emit page-number-archetype rows for base 0e47de2a.

    Evaluator is single-arg has_page_numbers_in_footers (no expected file).
    Oracle modifies file_path's footer to contain a digit; runs after
    postconfig so the agent's saved state is overridden by the gold state.
    """
    evaluator = {
        "func": "has_page_numbers_in_footers",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
    }
    out: list[dict] = []
    seen: set[str] = set()
    for _ in range(max_rows * 5):
        if len(out) >= max_rows:
            break
        instruction, py_code = _make_page_number(rng)
        if instruction in seen:
            continue
        seen.add(instruction)
        py_code_filled = py_code.format(file_path=file_path)
        # Oracle = single python config step rewriting file_path's footer.
        oracle = [_make_config_step(py_code_filled)]
        out.append(make_perturb_row(
            eval_row=eval_row,
            knob_assignment={
                "type": PERTURB_TYPE_3,
                "archetype": "page_number",
                "variant": instruction,
            },
            new_instruction=instruction,
            new_oracle=oracle,
            new_evaluator=evaluator,
            perturb_config_step=None,
            oracle_after_postconfig=True,
        ))
    return out


def _emit_page_break_rows(
    eval_row: dict, short: str, file_path: str, max_rows: int, rng: random.Random,
    expected_break_count: int,
) -> list[dict]:
    """Emit page-break-archetype rows for base ecc2413d.

    Evaluator is contains_page_break with rule-based expected count.
    Oracle adds a page break to paragraph 0; runs after postconfig.
    """
    evaluator = {
        "func": "contains_page_break",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "expected": {"type": "rule",
                     "rules": {"page_break_count": expected_break_count}},
    }
    out: list[dict] = []
    seen: set[str] = set()
    for _ in range(max_rows * 5):
        if len(out) >= max_rows:
            break
        instruction, py_code = _make_page_break(rng)
        if instruction in seen:
            continue
        seen.add(instruction)
        py_code_filled = py_code.format(file_path=file_path)
        oracle = [_make_config_step(py_code_filled)]
        out.append(make_perturb_row(
            eval_row=eval_row,
            knob_assignment={
                "type": PERTURB_TYPE_3,
                "archetype": "page_break",
                "variant": instruction,
            },
            new_instruction=instruction,
            new_oracle=oracle,
            new_evaluator=evaluator,
            perturb_config_step=None,
            oracle_after_postconfig=True,
        ))
    return out


# ---------------------------------------------------------------------------
# Main perturb function
# ---------------------------------------------------------------------------

def perturb_writer_per_task(
    eval_row: dict,
    rng: random.Random,
    max_type1: int = 1,
    max_type2: int = 2,
    max_type3: int = 3,
) -> list[dict]:
    file_path = _get_file_path(eval_row)
    # Skip .odt files: compare_docx_files returns 0 for extension mismatch (.odt vs .docx)
    if not file_path or not file_path.endswith('.docx'):
        return []

    tid = eval_row["task_id"]
    analysis = _load_analysis()
    short = tid.split("_")[-1]
    info = analysis.get(short, analysis.get(tid, {}))
    n_paras = info.get("n_paras", info.get("paragraphs", 5))
    # text_paras lists [(idx, text_prefix), ...] for non-empty paragraphs.
    # For paragraph-targeting ops we restrict the start idx so the runtime's
    # "advance past empty" loop can always reach a non-empty paragraph (even
    # if the chosen idx itself is empty, advance can land on a later non-empty
    # one). When the start idx exceeds the last non-empty paragraph index, the
    # advance loop has nowhere to go → either no-op (trivial_pass) or
    # sys.exit(42) (oracle FAIL). Cap idx at last_non_empty_idx + 1 so the
    # randint range stays within [0, last_non_empty_idx].
    text_paras = info.get("text_paras") or []
    if text_paras:
        last_non_empty_idx = max(t[0] for t in text_paras)
        n_paras_safe = max(1, min(n_paras, last_non_empty_idx + 1))
    else:
        n_paras_safe = n_paras

    expected_path = f"/tmp/perturb_expected_{tid[-8:]}.docx"

    rows: list[dict] = []

    # body_idxs for ordinal-safe bases: list of python-docx paragraph indices
    # corresponding to the user's natural visual paragraph counting (skipping
    # titles, headings, captions). When None, the base has ambiguous structure
    # and ordinal-paragraph ops (_PARA_IDX_OPS) are skipped from the T2 pool.
    body_idxs = _ORDINAL_SAFE_BASES.get(short)

    # TYPE_1: same op as eval, different params (only for tasks in _TYPE1_FNS)
    t1_entry = _TYPE1_FNS.get(short)
    if t1_entry is not None:
        t1_op_idx, t1_fn = t1_entry
        used_instrs: set[str] = set()
        for _ in range(max_type1 * 5):  # retry up to 5× to get distinct variants
            if len(rows) >= max_type1:
                break
            try:
                # TYPE_1 ops are spacing (0810415c, b21acd93 — para idx) and
                # font (0e763496, f178a4a9 — doc-wide). Use n_paras_safe for
                # idx ops; font ignores n_paras anyway so it's safe to pass.
                t1_np_arg = n_paras_safe if t1_op_idx in _PARA_IDX_OPS else n_paras
                instruction, py_code = t1_fn(rng, t1_np_arg, body_idxs)
            except Exception:
                break
            if instruction in used_instrs:
                continue
            used_instrs.add(instruction)
            py_code_filled = py_code.format(file_path=file_path, expected_path=expected_path)
            rows.append(_build_perturb_row(
                eval_row, t1_op_idx, instruction, py_code_filled,
                file_path, expected_path, PERTURB_TYPE_1,
            ))

    # TYPE_2: pick from explicit per-task candidate pool (see _T2_VARIANTS).
    # Pool is pre-designed to satisfy distribution match for any max_type2 value.
    t2_pool = list(_T2_VARIANTS.get(short, [0, 1, 3, 4]))
    rng.shuffle(t2_pool)

    t2_count = 0
    for op_idx in t2_pool:
        if t2_count >= max_type2:
            break
        if op_idx in _RUN_FORMAT_OP_INDICES and tid in _RUN_FORMAT_INFEASIBLE:
            continue
        # Skip ordinal-paragraph ops on bases without an _ORDINAL_SAFE_BASES
        # entry: the doc has ambiguous paragraph structure (interleaved
        # headings, banner+image+table, line-break-merged bullets, etc.) where
        # agent's visual counting diverges from python-docx <w:p> indexing.
        if op_idx in _PARA_IDX_OPS and body_idxs is None:
            continue
        np_arg = n_paras_safe if op_idx in _PARA_IDX_OPS else n_paras
        try:
            if op_idx in _PARA_IDX_OPS:
                if op_idx == 3:
                    instruction, py_code = _make_spacing(
                        rng,
                        np_arg,
                        body_idxs=body_idxs,
                        long_quote_limit=_SPACING_LONG_QUOTE_LIMITS.get(short),
                    )
                else:
                    instruction, py_code = _OP_POOL[op_idx](rng, np_arg, body_idxs=body_idxs)
            elif op_idx == 1 and short in _BASE_SOURCE_FONT:
                # Validation: avoid picking the source's current font for
                # TYPE_2 font op — otherwise the perturb is vacuous.
                instruction, py_code = _make_font(
                    rng, np_arg, exclude_fonts={_BASE_SOURCE_FONT[short]},
                )
            else:
                instruction, py_code = _OP_POOL[op_idx](rng, np_arg)
        except Exception:
            continue

        py_code_filled = py_code.format(file_path=file_path, expected_path=expected_path)
        rows.append(_build_perturb_row(
            eval_row, op_idx, instruction, py_code_filled,
            file_path, expected_path, PERTURB_TYPE_2,
        ))
        t2_count += 1

    # TYPE_3: per-base archetype rows targeting evaluators outside
    # compare_docx_strict's coverage. Only fires for the 3 specific bases in
    # _TYPE3_BASES; runs *in addition to* TYPE_1/TYPE_2 for those bases.
    if short in _TYPE3_BASES and max_type3 > 0:
        if short == "0a0faba3":
            rows.extend(_emit_tabstops_rows(
                eval_row, short, file_path, max_type3, rng,
            ))
        elif short == "0b17a146":
            # validation: 0b17a146 has visual ordinal ambiguity (FACT SHEET
            # banner + H2O Title + WATER IS VITAL... small-caps lead before
            # regular body paragraphs). Subscript instruction "first digit in
            # the {ord} paragraph" cannot be unambiguously interpreted by the
            # agent — disabled until base is added to _ORDINAL_SAFE_BASES.
            # Eval coverage of compare_subscript_contains is therefore from
            # the eval task only (no perturb archetype rows).
            pass
        elif short == "4bcb1253":
            rows.extend(_emit_pdf_export_rows(
                eval_row, short, file_path, max_type3, rng,
            ))
        elif short == "0e47de2a":
            rows.extend(_emit_page_number_rows(
                eval_row, short, file_path, max_type3, rng,
            ))
        elif short == "ecc2413d":
            # ecc2413d.docx already has 4 page breaks; eval expects 5
            # (oracle adds 1). Keep same expected count for the perturb so
            # contains_page_break(page_break_count=5) trivially aligns.
            rows.extend(_emit_page_break_rows(
                eval_row, short, file_path, max_type3, rng,
                expected_break_count=5,
            ))

    return rows
