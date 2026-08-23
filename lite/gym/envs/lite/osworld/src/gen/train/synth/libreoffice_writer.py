"""LibreOffice Writer synth generator (Batch, dataclass §I).

Per AGENTS.md:
every row builds the source docx in `pre_config_steps` via a python-docx
heredoc, then a SECOND heredoc writes the gold docx to
`/tmp/expected_<id>.docx` (Hard Constraint #2 — gold MUST land in
metadata.config). The oracle `cp`s the gold over the agent's file path.
The evaluator pulls both docx files (`vm_file`) and runs the mapped eval
func: compare_docx_strict / compare_docx_files / compare_docx_tables /
contains_page_break / compare_line_spacing / has_page_numbers_in_footers /
is_first_line_centered / find_default_font / compare_font_names /
evaluate_strike_through_last_paragraph / check_italic_font_size_14 /
compare_pdfs.

Postconfig is `LO_SAVE_POSTCONFIG` and `oracle_after_postconfig=True` by
default, so validate.py runs Save first, kills LO, then `cp` over the
result file before the evaluator pulls. The one exception is
`eval_kind == "highlighted_words"` (F_WRITER_114/115 .odt fixtures):
LO's Ctrl+S drops the `fo:background-color="#ffff00"` span on .odt
roundtrip, so those rows skip postconfig entirely (`postconfig=[]` and
`oracle_after_postconfig=False`) — see `_build_highlighted_words_evaluator`
and `_to_synth_template._params` for the bypass wiring.

This file is the §I dataclass form ONLY. Legacy `_make_*_row` factories
have been deleted as part of the Batch file-first refactor — File.src +
FileTask.gold callables now wire source / gold / evaluator construction
directly. §I expands across ~30 Files spanning all 11 writer eval_class
buckets (bold_text, change_font, change_line_spacing, add_header_footer,
highlight_text, insert_image, color_table_text, pdf_export, text_to_table,
blank_table_insert, footnote_citation).

Asset infra retained between line 444 (TEMPLATES decl) and line 800 (§I
header) consists of pure heredoc generators + the `_GUTENBERG_BOOKS`
catalog. They have no top-level emission — FileTask callables compose
them.

Usage:
    uv run python -m lite.gym.envs.lite.osworld.src.gen.train \\
        --track synth --domain libreoffice_writer
"""

from __future__ import annotations

import os
import pathlib
import re
import textwrap
from html.parser import HTMLParser

from lite.gym.envs.lite.osworld.src.gen.common import (
    LO_SAVE_POSTCONFIG,
)
from lite.gym.envs.lite.osworld.src.gen.train.synth._utils import (
    SynthTemplate,
    _stage_asset,
)
from lite.gym.envs.lite.osworld.src.utils.assets import asset_root

_DESKTOP = "/home/user/Desktop"


# ---------------------------------------------------------------------------
# Body-focus post-open step — runs AFTER launch + activate_window so the
# agent's first Ctrl+A actually selects body text (not the first-run Release
# Notes banner). See common.py post_open_config_steps hook for invocation.
# ---------------------------------------------------------------------------

# LO post-launch settle (validation — Trigger F race fix). Owned by the LO
# domain per the per-domain setup convention; previously lived in
# common.py and was applied via a `template.domain == "libreoffice_*"`
# conditional.
_LO_POSTLAUNCH_SETTLE: list[dict] = [
    {"type": "sleep", "parameters": {"seconds": 6}},
    {"type": "activate_window", "parameters": {"window_name": "LibreOffice"}},
]


# Batch lesson: LO Writer launches with focus on the Release Notes banner
# (first-run notification). Agent's first Ctrl+A doesn't select body — it
# selects within the banner. Pure-keyboard focus shift so we don't depend
# on resolution-specific screen coordinates: activate the window, then F6
# cycles focus to the document pane, and Ctrl+End anchors the caret in
# body. (The Release Notes infobar itself should be suppressed at the
# Dockerfile / xcu seed level — deferred.)
# validation: prefixed with `_LO_POSTLAUNCH_SETTLE` so the LO domain owns its
# own post-launch sequence (no `common.py` LO-special-casing).
#
# UNCERTAIN: the pure-keyboard fix shifts caret focus to
# body but does NOT dismiss the LO 7.3 Release Notes infobar itself, which
# sits at y≈155-190 and intercepts any subsequent toolbar click at
# y∈[170,180]. ≥9 Writer rows in Validation showed infobar still visible
# at the agent's first screenshot. The proper fix is Dockerfile xcu seed
# (`/etc/skel/.config/libreoffice/4/user/registrymodifications.xcu` or
# `bootstraprc ReleaseNotesURL=`) — deferred.
# Until then, agents that rely on coord-targeted toolbar clicks may still
# fail with this trigger.
_WRITER_BODY_FOCUS_STEPS: list[dict] = [
    *_LO_POSTLAUNCH_SETTLE,
    {"type": "execute", "parameters": {
        "command": (
            "WID=$(xdotool search --name 'LibreOffice Writer' 2>/dev/null | head -1); "
            "if [ -n \"$WID\" ]; then xdotool windowactivate --sync \"$WID\"; fi; true"
        ),
        "shell": True,
    }},
    {"type": "key", "parameters": {"key": "Escape"}},
    {"type": "key", "parameters": {"key": "F6"}},
    {"type": "sleep", "parameters": {"seconds": 0.3}},
    {"type": "key", "parameters": {"key": "ctrl+End"}},
    {"type": "sleep", "parameters": {"seconds": 0.3}},
]


# ---------------------------------------------------------------------------
# Pre-config helpers — heredoc into execute step
# ---------------------------------------------------------------------------

def _execute(command: str, *, shell: bool = True, **extra) -> dict:
    return {"type": "execute", "parameters": {"command": command, "shell": shell, **extra}}


def _make_config_step(py_code: str) -> dict:
    """Wrap `py_code` in a python3 heredoc executed inside the container."""
    return _execute(f"python3 << 'PYEOF'\n{py_code}\nPYEOF")


# ---------------------------------------------------------------------------
# Source-file heredoc bodies — small inline-text helpers that emit a docx
# with a fixed body (no Title, no Heading) so body_idxs[i] == i.
# ---------------------------------------------------------------------------

# Inline body pool — short, neutral paragraphs whose text content is largely
# task-agnostic. Each tuple is (genre_tag, paragraphs_list). The opening
# clause of paragraph 0 is used as the {anchor} substitution in
# quote-anchored row instructions.
# Each genre is exactly 6 paragraphs (idx 0..5). Paragraph count is the
# contract every FileTask param_idx depends on. Within each paragraph we
# write 2-3 sentences in the voice of a real document excerpt (a memo,
# essay draft, service procedure, recipe page, briefing note, runbook)
# rather than one isolated declarative sentence, so the rendered docx
# reads like genuine prose. Token anchors used by find_replace and
# first_centered tasks are preserved verbatim: 'training' and 'Managers'
# (capital M, sentence-initial in para 3) appear in policy;
# 'Phillips' and 'ribbon cable' appear in manual;
# recipe paragraph 0 still leads with the preheat step.
_BODY_POOL: list[tuple[str, list[str]]] = [
    ("policy", [
        "Effective immediately, all employees must complete the annual data-security training before the close of the current quarter. The course is mandatory for full-time staff, contractors, and any vendor with persistent access to internal systems.",
        "Course materials are hosted on the internal learning portal under the Compliance section. The portal will accept submissions for two weeks from this notice; manual extensions are reviewed only when supported by a documented scheduling conflict.",
        "Managers are responsible for monitoring completion rates for their direct reports and for escalating any pattern of missed deadlines to the people-operations partner assigned to their division. A weekly dashboard is circulated each Monday.",
        "Questions about scheduling or content scope can be directed to the people-operations mailbox at any time. The team aims to respond within one business day; tickets opened on Friday afternoon may roll over to the following Monday.",
        "Failure to complete the training on schedule will result in a temporary suspension of system access until the certificate is on file. Suspended accounts are restored automatically within one hour of the portal recording completion.",
        "We thank you in advance for the time you set aside to keep our customers' data safe and the company compliant with its regulatory obligations. Your cooperation is what makes our annual SOC-2 attestation possible.",
    ]),
    ("essay", [
        "The morning light filtered through the high windows of the reading room and fell across the worn oak desks in long, slanted bars. Dust drifted lazily through it, the way it always seems to in old libraries on quiet weekdays.",
        "She set her stack of books down on the corner of the nearest desk and walked over to the catalog cabinet, fingers running along the wood of the drawer fronts as she searched out the right letter.",
        "Outside, a gentle rain had begun to fall, and the sound of it striking the slate roof three stories up was steady and soothing — the kind of weather, she thought, that argued strongly against ever leaving the building.",
        "Most of the other patrons had already drifted away to drier corners of the city, and the long tables that had been full an hour ago were empty now save for an elderly man bent over a stack of bound newspapers near the window.",
        "She had come in meaning to finish a chapter she had been struggling with for weeks, but found herself drawn instead to the slim, unfamiliar volumes on the shelves nearest the catalog, the ones that nobody seemed to take down very often.",
        "Hours passed before she finally chose one and carried it up to the front desk to be stamped. The librarian smiled without looking up from her own book, the way librarians sometimes do when they recognise a fellow afternoon escapee.",
    ]),
    ("manual", [
        "Begin by removing the cover plate using a Phillips screwdriver and setting the four screws aside on the anti-static mat. The cover plate is a friction fit once the screws are out; lift it straight up rather than prying from one corner.",
        "Disconnect the ribbon cable carefully from its socket on the main board. The connector is mechanically keyed and will only seat one way, so note the orientation of the red stripe before withdrawing the cable for later reference.",
        "Inspect the internal fan blades for visible dust accumulation and clean with a short burst of compressed air. Hold the fan hub with a fingertip while you spray so the bearings are not driven beyond their rated speed.",
        "Reseat the ribbon cable, taking care that the keyed end engages fully before any downward pressure is applied. The cable should bottom out with light thumb pressure; if it resists, withdraw and check the orientation again before retrying.",
        "Replace the cover plate and tighten each of the four screws to roughly hand tight. Do not overtighten — the standoffs are pressed-fit into a polymer chassis and stripped threads here are the leading cause of warranty returns.",
        "Power the unit on at the line switch and confirm that the front-panel indicator LED transitions from amber to green within ten seconds. If the LED remains amber for longer than thirty seconds, return to step two and re-verify the ribbon cable seat.",
    ]),
    ("recipe", [
        "Preheat the oven to 350 degrees Fahrenheit and lightly butter a nine-inch round cake pan, then dust the buttered surface with a thin layer of flour and tap out the excess over the sink.",
        "In a medium bowl, whisk together the flour, baking powder, baking soda, and salt until evenly combined. Sifting is not required for this cake, but a brisk whisk for ten or fifteen seconds helps break up any clumps.",
        "In a separate larger bowl, cream the softened butter with the sugar at medium speed until the mixture noticeably lightens in colour and the volume nearly doubles, scraping the sides of the bowl once or twice along the way.",
        "Add the eggs one at a time, beating thoroughly between additions and scraping the bowl as needed. The batter may briefly look curdled after the second egg; this is normal and resolves once the dry ingredients are folded in.",
        "Fold the dry ingredients into the wet in three additions, alternating with the buttermilk in two additions, and beginning and ending with the dry. Stop mixing as soon as the last streaks of flour disappear.",
        "Pour the batter into the prepared pan, smooth the top with the back of a spoon, and bake for thirty-five minutes or until a tester inserted in the centre comes out clean. Cool in the pan for ten minutes before turning out onto a wire rack.",
    ]),
    ("brief", [
        "The committee convened on Tuesday morning to review the second-quarter financial summary and confirm the working assumptions for the half-year reforecast. All voting members were present in person; two non-voting observers joined by video.",
        "Revenue exceeded the internal forecast by nearly four percent, driven primarily by stronger renewals in the mid-market subscription tier. Three of the five geographic regions came in ahead of plan; one region was flat and one slightly behind.",
        "Operating expenses came in slightly above budget owing to one-time legal fees associated with the Q1 acquisition. Excluding those fees, run-rate expenses tracked within one percent of plan, in line with prior quarters.",
        "Cash reserves remain comfortable for the planned investment programme. The treasury team has identified two additional short-term investment vehicles offering yields modestly above the current sweep account and recommends moving twenty million dollars by month end.",
        "Hiring is broadly on track for the calendar year, with twelve open requisitions concentrated in engineering and customer success. Two roles in regulated-markets sales remain unfilled after a second search round and will be re-scoped by talent partners.",
        "The next review is scheduled for the third week of the following month, with materials due to the committee two business days in advance. Any pre-reads or motions for that meeting should be submitted to the corporate secretary by close of business Thursday.",
    ]),
    ("guide", [
        "This runbook explains how to configure the new shared inbox for the customer-support team. It assumes you already have administrator rights in the directory service and that the routing-rule console is reachable from your workstation.",
        "First, create the distribution group in the directory service and add every support agent as a member. Use the naming convention `support-<region>-shared` so that downstream automation can identify the group by prefix without an explicit allow-list.",
        "Next, grant send-as permission to the group so that outgoing messages appear to originate from the customer-facing support address rather than from any individual agent's mailbox. Permission changes can take up to fifteen minutes to propagate to the mail flow service.",
        "Then, configure the routing rule that forwards every incoming ticket from the public address to the shared inbox. Keep the rule strictly additive — leaving the original mailbox as a recipient guarantees a recoverable copy if the routing rule is later disabled.",
        "Finally, verify the setup by sending a test message from an external account and confirming receipt in the agent dashboards within two minutes. The test message should appear under the shared label and should not trigger any spam-bin classifier.",
        "Document the change in the operations runbook so that the next on-call engineer can audit the configuration during the quarterly review. Include the directory-service group ID, the routing-rule UID, and a screenshot of the verification email.",
    ]),
]


# ---------------------------------------------------------------------------
# Wikipedia paragraph extractor (real-source content for §I FileTasks).
#
# Loads cleaned `<p>` text from each HTML article in
# `assets/synth/html/wikipedia/`. The result is a module-level cache keyed by
# article slug ("coffee" → list[str]). Used by `_src_wiki_*` builders below
# AND by the Batch length-variation pass on `_LONG_BODY_*` constants
# (paragraph[3] of six selected long-body lists is now a real Wikipedia
# paragraph) — the cache must therefore be loaded before those constants.
#
# Stdlib HTMLParser only — no external deps. We skip <script>, <style>,
# <sup> (footnote refs), <table>, <figure>, <figcaption>, <noscript>; strip
# `[123]` footnote markers; collapse whitespace; drop paragraphs <80 chars.
# ---------------------------------------------------------------------------

_WIKIPEDIA_DIR = asset_root() / "html" / "wikipedia"


class _WikiParagraphParser(HTMLParser):
    """Extract clean <p> text, skipping non-prose containers."""
    _SKIP_TAGS = {"script", "style", "sup", "table", "figure",
                  "figcaption", "noscript"}

    def __init__(self) -> None:
        super().__init__()
        self._skip_depth = 0
        self._in_p_depth = 0
        self._cur_chunks: list[str] = []
        self.paragraphs: list[str] = []

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP_TAGS:
            self._skip_depth += 1
            return
        if self._skip_depth:
            return
        if tag == "p":
            self._in_p_depth += 1

    def handle_endtag(self, tag):
        if tag in self._SKIP_TAGS:
            if self._skip_depth:
                self._skip_depth -= 1
            return
        if self._skip_depth:
            return
        if tag == "p" and self._in_p_depth:
            self._in_p_depth -= 1
            if self._in_p_depth == 0:
                txt = "".join(self._cur_chunks)
                txt = re.sub(r"\[\d+\]", "", txt)
                txt = re.sub(r"\s+", " ", txt).strip()
                if len(txt) >= 80:
                    self.paragraphs.append(txt)
                self._cur_chunks = []

    def handle_data(self, data):
        if self._skip_depth or not self._in_p_depth:
            return
        self._cur_chunks.append(data)


def _load_wiki_paragraphs() -> dict[str, list[str]]:
    """Parse every Wikipedia HTML in `_WIKIPEDIA_DIR` once at import time."""
    out: dict[str, list[str]] = {}
    for path in sorted(_WIKIPEDIA_DIR.glob("*.html")):
        slug = path.stem
        parser = _WikiParagraphParser()
        parser.feed(path.read_text(encoding="utf-8"))
        out[slug] = parser.paragraphs
    return out


_WIKI_ARTICLES: dict[str, list[str]] = _load_wiki_paragraphs()


# ---------------------------------------------------------------------------
# Long-form realistic bodies — each paragraph is 50-150 words, not the short
# 1-line snippets in _BODY_POOL. Used by §I files that benefit from real
# document length (research notes, technical reports, novel-style prose,
# regulatory text).
# ---------------------------------------------------------------------------

_LONG_BODY_RESEARCH_REPORT: list[str] = [
    "Quarterly research progress report — the cosmic ray spectrum group has completed the first analysis pass on the 2025 detector calibration data. Initial fits to the proton flux above 10 GeV agree with the published AMS-02 reference within statistical uncertainty, confirming that the new gain-correction pipeline preserves spectral shape across the full rigidity range. Two issues remain open: a 1.4 percent normalization offset in the helium channel that correlates with the silicon-tracker temperature, and a small but persistent excess in the deuteron arrival fraction below 5 GeV which we are tracking as item TR-217 in the issue log.",
    "On the instrumentation side, the three new front-end boards passed thermal-vacuum qualification last month and have been integrated into the spare flight unit. We found that the threshold drift after a full thermal cycle stays within twelve millivolts, which is comfortably inside the requirement of fifty millivolts. The remaining concern is the spike rate observed during the warm-up transient, which appears to be related to charge injection on the hold capacitor; the analog group will run a dedicated bench test next week with a slow ramp profile to characterize it.",
    "The simulation team has finalized the new geometry description for the upgraded calorimeter and the first-pass material budget shows a 6 percent reduction in shower leakage compared with the legacy build. We are now running a one-million-event Monte Carlo sample to confirm that the energy resolution at one TeV remains below 2 percent. The output will feed directly into the response matrices used by the analysis pipeline. We expect the sample to complete by the end of next week, and a preliminary plot will be circulated at the Tuesday collaboration meeting.",
    "Outreach activities continue at the planned cadence. The lab open day in early April drew approximately 320 visitors, the highest attendance since the pre-pandemic series, and the student questionnaire returned a strong positive response on the cleanroom tour. Two undergraduate research projects are starting this term, both supervised jointly with the physics department, and a draft of the freshman-seminar reading list has been shared with the curriculum committee. We propose to repeat the open day in October and to expand the volunteer roster by reaching out to the postdoc community in advance.",
    "Looking ahead to the next quarter, the priorities are: complete the calibration write-up, file the upgrade proposal for the replacement front-end boards, and begin the formal review of the analysis pipeline. We will also revisit the budget projection for the 2026 detector campaign in light of the new vendor quotes received last month, which were on the whole favourable. The detailed schedule and the revised milestone table are attached as appendices for review at the next steering meeting.",
]

_LONG_BODY_TRAVEL_GUIDE: list[str] = [
    "Welcome to the Pacific Coast itinerary. This printed guide accompanies the seven-day driving tour from Monterey to the Olympic Peninsula and is intended as a companion to the route map and the booking confirmations included in your travel folder. Each section that follows describes one travel day: an overview of the route, recommended stops, suggested timing, and a short food note for lunch and dinner. Distances are given in miles, drive times assume conservative weekday traffic, and all suggested stops are accessible without advance reservation unless explicitly noted in the body of the entry.",
    "Day one — Monterey to Big Sur. This is a short driving day designed to ease into the trip and let you spend the afternoon walking on the cliffs above Pfeiffer Beach. Plan to depart Monterey by mid-morning, take Highway 1 south through Carmel-by-the-Sea, and stop briefly at the Bixby Bridge viewpoint where the parking turnout fills quickly after eleven. Continue south to the inn, drop bags, and walk down to the beach for the afternoon. Dinner is at the inn restaurant; the seasonal tasting menu is recommended over the standard à-la-carte selection.",
    "Day two — Big Sur to San Simeon. The drive south from Big Sur threads through some of the most photographed coastline on the route, and the road narrows in several places where storm damage has reduced shoulder width. Drive defensively, allow approximately three hours including stops, and plan a short walk at Julia Pfeiffer Burns State Park to see McWay Falls. Arrive at San Simeon in time for the afternoon Hearst Castle tour; the Grand Rooms tour is the most accessible introduction if this is your first visit. Evening: a casual dinner at the harbour-side seafood restaurant, walking distance from the inn.",
    "Day three — San Simeon to Cambria and inland to Paso Robles. Today the route turns inland for the first time, leaving the coast at Cambria and climbing into the wine country east of the highway. The drive is approximately ninety minutes, with several worthwhile stops along the way. Recommended winery visits are listed on the next page; please call ahead for tasting reservations as many of the smaller producers no longer accept walk-ins on weekends. Plan to arrive at the lodge by late afternoon. Dinner is the chef's-table prix-fixe menu at the lodge restaurant — confirmation included in the folder.",
    "Day four is a rest day in Paso Robles. The town centre has a walkable historic district with several independent bookshops and coffee roasters, and the lodge concierge can arrange a half-day cycling tour through the eastern vineyards if you want a structured activity. Otherwise, please feel free to spend the day at your own pace. The pool and the spa are open from eight in the morning until ten at night, and a light lunch menu is available at the poolside bar from eleven to four. Dinner is on your own; the concierge has a list of recommended restaurants in the historic district.",
    "Days five through seven take the route north along the inland Highway 101 corridor, with overnight stops in Carmel Valley, Mendocino, and Crescent City. Detailed driving notes for each leg are on the next three pages. The final leg from Crescent City to Port Angeles is the longest single drive of the trip — please leave early and plan a substantial lunch stop in Coos Bay. The ferry across the Columbia River runs hourly during daylight; weekday morning crossings are reliable. Welcome to your trip, and please contact the trip coordinator with any questions before departure.",
]

_LONG_BODY_PRODUCT_MANUAL: list[str] = [
    "Product manual — Model TR-7 portable thermal recorder. Read this document in full before the first use of the instrument. The TR-7 is a battery-powered four-channel temperature data logger intended for laboratory and field use in the range of negative forty to one hundred and fifty degrees Celsius. The instrument is shipped factory-calibrated and meets the published accuracy specification of plus or minus one tenth of a degree across the full range. Annual recalibration is recommended for laboratories operating under accredited quality programs.",
    "Section one — unpacking and inspection. The shipping carton contains the recorder itself, a quick-start card, a USB-C cable, four type-K thermocouple probes with two-metre leads, a soft carrying case, and the printed warranty card. Inspect each item for shipping damage immediately upon receipt and report any visible defects to the supplier within ten business days; later claims may not be honoured. Retain the original carton and inner packaging for at least sixty days in case the unit needs to be returned during the initial commissioning period.",
    "Section two — power and battery. The TR-7 ships with a fully charged lithium-ion battery rated for approximately seventy-two hours of continuous logging at the default ten-second sample interval. To charge the unit, connect the included USB-C cable to any standard 5V USB power source rated at least 1.5 amperes. The charge indicator on the front panel will pulse amber while charging and remain solid green when full. Do not operate the instrument outside the published temperature range; doing so may permanently damage the battery cells and will void the warranty.",
    "Section three — connecting probes. Each of the four channels accepts a standard miniature thermocouple connector. The polarity is keyed and the connector will only seat fully in the correct orientation. Insert each probe gently and confirm that the channel indicator on the display turns green before starting a logging session. If a channel reads as 'open' or shows an obviously incorrect value, check the connector seat first and the probe continuity second. Spare probes are available from the supplier under part number TR-K2 and ship within two business days when ordered before noon Pacific time.",
    "Section four — starting and stopping a logging session. To begin a session, press and hold the centre button for two seconds. The display will show the current sample interval, the channel count, and an estimated battery runtime; confirm with a short press to begin recording. To stop a session, press and hold the centre button again for two seconds. The instrument will save the recorded data to internal flash memory and return to the idle screen. Recorded sessions can be downloaded over USB using the free TR-7 desktop application available from the supplier's website.",
    "Section five — data download and analysis. Connect the TR-7 to a host computer using the included USB-C cable. The desktop application will detect the unit automatically and present a list of completed sessions. Select one or more sessions and click 'Export' to save the data as CSV or as a binary capture file. The CSV format is human-readable and compatible with spreadsheet software; the binary format preserves full sample resolution and is recommended for downstream signal-processing pipelines. Detailed format specifications are documented in appendix B of this manual.",
]

_LONG_BODY_NOVEL_OPENING: list[str] = [
    "The morning train from Edinburgh was running late, as it almost always was on the first Monday after a holiday weekend, and Marianne had settled into her seat by the window with the resigned patience of someone who had made the same journey many times before. She watched the low Scottish hills slide past in the pale autumn light and thought, not for the first time, about how strange it was to be travelling south on a day that should by rights have been a working day in town. The conductor had not yet appeared. A thin rain began to fall, leaving long diagonal streaks on the carriage window.",
    "She had not expected to be travelling at all this week. The telegram had arrived on Saturday afternoon, brief and formal in the way that telegrams from her uncle's solicitors had always been, and she had spent most of Sunday packing and re-packing the small leather valise that now sat on the rack above her head. The funeral was on Wednesday, the reading of the will the morning after, and her uncle's house — a gloomy brick affair in a Yorkshire village she had not visited since she was a child — would be hers to dispose of by the end of the week. She did not yet know what she felt about any of it.",
    "Across the carriage from her sat a thin man of perhaps sixty, with a closely trimmed white beard and the air of a country doctor, reading a newspaper that he held very close to his face. He had not spoken since she boarded at Waverley Station, and she had not encouraged conversation. A young woman in a green coat was knitting in the seat opposite, glancing up occasionally at the rain. The conductor came at last, examined her ticket without comment, and continued on his way down the carriage. The train picked up speed as it left the last of the Borders behind.",
    "Marianne unwrapped the cheese sandwich she had bought at the station and ate it slowly, watching the slow change of the landscape from soft hill country to the flatter, more cultivated land of the north of England. She tried to remember her uncle's face and found, with a small twinge of guilt, that she could not. He had been a remote, bookish figure on the edge of her childhood, present at one or two family Christmases and then absent for the rest of her growing up. The last letter she had received from him was perhaps three years old, and she could not recall what it had said.",
    "By the time the train crossed the Tyne and rolled into Newcastle the rain had stopped, and a thin sun was breaking through the cloud over the river. Marianne checked her watch, calculated that she would still make the connecting train south with twenty minutes to spare, and gathered her things. The thin man with the beard folded his newspaper and stood up to leave at the same time, nodded politely as he passed her, and disappeared into the bustle of the platform. She lifted her valise from the rack and stepped down onto the long curved platform of Newcastle Central, into the noise and motion of a busy junction station.",
]

_LONG_BODY_POLICY_LONGFORM: list[str] = [
    "Information security policy — version 4.2, effective from the first of June 2026. This document supersedes all previous versions of the company information security policy and applies without exception to every employee, contractor, and authorised third party who accesses company systems or handles company data. Failure to comply with the provisions described in this document may result in disciplinary action up to and including termination of employment or contract, and where applicable may also be referred to law enforcement for investigation under the relevant national and international statutes. Questions about interpretation should be directed to the chief information security officer.",
    "Section one — scope and definitions. For the purposes of this policy, 'company data' means any information generated, received, or stored by the company in the course of its business activities, regardless of the medium on which it is recorded or the location at which it is stored. 'Authorised third party' means any external organisation or individual that has been granted access to company systems or data under a written agreement signed by an officer of the company. 'Sensitive data' means a defined subset of company data identified in appendix A as requiring elevated handling controls.",
    "Section two — access control. All access to company systems must be authenticated using a unique user identifier issued by the IT department and a strong password that meets the complexity requirements published in the password standard. Multi-factor authentication is mandatory for all remote access and for all access to systems handling sensitive data. Shared accounts are prohibited without explicit written approval from the chief information security officer. Access privileges must be reviewed at least quarterly by the relevant system owner and must be revoked promptly when no longer required for the user's role.",
    "Section three — data classification and handling. All company data must be classified at the time of creation according to the four-level scheme described in appendix A: public, internal, confidential, and restricted. Each classification level is associated with a defined set of handling requirements covering storage, transmission, retention, and disposal. Confidential and restricted data must be encrypted at rest and in transit using algorithms and key lengths approved by the cryptographic standards committee. Personal data is subject to additional controls as set out in the company privacy policy and the applicable data protection regulations.",
    "Section four — incident response. Any actual or suspected security incident must be reported to the security operations centre within one hour of discovery, using the contact details published on the company intranet. A security incident includes, but is not limited to, the unauthorised access to or disclosure of company data, the loss or theft of any device containing company data, and the discovery of malware or other unauthorised software on company systems. Employees must cooperate fully with any subsequent investigation. Retaliation against any individual reporting an incident in good faith is strictly prohibited.",
    "Section five — review and revision. This policy will be reviewed at least annually by the chief information security officer in consultation with the executive committee, and revised as necessary to reflect changes in the threat landscape, applicable regulations, and the company's business environment. Significant revisions will be communicated to all affected personnel through the company-wide announcement channels at least thirty days before they take effect. The current version of the policy is always available on the security portal of the company intranet, alongside the supporting standards and procedures referenced in this document.",
]

# ---------------------------------------------------------------------------
# Catch-up: additional long-form bodies for genre diversity.
# Each list is 4-6 paragraphs, each paragraph 50-150 words. These feed the
# structured-source files (with title + headings) used by the 80-file target.
# ---------------------------------------------------------------------------

# Length-variation pass: paragraph[1] replaced with a short
# (30-50 word) note, paragraph[3] replaced with a real Wikipedia paragraph
# from the internet-of-things article (data/computing topic adjacency).
# Originals at indices 0, 2, 4 preserved (paragraph[0] anchors the
# find_replace token 'lakehouse').
_LONG_BODY_WHITE_PAPER: list[str] = [
    "Executive summary — this white paper examines the operational impact of the proposed shift from a centralised data-warehouse architecture to a federated lakehouse model across the four largest divisions of the firm. Drawing on six months of measured workload data from the production reporting cluster and a parallel proof-of-concept lakehouse environment, we conclude that the federated model offers material query-latency improvements for the analyst workload but introduces new operational risk around schema governance that must be addressed before any production rollout is recommended.",
    "Scope note: this paper covers the four largest divisions only; smaller divisions remain on the legacy stack. A separate planning paper will revisit the smaller-division roadmap once the headline rollout is complete.",
    "Findings from the proof-of-concept are summarised in section three. The headline result is a thirty-eight percent reduction in median query latency for the analyst-workload sample, accompanied by a comparable reduction in storage cost per terabyte once the lakehouse object-store tier is fully amortised. These savings are partially offset by an estimated fifteen percent increase in compute cost for the heavy ETL workload, which we attribute to the duplication of effort between the legacy ingest path and the new lakehouse staging layer during the transition period.",
    _WIKI_ARTICLES["internet-of-things"][0],
    "Recommendations and next steps are detailed in section five. We propose a twelve-month rollout plan with three explicit decision gates, each tied to a measurable readiness criterion. The total estimated investment over the rollout period is within the planning envelope agreed in the Q3 strategy review, and we expect the model to reach steady-state operating cost parity with the legacy platform within twenty-four months of the production cutover. The detailed financial model and the readiness-criteria tracking sheet are attached as appendices D and E.",
]

# Batch length-variation pass: paragraph[1] replaced with a short note,
# paragraph[3] replaced with a real Wikipedia paragraph from the bicycle
# article (urban-transport topic adjacency). Originals preserved at 0, 2, 4
# (paragraph[0] anchors the find_replace token 'Northbridge').
_LONG_BODY_CASE_STUDY: list[str] = [
    "Background — the city of Northbridge had operated its public-transit fare-collection system on the same magnetic-stripe technology since the mid-nineteen-nineties, and by twenty-twenty the operating cost of maintaining the legacy stripe readers had grown to consume approximately four percent of the transit authority's annual operating budget. This case study describes the eighteen-month program that replaced the stripe system with a contactless smart-card and mobile-payment platform, and analyses the lessons learned for transit authorities considering a similar transition.",
    "Procurement note: the RFP was issued in March and the contract was awarded in October the same year. Detailed procurement metrics are out of scope for this study and are reported separately in the procurement office annual review.",
    "Implementation proceeded in three phases over the twelve months following contract award. Phase one delivered the back-office platform and ran in parallel with the legacy system for ninety days to allow side-by-side reconciliation of every transaction. Phase two replaced the gate hardware on the busiest fifteen stations during a four-week overnight campaign, with passenger volumes monitored continuously and a defined rollback plan for any station whose passenger throughput dropped below ninety-five percent of the historical baseline. Phase three completed the rollout on the remaining sixty-eight stations on a station-by-station schedule.",
    _WIKI_ARTICLES["bicycle"][1],
    "Lessons learned for similar transitions are summarised in the closing section. The single most important takeaway is the value of the parallel-running phase: although it added three months to the project timeline, the side-by-side reconciliation surfaced a calibration issue in one ticket-type mapping that would have caused substantial revenue leakage if the legacy system had been switched off on the original schedule. We strongly recommend that any comparable transition include an explicit parallel-run phase of no less than ninety days.",
]

# Batch length-variation pass: paragraph[1] replaced with a short note,
# paragraph[3] replaced with a real Wikipedia paragraph from the
# renewable-energy article (industrial-strategy topic adjacency).
# Originals preserved at 0, 2, 4 (paragraph[0] anchors 'Eastfield').
_LONG_BODY_PRESS_BRIEFING: list[str] = [
    "Press briefing — quarterly results announcement, Eastfield Manufacturing Group, fourth quarter and full year ending thirty-first December. This document is provided for accredited financial press only and is embargoed until eight o'clock GMT on the publication date printed on the cover. All figures quoted are unaudited and prepared on the same accounting basis as previous quarterly disclosures. A detailed reconciliation to the audited annual statements will be published in the annual report scheduled for release in the first week of March.",
    "Note to editors: full-year revenue, margin, and cash-flow tables are appended to this briefing as exhibits one through three. Selected exhibits are available on request in spreadsheet form via the investor-relations team.",
    "Divisional performance — the industrial-equipment division had its strongest year on record, with revenue growth of eleven percent and a margin expansion of one hundred and twenty basis points. The consumer-products division grew revenue at three percent with a flat margin, broadly in line with the guidance issued at the third-quarter update. The services division saw a small revenue decline of one percent against an exceptionally strong prior year, but maintained margin and continued to win significant new contract awards including the previously announced multi-year framework agreement with Ridgemont Logistics.",
    _WIKI_ARTICLES["renewable-energy"][2],
    "Closing remarks — the chief executive will host a webcast for accredited press at nine thirty GMT on the announcement morning to take questions on the results. Press kits, including high-resolution photographs of the senior leadership team and product-launch images for the new industrial-equipment range, are available on the investor-relations area of the corporate website. For media enquiries please contact the corporate communications team using the details printed on the cover sheet of this briefing document.",
]

# Batch length-variation pass: paragraph[1] replaced with a short note,
# paragraph[3] replaced with a real Wikipedia paragraph from
# renewable-energy (energy-market topic adjacency). Originals preserved at
# 0, 2, 4 (paragraph[0] anchors the find_replace token 'precision sensors').
_LONG_BODY_MARKET_ANALYSIS: list[str] = [
    "Market analysis — the global market for industrial precision sensors reached an estimated eighteen point six billion United States dollars in the most recent reporting year, growing at a compound annual rate of just over six percent over the preceding five-year period. Growth has been driven principally by the continued adoption of factory-floor automation across the developed industrial economies, supplemented by the more recent emergence of substantial industrial-automation programmes in the upper-middle-income economies of South-East Asia. We expect this growth profile to continue into the medium term.",
    "Methodology note: market sizing is based on a bottom-up product-segment build, cross-checked against published industry-association revenue data and the financial filings of the four largest international suppliers.",
    "Competitive landscape — the supplier base has consolidated considerably over the analysis period through a sustained pattern of acquisition activity by the four largest international groups. The top four suppliers now account for approximately forty-one percent of the total market by revenue, up from twenty-eight percent five years previously. The remaining market is highly fragmented, with several hundred specialist regional suppliers serving narrow application niches. We expect further consolidation activity over the medium term, particularly in the position-sensor sub-segment where scale advantages are most pronounced.",
    _WIKI_ARTICLES["renewable-energy"][1],
    "Risks and uncertainties — the principal downside risk to the forecast is a sharper-than-expected slowdown in the global capital-goods cycle, which would reduce demand for new factory-automation projects and consequently for the precision sensors that support them. The principal upside risk is a faster-than-expected adoption of advanced robotic assembly in the consumer-electronics manufacturing sector, which would lift demand for position sensors above our base-case projection. Detailed scenario analysis is provided in section six.",
]

# Batch length-variation pass: paragraph[1] replaced with a short note,
# paragraph[3] replaced with a real Wikipedia paragraph from
# internet-of-things (industrial networking topic adjacency). Originals
# preserved at 0, 2, 4 (paragraph[0] anchors the 'Aurora-3' find_replace).
_LONG_BODY_TECH_SPEC: list[str] = [
    "Technical specification — Aurora-3 industrial network gateway, hardware revision C, firmware version four point one. This document is the authoritative specification for the Aurora-3 gateway and supersedes all previous revisions. It is intended for use by integrators and operations engineers responsible for deploying and maintaining the gateway in production environments. Changes from the previous revision are summarised in appendix A and are flagged in the body of this document by a vertical change-bar in the right margin.",
    "Compliance summary: the unit conforms to the IEC class-A industrial-EMC directive and the published RoHS substance restrictions. Detailed compliance certificates are appended as appendix E of this specification.",
    "Section two — electrical interfaces. The unit is powered from a single redundant pair of input feeds, each accepting nominal one hundred to two hundred and forty volts AC at fifty or sixty hertz. Maximum power consumption under full load is one hundred and twenty watts; typical idle power consumption is approximately forty-five watts. The chassis ground stud on the rear panel must be bonded to the local equipment-room ground bus using a conductor of at least four square millimetre cross-section. Detailed power-supply derating curves are provided in appendix B.",
    _WIKI_ARTICLES["internet-of-things"][2],
    "Section four — software and protocol support. The Aurora-3 supports the full set of routing and tunnelling protocols listed in appendix D and is interoperable with all three of the major industry-standard network controllers. Configuration management is supported via the published REST and gRPC interfaces; the legacy CLI interface remains available for diagnostic and recovery purposes only and is not recommended for production change-control workflows. Firmware updates are delivered via the standard signed-image distribution channel; rollback is supported within a single major version boundary.",
]

_LONG_BODY_ORG_ANNOUNCEMENT: list[str] = [
    "All-hands announcement — organisational changes effective from the start of the next financial quarter. As communicated at the quarterly leadership review last week, the executive committee has approved a set of organisational changes intended to better align the engineering and product organisations with the company's medium-term strategic priorities. This note describes the changes, the rationale behind them, and the practical implications for individual employees, and it is being distributed to all staff simultaneously to ensure consistent information across the organisation.",
    "The principal change is the consolidation of the platform-engineering and infrastructure-engineering organisations into a single unified platform group reporting to the chief technology officer. The new group will be led by Marina Petrov, currently vice president of platform engineering, who will assume the new title of vice president of platform and infrastructure with effect from the first of the new quarter. The previous vice president of infrastructure engineering, Dr Hassan Akbari, will move to a newly created strategic advisor role reporting to the chief executive officer.",
    "The product organisation will see a corresponding realignment, with the previously separate product-management and product-design functions being brought together under a single chief product officer. We are pleased to confirm that Lila Bergstrom has accepted the role of chief product officer and will join the executive committee with effect from the same date. Lila brings over eighteen years of product leadership experience from her previous roles at three highly successful enterprise-software companies, and her formal biography is attached to this announcement as a separate document.",
    "Practical implications for individual employees are limited in the immediate term. Reporting lines for individual contributors will not change as part of this announcement; the changes operate at the leadership-team layer only. The detailed organisational chart, including the revised director-level reporting lines that will follow over the next thirty days, will be published on the people-portal as soon as it has been finalised. Individual transition discussions for the small number of director-level roles affected by the changes are already under way and will be completed before the effective date.",
    "Next steps — the executive committee will host an all-hands video conference at four o'clock local time on Friday to take questions on the announcement, with simultaneous translation provided in the languages listed on the invitation. The recording and a written summary of the questions and answers will be made available on the all-hands channel within forty-eight hours of the live session. Questions in advance can be submitted anonymously through the standard all-hands questions channel; we will prioritise those questions in the live session that have received the highest community vote.",
]

# Batch length-variation pass: paragraph[1] replaced with a short note,
# paragraph[3] replaced with a real Wikipedia paragraph from earth
# (planetary-science topic adjacency). Originals preserved at 0, 2, 4
# (paragraph[0] anchors any urban-heat-island-related find_replace).
_LONG_BODY_ACADEMIC_ABSTRACT: list[str] = [
    "Abstract — we report on a longitudinal study of urban-microclimate variability across three medium-sized European cities over a continuous five-year measurement period from twenty-twenty to twenty-twenty-four. Data were collected from a network of one hundred and fifty calibrated weather stations distributed across the three study cities, and supplemented by satellite-derived land-surface temperature observations at one-kilometre spatial resolution. The dataset captures three complete annual cycles of measurement at each station, after exclusion of station-years with greater than ten percent data loss.",
    "Funding declaration: this study was supported by an inter-institutional grant from the participating universities and by a complementary instrument grant from the regional research council.",
    "Results — across the three study cities the urban-heat-island effect averaged two point four degrees Celsius during summer night-time conditions, with substantial inter-city and intra-city variability. The strongest urban heat-island signal was observed in the most densely built central districts and was correlated most strongly with the local impervious-surface fraction. Proximity to large water bodies provided a measurable cooling effect of up to one degree Celsius within the first kilometre of shoreline, decaying approximately exponentially with distance inland.",
    _WIKI_ARTICLES["earth"][1],
    "Conclusions — the five-year measurement programme provides a high-quality multi-city dataset that we believe will be of value to the broader urban-climate research community. The full dataset is being prepared for public release through the standard climate-data archive and will be available under a Creative Commons licence. Future work will extend the analysis to investigate the relationship between observed urban heat-island intensity and the public-health outcomes recorded in parallel by the participating municipal health authorities, subject to ethical approval.",
]

_LONG_BODY_POLICY_MEMO: list[str] = [
    "Policy memo — proposed amendments to the company travel-and-expense policy, prepared for the executive committee review session scheduled for the second Tuesday of next month. This memo summarises the principal proposed amendments, the rationale behind each, and the expected operational impact. A redlined draft of the full revised policy is attached for detailed review; please send any substantive comments to the policy office by the close of business on the Friday before the executive review session so they can be incorporated into the final draft.",
    "The principal proposed amendments fall into three categories: alignment with the new corporate sustainability commitments, simplification of the per-diem rate structure, and clarification of the approval-threshold rules. The sustainability-alignment changes introduce a preferred-rail recommendation for European intercity journeys of less than four hours and require explicit business-case justification for any short-haul flight where a comparable rail option exists. These changes implement the carbon-reduction commitments published in the most recent corporate sustainability report.",
    "The per-diem simplification reduces the existing twenty-three-tier rate structure to a six-tier structure based on broad geographical regions rather than on individual city-by-city rates. The simplified structure preserves the overall expense envelope at the corporate level — modelled cost impact is plus or minus one percent on a like-for-like basis — but materially reduces the administrative burden on travelling employees and on the expense-processing team. A detailed cost-impact analysis is included as appendix B of the attached redline draft.",
    "The approval-threshold clarifications address an ambiguity in the current policy that has produced a small but persistent stream of escalations to the policy office over the past eighteen months. The proposed amendment establishes clear single-approver thresholds for routine expense categories and explicit dual-approver requirements for the small set of higher-risk categories identified by the internal-audit team during their twenty-twenty-three review. The clarifications do not change any approval threshold downwards from current practice.",
    "Implementation timeline and communications — assuming executive approval at the scheduled review session, the revised policy would take effect from the start of the new financial quarter. The policy office and the people-operations team will jointly run an awareness campaign during the four-week window between approval and effective date, including manager briefings, an updated self-service expense portal, and refreshed training materials in the standard learning-management system. Detailed communications timeline and owner assignments are appendix C of the attached redline draft.",
]

_LONG_BODY_LONG_ESSAY_A: list[str] = [
    "On the persistence of letters — there is a particular pleasure in receiving an unexpected letter that no electronic medium has ever quite been able to reproduce, and although our age is the most thoroughly connected in human history we still write letters, and other people still keep them. The reasons we keep them are sometimes obvious — the love letters from a marriage long settled, the formal notice of a long-awaited appointment — and sometimes obscure even to ourselves, but the keeping of them seems to be a deep and continuing human practice that resists the steady pressure of digital substitution.",
    "Part of the persistence, I think, lies in the simple physicality of the letter as object. A letter has a particular handwriting, a particular paper, perhaps a postmark from a particular city; it bears in its small way the stamp of a single human hand performing a single deliberate act on a single afternoon. Even the most carefully composed email cannot quite reproduce that quality of singularity. The same is true of a postcard, or a handwritten birthday card, or a thank-you note: the medium itself carries a kind of evidence that is difficult to forge and difficult to copy.",
    "Part of it, too, is the discipline of writing. To sit down to compose a letter, even a short one, is to commit to a particular sustained attention that the ordinary scroll of digital messaging does not require. There is no notification to interrupt the train of thought, no autocomplete to nudge the writer towards a more conventional turn of phrase, no opportunity to delete the first paragraph and try again without leaving evidence of the attempt. The discipline of the medium produces, in the better cases, a quality of writing that the easier medium does not consistently produce.",
    "And part of it is the temporal asymmetry. To send a letter is to commit a thought to a slow channel, and to receive one is to read a thought that the writer committed several days ago, perhaps a week ago, perhaps in another mood than the one the writer is now in. The asymmetry creates, in the better correspondences, a sustained register of considered thought that the immediate-response patterns of digital communication tend to dissolve. The exchange of letters is by its nature a slow art, and slowness in this case is not the absence of speed but a positive aesthetic of its own.",
    "None of this is to argue that the letter will or should displace the more efficient electronic substitutes for the routine business of human contact. The letter as a medium for ordering groceries or for confirming a meeting is plainly inferior, and nobody seriously proposes otherwise. But the letter as a medium for the sustained slow-attention exchange of considered thought between two people seems to me to have a continuing place even in the most relentlessly connected age, and I notice with something like satisfaction that the practice does indeed persist.",
]

_LONG_BODY_LONG_ESSAY_B: list[str] = [
    "On reading old maps — there is a section of my study where I keep, in a flat drawer designed for the purpose, a small collection of old folded road maps from the years before satellite navigation became universal. They are Michelin maps mostly, of various French regions, dating from the late nineteen-eighties through to the early years of this century. I bought most of them when they were new, used them for one driving holiday or another, and brought them home and put them away. They have been in the drawer ever since.",
    "I take them out perhaps twice a year, usually on a wet afternoon when the work I should be doing has become temporarily intolerable, and I spread one out on the kitchen table and look at it. The pleasure is partly that of remembering the particular journey for which the map was bought, and partly that of looking at the map as an object in its own right. A road map of this period is, I think, one of the most successful objects of practical graphic design ever produced. The hierarchy of information is established by colour and weight of line in a way that the eye reads almost without conscious effort.",
    "The maps record, of course, the road network as it existed at the date of printing. The motorways are there, drawn in the characteristic Michelin orange, but the network is sparser than the present one in interesting ways. The bypasses around the medium-sized towns are mostly absent, the high-speed orbital roads around the larger cities are partial, and several of the trans-regional motorway routes that I would now take for granted are shown still as the older single-carriageway national roads they were before upgrade. The maps are, in a small way, a record of the recent transport history of the country.",
    "They record other things too: the location of forests and large parks, the boundaries of the regional natural areas, the distribution of small villages and isolated hamlets in the rural districts. The cartographer has made a series of careful judgements about which features deserve which weight on the page, and the resulting object is dense with information without ever being visually crowded. The same data displayed on a contemporary screen-based map app would be either oversimplified at the macro scale or unreadably busy at the micro scale; the printed map handles the multi-scale problem with an elegance that the screen has so far not equalled.",
    "I do not propose, of course, to navigate by paper map on any actual journey I am likely to take. The convenience of the satellite-navigation app is too great, the real-time traffic information too useful, the integration with the destination booking too convenient. But the paper map remains in its drawer, and is taken out occasionally and looked at with pleasure, and this seems to me to be no more eccentric a use of an obsolete technology than the practice of keeping a small bookcase of physical books in a study where most reading is now done electronically.",
]


def _docx_body_py(out_path: str, paragraphs: list[str], *,
                  font_name: str | None = None,
                  font_size_pt: int | None = None,
                  line_spacing: float | None = None,
                  initial_table: tuple[int, int] | None = None) -> str:
    """Heredoc body that writes a fresh docx with the given body paragraphs.

    No Title, no Heading — every paragraph is Default style so body_idxs[i] == i
    when reading via python-docx. Optional doc-wide font/size/spacing knobs
    let callers shape the visual template. ``initial_table`` (rows, cols)
    appends an empty table after the body — used for files whose tasks DO
    NOT touch the table (table count is then a stable initial-state fixture).
    """
    parts = ["from docx import Document",
             "from docx.shared import Pt",
             "doc = Document()"]
    if font_name is not None:
        # Set doc-default font on the Normal style so LO toolbar / docDefaults
        # report the target font (rather than theme fallback `minorHAnsi` →
        # Cambria), enabling agents to verify font via the toolbar font picker.
        parts.append(f"doc.styles['Normal'].font.name = {font_name!r}")
    for p in paragraphs:
        parts.append(f"para = doc.add_paragraph({p!r})")
        if line_spacing is not None:
            parts.append(f"para.paragraph_format.line_spacing = {line_spacing!r}")
        if font_name is not None or font_size_pt is not None:
            parts.append("for r in para.runs:")
            if font_name is not None:
                parts.append(f"    r.font.name = {font_name!r}")
            if font_size_pt is not None:
                parts.append(f"    r.font.size = Pt({font_size_pt})")
    if initial_table is not None:
        rows, cols = initial_table
        parts.append(f"doc.add_table(rows={rows}, cols={cols})")
    parts.append(f"doc.save({out_path!r})")
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Evaluator + oracle helpers (mirrors perturb pattern)
# ---------------------------------------------------------------------------

def _lo_normalize_cmd(path: str, fmt: str = "docx") -> str:
    """Headless LO round-trip a docx in place. One-to-one with perturb's
    `_lo_normalize_cmd` at perturb/libreoffice_writer.py:88. Used by
    `_build_oracle` to symmetrize the eval-time normalize chain."""
    return (
        f"tmpd=$(mktemp -d) && "
        f"DISPLAY=:1 soffice --headless --norestore --nofirststartwizard "
        f"--convert-to {fmt} --outdir \"$tmpd\" '{path}' 2>/dev/null && "
        f"[ -f \"$tmpd/$(basename '{path}')\" ] && "
        f"cp \"$tmpd/$(basename '{path}')\" '{path}'; "
        f"rm -rf \"$tmpd\"; true"
    )


def _build_oracle(file_path: str, expected_path: str) -> list[dict]:
    """Oracle: ① normalize gold, ② plant at sink.

    oracle_after_postconfig=True kills LO before running oracle actions and
    sets _postconfig_done=True so LO_SAVE_POSTCONFIG (Ctrl+S) is skipped after
    oracle. Result = cp of norm(gold). Expected = norm(gold). They are
    byte-identical without a third normalize step.  A third normalize would
    produce result = norm(norm(gold)) which diverges when LO normalize is
    non-idempotent (bold, font-size, colour, alignment properties).
    """
    return [
        _execute(_lo_normalize_cmd(expected_path, "docx")),
        _execute(f"cp '{expected_path}' '{file_path}'"),
    ]


def _build_default_font_oracle(font_name: str) -> list[dict]:
    """Write registrymodifications.xcu so find_default_font returns 1.0.

    The default_font evaluator reads the xcu file via result=vm_file. The
    standard _build_oracle only plants a docx, which the evaluator ignores.
    This oracle writes the minimal xcu fragment directly to the container
    using a python heredoc — using `python3 -c "..."` corrupts the XML
    because XML attribute quotes (`version="1.0"`) collide with the outer
    shell-level double quotes.
    """
    xcu = "/home/user/.config/libreoffice/4/user/registrymodifications.xcu"
    py = (
        "import os, pathlib\n"
        f"p = {xcu!r}\n"
        "pathlib.Path(os.path.dirname(p)).mkdir(parents=True, exist_ok=True)\n"
        f"font = {font_name!r}\n"
        "xml = (\n"
        "    '<?xml version=\"1.0\" encoding=\"UTF-8\"?>'\n"
        "    '<oor:items xmlns:oor=\"http://openoffice.org/2001/registry\">'\n"
        "    '<item oor:path=\"/org.openoffice.Office.Writer/DefaultFont\">'\n"
        "    '<prop oor:name=\"Standard\"><value>' + font + '</value></prop>'\n"
        "    '</item>'\n"
        "    '</oor:items>'\n"
        ")\n"
        "open(p, 'w').write(xml)\n"
    )
    return [_execute(f"python3 << 'PYEOF'\n{py}\nPYEOF")]


def _build_strict_evaluator(
    file_path: str, expected_path: str, *,
    examine_font_name: bool = False,
    examine_font_size: bool = False,
    examine_color: bool = False,
    examine_highlight: bool = False,
    examine_images: bool = False,
) -> dict:
    """compare_docx_strict evaluator. examine_* defaults to False — LO's
    docx round-trip normalizes font name/size/color/highlight on untouched
    runs and can false-fail. Set the targeted field True at the call site.

    bold/italic/underline/strike are ALWAYS checked by compare_docx_strict
    (no flag controls them) — that's why this evaluator covers the bulk of
    per-paragraph format ops without per-call configuration.
    """
    return {
        "func": "compare_docx_strict",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "expected": {"type": "vm_file", "path": expected_path,
                     "dest": "expected_file"},
        "options": {
            "examine_font_name": examine_font_name,
            "examine_font_size": examine_font_size,
            "examine_color": examine_color,
            "examine_highlight": examine_highlight,
            "examine_images": examine_images,
        },
        "postconfig": LO_SAVE_POSTCONFIG,
    }


def _build_files_evaluator(file_path: str, expected_path: str, **extra_opts: object) -> dict:
    """compare_docx_files — paragraph-text-only comparison.

    Batch lesson: agents commonly press Enter twice before typing appended
    content. delete_empty_lines=True tolerates this without rejecting the
    semantically-correct append (find/replace + append_paragraph rows).

    Pass extra_opts to override or extend the default options dict (e.g.
    ignore_blanks=False for sentence-per-line tasks where paragraph splitting
    must be visible to the evaluator).
    """
    return {
        "func": "compare_docx_files",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "expected": {"type": "vm_file", "path": expected_path,
                     "dest": "expected_file"},
        "options": {"delete_empty_lines": True, **extra_opts},
        "postconfig": LO_SAVE_POSTCONFIG,
    }


def _build_tables_evaluator(file_path: str, expected_path: str) -> dict:
    return {
        "func": "compare_docx_tables",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "expected": {"type": "vm_file", "path": expected_path,
                     "dest": "expected_file"},
        "postconfig": LO_SAVE_POSTCONFIG,
    }


def _build_line_spacing_evaluator(file_path: str, expected_path: str) -> dict:
    return {
        "func": "compare_line_spacing",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "expected": {"type": "vm_file", "path": expected_path,
                     "dest": "expected_file"},
        "postconfig": LO_SAVE_POSTCONFIG,
    }


def _build_first_centered_evaluator(file_path: str) -> dict:
    return {
        "func": "is_first_line_centered",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "postconfig": LO_SAVE_POSTCONFIG,
    }


def _build_page_numbers_evaluator(file_path: str) -> dict:
    return {
        "func": "has_page_numbers_in_footers",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "postconfig": LO_SAVE_POSTCONFIG,
    }


def _build_page_break_evaluator(file_path: str, expected_count: int) -> dict:
    return {
        "func": "contains_page_break",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "expected": {"type": "rule",
                     "rules": {"page_break_count": expected_count}},
        "postconfig": LO_SAVE_POSTCONFIG,
    }


def _build_font_names_evaluator(file_path: str, font_name: str) -> dict:
    # Validation fix: switch from upstream strict `compare_font_names` to
    # lite-side `compare_font_names_loose` (resolves paragraph-style / doc
    # default inheritance). LO save normaliser strips redundant run-level
    # rFonts when the font matches the paragraph/Normal style → upstream
    # returns 0 even when agent's Ctrl+A + Format → Character correctly
    # changed the visible font. Eval-side jsonl tasks still use upstream
    # strict via their own evaluator dicts; this switch only affects
    # synth doc_font_* templates.
    return {
        "func": "compare_font_names_loose",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "expected": {"type": "rule",
                     "rules": {"font_name": font_name}},
        "postconfig": LO_SAVE_POSTCONFIG,
    }


def _build_default_font_evaluator(font_name: str) -> dict:
    """find_default_font reads LO Writer's registrymodifications.xcu and
    looks up the per-language Default font. Result type=vm_file pulls the
    XML file from the container; expected is a rule with the target font.
    """
    xcu = "/home/user/.config/libreoffice/4/user/registrymodifications.xcu"
    return {
        "func": "find_default_font",
        "result": {"type": "vm_file", "path": xcu,
                   "dest": "registrymodifications.xcu"},
        "expected": {"type": "rule", "rules": {"font_name": font_name}},
        "postconfig": LO_SAVE_POSTCONFIG,
    }


def _build_strike_last_para_evaluator(file_path: str, expected_path: str) -> dict:
    return {
        "func": "evaluate_strike_through_last_paragraph",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "expected": {"type": "vm_file", "path": expected_path,
                     "dest": "expected_file"},
        "postconfig": LO_SAVE_POSTCONFIG,
    }


def _build_italic_size14_evaluator(file_path: str, expected_path: str) -> dict:
    return {
        "func": "check_italic_font_size_14",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "expected": {"type": "vm_file", "path": expected_path,
                     "dest": "expected_file"},
        "postconfig": LO_SAVE_POSTCONFIG,
    }


def _build_colored_table_evaluator(file_path: str, expected_path: str) -> dict:
    """evaluate_colored_words_in_tables — checks per-word RGB colour inside
    table cells. Eval (1 task) uses this; synth previously had 0 coverage at
    the func level (the `color_table_text` bucket actually colored paragraph
    text, not table cells). Source pre-colours cells GREEN to avoid the
    eval's None-colour TypeError crash (per env-side `synth_inject_colored_
    table` lesson at lite/gym/envs/lite/osworld/src/gen/train/synth/libreoffice_writer.py).
    """
    return {
        "func": "evaluate_colored_words_in_tables",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "expected": {"type": "vm_file", "path": expected_path,
                     "dest": "expected_file"},
        "postconfig": LO_SAVE_POSTCONFIG,
    }


def _build_subscript_contains_evaluator(file_path: str, expected_path: str) -> dict:
    """compare_subscript_contains — verifies at least one (run1, run2) pair
    in some paragraph where both runs are subscript. Eval (1 task) checks
    chemistry-formula style subscript; synth previously had 0 coverage of
    character-level subscript (compare_docx_strict has no `examine_subscript`
    flag — only bold/italic/underline/strike are auto-checked, plus 5
    examine_* flags for font_name/size/color/highlight/images). The broader
    character-level format axis was flagged separately.
    """
    return {
        "func": "compare_subscript_contains",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "expected": {"type": "vm_file", "path": expected_path,
                     "dest": "expected_file"},
        "postconfig": LO_SAVE_POSTCONFIG,
    }


def _build_tabstops_evaluator(file_path: str, expected_path: str,
                              **opts) -> dict:
    """check_tabstops — verifies tab-stop positions match the gold doc. Eval
    (1 task) uses this; synth had 0 coverage. Compares paragraph_format.
    tab_stops (positions + alignments) between result and expected docx.
    """
    return {
        "func": "check_tabstops",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "expected": {"type": "vm_file", "path": expected_path,
                     "dest": "expected_file"},
        "options": opts,
        "postconfig": LO_SAVE_POSTCONFIG,
    }


def _build_unique_train_records_evaluator(file_path: str,
                                          gold_path: str,
                                          initial_path: str) -> dict:
    """compare_unique_train_records — verifies the processed doc has a
    correct set of unique train_ids (col-1 of CSV-ish rows) matching the
    gold, and every processed line was present in the initial doc. Eval
    (1 task) uses this; synth had 0 coverage. Three-file evaluator: gold +
    initial are passed as `expected[0]` and `expected[1]`.
    """
    return {
        "func": "compare_unique_train_records",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "expected": [
            {"type": "vm_file", "path": gold_path,
             "dest": "expected_gold.docx"},
            {"type": "vm_file", "path": initial_path,
             "dest": "expected_initial.docx"},
        ],
        "postconfig": LO_SAVE_POSTCONFIG,
    }


def _build_images_evaluator(file_path: str, expected_path: str) -> dict:
    """compare_docx_images — verifies the embedded images in the docx
    binary-match the gold's embedded images. Mirrors eval task
    osworld_libreoffice_writer_6ada715d ("copy screenshot 1.png at cursor").
    Synth previously had ZERO coverage of this upstream func — image
    insertion rows used compare_docx_strict + examine_images=True instead.
    """
    return {
        "func": "compare_docx_images",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "expected": {"type": "vm_file", "path": expected_path,
                     "dest": "expected_file"},
        "postconfig": LO_SAVE_POSTCONFIG,
    }


def _build_highlighted_words_evaluator(file_path: str, expected_path: str) -> dict:
    """check_highlighted_words — REAL upstream evaluator from osworld_libre
    office_writer_6a33f9b9 ("remove all highlight"). REQUIRES both files
    to be .odt because the evaluator calls `odf.opendocument.load()`
    internally; docx inputs make it raise. The evaluator first checks
    text equality via `compare_docx_files`, then walks every Span and
    fails if ANY span's automatic-style backgroundcolor == '#ffff00'.

    Use only for `.odt` source+gold pairs (see `_build_odt_oracle`).

    Postconfig is INTENTIONALLY EMPTY (not LO_SAVE_POSTCONFIG). LO Writer's
    Ctrl+S save drops the `fo:background-color="#ffff00"` highlight span
    when LO rewrites the .odt on roundtrip — pre-eval would then read a
    source.odt with no highlight visible and `check_highlighted_words`
    would return 1.0 (trivial_pass). Skipping postconfig keeps the
    zipfile-built source.odt byte-stable on disk so pre-eval sees the
    yellow highlights (returns 0 = not trivial); oracle then cp's the
    gold-no-highlight over and post-eval reads gold (returns 1.0 = PASS).
    Paired with `oracle_after_postconfig=False` in `_to_synth_template`.
    """
    return {
        "func": "check_highlighted_words",
        "result": {"type": "vm_file", "path": file_path,
                   "dest": file_path.split("/")[-1]},
        "expected": {"type": "vm_file", "path": expected_path,
                     "dest": "expected_file"},
        "postconfig": [],
    }


def _build_odt_oracle(file_path: str, expected_path: str) -> list[dict]:
    """Oracle for .odt source/gold pairs. Same pattern as `_build_oracle`
    but normalizes with `--convert-to odt`. Two steps only — see
    `_build_oracle` docstring for why a third normalize step is omitted.
    """
    return [
        _execute(_lo_normalize_cmd(expected_path, "odt")),
        _execute(f"cp '{expected_path}' '{file_path}'"),
    ]


def _build_compound_evaluator(file_path: str, atoms: list[dict]) -> dict:
    """Compound multi-property evaluator. `atoms` is a list of dicts each
    of shape {"func": str, "expected_path": str|None, "options": dict}.

    Builds a `func`-as-list / `expected`-as-list shape (eval's compound
    style). Eval has 5 compound rows (mostly same-fn ×N at multiple targets,
    plus 1 truly multi-fn — H2O subscript+files). conj=None means AND.
    All atoms share the SAME result file (the agent's edited docx) since
    every property is computed on the final saved doc.
    """
    funcs: list[str] = []
    expecteds: list[dict] = []
    options: list[dict] = []
    for atom in atoms:
        funcs.append(atom["func"])
        ep = atom.get("expected_path")
        if ep is not None:
            expecteds.append({"type": "vm_file", "path": ep,
                              "dest": f"expected_{len(expecteds)}.docx"})
        options.append(atom.get("options", {}))
    return {
        "func": funcs,
        "result": [
            {"type": "vm_file", "path": file_path,
             "dest": file_path.split("/")[-1]}
            for _ in funcs
        ],
        "expected": expecteds if expecteds else None,
        "options": options if any(options) else None,
        "postconfig": LO_SAVE_POSTCONFIG,
    }


# ---------------------------------------------------------------------------
# All row factories now live in the §I dataclass form at the bottom of this
# file. Real-asset infra (_GUTENBERG_BOOKS + heredoc helpers, image-insert
# heredoc helpers, PDF-export helpers) below this point provides the
# building blocks that §I FileTasks reference.
# ---------------------------------------------------------------------------

TEMPLATES: list[SynthTemplate] = []


# ===========================================================================
# Real-asset infra — `_stage_asset`-backed Gutenberg-text + photo helpers.
# These heredoc generators + asset catalogs are referenced by the §I
# FileTask entries at the bottom of this file (image-insert files, gutenberg
# excerpt files, double-image files, PDF-export files). The legacy
# `_make_*_row` factories that wrapped them have been deleted as part of
# the §I migration — File.src + FileTask.gold callables now provide that
# wiring directly.
# ===========================================================================

# Per-book anchors — the Gutenberg files have ~50-line headers (Project
# Gutenberg start markers, illustration blocks, TOC, preface). We locate
# the first body paragraph by substring search on a known opening clause,
# then take the next N paragraphs (split on blank lines).
_GUTENBERG_BOOKS: list[dict[str, str]] = [
    {
        "key": "alice",
        "asset_rel": "docs/gutenberg/alice-in-wonderland.txt",
        "anchor": "Alice was beginning to get very tired",
        "display": "Alice's Adventures in Wonderland",
    },
    {
        "key": "pride",
        "asset_rel": "docs/gutenberg/pride-and-prejudice.txt",
        "anchor": "It is a truth universally acknowledged",
        "display": "Pride and Prejudice",
    },
    {
        "key": "moby",
        "asset_rel": "docs/gutenberg/moby-dick.txt",
        "anchor": "Call me Ishmael",
        "display": "Moby Dick",
    },
    {
        "key": "frank",
        "asset_rel": "docs/gutenberg/frankenstein.txt",
        "anchor": "You will rejoice to hear",
        "display": "Frankenstein",
    },
    {
        "key": "suntzu",
        "asset_rel": "docs/gutenberg/art-of-war.txt",
        "anchor": "Sun Tzŭ said:",
        "display": "The Art of War",
    },
    {
        "key": "tale2",
        "asset_rel": "docs/gutenberg/tale-of-two-cities.txt",
        "anchor": "It was the best of times",
        "display": "A Tale of Two Cities",
    },
    {
        "key": "meta",
        "asset_rel": "docs/gutenberg/metamorphosis.txt",
        "anchor": "One morning, when Gregor Samsa",
        "display": "The Metamorphosis",
    },
    {
        "key": "sherlock",
        "asset_rel": "docs/gutenberg/sherlock-adventures.txt",
        "anchor": "To Sherlock Holmes she is always",
        "display": "The Adventures of Sherlock Holmes",
    },
    {
        "key": "treasure",
        "asset_rel": "docs/gutenberg/treasure-island.txt",
        "anchor": "Squire Trelawney, Dr. Livesey",
        "display": "Treasure Island",
    },
    {
        "key": "tom",
        "asset_rel": "docs/gutenberg/tom-sawyer.txt",
        "anchor": "The old lady pulled her spectacles down",
        "display": "The Adventures of Tom Sawyer",
    },
    {
        "key": "earnest",
        "asset_rel": "docs/gutenberg/importance-of-being-earnest.txt",
        "anchor": "Lane is arranging afternoon tea",
        "display": "The Importance of Being Earnest",
    },
]

_GUTENBERG_BOOKS_BY_KEY = {b["key"]: b for b in _GUTENBERG_BOOKS}


def _gutenberg_source_heredoc(
    *, book_txt_path: str, anchor: str, n_paras: int, out_docx: str,
) -> str:
    """Heredoc body that reads a staged Gutenberg .txt, finds the anchor
    paragraph, takes the next `n_paras` body paragraphs, and writes a
    fresh docx (no Title/Heading — every para is Default style).

    Mirrors `_docx_body_py` shape but pulls text from the staged file at
    container time instead of an inline list literal.
    """
    return textwrap.dedent(f"""\
        import re
        from docx import Document
        with open({book_txt_path!r}) as _f:
            _text = _f.read()
        _idx = _text.find({anchor!r})
        assert _idx >= 0, "anchor not found in " + {book_txt_path!r}
        _pstart = _text.rfind('\\n\\n', 0, _idx)
        if _pstart < 0:
            _pstart = 0
        _text = _text[_pstart:].strip()
        _raw = re.split(r'\\n\\s*\\n', _text)
        _paras = []
        for _p in _raw:
            _p = re.sub(r'\\s+', ' ', _p).strip()
            if not _p or len(_p) < 60 or '[Illustration' in _p:
                continue
            _paras.append(_p)
            if len(_paras) >= {n_paras}:
                break
        _doc = Document()
        for _para in _paras:
            _doc.add_paragraph(_para)
        _doc.save({out_docx!r})
        """)


def _gutenberg_gold_heredoc(*, source_docx: str, gold_docx: str, op: str) -> str:
    """Heredoc body that opens the staged source docx, applies one
    paragraph-0 format op, and saves to the gold path. Mirrors
    `_row_para_format._gold` op-switch but reads paragraphs from disk
    instead of regenerating from a list literal — guarantees source/gold
    paragraph text byte-equivalence even if the Gutenberg parser drops a
    paragraph differently between calls.
    """
    op_lines = {
        "bold":             "for r in p0.runs: r.font.bold = True",
        "italic":           "for r in p0.runs: r.font.italic = True",
        "underline":        "for r in p0.runs: r.font.underline = True",
        "strike":           "for r in p0.runs: r.font.strike = True",
        "size18":           "for r in p0.runs: r.font.size = Pt(18)",
        # training validation: agent's natural palette click for "red"
        # lands on LO's "Red" swatch (#CC0000), not pure #FF0000. Switch gold
        # to palette-matched value so the easy path passes; "Custom Color #FF0000"
        # path also still passes if agent goes that route. Same for blue.
        "color_red":        "for r in p0.runs: r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)",
        "color_blue":       "for r in p0.runs: r.font.color.rgb = RGBColor(0x00, 0x00, 0xCC)",
        "font_georgia":     "for r in p0.runs: r.font.name = 'Georgia'",
        "highlight_yellow": "for r in p0.runs: r.font.highlight_color = WD_COLOR_INDEX.YELLOW",
        "size14":           "for r in p0.runs: r.font.size = Pt(14)",
    }
    if op not in op_lines:
        raise ValueError(f"unknown gutenberg op: {op}")
    return textwrap.dedent(f"""\
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_COLOR_INDEX
        _doc = Document({source_docx!r})
        p0 = _doc.paragraphs[0]
        {op_lines[op]}
        _doc.save({gold_docx!r})
        """)


# Per-row spec for Gutenberg format-op rows.
# Each: (template_id, book_key, op, n_paras, examine_kwargs, instr_template).
# ---------------------------------------------------------------------------
# Image-insert helpers — heredoc generators for photo-staging FileTasks.
# ---------------------------------------------------------------------------

# Per-row spec: (template_id, asset_rel, body_kind, instr_subject, width_in)
def _image_source_heredoc(*, body_paras: list[str], src_docx: str) -> str:
    """Build a docx whose body is `body_paras` (no Title, no Heading)."""
    lines = ["from docx import Document", "_doc = Document()"]
    for p in body_paras:
        lines.append(f"_doc.add_paragraph({p!r})")
    lines.append(f"_doc.save({src_docx!r})")
    return "\n".join(lines)


def _image_gold_heredoc(
    *, src_docx: str, gold_docx: str, image_path: str, width_in: float,
    insert_after_idx: int,
) -> str:
    """Open source docx, insert the staged image as a new run on a new
    paragraph after `insert_after_idx`, save to gold. We use python-docx
    `add_picture` on a fresh document then re-emit body in the right
    order — simpler than mid-document insertion via XML splicing.

    add_picture appends to the end; to put the image AFTER paragraph i,
    we rebuild the doc with body[:i+1] + image + body[i+1:].
    """
    # Batch lesson: LO Insert → Image GUI re-encodes image bytes (drawingML
    # wrapping) vs python-docx add_picture (inline image, raw bytes).
    # examine_images=True via per-image sha256 rejects. Solution: soffice docx
    # round-trip on gold so its image bytes are also LO-encoded → symmetric eval.
    return textwrap.dedent(f"""\
        from docx import Document
        from docx.shared import Inches
        import os as _os_img, tempfile as _tf_img
        from PIL import Image as _PIL
        _src = Document({src_docx!r})
        _texts = [p.text for p in _src.paragraphs]
        _new = Document()
        _i = {insert_after_idx}
        for _t in _texts[:_i + 1]:
            _new.add_paragraph(_t)
        _tmp_png = _tf_img.mktemp(suffix='.png')
        try:
            _PIL.open({image_path!r}).save(_tmp_png, 'PNG')
            _new.add_picture(_tmp_png, width=Inches({width_in!r}))
        finally:
            if _os_img.path.exists(_tmp_png):
                _os_img.unlink(_tmp_png)
        for _t in _texts[_i + 1:]:
            _new.add_paragraph(_t)
        _new.save({gold_docx!r})
        import os as _os, subprocess as _sp, tempfile as _tf, shutil as _sh
        _td = _tf.mkdtemp()
        try:
            _r = _sp.run(['soffice', '--headless', '--norestore', '--nofirststartwizard',
                          '--convert-to', 'docx', '--outdir', _td, {gold_docx!r}],
                         capture_output=True, env=_os.environ, timeout=120)
            _conv = _os.path.join(_td, _os.path.basename({gold_docx!r}))
            if _os.path.exists(_conv):
                _sh.copy(_conv, {gold_docx!r})
        finally:
            _sh.rmtree(_td, ignore_errors=True)
        """)


# ---------------------------------------------------------------------------
# Double-image helper — gold heredoc for two-photo FileTasks.
# ---------------------------------------------------------------------------

# Per-row spec: (template_id, two assets, body_paras with two anchor captions,
# width_in, instructions). The gold inserts image_a after paragraph index 1
# and image_b after the (now-shifted) caption-b paragraph.
def _double_image_gold_heredoc(
    *, src_docx: str, gold_docx: str,
    image_path_a: str, image_path_b: str, width_in: float,
) -> str:
    """Open source docx, insert image_a after the paragraph whose text starts
    with 'Caption A' and image_b after 'Caption B', save to gold. Mirrors
    `_image_gold_heredoc` but rebuilds with two pictures at anchor positions
    instead of a single fixed index.
    """
    # Batch lesson: LO Insert → Image GUI re-encodes image bytes (drawingML
    # wrapping) vs python-docx add_picture (inline image, raw bytes).
    # examine_images=True via per-image sha256 rejects. Solution: soffice docx
    # round-trip on gold so its image bytes are also LO-encoded → symmetric eval.
    return textwrap.dedent(f"""\
        from docx import Document
        from docx.shared import Inches
        import os as _os_img, tempfile as _tf_img
        from PIL import Image as _PIL
        _src = Document({src_docx!r})
        _texts = [p.text for p in _src.paragraphs]
        _idx_a = next(i for i, t in enumerate(_texts) if t.startswith('Caption A'))
        _idx_b = next(i for i, t in enumerate(_texts) if t.startswith('Caption B'))
        assert _idx_a < _idx_b, 'Caption A must precede Caption B'
        def _to_png(src):
            p = _tf_img.mktemp(suffix='.png')
            _PIL.open(src).save(p, 'PNG')
            return p
        _tmp_a = _to_png({image_path_a!r})
        _tmp_b = _to_png({image_path_b!r})
        try:
            _new = Document()
            for _i, _t in enumerate(_texts):
                _new.add_paragraph(_t)
                if _i == _idx_a:
                    _new.add_picture(_tmp_a, width=Inches({width_in!r}))
                elif _i == _idx_b:
                    _new.add_picture(_tmp_b, width=Inches({width_in!r}))
            _new.save({gold_docx!r})
        finally:
            for _p in (_tmp_a, _tmp_b):
                if _os_img.path.exists(_p):
                    _os_img.unlink(_p)
        import os as _os, subprocess as _sp, tempfile as _tf, shutil as _sh
        _td = _tf.mkdtemp()
        try:
            _r = _sp.run(['soffice', '--headless', '--norestore', '--nofirststartwizard',
                          '--convert-to', 'docx', '--outdir', _td, {gold_docx!r}],
                         capture_output=True, env=_os.environ, timeout=120)
            _conv = _os.path.join(_td, _os.path.basename({gold_docx!r}))
            if _os.path.exists(_conv):
                _sh.copy(_conv, {gold_docx!r})
        finally:
            _sh.rmtree(_td, ignore_errors=True)
        """)


# ---------------------------------------------------------------------------
# Gutenberg paragraph-offset heredoc helpers — slice [skip:skip+n] then
# locate target paragraph by sub-anchor substring (F6 ordinal-safety).
# ---------------------------------------------------------------------------

def _gutenberg_offset_source_heredoc(
    *, book_txt_path: str, anchor: str, skip: int, n_paras: int, out_docx: str,
) -> str:
    """Mirrors `_gutenberg_source_heredoc` but slices [skip:skip+n_paras]
    after the anchor instead of [0:n_paras]. Asserts the slice yields at
    least n_paras paragraphs so the resulting docx is well-formed.
    """
    return textwrap.dedent(f"""\
        import re
        from docx import Document
        with open({book_txt_path!r}) as _f:
            _text = _f.read()
        _idx = _text.find({anchor!r})
        assert _idx >= 0, "anchor not found in " + {book_txt_path!r}
        _pstart = _text.rfind('\\n\\n', 0, _idx)
        if _pstart < 0:
            _pstart = 0
        _text = _text[_pstart:].strip()
        _raw = re.split(r'\\n\\s*\\n', _text)
        _paras = []
        for _p in _raw:
            _p = re.sub(r'\\s+', ' ', _p).strip()
            if not _p or len(_p) < 60 or '[Illustration' in _p:
                continue
            _paras.append(_p)
            if len(_paras) >= {skip} + {n_paras}:
                break
        _paras = _paras[{skip}:{skip} + {n_paras}]
        assert len(_paras) == {n_paras}, (
            f"offset slice yielded only {{len(_paras)}} paragraphs (expected {n_paras})"
        )
        _doc = Document()
        for _para in _paras:
            _doc.add_paragraph(_para)
        _doc.save({out_docx!r})
        """)


def _gutenberg_subanchor_gold_heredoc(
    *, source_docx: str, gold_docx: str, sub_anchor: str, op: str,
) -> str:
    """Open source docx, find the paragraph whose text contains
    `sub_anchor`, apply one format op to that paragraph (all runs), save.
    """
    op_lines = {
        "bold":             "for r in pT.runs: r.font.bold = True",
        "italic":           "for r in pT.runs: r.font.italic = True",
        "underline":        "for r in pT.runs: r.font.underline = True",
        "strike":           "for r in pT.runs: r.font.strike = True",
    }
    if op not in op_lines:
        raise ValueError(f"unknown gutenberg-offset op: {op}")
    return textwrap.dedent(f"""\
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_COLOR_INDEX
        _doc = Document({source_docx!r})
        pT = next(p for p in _doc.paragraphs if {sub_anchor!r} in p.text)
        {op_lines[op]}
        _doc.save({gold_docx!r})
        """)


# Per-row spec: (template_id, book_key, skip, n_paras, sub_anchor, op,
# examine_kwargs, instructions). The sub_anchor MUST appear in the
# [skip:skip+n_paras] slice; if not, the gold heredoc will raise StopIteration.
# ---------------------------------------------------------------------------
# PDF export helpers — soffice docx→pdf conversion + 4-dir target list.
# ---------------------------------------------------------------------------
# Pattern:
#   pre_config: build source docx via python-docx heredoc, then convert it
#     to a gold PDF via `soffice --headless --convert-to pdf` into
#     `/tmp/expected_<id>.pdf`. The agent's task is GUI File > Export As >
#     PDF saving to <docx-stem>.pdf in any of 4 standard dirs.
#   evaluator: list-form `["compare_pdfs"]*4` with `conj=or` and 4
#     result/expected pairs (mirrors eval 4bcb1253). Same gold PDF is
#     reused 4× on the expected side; result side is the 4 standard target
#     directories. compare_pdfs computes fuzz_ratio over extracted text;
#     same-source-same-converter → identical text → score 1.0.
#   oracle: convert the docx to PDF (already done in pre_config — we cp the
#     pre-built expected PDF) and copy to the Desktop sink path. The `or`
#     conjunction means a single matching pair passes.
#   trivial-pass guard: pre_config does NOT pre-create any of the agent's
#     4 sink PDF paths. Without oracle, eval finds 4 missing files → fails.
#   `oracle_after_postconfig=True`: LO_SAVE_POSTCONFIG kills LO first to
#     release any PDF lock; then oracle's `cp` lands the gold PDF.
# ---------------------------------------------------------------------------


_PDF_EXPORT_TARGET_DIRS = (
    "/home/user/Desktop",
    "/home/user/Documents",
    "/home/user/Downloads",
    "/home/user",
)


def _soffice_docx_to_pdf_step(src_docx: str, expected_pdf: str) -> dict:
    """Headless soffice docx→pdf conversion step. soffice writes
    `<basename>.pdf` to `--outdir`; we mv it to a deterministic location.
    Mirrors `multi_apps._soffice_convert_to_pdf_step`.
    """
    out_dir = os.path.dirname(expected_pdf) or "/tmp"
    src_stem = "$(basename '" + src_docx + "' | sed 's/\\.[^.]*$//')"
    return _execute(
        f"mkdir -p '{out_dir}' && "
        f"soffice --headless --norestore --nofirststartwizard "
        f"--convert-to pdf --outdir '{out_dir}' '{src_docx}' >/dev/null 2>&1 && "
        f"mv '{out_dir}/'{src_stem}'.pdf' '{expected_pdf}'"
    )


# ===========================================================================
# §I. File-task templates (Batch, dataclass form)
#
# Mirrors synth/libreoffice_calc.py + synth/libreoffice_impress.py §I.
# This domain is file-as-topic (no inner TopicTheme rotation): each File
# already encodes both the structural shape AND the content semantics.
#
# Symmetric layout (all synth/*.py):
#   §I.a  Caps                — SYNTH_CAP_TASKS_PER_FILE / _PARAMS_PER_TASK
#   §I.b  Dataclasses         — File / Param / FileTask (frozen)
#   §I.c  File instances      — define each File ONCE
#   §I.d  Gold helpers        — heredoc generators (per-task)
#   §I.e  Factory + emit      — _to_synth_template / _emit_templates
#   §I.f  FILE_TASKS          — flat list, one entry per (file, task) pair
#   §I.g  Emission            — TEMPLATES.extend(_emit_templates(FILE_TASKS))
# ===========================================================================

from dataclasses import dataclass as _I_dataclass, field as _I_field
from typing import Callable as _I_Callable


# §I.a — caps
SYNTH_CAP_TASKS_PER_FILE: int = 2
SYNTH_CAP_PARAMS_PER_TASK: int = 2


# §I.b — Dataclasses (writer-shaped Param: eval_kind discriminates the
# evaluator-builder used; eval_args are extra kwargs for that builder).

@_I_dataclass(frozen=True)
class File:
    """One structurally distinct source docx.

    `src(path, seed) -> list[dict]` returns the pre_config steps that
    materialize the source docx (typically one `_make_config_step` wrapping
    a `_docx_body_py(...)` heredoc).
    """
    id: str
    setup_class: str
    basename: str
    src: _I_Callable[[str, int], list[dict]]


@_I_dataclass(frozen=True)
class Param:
    gold_args: dict       # forwarded to FileTask.gold(src, gold, **gold_args)
    eval_kind: str        # "strict"|"files"|"tables"|"line_spacing"|"page_break"|"font_names"|"default_font"|"first_centered"|"page_numbers"
    eval_args: dict       # extra kwargs for the evaluator builder
    instr: str


@_I_dataclass(frozen=True)
class FileTask:
    file: File
    task_id: str
    eval_class: str
    gold: _I_Callable[..., str]    # (src_path, gold_path, **gold_args) -> heredoc body
    params: list[Param] = _I_field(default_factory=list)


# §I.c — File instances.
#
# Files split into three families by source-build pattern:
#
# (A) Inline-body files (_src_genre / _src_long_body / _src_with_table) —
#     paragraphs are literal strings baked into the heredoc. Most files.
# (B) Gutenberg files (_src_gutenberg) — stage a Project Gutenberg .txt via
#     `_stage_asset`, then slice the first N body paragraphs after a known
#     anchor line. Real-source content per PD (3a).
# (C) Image-host files (_src_image_host) — stage one or two photo jpgs via
#     `_stage_asset`, build a body docx whose tasks include "insert this
#     image". Real-source content per PD (3a).

def _src_genre(genre_idx: int, *,
               font_name: str | None = None,
               font_size_pt: int | None = None,
               line_spacing: float | None = None,
               initial_table: tuple[int, int] | None = None):
    """File.src callable: _BODY_POOL[genre_idx] paragraphs (short snippets)
    with optional doc-wide font/size/spacing/initial-table dressing."""
    paras = _BODY_POOL[genre_idx][1]
    def _src(path, _seed):
        return [_make_config_step(_docx_body_py(
            path, paras, font_name=font_name,
            font_size_pt=font_size_pt, line_spacing=line_spacing,
            initial_table=initial_table))]
    return _src


def _src_long_body(paragraphs: list[str], *,
                   font_name: str | None = None,
                   font_size_pt: int | None = None,
                   line_spacing: float | None = None,
                   initial_table: tuple[int, int] | None = None):
    """File.src callable: long-form paragraphs (each 50-150 words) with
    optional doc-wide font/size/spacing/initial-table dressing."""
    def _src(path, _seed):
        return [_make_config_step(_docx_body_py(
            path, paragraphs, font_name=font_name,
            font_size_pt=font_size_pt, line_spacing=line_spacing,
            initial_table=initial_table))]
    return _src


def _src_gutenberg(book_key: str, n_paras: int = 5):
    """File.src callable: stage a Gutenberg .txt + slice first N body
    paragraphs after the anchor line. Real-source content."""
    book = _GUTENBERG_BOOKS_BY_KEY[book_key]
    def _src(path, _seed):
        book_path = f"/tmp/_book_{book_key}_{os.path.basename(path)}.txt"
        return [
            _stage_asset(book["asset_rel"], book_path),
            _make_config_step(_gutenberg_source_heredoc(
                book_txt_path=book_path, anchor=book["anchor"],
                n_paras=n_paras, out_docx=path,
            )),
        ]
    return _src


def _src_image_host(asset_rel: str, image_basename: str,
                    body_paras: list[str]):
    """File.src callable: stage a real photo at /home/user/Desktop/<basename>
    + build a body-only docx referencing the image. Used by image-insert
    FileTasks whose gold heredoc appends `add_picture(image_path)`."""
    def _src(path, _seed):
        image_path = f"{_DESKTOP}/{image_basename}"
        return [
            _stage_asset(asset_rel, image_path),
            _make_config_step(_image_source_heredoc(
                body_paras=body_paras, src_docx=path,
            )),
        ]
    return _src


def _docx_structured_py(
    out_path: str, *,
    title: str | None = None,
    blocks: list[tuple[str, object]],
    font_name: str | None = None,
    font_size_pt: int | None = None,
    line_spacing: float | None = None,
) -> str:
    """Build a docx with structured blocks (title / headings / paragraphs /
    lists / table). `blocks` is a list of (kind, payload) tuples:

      ("h1", "Section title")             — Heading 1
      ("h2", "Subsection")                — Heading 2
      ("p",  "Body paragraph text...")    — Default paragraph
      ("bullet", "list item text")        — List Bullet
      ("number", "list item text")        — List Number
      ("table", [[r0c0, r0c1], [...]])    — Table with content

    The `title`, if given, is emitted FIRST as a styled title paragraph
    (Heading 0 / Title style). The optional doc-wide font/size/spacing
    knobs are applied to every paragraph block (not to the heading or
    title — those keep their style-defined font).

    Returns python-docx heredoc body that writes the file. Gold helpers
    that target a specific paragraph index must account for the title/
    heading positions in the resulting `doc.paragraphs` list.
    """
    parts = [
        "from docx import Document",
        "from docx.shared import Pt",
        "doc = Document()",
    ]
    if font_name is not None:
        # Set doc-default font on the Normal style so LO toolbar / docDefaults
        # report the target font (rather than theme fallback `minorHAnsi` →
        # Cambria), enabling agents to verify font via the toolbar font picker.
        parts.append(f"doc.styles['Normal'].font.name = {font_name!r}")
    if title is not None:
        parts.append(f"_t = doc.add_heading({title!r}, level=0)")
    for kind, payload in blocks:
        if kind == "h1":
            parts.append(f"doc.add_heading({payload!r}, level=1)")
        elif kind == "h2":
            parts.append(f"doc.add_heading({payload!r}, level=2)")
        elif kind == "p":
            parts.append(f"para = doc.add_paragraph({payload!r})")
            if line_spacing is not None:
                parts.append(f"para.paragraph_format.line_spacing = {line_spacing!r}")
            if font_name is not None or font_size_pt is not None:
                parts.append("for r in para.runs:")
                if font_name is not None:
                    parts.append(f"    r.font.name = {font_name!r}")
                if font_size_pt is not None:
                    parts.append(f"    r.font.size = Pt({font_size_pt})")
        elif kind == "short_p":
            # Same emission as "p" — the kind tag is purely a callsite hint
            # signaling 'this paragraph is intentionally short'. Useful for
            # _src_mixed_length_essay where length variance is the point.
            parts.append(f"para = doc.add_paragraph({payload!r})")
            if line_spacing is not None:
                parts.append(f"para.paragraph_format.line_spacing = {line_spacing!r}")
            if font_name is not None or font_size_pt is not None:
                parts.append("for r in para.runs:")
                if font_name is not None:
                    parts.append(f"    r.font.name = {font_name!r}")
                if font_size_pt is not None:
                    parts.append(f"    r.font.size = Pt({font_size_pt})")
        elif kind == "quote":
            # Italic single-paragraph quote (no special style — Default
            # paragraph with italic run formatting). Used by mixed-length
            # essays to break up dense Wikipedia prose with literary
            # excerpts from Gutenberg.
            parts.append(f"para = doc.add_paragraph({payload!r})")
            if line_spacing is not None:
                parts.append(f"para.paragraph_format.line_spacing = {line_spacing!r}")
            parts.append("for r in para.runs:")
            parts.append("    r.italic = True")
            if font_name is not None:
                parts.append(f"    r.font.name = {font_name!r}")
            if font_size_pt is not None:
                parts.append(f"    r.font.size = Pt({font_size_pt})")
        elif kind == "bullet":
            parts.append(f"doc.add_paragraph({payload!r}, style='List Bullet')")
        elif kind == "number":
            parts.append(f"doc.add_paragraph({payload!r}, style='List Number')")
        elif kind == "table":
            rows = payload  # list[list[str]]
            n_rows = len(rows)
            n_cols = len(rows[0]) if rows else 0
            parts.append(f"_tbl = doc.add_table(rows={n_rows}, cols={n_cols})")
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    parts.append(f"_tbl.rows[{ri}].cells[{ci}].text = {cell!r}")
        else:
            raise ValueError(f"unknown structured block kind: {kind!r}")
    parts.append(f"doc.save({out_path!r})")
    return "\n".join(parts)


def _src_structured(*,
                    title: str | None = None,
                    blocks: list[tuple[str, object]],
                    font_name: str | None = None,
                    font_size_pt: int | None = None,
                    line_spacing: float | None = None):
    """File.src callable: build a structured doc with title + heading +
    paragraphs (+ optional table / list blocks). Heading-aware files use
    this builder rather than `_src_long_body`."""
    def _src(path, _seed):
        return [_make_config_step(_docx_structured_py(
            path, title=title, blocks=blocks,
            font_name=font_name, font_size_pt=font_size_pt,
            line_spacing=line_spacing,
        ))]
    return _src


def _src_structured_with_photo(*,
                               title: str | None,
                               blocks: list[tuple[str, object]],
                               asset_rel: str,
                               image_basename: str,
                               font_name: str | None = None,
                               font_size_pt: int | None = None):
    """File.src callable: stage a real photo + build a structured doc. The
    photo is staged at /home/user/Desktop/<basename>; the doc itself does
    not yet contain the image — image-insert FileTasks add it via gold
    heredoc (consistent with the existing `_src_image_host` pattern)."""
    def _src(path, _seed):
        image_path = f"{_DESKTOP}/{image_basename}"
        return [
            _stage_asset(asset_rel, image_path),
            _make_config_step(_docx_structured_py(
                path, title=title, blocks=blocks,
                font_name=font_name, font_size_pt=font_size_pt,
            )),
        ]
    return _src


def _src_double_image_host(asset_rel_a: str, asset_rel_b: str,
                           image_basename_a: str, image_basename_b: str,
                           body_paras: list[str]):
    """File.src callable: stage TWO photos + build a body docx with
    'Caption A' / 'Caption B' anchor paragraphs."""
    def _src(path, _seed):
        path_a = f"{_DESKTOP}/{image_basename_a}"
        path_b = f"{_DESKTOP}/{image_basename_b}"
        return [
            _stage_asset(asset_rel_a, path_a),
            _stage_asset(asset_rel_b, path_b),
            _make_config_step(_image_source_heredoc(
                body_paras=body_paras, src_docx=path,
            )),
        ]
    return _src


# ---------------------------------------------------------------------------
# Real-source Wikipedia / structural-variant builders. Each pulls cleaned
# Wikipedia paragraphs from `_WIKI_ARTICLES` (or Gutenberg quotes) and lays
# them out in a structurally distinct shape (article-style, structured H1
# sections, short memo, Q&A, bullet reference, mixed-length essay). These
# break the "single-page 4-6 paragraph essay" mold of the original synth
# files and inject natural paragraph-length variance from real sources.
# ---------------------------------------------------------------------------

def _src_wiki_article(article_key: str, n_paras: int = 5, start: int = 0,
                      title: str | None = None,
                      font_name: str | None = None,
                      font_size_pt: int | None = None,
                      line_spacing: float | None = None):
    """File.src callable: Wikipedia paragraphs sliced [start:start+n_paras],
    optional title heading. Like `_src_gutenberg` but for Wikipedia — body
    paragraphs come from `_WIKI_ARTICLES[article_key]` so length varies
    naturally (40-300 words per paragraph).
    """
    paras = _WIKI_ARTICLES[article_key][start:start + n_paras]
    assert len(paras) == n_paras, (
        f"_src_wiki_article: only {len(paras)} paragraphs available for "
        f"{article_key!r} starting at {start} (need {n_paras})"
    )
    blocks: list[tuple[str, object]] = [("p", p) for p in paras]
    def _src(path, _seed):
        return [_make_config_step(_docx_structured_py(
            path, title=title, blocks=blocks,
            font_name=font_name, font_size_pt=font_size_pt,
            line_spacing=line_spacing,
        ))]
    return _src


def _src_wiki_structured(article_key: str, title: str,
                         sections: list[tuple[str, int]],
                         font_name: str | None = None,
                         font_size_pt: int | None = None,
                         line_spacing: float | None = None):
    """File.src callable: Wikipedia content split across H1 sections.

    `sections` is a list of (heading_text, n_paras) — paragraphs are pulled
    consecutively from `_WIKI_ARTICLES[article_key]` into each section.
    Total paragraph count must not exceed the available pool.
    """
    pool = _WIKI_ARTICLES[article_key]
    blocks: list[tuple[str, object]] = []
    cursor = 0
    for heading, n in sections:
        blocks.append(("h1", heading))
        for i in range(n):
            blocks.append(("p", pool[cursor + i]))
        cursor += n
    assert cursor <= len(pool), (
        f"_src_wiki_structured: requested {cursor} paragraphs from "
        f"{article_key!r} (only {len(pool)} available)"
    )
    def _src(path, _seed):
        return [_make_config_step(_docx_structured_py(
            path, title=title, blocks=blocks,
            font_name=font_name, font_size_pt=font_size_pt,
            line_spacing=line_spacing,
        ))]
    return _src


def _src_short_memo(title: str, paragraph_text: str,
                    font_name: str | None = None,
                    font_size_pt: int | None = None):
    """File.src callable: single-paragraph short memo (50-80 words). Title
    on top, one body paragraph below. Used to break the multi-paragraph
    essay default with a deliberately compact document shape.
    """
    blocks: list[tuple[str, object]] = [("p", paragraph_text)]
    def _src(path, _seed):
        return [_make_config_step(_docx_structured_py(
            path, title=title, blocks=blocks,
            font_name=font_name, font_size_pt=font_size_pt,
        ))]
    return _src


def _src_qa_format(title: str, qa_pairs: list[tuple[str, str]],
                   font_name: str | None = None,
                   font_size_pt: int | None = None):
    """File.src callable: title + alternating Q (bold) / A (normal) pairs.

    Each Q is rendered as a bold paragraph (the question), each A as a
    normal-weight paragraph below it (the answer drawn from a real Wikipedia
    source). The bold formatting is baked at codegen time so para_idx-based
    tasks see a deterministic initial bolded layout. Doc-wide ops only —
    individual Q/A targeting is ambiguous so callers MUST avoid para_idx
    tasks on this builder.
    """
    parts = ["from docx import Document",
             "from docx.shared import Pt",
             "doc = Document()"]
    if font_name is not None:
        # Set doc-default font on the Normal style so LO toolbar / docDefaults
        # report the target font (rather than theme fallback `minorHAnsi` →
        # Cambria), enabling agents to verify font via the toolbar font picker.
        parts.append(f"doc.styles['Normal'].font.name = {font_name!r}")
    parts.append(f"_t = doc.add_heading({title!r}, level=0)")
    for q, a in qa_pairs:
        # Question — bold paragraph (single bold run).
        parts.append(f"_q = doc.add_paragraph({q!r})")
        parts.append("for _r in _q.runs:")
        parts.append("    _r.bold = True")
        if font_name is not None:
            parts.append(f"    _r.font.name = {font_name!r}")
        if font_size_pt is not None:
            parts.append(f"    _r.font.size = Pt({font_size_pt})")
        # Answer — normal paragraph below.
        parts.append(f"_a = doc.add_paragraph({a!r})")
        if font_name is not None or font_size_pt is not None:
            parts.append("for _r in _a.runs:")
            if font_name is not None:
                parts.append(f"    _r.font.name = {font_name!r}")
            if font_size_pt is not None:
                parts.append(f"    _r.font.size = Pt({font_size_pt})")
    body = "\n".join(parts)
    def _src(path, _seed):
        full = body + f"\ndoc.save({path!r})\n"
        return [_make_config_step(full)]
    return _src


def _src_bullet_reference(title: str, intro_para: str, bullets: list[str],
                          conclusion_para: str | None = None,
                          font_name: str | None = None,
                          font_size_pt: int | None = None):
    """File.src callable: title + 1 short intro + many bullets (8-15) +
    optional 1 short conclusion. Reference-document shape — distinct from
    the multi-paragraph essay default. Use real Wikipedia paragraphs for
    intro/conclusion. Bullet/paragraph indices interleave ambiguously, so
    callers MUST avoid para_idx tasks on this builder.
    """
    blocks: list[tuple[str, object]] = [("p", intro_para)]
    for b in bullets:
        blocks.append(("bullet", b))
    if conclusion_para is not None:
        blocks.append(("p", conclusion_para))
    def _src(path, _seed):
        return [_make_config_step(_docx_structured_py(
            path, title=title, blocks=blocks,
            font_name=font_name, font_size_pt=font_size_pt,
        ))]
    return _src


def _src_mixed_length_essay(title: str,
                            blocks: list[tuple[str, object]],
                            font_name: str | None = None,
                            font_size_pt: int | None = None,
                            line_spacing: float | None = None):
    """File.src callable: like `_src_structured` but the call site is
    expected to mix short paragraphs (15-40 words: captions, one-line
    statements) with long paragraphs (200-400 words from Wikipedia /
    Gutenberg). `blocks` accepts every existing kind plus `("short_p",
    text)` and `("quote", text)` (italic single-paragraph quote).
    """
    def _src(path, _seed):
        return [_make_config_step(_docx_structured_py(
            path, title=title, blocks=blocks,
            font_name=font_name, font_size_pt=font_size_pt,
            line_spacing=line_spacing,
        ))]
    return _src


# ----- Inline-body files (A) — _BODY_POOL short fixtures -----------------
# These six files retain the historical naming so existing tooling that
# greps `F-WRITER-1..6` keeps resolving.
F_WRITER_1  = File("F-WRITER-1",  "writer_policy",   "policy.docx",
                   _src_genre(0))
F_WRITER_2  = File("F-WRITER-2",  "writer_essay",    "essay.docx",
                   _src_genre(1))
F_WRITER_3  = File("F-WRITER-3",  "writer_manual",   "manual.docx",
                   _src_genre(2))
F_WRITER_4  = File("F-WRITER-4",  "writer_recipe",   "recipe.docx",
                   _src_genre(3))
F_WRITER_5  = File("F-WRITER-5",  "writer_brief",    "brief.docx",
                   _src_genre(4))
F_WRITER_6  = File("F-WRITER-6",  "writer_guide",    "guide.docx",
                   _src_genre(5))

# Styled variants of the short fixtures — each carries a distinct initial
# font / size / spacing / initial-table so the source docx visually differs
# even when the underlying genre fixture is shared.
F_WRITER_7  = File("F-WRITER-7",  "writer_policy_serif", "policy_serif.docx",
                   _src_genre(0, font_name="Liberation Serif"))
F_WRITER_8  = File("F-WRITER-8",  "writer_essay_spaced", "essay_spaced.docx",
                   _src_genre(1, line_spacing=1.5))
F_WRITER_9  = File("F-WRITER-9",  "writer_manual_13pt",  "manual_13pt.docx",
                   _src_genre(2, font_size_pt=13))
F_WRITER_10 = File("F-WRITER-10", "writer_guide_dejavu", "guide_dejavu.docx",
                   _src_genre(5, font_name="DejaVu Sans"))

# ----- Inline long-body files (A) — realistic 50-150 word paragraphs ----
F_WRITER_11 = File("F-WRITER-11", "writer_research_report", "research_report.docx",
                   _src_long_body(_LONG_BODY_RESEARCH_REPORT,
                                  font_name="Liberation Serif", font_size_pt=12))
F_WRITER_12 = File("F-WRITER-12", "writer_travel_guide", "travel_guide.docx",
                   _src_long_body(_LONG_BODY_TRAVEL_GUIDE,
                                  font_name="DejaVu Serif", font_size_pt=11))
F_WRITER_13 = File("F-WRITER-13", "writer_product_manual_long", "product_manual.docx",
                   _src_long_body(_LONG_BODY_PRODUCT_MANUAL,
                                  font_name="Liberation Mono", font_size_pt=11))
F_WRITER_14 = File("F-WRITER-14", "writer_novel_opening", "novel_opening.docx",
                   _src_long_body(_LONG_BODY_NOVEL_OPENING,
                                  font_name="Liberation Serif", font_size_pt=12,
                                  line_spacing=1.5))
F_WRITER_15 = File("F-WRITER-15", "writer_policy_longform", "policy_longform.docx",
                   _src_long_body(_LONG_BODY_POLICY_LONGFORM,
                                  font_name="Liberation Sans", font_size_pt=11))

# ----- Inline files with initial table (A) — for tasks that DO NOT touch
# the table, leaving the initial table as document furniture. -----
F_WRITER_16 = File("F-WRITER-16", "writer_recipe_with_table", "recipe_card.docx",
                   _src_genre(3, font_name="DejaVu Sans", initial_table=(3, 3)))
F_WRITER_17 = File("F-WRITER-17", "writer_policy_with_table", "policy_with_table.docx",
                   _src_genre(0, font_name="Liberation Serif", initial_table=(4, 4)))

# ----- Gutenberg real-source files (B) -----------------------------------
F_WRITER_18 = File("F-WRITER-18", "writer_gutenberg_alice", "alice_excerpt.docx",
                   _src_gutenberg("alice", n_paras=5))
F_WRITER_19 = File("F-WRITER-19", "writer_gutenberg_pride", "pride_excerpt.docx",
                   _src_gutenberg("pride", n_paras=5))
F_WRITER_20 = File("F-WRITER-20", "writer_gutenberg_moby", "moby_excerpt.docx",
                   _src_gutenberg("moby", n_paras=5))
F_WRITER_21 = File("F-WRITER-21", "writer_gutenberg_frank", "frankenstein_excerpt.docx",
                   _src_gutenberg("frank", n_paras=5))
F_WRITER_22 = File("F-WRITER-22", "writer_gutenberg_sherlock", "sherlock_excerpt.docx",
                   _src_gutenberg("sherlock", n_paras=5))
F_WRITER_23 = File("F-WRITER-23", "writer_gutenberg_treasure", "treasure_excerpt.docx",
                   _src_gutenberg("treasure", n_paras=5))
F_WRITER_24 = File("F-WRITER-24", "writer_gutenberg_meta", "metamorphosis_excerpt.docx",
                   _src_gutenberg("meta", n_paras=5))
F_WRITER_25 = File("F-WRITER-25", "writer_gutenberg_earnest", "earnest_excerpt.docx",
                   _src_gutenberg("earnest", n_paras=5))

# ----- Real-photo image-host files (C) -----------------------------------
F_WRITER_26 = File("F-WRITER-26", "writer_image_host_earth", "earth_brief.docx",
                   _src_image_host(
                       "photos/space/earth-blue-marble-apollo17.jpg",
                       "earth-blue-marble-apollo17.jpg",
                       [
                           "This briefing accompanies the cover image, which shows Earth as captured by the Apollo 17 crew on their journey to the Moon. The photograph, taken on December 7 1972, is one of the most widely reproduced images of the planet ever made and remains a touchstone for environmental and educational materials half a century after its creation.",
                           "Insert the photograph below this paragraph so it appears between the description and the closing note. The intended position is at full width on the page; please retain the default aspect ratio and do not crop or rotate the source image during placement.",
                           "Closing note — please review the layout before circulating to the wider distribution list. Any final revisions to the caption or the surrounding text should be sent to the editorial inbox by the end of business on Friday so we can lock the layout ahead of the Monday morning briefing.",
                       ],
                   ))
F_WRITER_27 = File("F-WRITER-27", "writer_image_host_tiger", "wildlife_journal.docx",
                   _src_image_host(
                       "photos/wildlife/tiger-closeup.jpg",
                       "tiger-closeup.jpg",
                       [
                           "Field journal entry — northern reserve, morning patrol along the eastern fence line. The patrol team logged a sighting at approximately oh-six-fifteen local time, recorded a brief observation of feeding behaviour, and departed the area without disturbing the subject. Weather conditions were clear with a light easterly breeze, and visibility was excellent throughout the encounter.",
                           "Place the close-up below this paragraph for the photo plate of the printed journal. The subject was photographed from a distance of approximately forty metres using a long telephoto lens; no attempt was made to approach the animal at any point during the encounter, in accordance with the reserve's standing protocol for large carnivore observations.",
                           "Sighting metadata: GPS coordinates redacted from the printed copy for security reasons but available in the original digital field record; weather summary appended to the daily log; no abnormal behaviour observed; no further action recommended at this time beyond routine continuation of the patrol schedule.",
                       ],
                   ))
F_WRITER_28 = File("F-WRITER-28", "writer_image_host_pizza", "menu_insert.docx",
                   _src_image_host(
                       "photos/food/pizza-dish.jpg",
                       "pizza-dish.jpg",
                       [
                           "Featured menu item — wood-fired Margherita pizza, prepared in the restaurant's authentic Neapolitan oven at a working temperature of approximately four hundred and fifty degrees Celsius. The dough is made from a long-fermented starter and rests for at least forty-eight hours before being shaped to order. Toppings are kept deliberately minimal to allow the quality of the base ingredients to speak for itself.",
                           "Insert the pizza photograph below this paragraph for the printed menu insert. Final layout will set the image at approximately three and a half inches wide to fit alongside the price column on the right-hand side; please retain the source aspect ratio and avoid cropping the visible basil leaves on the upper-left corner of the image.",
                           "Available daily from five o'clock in the evening through to last orders at ten thirty; reservations recommended on Friday and Saturday evenings. The kitchen pairs the dish with a house-selected Sangiovese available by the glass or the carafe, and the seasonal salad selection makes a recommended starter for two diners.",
                       ],
                   ))
F_WRITER_29 = File("F-WRITER-29", "writer_image_host_andromeda", "astronomy_brief.docx",
                   _src_image_host(
                       "photos/nature/galaxy-andromeda.jpg",
                       "galaxy-andromeda.jpg",
                       [
                           "Astronomy briefing — the Andromeda galaxy (Messier 31) is the nearest large spiral galaxy to the Milky Way at a distance of approximately two and a half million light years. It is the most distant object visible to the naked eye from a dark sky site, and its angular diameter of roughly three degrees makes it a rewarding target even for modest binoculars and small refractor telescopes.",
                           "Insert the photograph of the galaxy below this paragraph as the cover plate of the printed handout. The image is a long-exposure composite taken from a mid-latitude observatory; please preserve the aspect ratio at the printed size, as cropping the outer disc would obscure the prominent dust lanes that the discussion in the following section refers to.",
                           "Caption text for the printed copy — Andromeda galaxy, Messier 31, the nearest spiral neighbour of the Milky Way at roughly two point five million light years. The two visible companion galaxies, Messier 32 and Messier 110, are also recorded in the same field. Image courtesy of the regional observatory archive, reproduced with permission for educational use only.",
                       ],
                   ))
F_WRITER_30 = File("F-WRITER-30", "writer_image_host_house", "architectural_review.docx",
                   _src_image_host(
                       "photos/architecture/house-modern.jpg",
                       "house-modern.jpg",
                       [
                           "Architectural review — a single-family residence completed last year in the late-modernist tradition, sited on a sloping wooded lot at the edge of a small mid-western university town. The building's principal facade faces south-east across a landscaped meadow, and the overall plan is organised around a continuous covered porch that runs the full length of the main living level.",
                           "Insert the exterior photograph of the building below this paragraph for the cover plate of the review. The chosen image is the south-east elevation taken in late-afternoon light, which shows the cantilevered upper bedroom wing and the integrated carport to good effect; an alternative dawn elevation is held in the editorial archive should the editor prefer it.",
                           "Programmatic notes — open-plan main level with a single large multi-purpose living space, a cantilevered upper bedroom wing supported by two steel-framed transfer beams, and an integrated carport at the north end of the building tucked under the projecting roof line. The detailed material schedule and the structural drawings appear as appendices to this review.",
                       ],
                   ))

# ----- Catch-up — Loop 1: long-form prose with headings,
# recipes with table+photo, manuals with numbered sections, travel guides
# with bullet lists, and Loop 2 additions (more Gutenberg books, letters,
# list-only files, photo essays). Each File below has a unique
# title+heading layout, varied font/size/spacing, and contextually
# appropriate artifacts (tables / lists / photos). -------------------------

# Loop 1.A — long-form prose with title + headings (no tables/photos)
F_WRITER_32 = File("F-WRITER-32", "writer_white_paper",
                   "lakehouse_white_paper.docx",
                   _src_structured(
                       title="Federated Lakehouse — A White Paper",
                       blocks=[
                           ("h1", "Executive Summary"),
                           ("p",  _LONG_BODY_WHITE_PAPER[0]),
                           ("h1", "Background"),
                           ("p",  _LONG_BODY_WHITE_PAPER[1]),
                           ("h1", "Findings"),
                           ("p",  _LONG_BODY_WHITE_PAPER[2]),
                           ("h2", "Operational Risks"),
                           ("p",  _LONG_BODY_WHITE_PAPER[3]),
                           ("h1", "Recommendations"),
                           ("p",  _LONG_BODY_WHITE_PAPER[4]),
                       ],
                       font_name="Liberation Serif", font_size_pt=12,
                   ))

F_WRITER_33 = File("F-WRITER-33", "writer_case_study",
                   "northbridge_case_study.docx",
                   _src_structured(
                       title="Northbridge Transit Fare-Collection Modernisation",
                       blocks=[
                           ("h1", "Background"),
                           ("p",  _LONG_BODY_CASE_STUDY[0]),
                           ("h1", "Procurement"),
                           ("p",  _LONG_BODY_CASE_STUDY[1]),
                           ("h1", "Implementation"),
                           ("p",  _LONG_BODY_CASE_STUDY[2]),
                           ("h1", "Outcomes"),
                           ("p",  _LONG_BODY_CASE_STUDY[3]),
                           ("h1", "Lessons Learned"),
                           ("p",  _LONG_BODY_CASE_STUDY[4]),
                       ],
                       font_name="DejaVu Serif", font_size_pt=11,
                   ))

F_WRITER_34 = File("F-WRITER-34", "writer_press_briefing",
                   "eastfield_press_briefing.docx",
                   _src_structured(
                       title="Eastfield Manufacturing — Quarterly Press Briefing",
                       blocks=[
                           ("h1", "Briefing Notes"),
                           ("p",  _LONG_BODY_PRESS_BRIEFING[0]),
                           ("h1", "Headline Results"),
                           ("p",  _LONG_BODY_PRESS_BRIEFING[1]),
                           ("h1", "Divisional Performance"),
                           ("p",  _LONG_BODY_PRESS_BRIEFING[2]),
                           ("h1", "Outlook"),
                           ("p",  _LONG_BODY_PRESS_BRIEFING[3]),
                           ("h1", "Closing Remarks"),
                           ("p",  _LONG_BODY_PRESS_BRIEFING[4]),
                       ],
                       font_name="Liberation Sans", font_size_pt=12,
                   ))

F_WRITER_35 = File("F-WRITER-35", "writer_market_analysis",
                   "precision_sensor_market.docx",
                   _src_structured(
                       title="Global Precision Sensor Market — Analysis",
                       blocks=[
                           ("h1", "Market Overview"),
                           ("p",  _LONG_BODY_MARKET_ANALYSIS[0]),
                           ("h1", "Segment Dynamics"),
                           ("p",  _LONG_BODY_MARKET_ANALYSIS[1]),
                           ("h1", "Competitive Landscape"),
                           ("p",  _LONG_BODY_MARKET_ANALYSIS[2]),
                           ("h1", "Regional Outlook"),
                           ("p",  _LONG_BODY_MARKET_ANALYSIS[3]),
                           ("h1", "Risks and Uncertainties"),
                           ("p",  _LONG_BODY_MARKET_ANALYSIS[4]),
                       ],
                       font_name="Liberation Serif", font_size_pt=11,
                       line_spacing=1.15,
                   ))

F_WRITER_36 = File("F-WRITER-36", "writer_tech_spec",
                   "aurora3_tech_spec.docx",
                   _src_structured(
                       title="Aurora-3 Gateway — Technical Specification",
                       blocks=[
                           ("h1", "Introduction"),
                           ("p",  _LONG_BODY_TECH_SPEC[0]),
                           ("h1", "Physical and Environmental"),
                           ("p",  _LONG_BODY_TECH_SPEC[1]),
                           ("h1", "Electrical Interfaces"),
                           ("p",  _LONG_BODY_TECH_SPEC[2]),
                           ("h1", "Network Interfaces"),
                           ("p",  _LONG_BODY_TECH_SPEC[3]),
                           ("h1", "Software and Protocols"),
                           ("p",  _LONG_BODY_TECH_SPEC[4]),
                       ],
                       font_name="Liberation Mono", font_size_pt=10,
                   ))

F_WRITER_37 = File("F-WRITER-37", "writer_org_announcement",
                   "org_announcement.docx",
                   _src_structured(
                       title="Organisational Changes — All-Hands Announcement",
                       blocks=[
                           ("h1", "Summary"),
                           ("p",  _LONG_BODY_ORG_ANNOUNCEMENT[0]),
                           ("h1", "Platform Engineering"),
                           ("p",  _LONG_BODY_ORG_ANNOUNCEMENT[1]),
                           ("h1", "Product Organisation"),
                           ("p",  _LONG_BODY_ORG_ANNOUNCEMENT[2]),
                           ("h1", "Practical Implications"),
                           ("p",  _LONG_BODY_ORG_ANNOUNCEMENT[3]),
                           ("h1", "Next Steps"),
                           ("p",  _LONG_BODY_ORG_ANNOUNCEMENT[4]),
                       ],
                       font_name="DejaVu Sans", font_size_pt=11,
                   ))

F_WRITER_38 = File("F-WRITER-38", "writer_academic_abstract",
                   "urban_microclimate_abstract.docx",
                   _src_structured(
                       title="Urban Microclimate Variability — Five-Year Study",
                       blocks=[
                           ("h1", "Abstract"),
                           ("p",  _LONG_BODY_ACADEMIC_ABSTRACT[0]),
                           ("h1", "Methods"),
                           ("p",  _LONG_BODY_ACADEMIC_ABSTRACT[1]),
                           ("h1", "Results"),
                           ("p",  _LONG_BODY_ACADEMIC_ABSTRACT[2]),
                           ("h1", "Discussion"),
                           ("p",  _LONG_BODY_ACADEMIC_ABSTRACT[3]),
                           ("h1", "Conclusions"),
                           ("p",  _LONG_BODY_ACADEMIC_ABSTRACT[4]),
                       ],
                       font_name="Liberation Serif", font_size_pt=11,
                       line_spacing=1.5,
                   ))

F_WRITER_39 = File("F-WRITER-39", "writer_policy_memo",
                   "travel_policy_memo.docx",
                   _src_structured(
                       title="Travel Policy Memo — Proposed Amendments",
                       blocks=[
                           ("h1", "Introduction"),
                           ("p",  _LONG_BODY_POLICY_MEMO[0]),
                           ("h1", "Proposed Amendments"),
                           ("p",  _LONG_BODY_POLICY_MEMO[1]),
                           ("h2", "Per-Diem Simplification"),
                           ("p",  _LONG_BODY_POLICY_MEMO[2]),
                           ("h2", "Approval-Threshold Clarifications"),
                           ("p",  _LONG_BODY_POLICY_MEMO[3]),
                           ("h1", "Implementation"),
                           ("p",  _LONG_BODY_POLICY_MEMO[4]),
                       ],
                       font_name="Liberation Serif", font_size_pt=12,
                   ))

F_WRITER_40 = File("F-WRITER-40", "writer_long_essay_letters",
                   "essay_on_letters.docx",
                   _src_structured(
                       title="On the Persistence of Letters",
                       blocks=[
                           ("p",  _LONG_BODY_LONG_ESSAY_A[0]),
                           ("p",  _LONG_BODY_LONG_ESSAY_A[1]),
                           ("p",  _LONG_BODY_LONG_ESSAY_A[2]),
                           ("p",  _LONG_BODY_LONG_ESSAY_A[3]),
                           ("p",  _LONG_BODY_LONG_ESSAY_A[4]),
                       ],
                       font_name="Liberation Serif", font_size_pt=12,
                       line_spacing=1.5,
                   ))

F_WRITER_41 = File("F-WRITER-41", "writer_long_essay_maps",
                   "essay_on_maps.docx",
                   _src_structured(
                       title="On Reading Old Maps",
                       blocks=[
                           ("p",  _LONG_BODY_LONG_ESSAY_B[0]),
                           ("p",  _LONG_BODY_LONG_ESSAY_B[1]),
                           ("p",  _LONG_BODY_LONG_ESSAY_B[2]),
                           ("p",  _LONG_BODY_LONG_ESSAY_B[3]),
                           ("p",  _LONG_BODY_LONG_ESSAY_B[4]),
                       ],
                       font_name="Liberation Serif", font_size_pt=12,
                   ))

# Loop 1.B — recipes with title + ingredient table + step body + photo
_RECIPE_SALAD_BODY = [
    "Welcome to this seasonal mixed-leaf salad with toasted hazelnuts and a sharp citrus dressing — a light, fresh starter or a comfortable summer-evening main course on its own. The recipe serves four as a starter or two as a main, and takes about twenty minutes from start to finish provided the leaves have been washed and dried in advance. Plan to dress the salad at the very last moment so the leaves retain their crisp texture; a wilted leaf is the most easily avoided fault in a recipe of this kind.",
    "Begin by toasting the hazelnuts in a dry frying pan over a medium heat, shaking the pan continuously, for approximately five minutes or until the skins are visibly darkened and the nuts smell distinctly toasted. Tip them out onto a clean tea towel and rub vigorously to remove the loose skins. The nuts can be prepared up to a day in advance and stored in an airtight jar at room temperature; do not refrigerate them as the cold reintroduces moisture into the toasted nut and softens the texture you have just developed.",
    "While the nuts cool, whisk together the dressing ingredients in a small bowl until they form a smooth emulsion; taste and adjust the seasoning carefully, remembering that the leaves will dilute the dressing markedly once tossed. Assemble the salad in a wide shallow bowl, scatter the cooled hazelnuts over the top, and dress only at the moment of serving. A small grating of hard sheep's-milk cheese over the top makes a pleasant final touch but is not essential to the recipe; serve immediately while the leaves are still crisp.",
]
F_WRITER_42 = File("F-WRITER-42", "writer_recipe_salad",
                   "salad_recipe.docx",
                   _src_structured_with_photo(
                       title="Mixed-Leaf Salad with Toasted Hazelnuts",
                       blocks=[
                           ("h2", "Serves 4 starter / 2 main — 20 minutes"),
                           ("h1", "Ingredients"),
                           ("table", [
                               ["Ingredient", "Quantity"],
                               ["Mixed salad leaves", "200 g"],
                               ["Hazelnuts, blanched", "60 g"],
                               ["Olive oil", "4 tbsp"],
                               ["Lemon juice", "1 tbsp"],
                               ["Dijon mustard", "1 tsp"],
                               ["Sea salt and pepper", "to taste"],
                           ]),
                           ("h1", "Method"),
                           ("p", _RECIPE_SALAD_BODY[0]),
                           ("p", _RECIPE_SALAD_BODY[1]),
                           ("p", _RECIPE_SALAD_BODY[2]),
                       ],
                       asset_rel="photos/food/salad-bowl.jpg",
                       image_basename="salad-bowl.jpg",
                       font_name="Liberation Serif", font_size_pt=12,
                   ))

_RECIPE_PIZZA_BODY = [
    "This Neapolitan-style Margherita is the foundation recipe of our home pizza repertoire and rewards patience at every stage of the process. The dough is a long-fermented water-flour-yeast-salt mixture that develops its characteristic open crumb structure during a twenty-four hour cold proof in the refrigerator. The toppings are deliberately spare — a single layer of crushed San Marzano tomatoes, torn fresh mozzarella, a few leaves of basil, and a thread of good olive oil. Serve straight from the oven onto a warm plate.",
    "Mix the dough ingredients to a shaggy mass and rest for fifteen minutes; then knead for ten minutes by hand or six minutes in a stand mixer until smooth and elastic. Divide into balls of approximately two hundred and seventy grams each, place into individual lightly oiled containers, and refrigerate for at least twenty-four and up to forty-eight hours. The slow cold fermentation develops both flavour and the structural strength needed for the very thin Neapolitan stretch.",
    "About two hours before service, remove the dough balls from the refrigerator and let them come to room temperature, covered. Preheat your hottest oven for at least forty-five minutes with a baking stone on the lowest rack. Stretch each ball by hand to a thin disc of approximately thirty centimetres, top quickly and sparingly, and bake on the stone for six to eight minutes — the pizza is ready when the crust is leopard-spotted and the cheese is just beginning to colour. Slide onto a warm plate, finish with the basil, and serve.",
]
F_WRITER_43 = File("F-WRITER-43", "writer_recipe_pizza",
                   "pizza_recipe.docx",
                   _src_structured_with_photo(
                       title="Neapolitan Margherita Pizza",
                       blocks=[
                           ("h2", "Serves 4 — 24 hours plus 15 minutes active"),
                           ("h1", "Ingredients"),
                           ("table", [
                               ["Ingredient", "Quantity"],
                               ["Tipo-00 flour", "500 g"],
                               ["Water, cool", "320 ml"],
                               ["Sea salt", "12 g"],
                               ["Fresh yeast", "2 g"],
                               ["San Marzano tomatoes", "1 tin"],
                               ["Fresh mozzarella", "250 g"],
                               ["Fresh basil leaves", "1 small bunch"],
                               ["Extra-virgin olive oil", "to finish"],
                           ]),
                           ("h1", "Method"),
                           ("p", _RECIPE_PIZZA_BODY[0]),
                           ("p", _RECIPE_PIZZA_BODY[1]),
                           ("p", _RECIPE_PIZZA_BODY[2]),
                       ],
                       asset_rel="photos/food/pizza-dish.jpg",
                       image_basename="pizza-dish.jpg",
                       font_name="DejaVu Serif", font_size_pt=11,
                   ))

_RECIPE_COFFEE_BODY = [
    "A successful home cappuccino is essentially a matter of three independent skills practised in sequence — extracting a balanced double-shot espresso, steaming milk to a tight micro-foam texture, and pouring the two together to produce a stable surface pattern. Each of the three is learnable in an evening of focused practice; combining them reliably into a single drink takes rather longer. The recipe below describes the process in three numbered passes and assumes a single-boiler home machine with a fifty-eight-millimetre portafilter.",
    "Pull the espresso shot first. Dose between eighteen and nineteen grams of freshly ground coffee into the portafilter, distribute and tamp to a level puck, and lock into the group head. Pull for approximately twenty-eight to thirty seconds; the resulting double-shot should weigh thirty-six to thirty-eight grams in the cup. If the shot runs significantly faster than this, grind a step finer; if it stalls, grind a step coarser. Adjust one variable at a time and pull a second test shot before changing anything else in the process.",
    "Steam the milk in parallel. Use a small stainless-steel jug, fill to about a third full with cold whole milk, and submerge the steam wand tip just below the surface. Open the steam valve fully and lower the jug slowly to draw air into the milk for the first two or three seconds, then submerge the tip deeper to swirl the milk into a tight micro-foam without further air introduction. Stop steaming when the jug is just too hot to hold comfortably. Pour the milk immediately, low and steady, through the centre of the espresso surface.",
]
F_WRITER_44 = File("F-WRITER-44", "writer_recipe_coffee",
                   "coffee_recipe.docx",
                   _src_structured_with_photo(
                       title="Home Cappuccino — A Reliable Method",
                       blocks=[
                           ("h2", "Serves 1 — 10 minutes including warm-up"),
                           ("h1", "Equipment and Ingredients"),
                           ("table", [
                               ["Item", "Specification"],
                               ["Espresso machine", "single-boiler, 58 mm portafilter"],
                               ["Coffee, freshly roasted", "18-19 g per double shot"],
                               ["Whole milk, cold", "120 ml"],
                               ["Cup, pre-warmed", "150-180 ml"],
                           ]),
                           ("h1", "Method"),
                           ("p", _RECIPE_COFFEE_BODY[0]),
                           ("p", _RECIPE_COFFEE_BODY[1]),
                           ("p", _RECIPE_COFFEE_BODY[2]),
                       ],
                       asset_rel="photos/food/coffee-latte.jpg",
                       image_basename="coffee-latte.jpg",
                       font_name="Liberation Serif", font_size_pt=11,
                   ))

_RECIPE_PASTA_BODY = [
    "A weekday tagliatelle al ragù is the canonical recipe for using up odds and ends in the refrigerator on a Wednesday evening, and rewards patience at the simmer rather than complexity in the ingredient list. The recipe here is the standard Bolognese ragù as taught in the home kitchens of Emilia-Romagna, scaled to a comfortable family quantity, and finished with hand-cut tagliatelle either fresh from a pasta-machine or shop-bought from a producer you trust. Allow three hours for the long simmer.",
    "Sweat the soffritto — finely chopped onion, carrot, and celery in roughly equal proportions — in olive oil and butter over a gentle heat for at least fifteen minutes, until softened and just barely golden. Add the minced beef and pork, raise the heat, and brown thoroughly, breaking up the lumps with a wooden spoon. When the meat is fully browned, deglaze the pan with a generous splash of dry white wine and reduce until the alcohol has burned off, which takes two or three minutes.",
    "Add the milk and bring to a bare simmer, then add the tomato passata and a small amount of stock if the pan looks dry. Season generously and reduce the heat to the lowest setting that maintains a slow steady bubble. Cook uncovered for at least two and a half hours, stirring occasionally and adding a splash of stock or water if the pan threatens to dry. Cook the tagliatelle in well-salted water until al dente, drain, and finish in the pan with the ragù for a final minute over a gentle heat.",
]
F_WRITER_45 = File("F-WRITER-45", "writer_recipe_pasta",
                   "pasta_recipe.docx",
                   _src_structured_with_photo(
                       title="Tagliatelle al Ragù — Bolognese-Style",
                       blocks=[
                           ("h2", "Serves 4 — 3 hours total"),
                           ("h1", "Ingredients"),
                           ("table", [
                               ["Ingredient", "Quantity"],
                               ["Onion, finely chopped", "1 medium"],
                               ["Carrot, finely chopped", "1 medium"],
                               ["Celery, finely chopped", "2 sticks"],
                               ["Minced beef", "300 g"],
                               ["Minced pork", "150 g"],
                               ["Dry white wine", "150 ml"],
                               ["Whole milk", "200 ml"],
                               ["Tomato passata", "400 g"],
                               ["Fresh tagliatelle", "500 g"],
                           ]),
                           ("h1", "Method"),
                           ("p", _RECIPE_PASTA_BODY[0]),
                           ("p", _RECIPE_PASTA_BODY[1]),
                           ("p", _RECIPE_PASTA_BODY[2]),
                       ],
                       asset_rel="photos/food/restaurant-meal.jpg",
                       image_basename="restaurant-meal.jpg",
                       font_name="DejaVu Serif", font_size_pt=12,
                   ))

_RECIPE_DESSERT_BODY = [
    "A classic flourless chocolate torte is, despite its visually impressive appearance, one of the simplest dessert recipes in the pastry-chef repertoire. The recipe rewards good chocolate above everything else; use the best dark chocolate you can comfortably afford, ideally at sixty-five to seventy-percent cocoa solids. Allow at least four hours from start to finish, of which only thirty minutes is active work — the remaining time is the long slow cooling in the oven that gives the torte its characteristic dense fudgy texture.",
    "Melt the chocolate and butter together gently over a pan of barely simmering water, stirring occasionally until smooth. Remove from the heat and allow to cool to body temperature. In a separate bowl, whisk the egg yolks with three quarters of the sugar until pale and thick — the ribbon stage. Fold the cooled chocolate mixture carefully into the egg-yolk mixture. In a third bowl, whisk the egg whites with the remaining sugar to soft peaks, and fold this meringue into the chocolate base in three additions to preserve the volume.",
    "Pour the batter into a buttered and lined twenty-three-centimetre springform tin, smooth the surface, and bake for approximately twenty-five minutes in an oven preheated to one hundred and seventy degrees Celsius. The torte is ready when the surface has set with a fine crackled crust but the centre is still very slightly unset to the touch. Turn off the oven and allow the torte to cool inside for at least one hour with the door propped slightly ajar. Refrigerate for at least three hours before turning out, dust with cocoa, and serve in slim slices.",
]
F_WRITER_46 = File("F-WRITER-46", "writer_recipe_dessert",
                   "dessert_recipe.docx",
                   _src_structured_with_photo(
                       title="Flourless Chocolate Torte",
                       blocks=[
                           ("h2", "Serves 10-12 — 4 hours (30 min active)"),
                           ("h1", "Ingredients"),
                           ("table", [
                               ["Ingredient", "Quantity"],
                               ["Dark chocolate, 65-70%", "250 g"],
                               ["Unsalted butter", "200 g"],
                               ["Eggs, separated", "6 large"],
                               ["Caster sugar", "180 g"],
                               ["Cocoa powder, to dust", "2 tbsp"],
                           ]),
                           ("h1", "Method"),
                           ("p", _RECIPE_DESSERT_BODY[0]),
                           ("p", _RECIPE_DESSERT_BODY[1]),
                           ("p", _RECIPE_DESSERT_BODY[2]),
                       ],
                       asset_rel="photos/event/birthday-table.jpg",
                       image_basename="birthday-table.jpg",
                       font_name="Liberation Serif", font_size_pt=11,
                   ))

# Loop 1.C — manuals with title + numbered sections + spec table
_MANUAL_INSTALL_BODY = [
    "This installation guide describes the recommended procedure for the rack-mount installation of the WR-200 series wireless router in a standard nineteen-inch equipment cabinet. The procedure assumes that the rack is already grounded, that the cabinet has at least one available U of contiguous vertical rack space, and that the installer has access to a Phillips-head screwdriver, a torque wrench calibrated for the cabinet's cage-nut hardware, and the rear-rack power and network connectivity required by the appliance. Read the full procedure before beginning the physical install.",
    "Step one — unpack and inspect the unit. Verify against the shipping manifest that the carton contains the WR-200 chassis, two rack-mount ears, a single power cord matched to the destination region's electrical standard, the printed quick-start card, and the warranty registration leaflet. Inspect the chassis for any visible shipping damage, including dents to the chassis, bent mounting holes, or loose internal components audible when the chassis is gently tilted. Report any damage to the supplier within five business days; later claims may be declined.",
    "Step two — install the rack-mount ears. Place the chassis upside-down on a soft, clean surface to avoid scratching the top cover during installation. Align each rack-mount ear with the four threaded holes on the corresponding side of the chassis. Insert the supplied M4 screws and tighten progressively in a cross-pattern to a final torque of two newton-metres. Do not exceed the specified torque; over-tightening can strip the chassis threads and will void the warranty on the chassis frame.",
    "Step three — install the chassis into the rack. With a second person supporting the rear of the chassis, lift the unit into the rack at the chosen mounting position and engage the front rack-mount ears with the cabinet's vertical rails. Hand-thread two cage-nut bolts on each side to support the weight of the chassis. Once the chassis is supported by all four front bolts, install the rear mounting bracket if used. Tighten all eight mounting bolts to the cabinet manufacturer's specified torque, typically four to six newton-metres for standard cage-nut hardware.",
]
F_WRITER_47 = File("F-WRITER-47", "writer_manual_install",
                   "wr200_install_guide.docx",
                   _src_structured(
                       title="WR-200 Wireless Router — Rack-Mount Installation Guide",
                       blocks=[
                           ("h1", "Specifications"),
                           ("table", [
                               ["Parameter", "Value"],
                               ["Form factor", "1U rack-mount"],
                               ["Weight", "4.2 kg"],
                               ["Depth", "370 mm"],
                               ["Power input", "100-240 V AC, 50/60 Hz"],
                               ["Max power draw", "120 W"],
                               ["Operating temperature", "0-45 °C"],
                           ]),
                           ("h1", "1. Preparation"),
                           ("p",  _MANUAL_INSTALL_BODY[0]),
                           ("h1", "2. Unpack and Inspect"),
                           ("p",  _MANUAL_INSTALL_BODY[1]),
                           ("h1", "3. Install Rack Ears"),
                           ("p",  _MANUAL_INSTALL_BODY[2]),
                           ("h1", "4. Mount into Rack"),
                           ("p",  _MANUAL_INSTALL_BODY[3]),
                       ],
                       font_name="Liberation Mono", font_size_pt=11,
                   ))

_MANUAL_SAFETY_BODY = [
    "This safety protocol governs the operation of the laboratory's chemical fume hood and applies without exception to every researcher, visitor, and contractor entering the laboratory space. Before performing any operation under the hood, read this protocol in full and confirm that you have completed the laboratory's standard safety induction within the past twelve months. If your induction has lapsed, contact the laboratory safety officer to schedule a refresher session before proceeding with any work in the hood.",
    "Pre-work inspection — at the start of each working session, verify that the hood's airflow indicator reads within the green band on the front panel. The acceptable face-velocity range is between zero point four and zero point six metres per second; readings outside this band indicate that the hood is not providing adequate containment and the hood must not be used for hazardous work until the airflow has been restored to specification. Report any out-of-range readings to the laboratory safety officer immediately.",
    "Operating practice — keep the sash at the marked working height of approximately forty-five centimetres above the work surface. Working with the sash raised significantly above this height reduces the hood's containment performance and increases the risk of vapour escape into the laboratory air space. Position equipment at least fifteen centimetres back from the front edge of the work surface to remain within the protected airflow envelope. Do not block the rear baffle slots with large items of equipment.",
    "Emergency procedures — in the event of a chemical spill within the hood, lower the sash to the closed position immediately, evacuate the laboratory, and trigger the laboratory's chemical-spill alarm using the wall-mounted call point nearest the exit door. Do not attempt to clean up a significant spill yourself; the laboratory's hazardous-material response team will assess the spill on arrival and direct the cleanup. Annual training in spill response is mandatory for all laboratory personnel and is delivered jointly by the safety office and the response team.",
]
F_WRITER_48 = File("F-WRITER-48", "writer_manual_safety",
                   "fume_hood_safety.docx",
                   _src_structured(
                       title="Chemical Fume Hood — Safety Protocol",
                       blocks=[
                           ("h1", "Operating Limits"),
                           ("table", [
                               ["Parameter", "Specification"],
                               ["Face velocity", "0.4-0.6 m/s"],
                               ["Sash working height", "45 cm"],
                               ["Minimum setback", "15 cm from edge"],
                               ["Annual training", "mandatory, all personnel"],
                           ]),
                           ("h1", "1. Pre-Work Authorisation"),
                           ("p",  _MANUAL_SAFETY_BODY[0]),
                           ("h1", "2. Pre-Work Inspection"),
                           ("p",  _MANUAL_SAFETY_BODY[1]),
                           ("h1", "3. Operating Practice"),
                           ("p",  _MANUAL_SAFETY_BODY[2]),
                           ("h1", "4. Emergency Procedures"),
                           ("p",  _MANUAL_SAFETY_BODY[3]),
                       ],
                       font_name="Liberation Sans", font_size_pt=11,
                   ))

_MANUAL_TROUBLE_BODY = [
    "This troubleshooting guide addresses the most commonly reported issues with the office printing system and is intended for use by frontline IT support staff. Each numbered section describes the symptom, the most likely root causes in order of probability, and the suggested resolution sequence. If the suggested resolution does not restore service within fifteen minutes of attempting the documented steps, escalate to the printer-services team via the standard ticketing system rather than continuing further investigation locally.",
    "Symptom one — print job submitted but no output produced. The most common cause by some margin is that the print queue has paused on the affected workstation, typically because the user has dismissed an earlier paper-jam notification without clearing the underlying queue state. Open the print-queue dialogue on the affected workstation, look for jobs marked 'paused' or 'error', and either delete the stuck job or restart the queue. Second-most-common cause is a network connectivity loss to the print server itself; verify with ping before deeper investigation.",
    "Symptom two — print output is faded or shows streaks. Almost always indicates that one or more toner cartridges is approaching end of life and is failing to deliver toner uniformly across the page. Check the printer's front-panel status display for any toner-low warning and replace the affected cartridge if so. Less commonly, a faded output indicates a worn imaging drum; if cartridge replacement does not resolve the symptom, the next likely component is the drum unit, which is a more involved service call typically requiring a printer-services visit.",
    "Symptom three — printer reports a paper jam but no jam is visible. Frequently caused by a small piece of torn paper retained in one of the secondary paper-path zones — particularly the duplex unit or the rear pickup roller area. Open every access door in sequence as described in the printer's quick-reference card and inspect each zone thoroughly. If no debris is visible, power-cycle the printer to clear the paper-path sensor latches; if the symptom persists after a clean power cycle, escalate to printer services for a sensor diagnostic.",
]
F_WRITER_49 = File("F-WRITER-49", "writer_manual_trouble",
                   "printer_troubleshoot.docx",
                   _src_structured(
                       title="Office Printer Troubleshooting Guide",
                       blocks=[
                           ("h1", "Escalation Targets"),
                           ("table", [
                               ["Severity", "Response time"],
                               ["P1 — service down", "30 minutes"],
                               ["P2 — degraded", "2 business hours"],
                               ["P3 — single user", "1 business day"],
                               ["P4 — cosmetic", "best effort"],
                           ]),
                           ("h1", "1. Scope and Escalation"),
                           ("p",  _MANUAL_TROUBLE_BODY[0]),
                           ("h1", "2. No Output Produced"),
                           ("p",  _MANUAL_TROUBLE_BODY[1]),
                           ("h1", "3. Faded or Streaky Output"),
                           ("p",  _MANUAL_TROUBLE_BODY[2]),
                           ("h1", "4. Phantom Paper Jam"),
                           ("p",  _MANUAL_TROUBLE_BODY[3]),
                       ],
                       font_name="DejaVu Sans", font_size_pt=11,
                   ))

_MANUAL_OPERATOR_BODY = [
    "This operator manual describes the routine operating procedures for the BX-50 laboratory bench-top centrifuge and is intended for daily reference by qualified laboratory personnel. The BX-50 is a fixed-angle bench-top centrifuge with a maximum rotor speed of fifteen thousand revolutions per minute and a maximum sample mass of two hundred and forty grams across the standard six-position fixed rotor. All operators must complete the laboratory's centrifuge-specific safety briefing before independent operation.",
    "Pre-run checks — before each operating session, inspect the rotor for any visible signs of corrosion, cracking, or stress whitening, and confirm that the rotor lock-down knob is fully tightened to its hand-tight stop. Verify that the lid latch engages positively when the lid is closed and that the safety interlock prevents the rotor from spinning with the lid open. Any visible defect in the rotor or any failure of the safety interlock must be reported to the laboratory equipment officer immediately, and the centrifuge must be taken out of service pending inspection.",
    "Loading the rotor — load samples in diametrically opposite pairs to maintain rotor balance. The mass of each sample-pair must match within plus or minus zero point five grams; use the balance on the laboratory bench to confirm before loading. For odd-numbered sample counts, prepare a counterweight tube of equivalent mass and load in the diametrically opposite position. Do not exceed the marked fill line on the sample tubes — overfilling causes leakage during operation and may damage the rotor or the centrifuge chamber.",
    "Operating the centrifuge — close the lid firmly and confirm that the interlock indicator is green on the front-panel display. Select the desired rotor speed using the up and down arrows; the unit will not allow speeds above the rotor's rated maximum. Enter the desired run time and press the start button. The centrifuge will accelerate to the set speed in approximately ninety seconds, hold for the programmed time, and decelerate to a complete stop. Do not attempt to open the lid until the rotor has come to a complete stop and the indicator turns blue.",
]
F_WRITER_50 = File("F-WRITER-50", "writer_manual_operator",
                   "bx50_operator_manual.docx",
                   _src_structured(
                       title="BX-50 Bench-Top Centrifuge — Operator Manual",
                       blocks=[
                           ("h1", "Specifications"),
                           ("table", [
                               ["Parameter", "Value"],
                               ["Max rotor speed", "15 000 RPM"],
                               ["Max sample mass", "240 g (6 × 40 g)"],
                               ["Rotor type", "fixed-angle, 6-position"],
                               ["Acceleration time", "~90 s to max speed"],
                               ["Mass balance tolerance", "±0.5 g per pair"],
                           ]),
                           ("h1", "1. Scope"),
                           ("p",  _MANUAL_OPERATOR_BODY[0]),
                           ("h1", "2. Pre-Run Checks"),
                           ("p",  _MANUAL_OPERATOR_BODY[1]),
                           ("h1", "3. Loading the Rotor"),
                           ("p",  _MANUAL_OPERATOR_BODY[2]),
                           ("h1", "4. Operating the Centrifuge"),
                           ("p",  _MANUAL_OPERATOR_BODY[3]),
                       ],
                       font_name="Liberation Serif", font_size_pt=11,
                   ))

_MANUAL_DEVICE_BODY = [
    "This device manual describes the use of the HM-12 hand-held multimeter for routine laboratory and field measurements. The HM-12 is a four-and-a-half-digit auto-ranging multimeter rated for category III six-hundred-volt measurement environments and is supplied calibrated to traceable standards from the manufacturer's accredited calibration laboratory. Annual recalibration is recommended for laboratories operating under accredited quality programmes; the calibration interval may be extended in less critical use cases at the user's discretion.",
    "Battery and power — the unit ships with one nine-volt alkaline battery installed in the rear compartment. Typical battery life under continuous use is approximately two hundred and forty hours; the front-panel battery indicator will display a low-battery icon when the remaining capacity falls below twenty percent. Replace the battery before the indicator reaches the empty mark to avoid measurement errors as the battery approaches its discharge knee. Always remove the battery if the instrument is to be stored for more than thirty days.",
    "Voltage measurement — to measure AC or DC voltage, insert the black test lead into the COM jack and the red test lead into the V-ohm-Hz jack. Rotate the function selector to the desired voltage range; the auto-ranging feature will select the appropriate measurement range automatically within the chosen function. Probe the circuit under test, observing standard safe-probing practice — keep the non-probing hand away from the circuit and behind the finger guard on the probe handle.",
    "Resistance measurement — disconnect the circuit under test from any power source before attempting a resistance measurement. The HM-12 will display a warning if it detects voltage in the resistance function, but the warning is not a substitute for the standard safe-isolation practice. Insert the test leads as for voltage measurement, rotate the function selector to the ohms position, and probe the resistive element. For very low resistances (below ten ohms), use the four-wire Kelvin connection accessory available separately under part number HM-12-K.",
]
F_WRITER_51 = File("F-WRITER-51", "writer_manual_device",
                   "hm12_device_manual.docx",
                   _src_structured(
                       title="HM-12 Hand-Held Multimeter — User Manual",
                       blocks=[
                           ("h1", "Specifications"),
                           ("table", [
                               ["Parameter", "Value"],
                               ["Display", "4½-digit, auto-ranging"],
                               ["Measurement category", "CAT III 600 V"],
                               ["Battery", "1 × 9 V alkaline"],
                               ["Typical battery life", "~240 hours"],
                               ["Recommended recalibration", "12 months"],
                           ]),
                           ("h1", "1. Overview"),
                           ("p",  _MANUAL_DEVICE_BODY[0]),
                           ("h1", "2. Power Source"),
                           ("p",  _MANUAL_DEVICE_BODY[1]),
                           ("h1", "3. Voltage Measurement"),
                           ("p",  _MANUAL_DEVICE_BODY[2]),
                           ("h1", "4. Resistance Measurement"),
                           ("p",  _MANUAL_DEVICE_BODY[3]),
                       ],
                       font_name="Liberation Mono", font_size_pt=10,
                   ))

# Loop 1.D — travel guides with title + intro + bullets + photo
_TRAVEL_KYOTO_BODY = [
    "Welcome to Kyoto, the cultural heart of Japan and a city where every neighbourhood seems to layer centuries of history into a few streets of carefully tended temples, gardens, and townhouses. This printed guide is the companion document to the four-day itinerary booked through the travel office, and it covers the recommended walking routes, the must-visit sites, and a handful of food and shopping notes for each day. Distances quoted are walking distances; the city is comfortably navigated on foot supplemented by occasional local-bus or taxi journeys between districts.",
    "Day one is reserved for the Higashiyama district on the eastern slope of the city, beginning at the famous Kiyomizu-dera temple and working slowly downhill through the preserved historic streets of Sannenzaka and Ninenzaka. Allow a full leisurely day; the area rewards slow exploration. Avoid the densest tourist hours by starting before half past eight in the morning, when the temple opens, and you will have the temple's main hall almost to yourself for the first half hour. Lunch at one of the family-run noodle restaurants in the Yasaka area.",
]
F_WRITER_52 = File("F-WRITER-52", "writer_travel_kyoto",
                   "kyoto_guide.docx",
                   _src_structured_with_photo(
                       title="Kyoto — A Four-Day Walking Guide",
                       blocks=[
                           ("p", _TRAVEL_KYOTO_BODY[0]),
                           ("p", _TRAVEL_KYOTO_BODY[1]),
                           ("h1", "Highlights"),
                           ("bullet", "Kiyomizu-dera temple — main hall, eastern slope, early morning"),
                           ("bullet", "Sannenzaka preserved street — traditional townhouses"),
                           ("bullet", "Yasaka shrine — sunset visit recommended"),
                           ("bullet", "Gion district — narrow lantern-lit streets"),
                           ("bullet", "Philosopher's Path — cherry-blossom season"),
                           ("h1", "Practical Notes"),
                           ("p", "Most temples charge a small entrance fee in the range of three hundred to five hundred yen, payable in cash at the entrance gate. Cash is preferred everywhere outside the central commercial district; carry small notes for the temple fees and for the smaller traditional restaurants in the historic areas where card payment is not always accepted."),
                       ],
                       asset_rel="photos/city/asian-skyline.jpg",
                       image_basename="kyoto-skyline.jpg",
                       font_name="Liberation Serif", font_size_pt=12,
                   ))

_TRAVEL_HIKE_BODY = [
    "This hike journal records the three-day traverse of the Cathedral Range, completed in late June and recorded here for the trip-report archive. The route is the standard south-to-north traverse beginning at the Tuolumne Meadows trailhead and ending at the Mono Pass exit. Total distance is approximately thirty-eight kilometres with cumulative elevation gain of roughly fourteen hundred metres, distributed over the three days as eleven, fifteen, and twelve kilometres respectively. Water is plentiful throughout; bear-resistant food storage is mandatory for overnight camps.",
    "Day one — Tuolumne to Cathedral Lakes. A short and gentle warm-up day designed to ease into the elevation. The trail climbs steadily through open lodgepole-pine forest with periodic views back across the meadows. Cathedral Lakes is reached by mid-afternoon and offers excellent established campsites on the eastern shore. Plan for a swim if the weather is settled; the lake water is cold but bearable in late afternoon sunshine. Cathedral Peak rises dramatically across the lake and makes a memorable backdrop for an early-evening photograph.",
]
F_WRITER_53 = File("F-WRITER-53", "writer_travel_hike",
                   "cathedral_range_hike.docx",
                   _src_structured_with_photo(
                       title="Cathedral Range — Three-Day Traverse Journal",
                       blocks=[
                           ("p", _TRAVEL_HIKE_BODY[0]),
                           ("p", _TRAVEL_HIKE_BODY[1]),
                           ("h1", "Itinerary"),
                           ("bullet", "Day 1 — Tuolumne Meadows to Cathedral Lakes (11 km)"),
                           ("bullet", "Day 2 — Cathedral Lakes to Vogelsang High Camp (15 km)"),
                           ("bullet", "Day 3 — Vogelsang to Mono Pass trailhead (12 km)"),
                           ("h1", "Notes"),
                           ("bullet", "Bear canister required; rent at trailhead ranger station"),
                           ("bullet", "Water filter recommended; lakes and creeks throughout"),
                           ("bullet", "Cell coverage absent for entire route — file plan with ranger"),
                       ],
                       asset_rel="photos/landscape/mountain-range.jpg",
                       image_basename="cathedral-range.jpg",
                       font_name="Liberation Serif", font_size_pt=11,
                   ))

_TRAVEL_FOOD_BODY = [
    "This restaurant review covers six recently visited places in the old quarter of the city, recorded during a long-weekend visit in early spring. The selection is biased towards family-run trattorias and away from the more obvious tourist routes; the recommendations should hold over the medium term but are subject to the usual caveats around restaurant ownership and head-chef stability. Where reservations are recommended this is explicitly noted; otherwise walk-ins are reliable outside the local public-holiday weekends.",
    "Trattoria Romolo, on the small piazza behind the cathedral, is the standout of the visit and is the recommendation of choice for a first-night dinner. The menu is short, hand-written, and changes daily on the basis of the morning's market. The signature dish is a hand-cut tonnarelli with bottarga that is worth ordering on the strength of the chef's reputation alone. Bookings are essential on Friday and Saturday evenings; for weekday dinners a walk-in arrival before half past seven in the evening is usually reliable.",
]
F_WRITER_54 = File("F-WRITER-54", "writer_travel_food",
                   "old_quarter_review.docx",
                   _src_structured_with_photo(
                       title="Old Quarter — Restaurant Review Notes",
                       blocks=[
                           ("p", _TRAVEL_FOOD_BODY[0]),
                           ("p", _TRAVEL_FOOD_BODY[1]),
                           ("h1", "Recommendations"),
                           ("bullet", "Trattoria Romolo — first-night dinner, tonnarelli with bottarga"),
                           ("bullet", "Osteria della Quercia — casual lunch, daily specials"),
                           ("bullet", "Bar Centrale — morning espresso and pastry, no seats"),
                           ("bullet", "Pizzeria San Pietro — wood-fired, family-friendly"),
                           ("bullet", "Caffè della Posta — late-afternoon aperitivo, terrace"),
                           ("bullet", "Forno Antico — bakery, pick up bread for the train home"),
                       ],
                       asset_rel="photos/city/european-street.jpg",
                       image_basename="old-quarter.jpg",
                       font_name="DejaVu Serif", font_size_pt=11,
                   ))

_TRAVEL_BEACH_BODY = [
    "The Atlantic-coast section of this guide covers a four-hundred-kilometre stretch of coastline running south from Lisbon to the Algarve, with twelve recommended beach stops, six recommended cliff-top walks, and three town-based recommendations for inland breaks if the weather turns. The driving route is straightforward and well-signed throughout; expect approximately five hours of total driving for the full length, plus whatever time you choose to spend at individual stops. Hire-car insurance with full glass cover is strongly recommended for the cliff-road sections.",
    "The signature beach of the route, by some margin, is the wide sandy crescent at Praia da Bordeira, reached by a short dirt road from the village of the same name. The beach faces directly west and is at its best on a calm late-spring afternoon when the surf has settled and the offshore wind is light. The walk along the cliffs above the beach to the south is one of the most rewarding short walks in the country, with a clear path and the dramatic coastline laid out below you for the full forty minutes of the walk.",
]
F_WRITER_55 = File("F-WRITER-55", "writer_travel_beach",
                   "atlantic_coast_guide.docx",
                   _src_structured_with_photo(
                       title="Atlantic Coast — Lisbon to the Algarve",
                       blocks=[
                           ("p", _TRAVEL_BEACH_BODY[0]),
                           ("p", _TRAVEL_BEACH_BODY[1]),
                           ("h1", "Top Beach Stops"),
                           ("bullet", "Praia da Bordeira — signature beach, west-facing crescent"),
                           ("bullet", "Praia do Amado — surf beach, beginner-friendly"),
                           ("bullet", "Praia da Arrifana — small bay, swimming"),
                           ("bullet", "Praia do Beliche — wild and exposed, photography"),
                           ("h1", "Practical"),
                           ("bullet", "Hire-car full glass cover recommended"),
                           ("bullet", "Cash for parking at smaller beaches"),
                           ("bullet", "Sun protection essential — limited shade"),
                       ],
                       asset_rel="photos/landscape/beach-sunset.jpg",
                       image_basename="beach-sunset.jpg",
                       font_name="Liberation Serif", font_size_pt=12,
                   ))

_TRAVEL_DESERT_BODY = [
    "This desert tour itinerary covers a six-day camel-supported traverse of the western Sahara, beginning at the small caravan town of Merzouga and ending at the oasis village of M'Hamid via the high dune fields and the dry lake bed of Iriqui. The route is the standard guided traverse with overnight camps at five established bivouac sites; the daily walking distance averages approximately fifteen kilometres of relatively flat terrain. Temperatures range from around five degrees Celsius at dawn to forty in mid-afternoon during the spring season.",
    "The high dunes of Erg Chebbi are the visual highlight of the first two days of the tour, particularly at sunrise and sunset when the low-angle light brings out the sand patterns to dramatic effect. Allow ample time for photography during the golden-hour windows; the light changes rapidly and the best moments often last only ten or fifteen minutes. The mid-day hours are deliberately scheduled for shade and rest under the camp awnings, in line with the long-established practical wisdom of the region's traditional travellers.",
]
F_WRITER_56 = File("F-WRITER-56", "writer_travel_desert",
                   "sahara_tour.docx",
                   _src_structured_with_photo(
                       title="Western Sahara — Six-Day Desert Tour",
                       blocks=[
                           ("p", _TRAVEL_DESERT_BODY[0]),
                           ("p", _TRAVEL_DESERT_BODY[1]),
                           ("h1", "Camp Schedule"),
                           ("bullet", "Day 1 — Merzouga to Erg Chebbi north bivouac"),
                           ("bullet", "Day 2 — Erg Chebbi to Hassi Labied"),
                           ("bullet", "Day 3 — Hassi Labied to Iriqui dry lake"),
                           ("bullet", "Day 4 — Iriqui to Foum Larjamne"),
                           ("bullet", "Day 5 — Foum Larjamne to Erg Chigaga"),
                           ("bullet", "Day 6 — Erg Chigaga to M'Hamid"),
                           ("h1", "Packing"),
                           ("bullet", "Layered clothing — cold dawns, hot afternoons"),
                           ("bullet", "Long-brim hat plus head scarf for sun and sand"),
                           ("bullet", "Sturdy walking boots, broken-in"),
                       ],
                       asset_rel="photos/landscape/desert-dunes.jpg",
                       image_basename="desert-dunes.jpg",
                       font_name="DejaVu Serif", font_size_pt=11,
                   ))


# ----- Catch-up — Loop 2: more Gutenberg books, letters,
# list-only docs, photo essays. ---------------------------------------------

# Loop 2.A — additional Gutenberg books (one new File per book; pairs of
# different format-op FileTasks below). Mixes already-catalogued books at
# alternate openings/ops to expand the format-op variants without adding
# new asset files.

F_WRITER_57 = File("F-WRITER-57", "writer_gutenberg_alice_v2", "alice_chapter.docx",
                   _src_gutenberg("alice", n_paras=6))
F_WRITER_58 = File("F-WRITER-58", "writer_gutenberg_pride_v2", "pride_chapter.docx",
                   _src_gutenberg("pride", n_paras=6))
F_WRITER_59 = File("F-WRITER-59", "writer_gutenberg_moby_v2", "moby_chapter.docx",
                   _src_gutenberg("moby", n_paras=6))
F_WRITER_60 = File("F-WRITER-60", "writer_gutenberg_frank_v2", "frankenstein_chapter.docx",
                   _src_gutenberg("frank", n_paras=6))
F_WRITER_61 = File("F-WRITER-61", "writer_gutenberg_suntzu", "suntzu_excerpt.docx",
                   _src_gutenberg("suntzu", n_paras=5))
F_WRITER_62 = File("F-WRITER-62", "writer_gutenberg_tale2", "tale_of_two_cities.docx",
                   _src_gutenberg("tale2", n_paras=6))
F_WRITER_63 = File("F-WRITER-63", "writer_gutenberg_meta_v2", "metamorphosis_chapter.docx",
                   _src_gutenberg("meta", n_paras=6))
F_WRITER_64 = File("F-WRITER-64", "writer_gutenberg_sherlock_v2", "sherlock_chapter.docx",
                   _src_gutenberg("sherlock", n_paras=6))
F_WRITER_65 = File("F-WRITER-65", "writer_gutenberg_treasure_v2", "treasure_chapter.docx",
                   _src_gutenberg("treasure", n_paras=6))
F_WRITER_66 = File("F-WRITER-66", "writer_gutenberg_tom", "tom_sawyer.docx",
                   _src_gutenberg("tom", n_paras=5))

# Loop 2.B — letter-genre files (cover letters, recommendation, business
# correspondence). Pure prose, no tables/photos. Use `_src_long_body` with
# realistic letter-length paragraphs.

_LETTER_COVER = [
    "Dear Dr Mitchell,",
    "I am writing to apply for the Senior Research Engineer position advertised on the institute's careers portal last week. My background in distributed signal processing and the recent shift in my professional interests towards real-time perception systems make the role an excellent fit, and I would welcome the opportunity to bring my experience to your team. The role description aligns closely with the work I have been leading at my current organisation, and the institute's published research on adaptive sensor fusion has been a particular reference point for our team over the past two years.",
    "My current position is as a principal engineer at Vantage Systems, where I have led the perception-fusion sub-team for the past four years. In that role I have been responsible for the architectural design of the production sensor-fusion stack now deployed across the firm's industrial-automation product line, and for the line management of a team of six engineers spanning algorithm development, embedded implementation, and field calibration. Before joining Vantage I spent five years at Apex Robotics in a more individual-contributor capacity working on the perception-stack research programme.",
    "My doctoral work, completed in 2017 at the University of Edinburgh, focused on robust state estimation for partially-observable systems under sensor degradation. The thesis introduced a constraint-based estimation framework that has since been cited approximately eighty times in the literature and remains in use as the reference design in two of the major open-source robotics middleware stacks. Since the thesis I have published seven peer-reviewed papers in the relevant venues; selected papers are listed in the attached curriculum vitae.",
    "I would be available to start in the role within twelve weeks of an offer, accounting for the standard notice period in my current contract. I am happy to attend interviews in person at the institute's site at any reasonable time, and I can be available for an initial video call at short notice. I have attached my full curriculum vitae and the contact details for three professional references. Thank you for considering my application; I look forward to hearing from you.",
    "Yours sincerely,",
    "Alex Reyes",
]
F_WRITER_67 = File("F-WRITER-67", "writer_letter_cover", "cover_letter.docx",
                   _src_long_body(_LETTER_COVER,
                                  font_name="Liberation Serif", font_size_pt=12))

_LETTER_RECO = [
    "To Whom It May Concern,",
    "I am writing in support of Dr Priya Subramanian's application for the postdoctoral fellowship at your institute. I have known Priya for the past four years, first as the supervisor of her doctoral research and more recently as a colleague on a joint grant proposal that is now in its second year of funded work. In every respect I have found her to be one of the most capable young researchers I have worked with in my twenty-year academic career, and I recommend her application without reservation.",
    "Priya's doctoral research developed a novel inference framework for high-dimensional time-series data that has since been adopted as a standard tool in the climate-monitoring community. The technical depth of the work was exceptional, but it is the breadth of her contribution that distinguishes her from her peers. She built the underlying mathematical theory, implemented the production reference codebase, and led the multi-institution validation programme that established the method's empirical performance on three independent benchmark datasets. The combined contribution led to the best-paper award at the major conference in the field in 2024.",
    "Beyond the technical work, Priya has been an outstanding member of the laboratory community. She has mentored three masters students through their dissertation projects, contributed substantially to the laboratory's outreach programme with local secondary schools, and acted as the laboratory's representative on the departmental equality and diversity committee for the past two years. Her capacity to balance the deep individual research effort with the wider community contribution is, in my experience, an unusual combination at this career stage and is a strong indicator of long-term academic potential.",
    "In summary, Priya combines deep technical capability with broad scholarly judgement and an established record of community contribution. She would be an outstanding addition to your institute and I recommend her application to you in the strongest possible terms. Please do not hesitate to contact me if any further information would be helpful in your evaluation.",
    "Yours sincerely,",
    "Professor Margaret Holloway, FRS",
]
F_WRITER_68 = File("F-WRITER-68", "writer_letter_reco", "recommendation_letter.docx",
                   _src_long_body(_LETTER_RECO,
                                  font_name="Liberation Serif", font_size_pt=12,
                                  line_spacing=1.15))

_LETTER_BUSINESS = [
    "Dear Mr Andersson,",
    "I am writing to follow up on our conversation of the second of this month regarding the proposed renewal of the framework supply agreement between Northbridge Manufacturing and your firm. As discussed, we have now completed the internal review of the technical specifications and we are pleased to confirm our intention to proceed with the renewal on terms broadly equivalent to those set out in your initial proposal letter of the fifteenth of last month, subject to the small number of commercial clarifications described in the body of this letter.",
    "The first clarification concerns the volume commitment for the second contract year. Our internal capacity forecast indicates that the proposed minimum order quantity of two thousand four hundred units in year two should be reduced to two thousand units to better match the production-planning envelope we are able to commit at this stage of the planning cycle. We would propose to address the volume question for year three at the standard mid-term review, which falls in month eighteen of the renewed agreement, by which point both parties will have substantially better visibility of the medium-term demand.",
    "The second clarification concerns the price-escalation mechanism. We accept the proposed mechanism in principle but would propose that the escalation cap be raised from three percent to four percent annually, in recognition of the more volatile commodity-price environment that has prevailed since the original framework agreement was signed. This change, in conjunction with the volume adjustment described above, should produce a contract envelope that is broadly cost-neutral on a like-for-like basis relative to the proposal letter of last month.",
    "If these two clarifications are acceptable to your firm, we would propose to proceed directly to a formal contract draft on this basis, with the aim of signing before the expiry of the current framework agreement on the thirty-first of the month. I have copied our procurement counsel on this correspondence so that the legal draft can be initiated in parallel with your acknowledgement of the commercial points above. I look forward to your response.",
    "Yours sincerely,",
    "Catherine Lemarchand, Director of Procurement",
]
F_WRITER_69 = File("F-WRITER-69", "writer_letter_business", "business_letter.docx",
                   _src_long_body(_LETTER_BUSINESS,
                                  font_name="Liberation Serif", font_size_pt=11))

_LETTER_RESIGN = [
    "Dear Sarah,",
    "I am writing formally to give notice of my resignation from my position as Lead Software Engineer with effect from the thirty-first of next month, in line with the eight-week notice period set out in my employment contract. The decision has not been an easy one to reach, and I would like to take this opportunity to thank you and the wider leadership team for the support I have received during my five years with the firm. The professional growth I have experienced in this time has been substantial, and the move has been made for entirely positive reasons connected to my next-career opportunity.",
    "Over the coming eight weeks I am committed to ensuring a clean and well-managed handover of my current responsibilities. I have drafted a detailed handover document covering the platform-team work-streams I currently lead, including the open architectural decisions, the in-flight technical-debt programme, and the personnel-management responsibilities I carry as line manager for four engineers. I would welcome the opportunity to walk you through the document at our next one-to-one so we can agree on the appropriate transition plan for each of these areas.",
    "On a personal note, I would like to record my appreciation for the working environment that you and the senior leadership team have built at the firm. The combination of technical ambition and a genuinely supportive engineering culture is unusual in this industry and is the principal reason I have stayed for as long as I have. I will be sorry to leave the team I have worked with, and I will retain warm memories of the colleagues with whom I have shared the past five years.",
    "I expect to be entirely available for any handover-related activity throughout the notice period and would be happy to accommodate an exit interview at a time of your convenience. I will of course remain a champion for the firm and its mission after my departure, and would be glad to remain in informal contact with the team in whatever way is most appropriate after the formal end of my employment.",
    "Yours sincerely,",
    "Jordan Park",
]
F_WRITER_70 = File("F-WRITER-70", "writer_letter_resign", "resignation_letter.docx",
                   _src_long_body(_LETTER_RESIGN,
                                  font_name="Liberation Serif", font_size_pt=12,
                                  line_spacing=1.5))

_LETTER_THANKS = [
    "Dear Professor Lindgren,",
    "I am writing to thank you, very belatedly, for the time you so generously gave me during my visit to the laboratory last month. I had expected the meeting to be a brief introductory exchange and was both surprised and grateful that you set aside the better part of an afternoon to walk me through the laboratory and to discuss the work in such depth. The conversation was easily the most useful single hour I have spent on the topic this year and has materially shaped my thinking on the direction of my own research programme over the coming months.",
    "In particular, your observations on the calibration-stability issues with the optical bench have been valuable, and I have already begun discussions with my own colleagues here in London about adopting a similar two-stage calibration protocol on our equipment. I will of course attribute the suggestion to your group in any written description of the work, and I will be sure to share the results once we have established whether the technique transfers cleanly to our slightly different optical configuration.",
    "I hope very much that the broader visit arrangements will permit me to host you at our laboratory in due course, and I would be glad to extend a similarly thorough invitation when the calendar permits. Several of my colleagues have already expressed strong interest in hearing your perspective on the latest experimental results from your group, and there are one or two of our own ongoing experiments that I think you would find of interest in return.",
    "Thank you again for the visit and for your hospitality. I look forward to continuing the conversation in due course.",
    "Yours sincerely,",
    "Dr Naomi Brewer",
]
F_WRITER_71 = File("F-WRITER-71", "writer_letter_thanks", "thank_you_letter.docx",
                   _src_long_body(_LETTER_THANKS,
                                  font_name="Liberation Serif", font_size_pt=12))

# Loop 2.C — bullet / numbered list artifact files (to-do / agenda / action
# items / checklists). Body is structured-list-heavy with short intro and
# closing paragraphs.

F_WRITER_72 = File("F-WRITER-72", "writer_list_project_todo",
                   "project_todo.docx",
                   _src_structured(
                       title="Project Pegasus — Outstanding Action Items",
                       blocks=[
                           ("p", "This consolidated action-item list captures the outstanding work items from the most recent project-team sync, organised by work-stream owner and ordered within each work-stream by current priority ranking. The list is the authoritative reference for the next two weeks and will be revised at the regular Monday morning sync; please send any out-of-band updates to the project office before the close of business on the preceding Friday so the revised list is current at the Monday meeting."),
                           ("h1", "Platform Engineering — Owner: Marina"),
                           ("bullet", "Finalise the upgrade runbook for the production database cluster"),
                           ("bullet", "Validate the failover test results from last Tuesday's drill"),
                           ("bullet", "Schedule the production change window with the operations team"),
                           ("bullet", "Update the on-call rota to reflect the new team members"),
                           ("h1", "Product — Owner: Lila"),
                           ("bullet", "Complete the user-research synthesis for the autumn release"),
                           ("bullet", "Review and approve the revised information architecture proposal"),
                           ("bullet", "Schedule the product-review session with the executive team"),
                           ("h1", "Engineering — Owner: Jordan"),
                           ("bullet", "Resolve the open architectural review items from last sprint"),
                           ("bullet", "Validate the integration test results on the staging environment"),
                           ("bullet", "Update the engineering on-call documentation"),
                           ("p", "Closing reminder — the project office will lock the action-item list at noon on Friday for the Monday sync. Any updates received after the lock will be deferred to the next iteration; please prioritise accordingly."),
                       ],
                       font_name="Liberation Sans", font_size_pt=11,
                   ))

F_WRITER_73 = File("F-WRITER-73", "writer_list_meeting_agenda",
                   "meeting_agenda.docx",
                   _src_structured(
                       title="Executive Committee — Monthly Meeting Agenda",
                       blocks=[
                           ("p", "This is the consolidated agenda for the executive committee monthly meeting scheduled for the second Tuesday of the month, fourteen hundred to seventeen hundred hours, in the boardroom on the fourth floor. Pre-read papers for the substantive items have been distributed via the standard executive-papers channel and are available on the executive portal; please confirm receipt with the committee secretary by close of business on the preceding Friday."),
                           ("h1", "Standing Items"),
                           ("number", "Approval of minutes from the previous meeting"),
                           ("number", "Matters arising from the previous minutes"),
                           ("number", "Chief executive's report — verbal update"),
                           ("number", "Chief financial officer's report — written paper distributed"),
                           ("h1", "Substantive Items"),
                           ("number", "Quarterly results review — pre-read paper Q4-EC-23"),
                           ("number", "Annual strategic plan refresh — pre-read paper Q4-EC-24"),
                           ("number", "Acquisition pipeline review — pre-read paper Q4-EC-25"),
                           ("number", "Risk register quarterly review — pre-read paper Q4-EC-26"),
                           ("h1", "Standing Closing Items"),
                           ("number", "Any other business"),
                           ("number", "Date and venue of next meeting"),
                           ("p", "Logistics — refreshments will be served on arrival from thirteen forty-five. Please notify the committee secretary of any dietary requirements no later than the Friday preceding the meeting."),
                       ],
                       font_name="Liberation Serif", font_size_pt=12,
                   ))

F_WRITER_74 = File("F-WRITER-74", "writer_list_action_items",
                   "action_items.docx",
                   _src_structured(
                       title="Quarterly Strategy Review — Action Items",
                       blocks=[
                           ("p", "Action items arising from the quarterly strategy review held on the eighteenth of the month, organised by responsible executive and presented here in the format required for the standard action-tracker workflow. Each item carries an explicit owner, an explicit completion date, and a category tag drawn from the standard action-tracker taxonomy."),
                           ("h1", "Strategic Actions"),
                           ("number", "Finalise the medium-term capital allocation framework — CFO, end of next month"),
                           ("number", "Refresh the acquisition target shortlist for board review — Head of M&A, six weeks"),
                           ("number", "Approve the revised three-year product roadmap — CPO, four weeks"),
                           ("h1", "Operational Actions"),
                           ("number", "Complete the platform-engineering reorganisation — CTO, end of quarter"),
                           ("number", "Roll out the revised expense policy to all employees — CHRO, eight weeks"),
                           ("number", "Publish the updated information-security policy — CISO, end of next month"),
                           ("h1", "Governance Actions"),
                           ("number", "Refresh the executive committee charter — Company Secretary, six weeks"),
                           ("number", "Schedule the annual board strategy day — Chair's office, twelve weeks"),
                           ("p", "Closing note — the action tracker will be reviewed at the regular monthly executive committee meeting; ensure that any updates are recorded in the tracker no less than two business days before the scheduled review."),
                       ],
                       font_name="Liberation Serif", font_size_pt=11,
                       line_spacing=1.15,
                   ))

F_WRITER_75 = File("F-WRITER-75", "writer_list_checklist",
                   "release_checklist.docx",
                   _src_structured(
                       title="Production Release — Pre-Deployment Checklist",
                       blocks=[
                           ("p", "This pre-deployment checklist is the authoritative reference for any production release of the platform stack and must be completed and signed off by the named release manager before any production change window. The checklist is owned by the platform-engineering quality team and is revised quarterly; the current version is always the one available on the platform-engineering wiki and supersedes any printed copies in circulation."),
                           ("h1", "Code and Build"),
                           ("number", "All commits in the release branch have passed CI"),
                           ("number", "Static analysis has reported zero new critical findings"),
                           ("number", "Release notes have been written and reviewed by product"),
                           ("number", "Version tag has been applied to the release branch"),
                           ("h1", "Test and Validation"),
                           ("number", "Integration test suite is green on the staging environment"),
                           ("number", "Performance regression suite shows no negative trends"),
                           ("number", "Security scan is clean against the supply-chain registry"),
                           ("number", "Database migration has been dry-run on the staging replica"),
                           ("h1", "Operational Readiness"),
                           ("number", "Change-management ticket is approved and scheduled"),
                           ("number", "Rollback plan is documented and reviewed by the on-call"),
                           ("number", "Monitoring dashboards have been refreshed and verified"),
                           ("number", "Stakeholder notification has been sent two business days ahead"),
                           ("p", "Sign-off — the release manager and the on-call engineer must both sign the bottom of this checklist before the deployment proceeds. Retain the signed copy in the release archive for the standard retention period."),
                       ],
                       font_name="Liberation Mono", font_size_pt=10,
                   ))

F_WRITER_76 = File("F-WRITER-76", "writer_list_minutes",
                   "meeting_minutes.docx",
                   _src_structured(
                       title="Engineering All-Hands — Meeting Minutes",
                       blocks=[
                           ("p", "Minutes of the engineering all-hands meeting held on the twentieth of the month, fifteen hundred to sixteen hundred hours, in the engineering town-hall room and broadcast simultaneously to the remote-engineering teams. The recording is available on the standard internal video portal for ninety days. The minutes below capture the substantive discussion items and the agreed actions; the verbatim transcript is available on request from the engineering operations office."),
                           ("h1", "Attendance"),
                           ("p", "Approximately one hundred and forty engineers in person plus an additional eighty connected remotely from the European, North-American, and East-Asian offices. Apologies received from twelve individuals; the apologies list is attached to the master minutes file in the engineering operations archive."),
                           ("h1", "Substantive Discussion"),
                           ("number", "Quarterly engineering KPIs — presented by the CTO; principal observations on time-to-resolution metrics; broadly on plan"),
                           ("number", "Platform-engineering reorganisation update — presented by the VP Platform; structural changes effective from the start of the next quarter"),
                           ("number", "Product roadmap refresh — presented by the CPO; significant adjustment to the autumn release schedule announced"),
                           ("number", "Open Q&A — covered topics included compensation review timing, the office return-to-work policy, and the next round of hiring"),
                           ("h1", "Actions"),
                           ("bullet", "Publish updated org chart on the people portal — CTO office, two weeks"),
                           ("bullet", "Share roadmap detail with engineering managers — CPO office, one week"),
                           ("bullet", "Schedule follow-up Q&A on compensation — HR, three weeks"),
                       ],
                       font_name="Liberation Sans", font_size_pt=11,
                   ))

# Loop 2.D — photo-essay files (single relevant image; mostly insert_image)
F_WRITER_77 = File("F-WRITER-77", "writer_photo_forest",
                   "forest_essay.docx",
                   _src_image_host(
                       "photos/landscape/forest-trail.jpg",
                       "forest-trail.jpg",
                       [
                           "Photo essay — the high pine-and-larch forest of the western Cairngorms in late autumn, photographed in the soft horizontal light of a clear afternoon in early October. The forest at this altitude is at its visual peak for perhaps ten days each year, and the timing of a successful photographic visit is correspondingly difficult to plan in advance. This year the weather and the foliage cooperated unusually well, and the resulting set of photographs is the strongest from this location since I first began working the area five seasons ago.",
                           "Insert the cover photograph below this paragraph for the printed essay. The chosen image is the wide path-and-larch composition taken at approximately three o'clock in the afternoon when the low autumn sun was illuminating the path from behind the photographer. The original digital negative is held in the archive under reference cairngorms-twenty-twenty-five-zero-four-three-eight; please use the high-resolution master file for any reproduction at print size larger than A5.",
                           "Caption — Western Cairngorms, October — Scots pine and larch on a recovering grazing meadow. The photograph was taken in horizontal natural light shortly before sunset, with no fill light or post-processing colour correction beyond the standard print preparation workflow. The path is the trace of a long-disused estate road, now substantially recovered by the natural regeneration that has followed the change in land management since the late nineteen-eighties.",
                       ],
                   ))

F_WRITER_78 = File("F-WRITER-78", "writer_photo_portrait",
                   "portrait_essay.docx",
                   _src_image_host(
                       "photos/portrait/person-headshot-1.jpg",
                       "person-headshot-1.jpg",
                       [
                           "Portrait essay — this collection of recent studio portraits is the result of a six-month project carried out in collaboration with the social-history department of the regional museum. The subjects are residents of a small mining town in the north-east of the country, photographed in their everyday clothes against a neutral grey backdrop, and the project's intention has been to record a representative slice of the town's working population at a particular moment in its long industrial decline.",
                           "Insert the cover-plate portrait below this paragraph for the printed exhibition catalogue. The chosen image is the head-and-shoulders study of a retired pit deputy, photographed in soft natural light from a north-facing studio window on a clear morning in late spring. The portrait is reproduced here at the catalogue's standard cover size; alternative crops and the alternative wider-format composition from the same session are available in the catalogue archive.",
                           "Caption — Sitter A, retired pit deputy, sixty-eight years old, in his everyday clothes and against the standard project backdrop. The portrait is one of forty-two in the complete series; the full set of portraits will be exhibited at the regional museum from the first of next month for an initial run of three months, with an accompanying catalogue available from the museum shop and the museum's online publication channel.",
                       ],
                   ))

F_WRITER_79 = File("F-WRITER-79", "writer_photo_night",
                   "night_essay.docx",
                   _src_image_host(
                       "photos/city/street-night.jpg",
                       "street-night.jpg",
                       [
                           "Photo essay — the small commercial high street of a working district in the north of the city, photographed after midnight on a quiet weekday in mid-winter. The series is part of a continuing project to record the visual character of the city's late-night working environments before they are progressively transformed by the redevelopment programme now in its third year of implementation. The light is provided entirely by the existing public-realm lighting and a small number of late-trading commercial premises.",
                           "Insert the lead image below this paragraph for the printed plate of the essay. The chosen photograph is a wide angle composition looking south along the high street from the bus stop opposite the late-night convenience store, taken with a thirty-second exposure to capture the residual movement of late-night pedestrians and the slow trail of an approaching night-bus. The original digital negative is reference late-night-twenty-twenty-five-zero-two-one-five in the city-record archive.",
                           "Caption — High street, north district, after midnight on a winter weekday — the late-night convenience store and the closed shop fronts of the older small-trader premises in the foreground, with the bus stop and the early-morning bus route nearer the camera. The photograph is one of twenty-eight in the series and will be exhibited at the regional photographic society's spring exhibition under the project working title 'After Midnight.'",
                       ],
                   ))

F_WRITER_80 = File("F-WRITER-80", "writer_photo_bird",
                   "bird_essay.docx",
                   _src_image_host(
                       "photos/wildlife/bird-perch.jpg",
                       "bird-perch.jpg",
                       [
                           "Wildlife essay — the small woodland songbird population of the southern oak-and-hazel coppice, recorded over a complete spring breeding season in a continuing collaboration with the regional ornithological society. The photographs were taken from a small permanent hide established in the coppice for the duration of the season, using a long focal-length lens and a remote shutter trigger to minimise disturbance to the resident birds during the most sensitive nesting period of the year.",
                           "Insert the cover photograph below this paragraph for the printed report. The chosen image is the close perch study of a male great tit photographed in early-morning light shortly after dawn, on a low coppice branch within approximately six metres of the hide. The bird showed no awareness of the camera presence throughout the photographic session, in line with the established protocol for the hide's continuing operation through the breeding season.",
                           "Caption — Great tit (Parus major), male, southern oak-and-hazel coppice, mid-April — photographed at dawn from the permanent hide on the western edge of the coppice. The complete photographic record for the season is held in the regional ornithological society's image archive, accessible to society members through the standard member portal. Reproduction outside the society's own publications requires explicit prior permission from the society's image-archive committee.",
                       ],
                   ))

F_WRITER_81 = File("F-WRITER-81", "writer_photo_mars",
                   "mars_essay.docx",
                   _src_image_host(
                       "photos/space/mars-curiosity-panorama.jpg",
                       "mars-curiosity-panorama.jpg",
                       [
                           "Planetary-imaging essay — a high-resolution composite panorama of the Gale Crater floor as imaged by the Curiosity rover during sol three thousand two hundred and four of its surface mission, processed from the rover's raw image archive and reproduced here at the highest reproduction quality the printed page permits. The composite represents approximately eight hours of in-situ imaging time on the rover's mast-mounted camera array and was assembled from one hundred and forty-three individual frames by the imaging team at the mission's home institution.",
                           "Insert the panoramic photograph below this paragraph for the printed plate of the essay. The chosen image is the final processed composite at the catalogue's standard double-page reproduction size. Owing to the very high pixel resolution of the original mosaic, careful attention has been paid to the screening and dot-pattern selection in the printed reproduction; the digital master file is available from the mission archive under the standard public-release licence terms.",
                           "Caption — Gale Crater floor, sol three thousand two hundred and four — composite mosaic from the Curiosity mast camera, looking west towards the lower slopes of Aeolis Mons. The dark dune fields in the middle distance are part of the Bagnold Dunes complex, sampled in detail by the rover during the relevant phase of the surface mission. The image is reproduced courtesy of the imaging team and is in the public domain under the standard public-release licence terms.",
                       ],
                   ))


# ---------------------------------------------------------------------------
# Catch-up — real-source content + paragraph length variation +
# diverse document structures. Twelve new Files (F-WRITER-82..93) layered on
# top of the original 81 to address the user feedback that synth docs all
# look like single-page 4-6 paragraph essays. Each new File pulls real text
# from `_WIKI_ARTICLES` (Wikipedia paragraphs) or `_GUTENBERG_BOOKS_BY_KEY`
# (Project Gutenberg literary excerpts) — no new synthetic prose.
# ---------------------------------------------------------------------------

# F-WRITER-82 — Wikipedia structured article (Coffee, 4 H1 sections).
F_WRITER_82 = File("F-WRITER-82", "writer_wiki_coffee",
                   "coffee_wiki.docx",
                   _src_wiki_structured(
                       article_key="coffee",
                       title="Coffee — A Reference Article",
                       sections=[
                           ("Overview", 2),
                           ("History and Cultivation", 3),
                           ("Production and Trade", 1),
                           ("Brewing and Consumption", 2),
                       ],
                       font_name="Liberation Serif", font_size_pt=12,
                   ))

# F-WRITER-83 — Wikipedia article-style (Eiffel Tower, no headings, 8 paras
# of natural length variation).
F_WRITER_83 = File("F-WRITER-83", "writer_wiki_eiffel",
                   "eiffel_tower.docx",
                   _src_wiki_article(
                       article_key="eiffel-tower",
                       n_paras=8,
                       font_name="DejaVu Serif", font_size_pt=11,
                   ))

# F-WRITER-84 — Wikipedia structured article (Octopus, 3 H1 sections).
F_WRITER_84 = File("F-WRITER-84", "writer_wiki_octopus",
                   "octopus_wiki.docx",
                   _src_wiki_structured(
                       article_key="octopus",
                       title="Octopus — Biology and Behaviour",
                       sections=[
                           ("Introduction", 2),
                           ("Etymology and Naming", 3),
                           ("Habitat and Distribution", 2),
                       ],
                       font_name="Liberation Serif", font_size_pt=11,
                       line_spacing=1.15,
                   ))

# F-WRITER-85 — Wikipedia article (Mount Everest, 6 paras + title).
F_WRITER_85 = File("F-WRITER-85", "writer_wiki_everest",
                   "mount_everest.docx",
                   _src_wiki_article(
                       article_key="mount-everest",
                       n_paras=6,
                       title="Mount Everest",
                       font_name="Liberation Sans", font_size_pt=12,
                   ))

# F-WRITER-86 — short memo (1 paragraph from coffee Wikipedia, ~54 words).
F_WRITER_86 = File("F-WRITER-86", "writer_short_memo_coffee",
                   "coffee_memo.docx",
                   _src_short_memo(
                       title="Coffee Production — Briefing Note",
                       paragraph_text=_WIKI_ARTICLES["coffee"][23],
                       font_name="Liberation Serif", font_size_pt=12,
                   ))

# F-WRITER-87 — Q&A format on the Solar System (5 Q&A pairs, real wiki
# paragraphs as answers).
F_WRITER_87 = File("F-WRITER-87", "writer_qa_solar_system",
                   "solar_system_qa.docx",
                   _src_qa_format(
                       title="Solar System — Frequently Asked Questions",
                       qa_pairs=[
                           ("Q1: What is the Solar System?",
                            _WIKI_ARTICLES["solar-system"][0]),
                           ("Q2: What is the Sun's role in the system?",
                            _WIKI_ARTICLES["solar-system"][1]),
                           ("Q3: How many planets are there?",
                            _WIKI_ARTICLES["solar-system"][2]),
                           ("Q4: What are dwarf planets?",
                            _WIKI_ARTICLES["solar-system"][3]),
                           ("Q5: What are natural satellites?",
                            _WIKI_ARTICLES["solar-system"][4]),
                       ],
                       font_name="Liberation Serif", font_size_pt=11,
                   ))

# F-WRITER-88 — bullet reference (Yoga: short intro + 12 bullets sliced
# from Yoga wiki + conclusion). Bullets are short summaries derived from
# real wiki paragraphs (first ~12 words of each).
_YOGA_POOL = _WIKI_ARTICLES["yoga"]
F_WRITER_88 = File("F-WRITER-88", "writer_bullet_yoga",
                   "yoga_reference.docx",
                   _src_bullet_reference(
                       title="Yoga — Quick Reference",
                       intro_para=_YOGA_POOL[0],
                       bullets=[
                           "Yoga combines physical, mental, and spiritual practices originating in ancient India.",
                           "The earliest attested traditions date to the early first millennium BCE.",
                           "Hatha yoga is the principal source of modern Western asana practice.",
                           "The Sanskrit root yuj means 'to attach, join, harness, yoke'.",
                           "Yoga shares an Indo-European root with the English word 'yoke'.",
                           "Buswell and Lopez translate yoga as bond, restraint, and spiritual discipline.",
                           "Asana practice emphasises postures held with breath awareness.",
                           "Pranayama denotes the regulation of breath through specific techniques.",
                           "Dhyana refers to sustained meditative absorption on a single object.",
                           "Patanjali's Yoga Sutras codified eight limbs of classical yoga.",
                           "Modern postural yoga developed in the early twentieth century.",
                           "Therapeutic yoga programmes are now offered in many clinical settings.",
                       ],
                       conclusion_para=_YOGA_POOL[1],
                       font_name="Liberation Sans", font_size_pt=11,
                   ))

# F-WRITER-89 — mixed-length essay on Pizza, alternating short_p + p +
# Gutenberg italic quote blocks. At least 6 blocks of mixed lengths.
_PIZZA_POOL = _WIKI_ARTICLES["pizza"]
F_WRITER_89 = File("F-WRITER-89", "writer_mixed_pizza",
                   "pizza_essay.docx",
                   _src_mixed_length_essay(
                       title="Pizza — A Brief Essay with Mixed Notes",
                       blocks=[
                           ("p", _PIZZA_POOL[1]),
                           ("short_p", "A note on origins: the dish predates its name by many centuries."),
                           ("p", _PIZZA_POOL[2]),
                           ("quote", "It is a truth universally acknowledged, that a single man in possession of a good fortune, must be in want of a wife."),
                           ("short_p", "Pizza now circles the globe, naturalised in every cuisine it touches."),
                           ("p", _PIZZA_POOL[18]),
                           ("short_p", "Tradition guards the recipe; commerce expands it without apology."),
                           ("p", _PIZZA_POOL[10]),
                       ],
                       font_name="Liberation Serif", font_size_pt=12,
                   ))

# F-WRITER-90 — mixed-length essay on Volcano, with one Gutenberg literary
# quote block.
_VOLCANO_POOL = _WIKI_ARTICLES["volcano"]
F_WRITER_90 = File("F-WRITER-90", "writer_mixed_volcano",
                   "volcano_essay.docx",
                   _src_mixed_length_essay(
                       title="Volcanoes — Earth and Beyond",
                       blocks=[
                           ("p", _VOLCANO_POOL[0]),
                           ("short_p", "What we see at the surface is the cool record of immense heat below."),
                           ("p", _VOLCANO_POOL[1]),
                           ("quote", "It was the best of times, it was the worst of times, it was the age of wisdom, it was the age of foolishness."),
                           ("p", _VOLCANO_POOL[2]),
                           ("short_p", "Eruptions reshape coastlines and the chemistry of the air alike."),
                           ("p", _VOLCANO_POOL[4]),
                           ("quote", "Call me Ishmael."),
                       ],
                       font_name="Liberation Serif", font_size_pt=11,
                       line_spacing=1.5,
                   ))

# F-WRITER-91 — Wikipedia structured article (Lego, 3 H1 sections).
F_WRITER_91 = File("F-WRITER-91", "writer_wiki_lego",
                   "lego_wiki.docx",
                   _src_wiki_structured(
                       article_key="lego",
                       title="Lego — A Reference Article",
                       sections=[
                           ("Introduction", 2),
                           ("Origins of the Lego Group", 2),
                           ("Corporate Philosophy and Growth", 2),
                       ],
                       font_name="DejaVu Sans", font_size_pt=11,
                   ))

# F-WRITER-92 — Q&A on Internet of Things (5 Q&A from IoT wiki).
F_WRITER_92 = File("F-WRITER-92", "writer_qa_internet",
                   "iot_qa.docx",
                   _src_qa_format(
                       title="Internet of Things — Frequently Asked Questions",
                       qa_pairs=[
                           ("Q1: What is the Internet of Things?",
                            _WIKI_ARTICLES["internet-of-things"][0]),
                           ("Q2: How did the field evolve?",
                            _WIKI_ARTICLES["internet-of-things"][1]),
                           ("Q3: What does IoT look like in the consumer market?",
                            _WIKI_ARTICLES["internet-of-things"][2]),
                           ("Q4: What concerns surround the growth of IoT?",
                            _WIKI_ARTICLES["internet-of-things"][3]),
                           ("Q5: When did the earliest networked smart device appear?",
                            _WIKI_ARTICLES["internet-of-things"][5]),
                       ],
                       font_name="Liberation Mono", font_size_pt=11,
                   ))

# F-WRITER-93 — Origami bullet-reference document.
_ORIGAMI_POOL = _WIKI_ARTICLES["origami"]
F_WRITER_93 = File("F-WRITER-93", "writer_bullet_origami",
                   "origami_reference.docx",
                   _src_bullet_reference(
                       title="Origami — Quick Reference",
                       intro_para=_ORIGAMI_POOL[0],
                       bullets=[
                           "Origami transforms a flat square sheet into a finished sculpture by folding alone.",
                           "The practice is universal even though the word origami is Japanese in origin.",
                           "Modern origami eschews cuts, glue, and decorative marks on the paper.",
                           "Ceremonial origami (girei origami) is distinct from playful or recreational forms.",
                           "Basic folds combine into intricate designs through careful crease patterns.",
                           "The crane is among the best-known origami designs worldwide.",
                           "Origami has applications in mathematics, packaging, and aerospace deployment.",
                           "Wet-folding produces sculptural curves not achievable with dry paper.",
                           "Modular origami assembles many identical units into a larger structure.",
                           "Tessellation origami yields repeating geometric patterns from a single sheet.",
                           "Action origami produces models that move when manipulated by hand.",
                           "Origami diagrams use a standard set of symbols proposed by Akira Yoshizawa.",
                       ],
                       conclusion_para=_ORIGAMI_POOL[2],
                       font_name="Liberation Serif", font_size_pt=11,
                   ))


# Double-image file — two photos with caption anchors.
F_WRITER_31 = File("F-WRITER-31", "writer_double_image_galaxy", "galaxy_compare.docx",
                   _src_double_image_host(
                       "photos/nature/galaxy-andromeda.jpg",
                       "photos/nature/galaxy-hubble.jpg",
                       "galaxy-andromeda.jpg",
                       "galaxy-hubble.jpg",
                       [
                           "Comparative astronomy plate — this two-image reference is prepared for the introductory survey course on galactic morphology, and the accompanying discussion compares the resolved structure of a nearby spiral galaxy against the distant population sampled in a deep-field exposure.",
                           "Caption A — Andromeda (Messier 31), the nearest large spiral galaxy at approximately two and a half million light years, used here as the local-universe baseline against which the deeper sample is compared.",
                           "Caption B — Hubble deep-field, a long-exposure mosaic covering an apparently empty patch of sky and revealing thousands of distant galaxies of varied morphological type, from large ellipticals through to the irregular dwarf systems characteristic of the early universe.",
                           "Discussion prompt — contrast the well-resolved bulge-and-disc structure of the nearby spiral plate against the morphological zoo sampled in the deep field, and consider what selection effects shape the apparent distribution of galaxy types in each image.",
                       ],
                   ))


# §I.d — Gold helpers (heredoc generators for the gold docx).
#
# Each `_gold_*` function takes (src, gold, **gold_args) and returns the
# python-docx heredoc body that materializes the gold. Empty-string return
# means "no gold needed" (used for default_font where the gold lives in
# the registry, not the docx).

# ----- Per-paragraph format ops (eval = compare_docx_strict) -------------

def _gold_bold_para(src: str, gold: str, *, para_idx: int) -> str:
    return f"""\
from docx import Document
_doc = Document({src!r})
_p = _doc.paragraphs[{para_idx}]
for _r in _p.runs:
    _r.bold = True
_doc.save({gold!r})
"""


def _gold_italic_para(src: str, gold: str, *, para_idx: int) -> str:
    return f"""\
from docx import Document
_doc = Document({src!r})
_p = _doc.paragraphs[{para_idx}]
for _r in _p.runs:
    _r.italic = True
_doc.save({gold!r})
"""


def _gold_underline_para(src: str, gold: str, *, para_idx: int) -> str:
    return f"""\
from docx import Document
_doc = Document({src!r})
_p = _doc.paragraphs[{para_idx}]
for _r in _p.runs:
    _r.underline = True
_doc.save({gold!r})
"""


def _gold_strike_para(src: str, gold: str, *, para_idx: int) -> str:
    return f"""\
from docx import Document
_doc = Document({src!r})
_p = _doc.paragraphs[{para_idx}]
for _r in _p.runs:
    _r.font.strike = True
_doc.save({gold!r})
"""


def _gold_highlight_para(src: str, gold: str, *, para_idx: int) -> str:
    """Apply yellow highlight to all runs of paragraph K (eval needs
    examine_highlight=True)."""
    return f"""\
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
_doc = Document({src!r})
_p = _doc.paragraphs[{para_idx}]
for _r in _p.runs:
    _r.font.highlight_color = WD_COLOR_INDEX.YELLOW
_doc.save({gold!r})
"""


def _gold_unhighlight_all(src: str, gold: str) -> str:
    """Doc-wide: remove yellow highlight from EVERY run of EVERY paragraph.
    Mirrors osworld_libreoffice_writer_6a33f9b9 ("remove all highlight").
    The companion source builder pre-highlights selected paragraphs so
    the agent has a meaningful REMOVE operation; gold drops the highlight
    on every run (eval = compare_docx_strict + examine_highlight=True).
    """
    return f"""\
from docx import Document
_doc = Document({src!r})
for _p in _doc.paragraphs:
    for _r in _p.runs:
        _r.font.highlight_color = None
_doc.save({gold!r})
"""


def _gold_size_para(src: str, gold: str, *, para_idx: int, size_pt: int) -> str:
    """Set font size on paragraph K (eval needs examine_font_size=True)."""
    return f"""\
from docx import Document
from docx.shared import Pt
_doc = Document({src!r})
_p = _doc.paragraphs[{para_idx}]
for _r in _p.runs:
    _r.font.size = Pt({size_pt})
_doc.save({gold!r})
"""


# ----- Doc-wide spacing (eval = compare_line_spacing) --------------------

def _gold_doc_spacing(src: str, gold: str, *, value: float) -> str:
    """Set line_spacing on EVERY paragraph (doc-wide spacing change)."""
    return f"""\
from docx import Document
_doc = Document({src!r})
for _p in _doc.paragraphs:
    _p.paragraph_format.line_spacing = {value!r}
_doc.save({gold!r})
"""


def _gold_para_spacing(src: str, gold: str, *, para_idx: int, value: float) -> str:
    """Set line_spacing on paragraph K only."""
    return f"""\
from docx import Document
_doc = Document({src!r})
_p = _doc.paragraphs[{para_idx}]
_p.paragraph_format.line_spacing = {value!r}
_doc.save({gold!r})
"""


# ----- Doc-wide font (eval = compare_font_names) -------------------------

def _gold_doc_font(src: str, gold: str, *, font_name: str) -> str:
    """Set run.font.name on EVERY run of EVERY paragraph (whole-doc font).

    LO docx round-trip strips redundant rFonts and canonicalizes empty
    trailing paragraphs whose `compare_font_names` would otherwise see
    `name=None`. With the symmetric oracle pattern in `_build_oracle` the
    normalize now happens at eval time (step ①), so this builder is plain
    python-docx — no inline soffice tail.
    """
    return f"""\
from docx import Document
_doc = Document({src!r})
for _p in _doc.paragraphs:
    for _r in _p.runs:
        _r.font.name = {font_name!r}
_doc.save({gold!r})
"""


# ----- Default font in registry (eval = find_default_font; no gold docx) -

def _gold_default_font_noop(src: str, gold: str, *, font_name: str) -> str:
    """No-op — eval reads registrymodifications.xcu, not the docx. We still
    emit a trivial gold-docx file (cp of source) so the oracle's
    `cp expected→source` step doesn't fail with a missing path."""
    return f"""\
import shutil
shutil.copy({src!r}, {gold!r})
"""


# ----- Structural ops -----------------------------------------------------

def _gold_append_paragraph(src: str, gold: str, *, text: str) -> str:
    """Append a new paragraph at the end of the document (eval =
    compare_docx_files; text-only paragraph-by-paragraph comparison).

    LO docx round-trip absorbs subtle LO-on-save artifacts (autocorrect
    whitespace, invisible paragraph adjustments) that python-docx never
    produces. With the symmetric oracle pattern in `_build_oracle` the
    normalize now happens at eval time (step ①), so this builder is plain
    python-docx — no inline soffice tail.
    """
    return f"""\
from docx import Document
_doc = Document({src!r})
_doc.add_paragraph({text!r})
_doc.save({gold!r})
"""


def _gold_find_replace(src: str, gold: str, *, old: str, new: str) -> str:
    """Replace every occurrence of `old` with `new` across paragraph
    text. Operates per-run to preserve character formatting (eval =
    compare_docx_files)."""
    return f"""\
import re
from docx import Document
_doc = Document({src!r})
_old, _new = {old!r}, {new!r}
for _p in _doc.paragraphs:
    if _old not in _p.text:
        continue
    for _r in _p.runs:
        if _old in _r.text:
            _r.text = _r.text.replace(_old, _new)
_doc.save({gold!r})
"""


def _gold_insert_empty_table(src: str, gold: str, *, rows: int, cols: int) -> str:
    """Append an empty rows×cols table at the end (eval =
    compare_docx_tables; expects a specific table count + shape)."""
    return f"""\
from docx import Document
_doc = Document({src!r})
_doc.add_table(rows={rows}, cols={cols})
_doc.save({gold!r})
"""


def _gold_page_break(src: str, gold: str, *, para_idx: int) -> str:
    """Insert a page break before paragraph K. compare_pages returns 0 for
    the original doc and 1 for the gold (one page break inserted).
    Evaluator: contains_page_break with expected_count=1."""
    return f"""\
from docx import Document
from docx.enum.text import WD_BREAK
_doc = Document({src!r})
_p = _doc.paragraphs[{para_idx}]
_r = _p.runs[0] if _p.runs else _p.add_run('')
_r.add_break(WD_BREAK.PAGE)
_doc.save({gold!r})
"""


def _gold_first_centered(src: str, gold: str) -> str:
    """Set first-paragraph alignment to centre (eval =
    is_first_line_centered)."""
    return f"""\
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
_doc = Document({src!r})
if _doc.paragraphs:
    _doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
_doc.save({gold!r})
"""


def _gold_page_numbers_footer(src: str, gold: str) -> str:
    """Add a page-number field to the section footer (eval =
    has_page_numbers_in_footers — checks footer contains a digit)."""
    return f"""\
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
_doc = Document({src!r})
_section = _doc.sections[0]
_footer = _section.footer
_p = _footer.paragraphs[0] if _footer.paragraphs else _footer.add_paragraph()
_r = _p.add_run()
# Build a simple PAGE field — two-step (begin / instrText / end).
for _tag, _attrs, _text in [
    ('w:fldChar', {{qn('w:fldCharType'): 'begin'}}, None),
    ('w:instrText', {{qn('xml:space'): 'preserve'}}, ' PAGE '),
    ('w:fldChar', {{qn('w:fldCharType'): 'end'}}, None),
]:
    _el = OxmlElement(_tag)
    for _k, _v in _attrs.items():
        _el.set(_k, _v)
    if _text is not None:
        _el.text = _text
    _r._r.append(_el)
# Append a fallback literal digit so has_page_numbers_in_footers (which
# tests footer.text for any digit) passes even when the field is
# unevaluated in python-docx's emitted XML.
_p.add_run('1')
_doc.save({gold!r})
"""


def _gold_strike_last(src: str, gold: str) -> str:
    """Strike-through the LAST paragraph (eval =
    evaluate_strike_through_last_paragraph)."""
    return f"""\
from docx import Document
_doc = Document({src!r})
if _doc.paragraphs:
    _p = _doc.paragraphs[-1]
    for _r in _p.runs:
        _r.font.strike = True
_doc.save({gold!r})
"""


def _gold_italic_size14(src: str, gold: str) -> str:
    """Set wrong state (italic=True, size=12pt) on source, then derive gold
    (italic=True, size=14pt).  Source must have explicit italic=True so the
    evaluator's italic-run guard fires; without it check_italic_font_size_14
    trivially passes because it only checks runs where italic is already True.
    """
    return f"""\
from docx import Document
from docx.shared import Pt
_doc = Document({src!r})
for _p in _doc.paragraphs:
    for _r in _p.runs:
        _r.italic = True
        _r.font.size = Pt(12)
_doc.save({src!r})
_doc2 = Document({src!r})
for _p in _doc2.paragraphs:
    for _r in _p.runs:
        _r.font.size = Pt(14)
_doc2.save({gold!r})
"""


# ----- Gutenberg & image gold helpers (compose existing heredoc gens) ----

def _gold_gutenberg_p0_op(src: str, gold: str, *, op: str) -> str:
    """Apply a per-op format change to paragraph 0 of a Gutenberg-bodied
    source. Op switch lives in _gutenberg_gold_heredoc."""
    return _gutenberg_gold_heredoc(source_docx=src, gold_docx=gold, op=op)


def _gold_image_insert(src: str, gold: str, *, image_path: str,
                       width_in: float, insert_after_idx: int) -> str:
    """Insert image_path as inline picture after paragraph K and save to
    gold. Tail soffice round-trip aligns image-byte encoding with LO's
    emit path so examine_images succeeds."""
    return _image_gold_heredoc(
        src_docx=src, gold_docx=gold,
        image_path=image_path, width_in=width_in,
        insert_after_idx=insert_after_idx,
    )


def _gold_double_image_insert(src: str, gold: str, *, image_path_a: str,
                              image_path_b: str, width_in: float) -> str:
    """Insert image_a after the 'Caption A' paragraph and image_b after
    the 'Caption B' paragraph."""
    return _double_image_gold_heredoc(
        src_docx=src, gold_docx=gold,
        image_path_a=image_path_a, image_path_b=image_path_b,
        width_in=width_in,
    )


# ----- Subscript gold helper (chemistry formulas digits → subscript) -----

def _gold_subscript_chemistry(src: str, gold: str) -> str:
    """Open source docx, find the formulas paragraph, split each digit into
    its own run with `font.subscript=True`, save as gold. Mirrors the env-
    side `synth_inject_subscript` helper but via heredoc so the §I dataclass
    can wire it as a normal FileTask gold.

    `compare_subscript_contains` requires at least one (run1, run2) pair in
    a paragraph where both runs are subscript.

    Validation PARAM_REDUCIBLE: reduced gold to subscript ONLY
    the '2' in H2O (single digit in single formula) rather than every digit
    across H2O/CO2/NaCl/H2SO4. The contains-based eval only needs one
    aligned subscript run pair, so a single H2O '2' suffices and the
    instruction (now naming H2O specifically) becomes a 1-formula task.
    """
    return textwrap.dedent(f"""\
        import re
        from docx import Document

        doc = Document({src!r})
        # Locate the formulas paragraph by content anchor — F-WRITER-94's
        # source has a known formulas line containing 'H2O' as the marker.
        target = None
        for p in doc.paragraphs:
            if 'H2O' in p.text and 'CO2' in p.text:
                target = p
                break
        if target is None:
            raise SystemExit('subscript: anchor paragraph not found')

        text = target.text
        # Drop existing runs and re-add. Only the '2' inside the *first*
        # H2O occurrence becomes subscript; all other text (including the
        # other digits in CO2 / H2SO4) stays as a normal run.
        idx = text.find('H2O')
        before = text[:idx + 1]      # up to and including 'H'
        sub = text[idx + 1: idx + 2] # the '2'
        after = text[idx + 2:]       # 'O' + remainder of paragraph

        for run in list(target.runs):
            run._element.getparent().remove(run._element)
        if before:
            target.add_run(before)
        sub_run = target.add_run(sub)
        sub_run.font.subscript = True
        if after:
            target.add_run(after)

        doc.save({gold!r})
    """).strip() + "\n"


# ----- Colored-table gold helper (re-color GREEN cells to RED/BLUE) ------

def _gold_colored_table(src: str, gold: str) -> str:
    """Open source docx, walk the 2x3 table whose cells are pre-coloured
    GREEN, and re-colour each cell's text RED for vowel-start words, BLUE
    for consonant-start. Source pre-coloring avoids the eval's None-colour
    crash (see env-side `synth_inject_colored_table` lesson).
    """
    return textwrap.dedent(f"""\
        from docx import Document
        from docx.shared import RGBColor

        RED = RGBColor(0xFF, 0x00, 0x00)
        BLUE = RGBColor(0x00, 0x00, 0xFF)
        VOWELS = set('AEIOUaeiou')

        doc = Document({src!r})
        if not doc.tables:
            raise SystemExit('colored_table: no table found')
        tbl = doc.tables[0]
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        text = (run.text or '').strip()
                        if not text:
                            continue
                        run.font.color.rgb = RED if text[0] in VOWELS else BLUE

        doc.save({gold!r})
    """).strip() + "\n"


# ----- Source builders for the two new Files ------------------------------

def _src_chemistry_notes():
    """Source docx for the subscript task — chemistry notes with a known
    formulas line ('H2O ... CO2 ... NaCl ... H2SO4') as the subscript
    anchor + an intro paragraph for context."""
    title = "Chemistry Notes — Common Formulas"
    intro = ("This brief covers four common chemical compounds taught in "
             "the introductory chemistry tutorial. Each formula uses the "
             "standard convention where atom counts are written as digits "
             "alongside the element symbols.")
    formulas = ("Common compounds include H2O for water, CO2 for carbon "
                "dioxide, NaCl for sodium chloride, and H2SO4 for sulfuric "
                "acid.")
    note = ("Note that the digits in each formula indicate how many atoms "
            "of the preceding element are in one molecule of the compound.")
    blocks = [("p", intro), ("p", formulas), ("p", note)]
    return _src_structured(title=title, blocks=blocks,
                           font_name="Liberation Serif", font_size_pt=12)


_VOWEL_WORDS = ["Apple", "Orange", "Egg", "Ice", "Umbrella", "Eagle"]
_CONSONANT_WORDS = ["Dog", "Banana", "Tree", "Cat", "Fish", "Mountain"]


def _src_colored_table():
    """Source docx for the colored-table task — paragraphs + a 2x3 table
    whose cells are all pre-coloured GREEN. Matches the env-side
    `synth_inject_colored_table` shape: agent must re-colour cell text RED
    for vowel-start, BLUE for consonant-start."""
    def _src(path: str, seed: int) -> list[dict]:
        import random as _rnd
        rng = _rnd.Random(seed)
        vw = rng.sample(_VOWEL_WORDS, 3)
        cw = rng.sample(_CONSONANT_WORDS, 3)
        rows = [vw, cw]
        body = textwrap.dedent(f"""\
            from docx import Document
            from docx.shared import Pt, RGBColor

            GREEN = RGBColor(0x00, 0x80, 0x00)

            doc = Document()
            doc.add_heading('Word Categorisation Worksheet', level=0)
            doc.add_paragraph(
                'Each cell of the table below contains a single word. '
                'Re-colour each word so that words beginning with a vowel '
                'appear in red and words beginning with a consonant appear '
                'in blue. The cells are currently coloured green as a '
                'placeholder; replace the colour on every cell.'
            )
            tbl = doc.add_table(rows=2, cols=3)
            rows_data = {rows!r}
            for ri, row in enumerate(rows_data):
                for ci, word in enumerate(row):
                    cell = tbl.rows[ri].cells[ci]
                    cell.text = word
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.color.rgb = GREEN

            doc.save({path!r})
        """).strip() + "\n"
        return [_make_config_step(body)]
    return _src


# ----- PDF export gold helper --------------------------------------------

def _gold_pdf_export(src: str, gold_pdf: str) -> str:
    """No-op python heredoc — the actual gold-pdf build is done by a
    separate `_soffice_docx_to_pdf_step` step. Returning empty string
    means _to_synth_template skips the gold-docx step entirely."""
    return ""


# ----- Tab-stop gold helper (set custom tabstops on the target paragraph) -

def _gold_set_tabstops(src: str, gold: str, *,
                       para_idx: int,
                       positions_in: list[float],
                       alignment: str = "left") -> str:
    """Set a list of explicit tab-stops on `paragraphs[para_idx]` and add
    `\\t`-joined tokens so the paragraph text reflects the columns. Eval
    `check_tabstops` walks paragraph.paragraph_format.tab_stops and
    compares positions + alignments to the gold.

    `positions_in` is a list of tabstop offsets in inches. `alignment` is
    one of {"left", "right", "center"} (mapped to WD_TAB_ALIGNMENT).
    """
    align_map = {"left": "LEFT", "right": "RIGHT", "center": "CENTER"}
    enum_name = align_map[alignment]
    return f"""\
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT
_doc = Document({src!r})
_p = _doc.paragraphs[{para_idx}]
_ts = _p.paragraph_format.tab_stops
for _pos in {positions_in!r}:
    _ts.add_tab_stop(Inches(_pos), WD_TAB_ALIGNMENT.{enum_name})
_doc.save({gold!r})
"""


# ----- Unique-train-records source + gold helpers ------------------------
#
# `compare_unique_train_records` evaluator semantics:
#   - processed_file paragraphs must each be a subset of initial paragraphs
#   - the col-1 token (train_id) across processed rows must form a UNIQUE set
#   - that set must equal the gold's train_id set
#   - and processed line-count == gold line-count
# Each CSV-shaped paragraph is exactly 4 comma-separated fields:
#   "<station>, <train_id>, <departure>, <arrival>"

_TRAIN_RECORDS_ROWS = [
    "King's Cross, T101, 07:15, 11:40",
    "Paddington, T102, 08:00, 13:25",
    "Euston, T103, 06:45, 12:10",
    "Liverpool Street, T101, 07:15, 11:40",        # duplicate of T101
    "Victoria, T104, 09:20, 14:55",
    "Marylebone, T102, 08:00, 13:25",              # duplicate of T102
    "St Pancras, T105, 10:05, 15:30",
    "Waterloo, T106, 11:30, 17:00",
    "Charing Cross, T104, 09:20, 14:55",           # duplicate of T104
    "Cannon Street, T107, 12:45, 18:10",
]


def _src_train_records():
    """File.src callable: builds a register docx whose body is a heading +
    short intro + the full _TRAIN_RECORDS_ROWS list (10 CSV-shaped rows
    with 3 duplicate train_ids). Also drops an identical snapshot at
    /tmp/initial_<basename>.docx so the evaluator can read the pre-edit
    state (the agent's task is to DELETE the duplicates from the user-
    visible docx; the snapshot stays at the /tmp path)."""
    def _src(path: str, _seed: int) -> list[dict]:
        title = "Departure Register — Manual Audit"
        intro = ("Each line below records one observed departure as a "
                 "comma-separated tuple of station, train id, departure "
                 "time, and arrival time. Several trains were logged from "
                 "more than one platform on the same day, so a handful of "
                 "rows duplicate the same train id; please remove those "
                 "duplicates so each train id appears at most once.")
        py = textwrap.dedent(f"""\
            import shutil
            from docx import Document
            _doc = Document()
            _doc.add_heading({title!r}, level=0)
            _doc.add_paragraph({intro!r})
            for _row in {_TRAIN_RECORDS_ROWS!r}:
                _doc.add_paragraph(_row)
            _doc.save({path!r})
            shutil.copyfile({path!r}, '/tmp/initial_' + {path!r}.split('/')[-1])
        """).strip() + "\n"
        return [_make_config_step(py)]
    return _src


def _gold_unique_train_records(src: str, gold: str) -> str:
    """Build the gold by writing only the first-seen line for each train_id
    (the CSV col-1 token). Body header + intro paragraphs are preserved.

    The evaluator requires: processed_lines ⊆ initial_lines (subset), train_id
    set uniqueness within processed file, and equality of train_id set with
    gold. Keeping the first-seen row for each id naturally satisfies all
    three.
    """
    return f"""\
from docx import Document
_src_doc = Document({src!r})
_dst_doc = Document()
_seen = set()
for _p in _src_doc.paragraphs:
    _text = _p.text
    _parts = [_t.strip() for _t in _text.split(',')]
    if len(_parts) == 4:
        _train_id = _parts[1]
        if _train_id in _seen:
            continue
        _seen.add(_train_id)
    _dst_doc.add_paragraph(_text)
_dst_doc.save({gold!r})
"""


# ----- Compound gold helpers --------------------------------------------
#
# Compound evaluator atoms share the SAME gold docx; building the gold is
# the same as building any single-atom gold (because every property is
# computed on the final saved doc). What changes is the evaluator dict —
# func is a list of fn names and result/expected are parallel lists. These
# helpers simply wrap an existing single-atom gold builder so the FileTask
# `gold` callable stays in the single-helper shape.

def _gold_subscript_files_compound(src: str, gold: str) -> str:
    """Compound-property gold for [compare_docx_files,
    compare_subscript_contains]: agent must (a) preserve paragraph text
    semantics (text comparison passes after delete-empty-lines tolerance)
    AND (b) subscript the '2' in H2O. Identical to `_gold_subscript_
    chemistry` since the gold doc is the same file used by both atoms."""
    return _gold_subscript_chemistry(src, gold)


# ----- Append empty-table-with-text gold helper -------------------------

def _gold_insert_filled_table(src: str, gold: str, *,
                              rows_data: list[list[str]]) -> str:
    """Append a table at the end of the document filled with `rows_data`
    (a list of row-lists of cell strings). Eval = compare_docx_tables
    checks table count, dims and per-cell text equality.

    Distinct from `_gold_insert_empty_table` (which adds a blank n×m
    placeholder) — agents must actually type text into each cell.
    """
    return f"""\
from docx import Document
_doc = Document({src!r})
_rows_data = {rows_data!r}
_n_rows = len(_rows_data)
_n_cols = len(_rows_data[0]) if _rows_data else 0
_tbl = _doc.add_table(rows=_n_rows, cols=_n_cols)
for _i, _row in enumerate(_rows_data):
    for _j, _cell in enumerate(_row):
        _tbl.rows[_i].cells[_j].text = _cell
_doc.save({gold!r})
"""


# ----- Mixed-alignment split (eval skill: split one paragraph into two
#       paragraphs with different alignment — first N words LEFT, rest RIGHT).
#       Gold uses `compare_docx_strict` which already checks paragraph
#       alignment via `_paragraph_format_signature`. The agent must split the
#       paragraph at the Nth word and apply alignment to each half. ----------

def _gold_mixed_alignment_split(src: str, gold: str, *,
                                para_idx: int, split_word: int,
                                first_align: str = "left",
                                second_align: str = "right") -> str:
    """Split paragraphs[para_idx] at the `split_word`-th word into two
    paragraphs; align the first paragraph `first_align` and the second
    `second_align`. Other paragraphs are unchanged.

    In-place mutation: clone src to gold, replace para_idx text with the
    first half + alignment, inject a sibling paragraph for the second
    half. Preserves all surrounding paragraph format / sectPr / styles.
    The earlier `_dst = Document()` empty-rebuild lost every other
    paragraph's runs and formatting beyond plain text.
    """
    align_map = {"left": "LEFT", "right": "RIGHT", "center": "CENTER"}
    a1 = align_map[first_align]
    a2 = align_map[second_align]
    return f"""\
import shutil
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
shutil.copyfile({src!r}, {gold!r})
_doc = Document({gold!r})
_target = _doc.paragraphs[{para_idx}]
_words = _target.text.split(' ')
_first = ' '.join(_words[:{split_word}])
_rest = ' '.join(_words[{split_word}:])
_target.text = _first
_target.alignment = WD_ALIGN_PARAGRAPH.{a1}
_style = _target.style
_new = _doc.add_paragraph(_rest, style=_style)
_new.alignment = WD_ALIGN_PARAGRAPH.{a2}
_target._element.addnext(_new._element)
_doc.save({gold!r})
"""


# ----- Doc-wide case conversion (eval = compare_docx_files; text-only). ----

def _gold_doc_case_convert(src: str, gold: str, *, mode: str) -> str:
    """Re-render every paragraph's text in `mode` (`upper` or `lower`).
    Eval kind `files` (compare_docx_files) compares text only — but the
    earlier `_dst = Document()` empty-rebuild ALSO stripped paragraph
    styles, and LO-normalize round-trip on the agent's UI-edited docx
    preserves them, so the diff in style/formatting still surfaces. The
    safe form is in-place: clone src, mutate each paragraph's text.
    """
    assert mode in ("upper", "lower"), f"unknown case mode: {mode!r}"
    method = "upper" if mode == "upper" else "lower"
    return f"""\
import shutil
from docx import Document
shutil.copyfile({src!r}, {gold!r})
_doc = Document({gold!r})
for _p in _doc.paragraphs:
    _p.text = _p.text.{method}()
_doc.save({gold!r})
"""


# ----- First / last sentence-level format (eval = compare_docx_strict).
#       Splits the target paragraph by sentence-ending '.' and applies bold
#       or italic to the runs of the first / last sentence. Gold paragraph
#       text is unchanged; only run-format inside that paragraph differs.

def _gold_sentence_op(src: str, gold: str, *,
                      para_idx: int, which: str, op: str) -> str:
    """Apply `op` (`bold` or `italic`) to the `which` (`first` or `last`)
    sentence of paragraphs[para_idx]. The paragraph is rebuilt as 2 runs —
    the targeted-sentence run carries the format, the remainder run is
    plain. Sentence boundary: literal '.' followed by space; the trailing
    period is included in the targeted sentence.
    """
    assert which in ("first", "last"), which
    assert op in ("bold", "italic"), op
    attr = "bold" if op == "bold" else "italic"
    return f"""\
from docx import Document
_doc = Document({src!r})
_p = _doc.paragraphs[{para_idx}]
_text = _p.text
# Sentence split on '. ' — keep period with leading sentence.
_parts = _text.split('. ')
if len(_parts) == 1:
    # No internal split; treat the whole paragraph as one sentence.
    _target = _parts[0]
    _rest = ''
else:
    if {which!r} == 'first':
        _target = _parts[0] + '.'
        _rest = ' ' + '. '.join(_parts[1:])
    else:
        # last sentence
        _target = _parts[-1]
        _rest = '. '.join(_parts[:-1]) + '. '
# Drop existing runs and re-add two runs (rest order depends on `which`).
for _r in list(_p.runs):
    _r._element.getparent().remove(_r._element)
if {which!r} == 'first':
    _r1 = _p.add_run(_target)
    setattr(_r1, {attr!r}, True)
    if _rest:
        _p.add_run(_rest)
else:
    if _rest:
        _p.add_run(_rest)
    _r2 = _p.add_run(_target)
    setattr(_r2, {attr!r}, True)
_doc.save({gold!r})
"""


# F-WRITER-94 — chemistry notes source for the subscript task.
F_WRITER_94 = File("F-WRITER-94", "writer_chemistry_notes",
                   "chemistry_notes.docx",
                   _src_chemistry_notes())

# F-WRITER-95 — coloured-table worksheet source.
F_WRITER_95 = File("F-WRITER-95", "writer_colored_table",
                   "color_worksheet.docx",
                   _src_colored_table())

# F-WRITER-96 — mixed-alignment split fixture. Body uses an essay genre so
# paragraph[0] is a single long sentence the agent can split + align.
F_WRITER_96 = File("F-WRITER-96", "writer_mixed_align", "mixed_align.docx",
                   _src_genre(1))  # essay

# F-WRITER-97 — doc-wide case-conversion fixture. Body is a brief genre so
# the agent applies the same case operation uniformly across paragraphs.
F_WRITER_97 = File("F-WRITER-97", "writer_doc_case", "doc_case.docx",
                   _src_genre(4))  # brief

# F-WRITER-98 — sentence-level ordinal fixture. Uses the long research-report
# body so each paragraph contains multiple sentences delimited by '. '.
F_WRITER_98 = File("F-WRITER-98", "writer_sentence_op", "sentence_op.docx",
                   _src_long_body(_LONG_BODY_RESEARCH_REPORT))


# ----- validation RESCALER (eval-coverage fill, 4 niche skills) ---------------
#
# These 4 helpers + 4 File instances close the residual eval gaps surfaced by
# the eval-vs-synth coverage audit:
#   eval row 6  — center-align heading              → F_WRITER_99
#   eval row 11 — delete table rows by condition    → F_WRITER_100
#   eval row 14 — split paragraph into sentences    → F_WRITER_101
#   eval row 15 — convert text to table             → F_WRITER_102


def _gold_center_para(src: str, gold: str, *, para_idx: int) -> str:
    """Set paragraph[para_idx]'s alignment to CENTER.

    Eval kind `strict` (compare_docx_strict) checks alignment via
    `_paragraph_format_signature`.
    """
    return f"""\
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
_doc = Document({src!r})
_doc.paragraphs[{para_idx}].alignment = WD_ALIGN_PARAGRAPH.CENTER
_doc.save({gold!r})
"""


def _gold_filter_table_rows_keep_range(src: str, gold: str, *,
                                       col_idx: int,
                                       start: str, end: str) -> str:
    """Keep table rows whose leading `col_idx` token is in [start, end]
    inclusive; drop all other body rows (header is preserved).

    Eval kind `files` (compare_docx_files) compares text only.
    """
    return f"""\
from docx import Document
_doc = Document({src!r})
_tbl = _doc.tables[0]
_to_delete = []
_start = {start!r}.strip().split()[0]
_end = {end!r}.strip().split()[0]
for _i, _r in enumerate(_tbl.rows):
    if _i == 0:
        continue
    _v = _r.cells[{col_idx}].text.strip().split()[0]
    if not (_start <= _v <= _end):
        _to_delete.append(_r)
for _r in _to_delete:
    _r._element.getparent().remove(_r._element)
_doc.save({gold!r})
"""


def _gold_sentence_per_line(src: str, gold: str, *, para_idx: int) -> str:
    """Split paragraph[para_idx] at '. ' boundaries; each sentence becomes
    a separate paragraph. Other paragraphs unchanged.

    Eval kind `files` (compare_docx_files). In-place mutation: clone src
    to gold, then for the target paragraph replace text in place with the
    first sentence and inject sibling paragraphs for each subsequent
    sentence (preserves surrounding paragraph format / sectPr). The
    earlier `_dst = Document()` empty-rebuild stripped every paragraph's
    style and runs, which compare_docx_files sometimes still detected via
    LO-normalize round-trip diffs."""
    return f"""\
import shutil
from docx import Document
shutil.copyfile({src!r}, {gold!r})
_doc = Document({gold!r})
_target = _doc.paragraphs[{para_idx}]
_parts = _target.text.split('. ')
# Build the per-sentence text list (re-append the period dropped by split)
_sentences = []
for _j, _s in enumerate(_parts):
    if not _s:
        _sentences.append('')
    elif _j < len(_parts) - 1:
        _sentences.append(_s + '.')
    else:
        _sentences.append(_s)
# Replace the target paragraph's text with the first sentence (preserves
# its style); inject the remaining sentences as new paragraphs immediately
# after, also preserving the target's style.
_target.text = _sentences[0]
_style = _target.style
_anchor = _target
for _s in _sentences[1:]:
    _new = _doc.add_paragraph(_s, style=_style)
    _anchor._element.addnext(_new._element)
    _anchor = _new
_doc.save({gold!r})
"""


def _gold_text_to_table(src: str, gold: str, *,
                        para_idx: int, sep: str = ",") -> str:
    """Convert paragraph[para_idx] (a `sep`-separated value list) into a
    1-row table with one column per token. All OTHER paragraphs (and
    document-level formatting / styles / sectPr) preserved verbatim.

    Eval kind `files` (compare_docx_files) is strict on body content
    including paragraph runs and table cell text. The earlier `_dst =
    Document()` empty-rebuild lost every source paragraph's formatting
    runs (only `.text` was copied), and the agent's LO Writer
    Table→Convert→Text-to-Table operation does NOT strip run formatting,
    so any seed with non-default run formatting would diverge. Clone
    `src` to `gold`, then splice the table in-place at para_idx and
    drop the original paragraph element — same XML shape LO Writer
    produces."""
    return f"""\
import shutil
from docx import Document
shutil.copyfile({src!r}, {gold!r})
_doc = Document({gold!r})
_target = _doc.paragraphs[{para_idx}]
_cells = [_c.strip() for _c in _target.text.split({sep!r})]
_tbl = _doc.add_table(rows=1, cols=len(_cells))
for _j, _v in enumerate(_cells):
    _tbl.rows[0].cells[_j].text = _v
# Move new table (currently appended at body end) into the slot where
# para_idx lived; then delete the original paragraph element.
_target._element.addprevious(_tbl._element)
_target._element.getparent().remove(_target._element)
_doc.save({gold!r})
"""


def _src_pre_highlighted_phone_script():
    """File.src callable: builds a docx mimicking the recruitment-phone-script
    style of osworld_libreoffice_writer_6a33f9b9 — every paragraph carries
    yellow highlight on every run, simulating the user's pre-flagged
    "needs rewriting" markup. The agent's task is to REMOVE all highlight
    (the opposite operation of the dominant ADD-highlight synth coverage).
    """
    _PARAS = [
        "Sample recruitment phone script",
        "Good morning, this is the talent team calling on behalf of the regional staffing office regarding the engineering opening you applied for last week through our online portal.",
        "Before we go further, I want to confirm a few details that were flagged for follow-up: your earliest availability, your current notice period, and whether the listed salary range still works for you.",
        "On the role itself, this is a mid-level individual contributor position embedded within the platform reliability group, reporting to the team lead who handles incident response and on-call rotation planning.",
        "If everything I have shared still aligns with what you are looking for, the next step is a one-hour technical screen with the hiring manager, scheduled at a time of mutual convenience.",
        "Thank you for your time today, and please reply by end of week to the email I sent earlier this morning so we can lock in the technical screen slot before the calendar fills up.",
    ]
    def _src(path, _seed):
        py = f"""\
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
doc = Document()
for _t in {_PARAS!r}:
    _p = doc.add_paragraph(_t)
    for _r in _p.runs:
        _r.font.highlight_color = WD_COLOR_INDEX.YELLOW
doc.save({path!r})
"""
        return [_make_config_step(py)]
    return _src


def _src_pre_highlighted_meeting_notes():
    """Variant pre-highlighted source (meeting-notes flavour) for a second
    REMOVE-highlight task. Every paragraph carries yellow highlight on every
    run; gold removes them all."""
    _PARAS = [
        "Weekly programme review — draft notes",
        "Attendees confirmed at the standing review: programme manager, two engineering leads, the design lead, and the rotating analytics liaison covering the cross-team data pipeline workstream this fortnight.",
        "Open items carried forward from last cycle include the dependency conflict on the rollout calendar, the unresolved budget variance on the staging environment, and the still-pending procurement approval for the load-test cluster.",
        "Decisions taken in today's session: defer the rollout window by one cycle pending the procurement decision, fold the staging variance into the next quarterly recast, and assign the dependency mitigation to the platform on-call rotation for short-burst follow-up.",
        "Action items captured for the next standing review: the programme manager will share an updated rollout calendar by Wednesday, the analytics liaison will reconcile the staging variance with finance by end of week, and the platform lead will publish the dependency-mitigation plan in the team wiki.",
    ]
    def _src(path, _seed):
        py = f"""\
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
doc = Document()
for _t in {_PARAS!r}:
    _p = doc.add_paragraph(_t)
    for _r in _p.runs:
        _r.font.highlight_color = WD_COLOR_INDEX.YELLOW
doc.save({path!r})
"""
        return [_make_config_step(py)]
    return _src


def _src_csv_paragraph():
    """File.src callable: a memo-shaped doc with two comma-separated value
    paragraphs (one contact row, one project-status row) embedded amid
    realistic multi-sentence prose. Each Param targets one of the CSV
    paragraphs by its unique semantic anchor ("Alice Chen" / "Project
    Apollo") rather than positional index, since the prose paragraphs
    around it would otherwise make "the Nth paragraph" ambiguous to the
    agent reading the live document. Paragraph indices below match the
    `para_idx` values referenced by FileTask params using this fixture."""
    _PARAS = [
        # 0 — section heading
        "Team Contact and Project Status Register",
        # 1 — introductory prose (multi-sentence, sets context)
        "This memo captures the current point-of-contact for each active "
        "engineering workstream alongside a quick snapshot of where the "
        "flagship project sits this week. Please use it as the canonical "
        "reference until the next quarterly refresh; route any corrections "
        "through the program-office mailbox so the registry stays clean.",
        # 2 — TARGET row (CSV contact entry — Param para_idx=2)
        "Alice Chen, Senior Engineer, alice@example.com, ext 4012",
        # 3 — separator prose between the two CSV anchors
        "Recent organisational moves may have shifted reporting lines for a "
        "handful of teams. If a name no longer matches the chart on your "
        "side of the org, ping HR-Operations to refresh the registry before "
        "the next standing meeting so onboarding documents stay accurate.",
        # 4 — sub-heading
        "Project Status Snapshot",
        # 5 — prose introducing the status row
        "The summary below reflects this week's stand-up notes; the full "
        "milestone dashboard with burndown charts and dependency graph "
        "lives in the project portal under the engineering programs space.",
        # 6 — TARGET row (CSV status entry — Param para_idx=6)
        "Project Apollo, Phase 2, On Track, due May 30",
        # 7 — closing prose
        "Please raise blockers in the #project-apollo channel rather than "
        "via direct message so the program manager can triage in context "
        "and surface cross-team dependencies during the daily sync.",
    ]
    def _src(path, _seed):
        py = f"""\
from docx import Document
_doc = Document()
for _p in {_PARAS!r}:
    _doc.add_paragraph(_p)
_doc.save({{path!r}})
""".replace("{path!r}", repr(path))
        return [_make_config_step(py)]
    return _src


# F-WRITER-102 — validation eval-coverage fill (text→table conversion).
# Only NEW File needed: existing fixtures don't have a paragraph that is
# itself comma-separated values. The other 3 niche skills (center heading,
# delete rows by condition, sentence-per-line) reuse existing real-source
# Files: F_WRITER_94 (chemistry notes), F_WRITER_49 (printer troubleshoot
# with severity table), F_WRITER_98 (multi-sentence research report).
F_WRITER_102 = File("F-WRITER-102", "writer_csv_text",
                    "csv_text.docx", _src_csv_paragraph())


# ----- Real-photo image-host files (C, validation expansion) ----------------
# Four additional photo-host fixtures leveraging the newly-added asset
# directories: photos/energy/, photos/industrial/, photos/education/.
# Mirrors the F_WRITER_26..30 pattern: 3-paragraph mini-document where the
# second paragraph hosts the inserted image.
F_WRITER_103 = File("F-WRITER-103", "writer_image_host_wind_turbine",
                    "renewable_brief.docx",
                    _src_image_host(
                        "photos/energy/wind-turbine.jpg",
                        "wind-turbine.jpg",
                        [
                            "Renewable energy briefing — onshore and offshore wind capacity has expanded dramatically over the past decade, with global installed nameplate capacity passing nine hundred gigawatts in the most recent annual review from the international energy agency. The current generation of three-bladed horizontal-axis turbines now routinely exceeds five megawatts per unit at favourable coastal sites, and the largest offshore prototypes under construction in the North Sea approach the fifteen-megawatt threshold once thought to be a theoretical ceiling for the technology.",
                            "Insert the wind-turbine photograph below this paragraph as the cover plate of the briefing document. Capacity factors on modern offshore installations now average between forty and fifty percent over a representative twelve-month period, comfortably above the thirty-percent figure that characterised onshore fleets a decade ago. Levelised cost of energy has fallen in parallel with these performance gains, and several recent auctions in northern Europe have cleared at strike prices below the prevailing wholesale electricity benchmark.",
                            "Adoption outlook — analysts at the major consulting firms now project that wind will supply between twenty and thirty percent of total electricity demand across the OECD economies by the middle of the next decade, contingent on continued investment in transmission infrastructure and storage capacity. The bottleneck for further growth is no longer turbine cost or siting opposition but rather the lead time on high-voltage interconnects between coastal wind zones and the inland load centres that consume the bulk of the generated power.",
                        ],
                    ))
F_WRITER_104 = File("F-WRITER-104", "writer_image_host_cargo_port",
                    "logistics_brief.docx",
                    _src_image_host(
                        "photos/industrial/cargo-port.jpg",
                        "cargo-port.jpg",
                        [
                            "Maritime logistics dossier — the container port photographed for the cover plate handles in excess of four million twenty-foot-equivalent units per year and ranks among the twenty busiest seaports in the region by throughput volume. The terminal operates around the clock across three eight-hour shifts, with rubber-tyred gantry cranes and ship-to-shore quay cranes coordinated through a yard management system that tracks each container from vessel arrival to inland gate-out in near real time.",
                            "Place the cargo-port photograph below this paragraph for the printed dossier. Recent investment in two additional super-post-Panamax cranes has lifted the terminal's peak hourly handling rate by roughly eighteen percent over the figures recorded at the same point last year, and the new automation pilot in the rail-mounted gantry yard has reduced average truck dwell time at the inland gate by close to twelve minutes per visit. The terminal operator expects further productivity gains as the automation rollout extends to the remaining yard blocks.",
                            "Global trade context — container shipping continues to handle roughly four fifths of world merchandise trade by volume despite the recent rerouting of long-haul services away from the Red Sea corridor. The major liner alliances are progressively retiring smaller panamax-class vessels in favour of twenty-four-thousand-TEU mega-ships, a shift that places sustained pressure on landside ports to deepen approach channels and lengthen quay walls to accommodate the new tonnage.",
                        ],
                    ))
F_WRITER_105 = File("F-WRITER-105", "writer_image_host_graduation",
                    "commencement_essay.docx",
                    _src_image_host(
                        "photos/education/graduation-ceremony.jpg",
                        "graduation-ceremony.jpg",
                        [
                            "Commencement reflection — four years ago the members of this graduating class arrived on campus as strangers, wheeling cardboard boxes up dormitory staircases and clutching maps of a place that has since become the most familiar landscape any of them has ever known. The intervening seasons have layered on top of one another in a way that is difficult to summarise without resorting to the familiar imagery of changed leaves, snowed-in libraries, and long evenings spent in conversation on the steps of the student union.",
                            "Insert the commencement ceremony photograph below this paragraph as the centrepiece of the printed essay. The photograph captures the moment at which the assembled class first stood together as graduates, an instant that the speakers earlier in the programme described as both an ending and a beginning. Those of us who have watched several such ceremonies will recognise the curious mixture of relief and apprehension that settles over the rows of caps and gowns once the conferral of degrees is formally complete.",
                            "Closing address — to the graduating class, the message from the faculty is straightforward and earnestly meant: the credentials you carry from this institution are valuable, but the friendships and the habits of mind that you have cultivated here are the more durable inheritance. Wherever the next chapter takes you, return to this campus from time to time, and remember that the doors of your former departments remain open to alumni in good standing in perpetuity.",
                        ],
                    ))
F_WRITER_106 = File("F-WRITER-106", "writer_image_host_solar_farm",
                    "solar_briefing.docx",
                    _src_image_host(
                        "photos/energy/solar-farm.jpg",
                        "solar-farm.jpg",
                        [
                            "Solar adoption note — utility-scale photovoltaic farms of the kind illustrated on the cover plate now constitute the single largest source of new generating capacity coming online each year across most major electricity markets. Module efficiency for the dominant crystalline-silicon technology has crept past twenty-two percent at the cell level in commercial production lines, and tandem perovskite-silicon prototypes under development at several research consortia have demonstrated laboratory efficiencies approaching thirty-three percent under standard test conditions.",
                            "Insert the solar farm photograph below this paragraph for the briefing's cover plate. Tracking-mounted installations of the type photographed here typically generate between fifteen and twenty percent more energy per installed watt over a full year than equivalent fixed-tilt arrays at the same latitude, and the additional mechanical complexity has been largely paid down by continued reductions in the per-watt cost of the underlying module hardware. Operations and maintenance costs on tracking systems have also fallen as second-generation single-axis drives have matured.",
                            "Adoption outlook — the cumulative cost reduction observed in photovoltaic modules over the past fifteen years remains the most consistently underestimated trend in the long-range energy forecasting literature. Several recent analyses now project that solar will supply between a quarter and a third of global electricity demand by the middle of the next decade, with the trajectory beyond that point depending heavily on the pace at which long-duration storage and demand-side flexibility are deployed alongside the new generating capacity.",
                        ],
                    ))


# ----- F_WRITER_107 (train-records duplicate-removal worksheet) -----------
# Source = 10 CSV-shaped paragraphs with 3 duplicate train_ids; agent's
# task is to delete the duplicate rows. Evaluator = compare_unique_train_
# records (3-file: result/gold/initial). Uses _src_train_records which
# also drops the initial snapshot at /tmp/initial_<basename>.docx.
F_WRITER_107 = File("F-WRITER-107", "writer_train_records",
                    "train_register.docx", _src_train_records())


# ----- F_WRITER_108 (chemistry brief — compound subscript+files target) ---
# Separate chemistry-style fixture so a 2nd subscript-bearing FileTask can
# wire the compound `[compare_docx_files, compare_subscript_contains]`
# evaluator (mirrors eval row 0b17a146 — H2O subscript). F_WRITER_94 is
# already at its 2-task cap.
def _src_chemistry_brief():
    title = "Lab Brief — Water and Carbon Dioxide"
    intro = ("This brief summarises the two simplest molecular formulas "
             "covered in the orientation session. Both involve a single "
             "subscript digit that distinguishes the molecular form from "
             "the individual atoms.")
    formulas = ("Water is written H2O and carbon dioxide is written CO2; "
                "in each case the digit immediately following the element "
                "symbol records how many atoms of that element occur in "
                "one molecule of the compound.")
    closing = ("Please apply the standard subscript convention to the "
               "digit inside the water formula before circulating this "
               "brief to the new lab cohort.")
    blocks = [("p", intro), ("p", formulas), ("p", closing)]
    return _src_structured(title=title, blocks=blocks,
                           font_name="Liberation Serif", font_size_pt=12)


F_WRITER_108 = File("F-WRITER-108", "writer_chemistry_brief",
                    "chem_brief.docx", _src_chemistry_brief())


# ----- ODT highlight-removal helpers (emulates osworld_libreoffice_writer_
# 6a33f9b9 — the real eval pulls a .odt and runs `check_highlighted_words`
# which calls `odf.opendocument.load()`). The synth pipeline was docx-only
# pre-validation; the helpers below add a parallel .odt path:
#
#   source: python-docx writes a docx with yellow highlight on every run,
#           then `soffice --convert-to odt` produces the editable .odt.
#   gold:   python-docx writes a docx with NO highlight (text identical),
#           then `soffice --convert-to odt` produces the .odt expected.
#
# `check_highlighted_words` first runs `compare_docx_files` (paragraph text
# equality with ignore_blanks=True) and then checks that no Span carries
# automatic-style backgroundcolor='#ffff00'. Both source and gold share
# the same paragraphs list so the text-equality atom is trivially satisfied
# and the highlight-presence atom is the only one that discriminates.
def _src_pre_highlighted_odt(paras: list[str]):
    """File.src callable that yields ONE .odt with yellow highlight on every
    paragraph. Built via stdlib zipfile (odfpy is not in the container) so the
    .odt carries `fo:background-color="#ffff00"` text-properties — the docx→odt
    soffice chain dropped highlights, leaving source identical to gold and
    trivial_pass'ing the eval.
    """
    def _src(path: str, _seed: int) -> list[dict]:
        py = f"""\
import zipfile, html
paras = {paras!r}
spans = ''.join(
    '<text:p text:style-name="P1"><text:span text:style-name="HL">'
    + html.escape(t) + '</text:span></text:p>'
    for t in paras
)
content_xml = (
    '<?xml version="1.0" encoding="UTF-8"?>'
    '<office:document-content '
    'xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" '
    'xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" '
    'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" '
    'xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" '
    'office:version="1.2">'
    '<office:automatic-styles>'
    '<style:style style:name="P1" style:family="paragraph">'
    '<style:paragraph-properties/>'
    '</style:style>'
    '<style:style style:name="HL" style:family="text">'
    '<style:text-properties fo:background-color="#ffff00"/>'
    '</style:style>'
    '</office:automatic-styles>'
    '<office:body><office:text>' + spans + '</office:text></office:body>'
    '</office:document-content>'
)
manifest_xml = (
    '<?xml version="1.0" encoding="UTF-8"?>'
    '<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0">'
    '<manifest:file-entry manifest:full-path="/" '
    'manifest:media-type="application/vnd.oasis.opendocument.text"/>'
    '<manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/>'
    '</manifest:manifest>'
)
with zipfile.ZipFile({path!r}, 'w', zipfile.ZIP_DEFLATED) as z:
    # mimetype must be the first entry and stored uncompressed.
    info = zipfile.ZipInfo('mimetype'); info.compress_type = zipfile.ZIP_STORED
    z.writestr(info, 'application/vnd.oasis.opendocument.text')
    z.writestr('META-INF/manifest.xml', manifest_xml)
    z.writestr('content.xml', content_xml)
"""
        return [_make_config_step(py)]
    return _src


def _gold_unhighlight_all_odt(src_odt: str, gold_odt: str, *,
                              paras: list[str]) -> str:
    """Gold heredoc producing the unhighlighted .odt twin. Built via stdlib
    zipfile (odfpy is not in the container) so the gold has plain paragraphs
    without any highlight span — the eval's `compare_docx_files` atom passes
    (same text) and the highlight-presence atom discriminates.
    """
    return f"""\
import zipfile, html
paras = {paras!r}
ps = ''.join(
    '<text:p text:style-name="P1">' + html.escape(t) + '</text:p>'
    for t in paras
)
content_xml = (
    '<?xml version="1.0" encoding="UTF-8"?>'
    '<office:document-content '
    'xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" '
    'xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" '
    'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" '
    'xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" '
    'office:version="1.2">'
    '<office:automatic-styles>'
    '<style:style style:name="P1" style:family="paragraph">'
    '<style:paragraph-properties/>'
    '</style:style>'
    '</office:automatic-styles>'
    '<office:body><office:text>' + ps + '</office:text></office:body>'
    '</office:document-content>'
)
manifest_xml = (
    '<?xml version="1.0" encoding="UTF-8"?>'
    '<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0">'
    '<manifest:file-entry manifest:full-path="/" '
    'manifest:media-type="application/vnd.oasis.opendocument.text"/>'
    '<manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/>'
    '</manifest:manifest>'
)
with zipfile.ZipFile({gold_odt!r}, 'w', zipfile.ZIP_DEFLATED) as z:
    info = zipfile.ZipInfo('mimetype'); info.compress_type = zipfile.ZIP_STORED
    z.writestr(info, 'application/vnd.oasis.opendocument.text')
    z.writestr('META-INF/manifest.xml', manifest_xml)
    z.writestr('content.xml', content_xml)
"""


_PARAS_PHONE_ODT = [
    "Sample recruitment phone script",
    "Good morning, this is the talent team calling on behalf of the regional staffing office regarding the engineering opening you applied for last week through our online portal.",
    "Before we go further, I want to confirm a few details that were flagged for follow-up: your earliest availability, your current notice period, and whether the listed salary range still works for you.",
    "On the role itself, this is a mid-level individual contributor position embedded within the platform reliability group, reporting to the team lead who handles incident response and on-call rotation planning.",
    "If everything I have shared still aligns with what you are looking for, the next step is a one-hour technical screen with the hiring manager, scheduled at a time of mutual convenience.",
    "Thank you for your time today, and please reply by end of week to the email I sent earlier this morning so we can lock in the technical screen slot before the calendar fills up.",
]


_PARAS_RECIPE_ODT = [
    "Slow-cooked tomato ragu — kitchen draft",
    "Warm two tablespoons of olive oil in a heavy-bottomed pan over medium heat until the oil shimmers but has not yet begun to smoke, then add the diced onion and a pinch of fine salt.",
    "Sweat the onion for roughly eight minutes, stirring occasionally, until the pieces turn translucent and just barely start to take on colour around the edges of the pan.",
    "Add the minced garlic and stir for forty-five seconds, taking care not to let the garlic catch and turn bitter before the tomatoes go in to bring the temperature back down.",
    "Pour the tinned tomatoes into the pan, crush them gently with the back of a wooden spoon, season with the dried oregano and freshly ground black pepper, and bring the sauce up to a gentle simmer.",
    "Cover the pan loosely, reduce the heat to low, and leave the ragu to cook for forty minutes, stirring every ten or so to keep the bottom of the pan from catching as the sauce thickens.",
    "Taste once near the end and adjust the salt before serving over freshly drained pasta with a generous shower of grated parmesan and a few torn basil leaves on top.",
]


# ----- F_WRITER_109..110 (pre-highlighted docs — REMOVE-highlight target) -
# Each fixture's source builder bakes yellow highlight onto every run of
# every paragraph; the agent's task is to REMOVE all highlights. Mirrors
# osworld_libreoffice_writer_6a33f9b9 ("remove all highlight"). Synth had
# zero REMOVE-highlight coverage prior to this fixture — all existing
# highlight rows ADD highlight.
F_WRITER_109 = File("F-WRITER-109", "writer_pre_highlighted_phone",
                    "phone_script.docx",
                    _src_pre_highlighted_phone_script())
F_WRITER_110 = File("F-WRITER-110", "writer_pre_highlighted_notes",
                    "meeting_notes.docx",
                    _src_pre_highlighted_meeting_notes())


# ----- F_WRITER_114..115 (.odt twins — the eval's `check_highlighted_words`
# REQUIRES .odt because the upstream evaluator calls `odf.opendocument.load`.
# validation added F_WRITER_109/110 as the docx-side REMOVE-highlight
# fixtures; validation below adds the .odt twins that actually wire the upstream
# `check_highlighted_words` function. Two fixtures span distinct domains
# (recruitment phone script vs. tomato ragu recipe) to keep param-pair
# Jaccard well under 0.3.
F_WRITER_114 = File("F-WRITER-114", "writer_pre_highlighted_phone_odt",
                    "sample-recruitment-phone-script.odt",
                    _src_pre_highlighted_odt(_PARAS_PHONE_ODT))
F_WRITER_115 = File("F-WRITER-115", "writer_pre_highlighted_recipe_odt",
                    "tomato-ragu-draft.odt",
                    _src_pre_highlighted_odt(_PARAS_RECIPE_ODT))


# ----- F_WRITER_111 (second comma→table fixture) --------------------------
# Mirrors osworld_libreoffice_writer_936321ce more directly — the eval task
# uses a teaching worksheet with a single comma-separated list. Provides a
# second variant beyond F_WRITER_102 for added doc-wide table-conversion
# coverage.
def _src_phoneme_csv():
    """File.src callable: a single-list teaching worksheet whose body has
    one comma-separated phoneme line, plus surrounding intro/closing prose.
    The CSV row is paragraph index 1."""
    _PARAS = [
        "Phoneme reference list — vowel digraphs",
        "ai, ay, ea, ee, ie, oa, ow, oo, ue, ui",
        "Tutors may use this reference table in lesson planning for the early-readers cohort, particularly during the phonics review block at the start of each session.",
        "When circulating this list to families, please attach the most recent edition of the home-reading guidance so caregivers can apply the digraph cues consistently outside of class time.",
    ]
    def _src(path, _seed):
        py = f"""\
from docx import Document
doc = Document()
for _p in {_PARAS!r}:
    doc.add_paragraph(_p)
doc.save({path!r})
"""
        return [_make_config_step(py)]
    return _src


F_WRITER_111 = File("F-WRITER-111", "writer_phoneme_csv",
                    "phoneme_worksheet.docx", _src_phoneme_csv())


# ----- F_WRITER_112 (class-schedule image host — mirrors osworld_libre
# office_writer_6ada715d "Viewing Your Class Schedule and Textbooks").
# A fresh image-host docx whose body is registrar-style "where your
# cursor is located" so the compare_docx_images eval has a fixture
# dedicated to the cursor-anchored insertion idiom.
F_WRITER_112 = File("F-WRITER-112", "writer_class_schedule",
                    "class_schedule.docx",
                    _src_image_host(
                        "photos/education/graduation-ceremony.jpg",
                        "schedule_screenshot.jpg",
                        [
                            "Viewing your class schedule and textbooks — this short guide walks new students through pulling the latest term schedule from the registrar portal, exporting the PDF, and saving a screenshot to the desktop for offline reference during the orientation week.",
                            "Once the screenshot is saved on your desktop as 1.png (or with whatever filename your screenshot tool used), copy it into this guide at the cursor anchor below so the printed packet has the schedule reference inline with the textbook list.",
                            "After the schedule image is inserted, scroll to the next section for the textbook checklist and the campus map. If anything in your schedule is missing a section number, return to the portal and re-export before the start-of-term meeting with your academic adviser.",
                        ],
                    ))


# ----- F_WRITER_113 (sentence-split host — mirrors osworld_libreoffice_
# writer_88fe4b2d "separate each sentence in the first paragraph"). A
# real-source-style fixture whose opening paragraph is a single long
# multi-sentence block, ideally suited to the sentence-per-line operation.
def _src_tutorial_guidelines():
    """File.src callable: a tutorial-guidelines fixture where the first
    three paragraphs are each multi-sentence blocks (suitable for the
    sentence-per-line split + the 3-paragraph compound variant). The
    sentence_per_line gold splits on '. ' boundaries, so every targeted
    paragraph must contain at least one '. ' internally."""
    _PARAS = [
        "Tutorial guidelines — please read carefully before the first session. Attendance is expected at every weekly tutorial in addition to the main lecture series. Late arrivals after the first ten minutes will be marked as absent for purposes of the participation grade. Mobile phones must be silenced for the duration of the session and stowed out of sight. Questions are encouraged at any point but please raise your hand and wait to be called on by the tutor.",
        "Tutorial readings are circulated by email at least seventy-two hours in advance of the corresponding session. Every student is expected to arrive having read the assigned material at least once. Skimming on the morning of the session is not sufficient preparation. The tutor will assume working familiarity with the readings from the very first question.",
        "Group assignments are weighted at thirty percent of the final tutorial grade. Groups will be assigned during the second week of term based on student preferences gathered via the orientation survey. Each group typically contains four students drawn from across the cohort. Self-organised groupings are discouraged in the first half of term to encourage cross-disciplinary mixing.",
        "If you are unable to attend a tutorial for documented medical or family reasons, please email the tutor in advance and arrange to attend a make-up session within the same teaching block.",
    ]
    def _src(path, _seed):
        py = f"""\
from docx import Document
doc = Document()
for _p in {_PARAS!r}:
    doc.add_paragraph(_p)
doc.save({path!r})
"""
        return [_make_config_step(py)]
    return _src


F_WRITER_113 = File("F-WRITER-113", "writer_tutorial_guidelines",
                    "tutorial_guidelines.docx", _src_tutorial_guidelines())


# §I.e — Factory + emit.

def _to_synth_template(ft: FileTask) -> SynthTemplate:
    """Compile one FileTask into a SynthTemplate. eval_kind dispatches to
    the matching `_build_*_evaluator` helper from earlier in this file.

    Special pdf_export branch: gold heredoc returns empty string and the
    factory injects a `_soffice_docx_to_pdf_step` to build the gold pdf
    instead of a gold docx; evaluator is the 4-way list-form
    `["compare_pdfs"]*4` with conj=or; oracle copies gold-pdf to the
    Desktop sink (one of the 4 dirs the `or` accepts)."""
    pool = ft.params[:SYNTH_CAP_PARAMS_PER_TASK]
    template_id = f"{ft.file.id.lower().replace('-', '_')}__{ft.task_id}"

    is_pdf_export = any(p.eval_kind == "pdf_export" for p in pool)

    def _params(seed: int) -> dict:
        variant = pool[seed % len(pool)]
        # validation: .odt fixtures (F_WRITER_114/115 for the
        # `check_highlighted_words` upstream evaluator) keep the .odt
        # extension on both source and gold so the evaluator's odfpy
        # `load()` succeeds. All other fixtures stay docx.
        is_odt = ft.file.basename.endswith(".odt")
        if is_odt:
            stem = ft.file.basename[:-len(".odt")]
            src_path = f"{_DESKTOP}/{stem}_{seed:04d}.odt"
        else:
            src_path = f"{_DESKTOP}/{ft.file.basename.replace('.docx', '')}_{seed:04d}.docx"
        if variant.eval_kind == "pdf_export":
            # Gold path is a PDF, sink path lands the converted docx PDF.
            gold_path = f"/tmp/expected_{template_id}_{seed:04d}.pdf"
        elif is_odt:
            gold_path = f"/tmp/expected_{template_id}_{seed:04d}.odt"
        else:
            gold_path = f"/tmp/expected_{template_id}_{seed:04d}.docx"
        pre = list(ft.file.src(src_path, seed))
        if variant.eval_kind == "pdf_export":
            # Build gold-pdf via headless soffice (mirrors the agent's
            # File > Export As > PDF chain → identical text → fuzz=1.0).
            pre.append(_soffice_docx_to_pdf_step(src_path, gold_path))
        else:
            gold_body = ft.gold(src_path, gold_path, **variant.gold_args)
            if gold_body:
                pre.append(_make_config_step(gold_body))
        # `highlighted_words` eval (F_WRITER_114/115 .odt fixtures) must
        # bypass LO_SAVE_POSTCONFIG entirely — LO Writer's Ctrl+S drops the
        # `fo:background-color="#ffff00"` span on .odt roundtrip, making
        # pre-eval read a source.odt with no highlights → trivial_pass.
        # Paired with `postconfig: []` in `_build_highlighted_words_evaluator`.
        # `oracle_after_postconfig=False` keeps oracle runs from being
        # gated on a postconfig (Ctrl+S) that never executes.
        oracle_after_pc = variant.eval_kind != "highlighted_words"
        return {
            "instr":             variant.instr,
            "out_path":          src_path,
            "expected_path":     gold_path,
            "_eval_kind":        variant.eval_kind,
            "_eval_args":        variant.eval_args,
            "_template_id":      template_id,
            "pre_config_steps":  pre,
            "open_command":      ["libreoffice", "--writer", src_path],
            "oracle_after_postconfig": oracle_after_pc,
            "post_open_config_steps": _WRITER_BODY_FOCUS_STEPS,
        }

    def _eval(p: dict) -> dict:
        kind = p["_eval_kind"]
        args = p["_eval_args"]
        src, gold = p["out_path"], p["expected_path"]
        if kind == "strict":
            return _build_strict_evaluator(src, gold, **args)
        if kind == "files":
            return _build_files_evaluator(src, gold, **args)
        if kind == "tables":
            return _build_tables_evaluator(src, gold)
        if kind == "line_spacing":
            return _build_line_spacing_evaluator(src, gold)
        if kind == "page_break":
            return _build_page_break_evaluator(src, args.get("expected_count", 1))
        if kind == "font_names":
            return _build_font_names_evaluator(src, args["font_name"])
        if kind == "default_font":
            return _build_default_font_evaluator(args["font_name"])
        if kind == "first_centered":
            return _build_first_centered_evaluator(src)
        if kind == "page_numbers":
            return _build_page_numbers_evaluator(src)
        if kind == "strike_last":
            return _build_strike_last_para_evaluator(src, gold)
        if kind == "italic_size14":
            return _build_italic_size14_evaluator(src, gold)
        if kind == "colored_table":
            return _build_colored_table_evaluator(src, gold)
        if kind == "images":
            # compare_docx_images — mirrors osworld_libreoffice_writer_6ada715d.
            return _build_images_evaluator(src, gold)
        if kind == "highlighted_words":
            # check_highlighted_words — REAL upstream func from osworld_
            # libreoffice_writer_6a33f9b9. Requires .odt source+gold.
            return _build_highlighted_words_evaluator(src, gold)
        if kind == "subscript":
            return _build_subscript_contains_evaluator(src, gold)
        if kind == "tabstops":
            return _build_tabstops_evaluator(src, gold, **args)
        if kind == "unique_train_records":
            # `unique_train_records` needs the *initial* (pre-edit) docx
            # as a third reference. The source builder pairs the agent's
            # editable copy with a sibling initial-snapshot at
            # /tmp/initial_<basename>.
            init_path = f"/tmp/initial_{os.path.basename(src)}"
            return _build_unique_train_records_evaluator(src, gold, init_path)
        if kind == "compound":
            # args["atoms"] = [{"func": "compare_docx_files", "expected_path": gold},
            #                  {"func": "compare_subscript_contains", "expected_path": gold}]
            # When an atom's expected_path is the same as the global gold, the
            # caller passes gold; otherwise a per-atom gold path (e.g. for
            # disjunctive same-fn ambiguity), but this synth keeps it simple:
            # AND-style multi-property checks all sharing the same gold doc.
            atoms_def = args["atoms"]
            atoms = []
            for a in atoms_def:
                ep = a.get("expected_path", gold)
                if isinstance(ep, str) and ep == "__GOLD__":
                    ep = gold
                atoms.append({"func": a["func"], "expected_path": ep,
                              "options": a.get("options", {})})
            return _build_compound_evaluator(src, atoms)
        if kind == "pdf_export":
            # 4-way list-form evaluator (mirrors eval task 4bcb1253) — `or`
            # conj means a single matching dir/pdf pair passes. Oracle
            # lands the gold pdf at the Desktop sink (1st of the 4 dirs).
            pdf_basename = os.path.basename(src).replace(".docx", ".pdf")
            return {
                "func": ["compare_pdfs"] * 4,
                "conj": "or",
                "result": [
                    {"type": "vm_file", "path": f"{d}/{pdf_basename}",
                     "dest": f"result_{i}.pdf"}
                    for i, d in enumerate(_PDF_EXPORT_TARGET_DIRS, 1)
                ],
                "expected": [
                    {"type": "vm_file", "path": gold,
                     "dest": f"expected_{i}.pdf"}
                    for i in range(1, 5)
                ],
                "postconfig": LO_SAVE_POSTCONFIG,
            }
        raise ValueError(f"unknown eval_kind: {kind!r}")

    def _oracle(p: dict) -> list[dict]:
        if p["_eval_kind"] == "pdf_export":
            # cp gold pdf to Desktop sink (1st dir; `or` accepts any of 4).
            pdf_basename = os.path.basename(p["out_path"]).replace(".docx", ".pdf")
            sink = f"{_PDF_EXPORT_TARGET_DIRS[0]}/{pdf_basename}"
            return [_execute(f"cp '{p['expected_path']}' '{sink}'")]
        if p["_eval_kind"] == "default_font":
            # The evaluator reads registrymodifications.xcu, not the docx.
            # Write the xcu entry directly instead of the standard docx oracle.
            return _build_default_font_oracle(p["_eval_args"]["font_name"])
        if p["out_path"].endswith(".odt"):
            # .odt-fixture oracle: normalize gold .odt,
            # plant at sink, normalize sink. Symmetric round-trip keeps
            # `check_highlighted_words` byte-stable across LO's Ctrl+S.
            return _build_odt_oracle(p["out_path"], p["expected_path"])
        return _build_oracle(p["out_path"], p["expected_path"])

    return SynthTemplate(
        template_id=template_id,
        domain="libreoffice_writer",
        instruction_fn=lambda p: p["instr"],
        evaluator_fn=_eval,
        oracle_fn=_oracle,
        postconfig_fn=lambda _p: None,
        param_fn=_params,
        n_rows=len(pool),
        eval_class=ft.eval_class,
        setup_class=ft.file.setup_class,
        open_command=None,
    )


def _emit_templates(file_tasks: list[FileTask]) -> list[SynthTemplate]:
    per_file: dict[str, int] = {}
    out: list[SynthTemplate] = []
    for ft in file_tasks:
        c = per_file.get(ft.file.id, 0)
        if c >= SYNTH_CAP_TASKS_PER_FILE:
            continue
        per_file[ft.file.id] = c + 1
        out.append(_to_synth_template(ft))
    return out


# §I.f — FILE_TASKS.
#
# Per-file the cap is ≤2 tasks × ≤2 params = ≤4 emitted rows. Headroom
# (>2 tasks listed) is preserved in the source for future ablations but
# only the first 2 tasks are emitted (cf. _emit_templates).
#
# Eval-class distribution targets the 11 writer taxonomy buckets:
#   bold_text, change_font, change_line_spacing, add_header_footer,
#   highlight_text, insert_image, color_table_text, pdf_export,
#   text_to_table, blank_table_insert, footnote_citation.
#
# Prefer compare_docx_files (eval_kind=
# "files") over compare_docx_strict where the task is text-equivalent
# (find_replace, append_paragraph) — writer eval is heavy on text-only
# comparison and light on strict (35 perturb / 0 eval).

FILE_TASKS: list[FileTask] = [
    # ====================================================================
    # F-WRITER-1 (policy short fixture) — bold + italic per-paragraph
    # ====================================================================
    FileTask(F_WRITER_1, "bold_para", "bold_text",
             _gold_bold_para, params=[
        Param({"para_idx": 0}, "strict", {},
              "I'm finalising this policy for the printed handout and would like the opening notice to stand out — please bold the first paragraph of the document."),
        Param({"para_idx": 2}, "strict", {}, "Bold the third paragraph of the document."),
    ]),
    FileTask(F_WRITER_1, "append_signoff", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Signed: People Operations, May 2026."}, "files", {},
              "I'd like to add a sign-off line before this policy goes out to staff — please append a new paragraph at the end of the document with the text 'Signed: People Operations, May 2026.'."),
        Param({"text": "Reviewed and approved by the Compliance team."}, "files", {},
              "I'd like to flag that compliance has signed off on this policy before it goes into the records system — could you add a new paragraph at the very end of the document reading 'Reviewed and approved by the Compliance team.'?"),
    ]),

    # ====================================================================
    # F-WRITER-2 (essay short fixture) — bold + line spacing
    # ====================================================================
    FileTask(F_WRITER_2, "bold_para", "bold_text",
             _gold_bold_para, params=[
        Param({"para_idx": 1}, "strict", {},
              "I'm peer-reviewing this short essay for a friend and want to flag the second paragraph for closer attention — could you bold the second paragraph for me?"),
        Param({"para_idx": 4}, "strict", {},
              "I'd like to highlight the closing reflection of this essay so the reader takes that thought with them — could you bold the fifth paragraph for me?"),
    ]),
    # Validation PARAM_REDUCIBLE: dropped 2.0-spacing Param;
    # kept 1.5 (the simpler/more common spacing value).
    FileTask(F_WRITER_2, "doc_spacing", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 1.5}, "line_spacing", {},
              "I'd like to leave a little more breathing room between the lines of this essay before I print it for the workshop — please change the line spacing of the whole document to 1.5."),
    ]),

    # ====================================================================
    # F-WRITER-3 (manual short fixture) — find/replace + page break
    # ====================================================================
    FileTask(F_WRITER_3, "find_replace", "find_replace",
             _gold_find_replace, params=[
        Param({"old": "Phillips", "new": "flat-head"}, "files", {},
              "I'm updating this manual to match the tool kit we now ship with the device — could you find every occurrence of the word 'Phillips' and replace it with 'flat-head' for me?"),
        Param({"old": "ribbon cable", "new": "data cable"}, "files", {},
              "I've been asked to align the manual terminology with the new harness naming we now use on the line — please replace every occurrence of 'ribbon cable' with 'data cable'."),
    ]),
    FileTask(F_WRITER_3, "page_break", "add_page_break",
             _gold_page_break, params=[
        Param({"para_idx": 3}, "page_break", {"expected_count": 1},
              "I'd like the inspection step in this manual to start on a fresh page when printed — please throw in a page break before the fourth paragraph so it begins at the top of a new page."),
        Param({"para_idx": 4}, "page_break", {"expected_count": 1},
              "Hey, can you throw in a page break right before the fifth paragraph of this document? I want that section to begin on its own page when the manual is printed."),
    ]),

    # ====================================================================
    # F-WRITER-4 (recipe short fixture) — highlight + first-line center
    # ====================================================================
    # Validation PARAM_REDUCIBLE: dropped the para_idx=2
    # variant — kept para_idx=0 (first paragraph is the unambiguous,
    # least-step target for highlight via the UI dropdown).
    FileTask(F_WRITER_4, "highlight_para", "highlight_text",
             _gold_highlight_para, params=[
        Param({"para_idx": 0}, "strict", {"examine_highlight": True},
              "I'm printing this recipe card for my mother who keeps forgetting the preheat step — could you highlight the first paragraph in yellow so the instruction jumps out on the page?"),
    ]),
    FileTask(F_WRITER_4, "first_centered", "bold_text",
             _gold_first_centered, params=[
        Param({}, "first_centered", {},
              "I'd like the printed recipe card to look symmetric on the page, so please centre-align the first paragraph (the preheat instruction) for me."),
    ]),

    # ====================================================================
    # F-WRITER-5 (brief short fixture) — bold + insert empty table
    # ====================================================================
    # Pruned-2 (writer low-value swap, task_id=bold_para):
    # FileTask(F_WRITER_5, "bold_para", "bold_text",
             # _gold_bold_para, params=[
        # Param({"para_idx": 0}, "strict", {},
              # "I'm circulating this Q2 brief to the committee tomorrow and want the executive summary to stand out — could you bold the first paragraph of the brief for me?"),
        # Param({"para_idx": 5}, "strict", {},
              # "I'd like the closing recommendation to stand out for the committee chair when they skim the brief — could you bold the last paragraph of the brief for me?"),
    # ]),
    FileTask(F_WRITER_5, "insert_table", "blank_table_insert",
             _gold_insert_empty_table, params=[
        Param({"rows": 3, "cols": 3}, "tables", {},
              "I'd like a placeholder for the action-items grid at the foot of the brief — please insert a 3-row by 3-column empty table at the end of the document and leave all cells blank, no header text."),
        Param({"rows": 4, "cols": 4}, "tables", {},
              "Append an empty 4x4 table at the end of the brief. Do not fill in any cells."),
    ]),

    # ====================================================================
    # F-WRITER-6 (guide short fixture) — underline + page numbers footer
    # ====================================================================
    FileTask(F_WRITER_6, "underline_para", "bold_text",
             _gold_underline_para, params=[
        Param({"para_idx": 0}, "strict", {},
              "I'm writing this configuration guide for the new on-call rotation and want the lead-in to stand out — please underline the first paragraph of the guide for me."),
        Param({"para_idx": 4}, "strict", {},
              "I want the verification step to draw the on-call engineer's eye when they're skimming for the smoke test — could you underline the fifth paragraph of the guide for me?"),
    ]),
    FileTask(F_WRITER_6, "page_numbers_footer", "add_header_footer",
             _gold_page_numbers_footer, params=[
        Param({}, "page_numbers", {},
              # validation: bare "Insert > Page Number" misdirects the
              # agent to insert at body cursor. F_WRITER_42 was fixed first;
              # generalized here to all 7 siblings (F_WRITER_6/15/36/51/71/85/90)
              # since they share the same procedural-hint root cause.
              "Add page numbers to the document footer — first enable the footer via Insert > Header and Footer > Footer > Default, then with the cursor inside the footer use Insert > Page Number."),
    ]),

    # ====================================================================
    # F-WRITER-7 (policy serif) — doc-wide font + strike last
    # ====================================================================
    FileTask(F_WRITER_7, "doc_font_arial", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Arial"}, "font_names", {"font_name": "Arial"},
              "I'd like to standardise the typography across all our policy documents — could you change the font of the entire document to Arial so it matches the corporate style guide?"),
        Param({"font_name": "Verdana"}, "font_names", {"font_name": "Verdana"},
              "I'd like to test how this policy reads in our alternate house typeface before the design committee picks a winner — please change the font of the entire document to Verdana."),
    ]),
    FileTask(F_WRITER_7, "strike_last", "bold_text",
             _gold_strike_last, params=[
        Param({}, "strike_last", {},
              "I'm peer-reviewing this policy and I think the final paragraph is redundant — could you apply strikethrough to the last paragraph so the author can see what I'd cut?"),
    ]),

    # ====================================================================
    # F-WRITER-8 (essay 1.5-spaced) — italic + per-para spacing
    # ====================================================================
    FileTask(F_WRITER_8, "italic_para", "bold_text",
             _gold_italic_para, params=[
        Param({"para_idx": 2}, "strict", {},
              "I'm marking up this essay for my writing workshop and I want to flag the third paragraph as a vivid passage worth discussing — please italicise the third paragraph for me."),
        Param({"para_idx": 5}, "strict", {},
              "I'd like to mark the resolving paragraph at the end of this excerpt as a stylistic aside for my workshop students — could you italicise the sixth paragraph for me?"),
    ]),
    FileTask(F_WRITER_8, "para_spacing", "change_line_spacing",
             _gold_para_spacing, params=[
        Param({"para_idx": 0, "value": 2.0}, "line_spacing", {},
              "Could you help me leave extra room above this essay's opening so my teacher can scrawl margin notes? Please set double line spacing on the first paragraph only."),
        Param({"para_idx": 1, "value": 2.0}, "line_spacing", {},
              "I'd like the second paragraph of this essay to read with extra room for margin notes before I print the workshop copy — could you apply double line spacing to that paragraph only?"),
    ]),

    # ====================================================================
    # F-WRITER-9 (manual 13pt) — bold + size-18 paragraph (highlight bucket)
    # ====================================================================
    # Pruned-2 (writer low-value swap, task_id=bold_para):
    # FileTask(F_WRITER_9, "bold_para", "bold_text",
             # _gold_bold_para, params=[
        # Param({"para_idx": 0}, "strict", {},
              # "I want technicians to notice the screwdriver step before they begin a repair — could you bold the first paragraph of the manual for me, please?"),
        # Param({"para_idx": 3}, "strict", {},
              # "I want technicians to spot the fan-cleaning step quickly when they reach the maintenance section — could you bold the fourth paragraph of the manual for me?"),
    # ]),
    FileTask(F_WRITER_9, "size_para", "highlight_text",
             _gold_size_para, params=[
        Param({"para_idx": 0, "size_pt": 18}, "strict", {"examine_font_size": True},
              "I'd like this opening warning to read clearly across the workshop floor when the printed manual is pinned to the wall — please set the font size of the first paragraph to 18 points."),
        Param({"para_idx": 2, "size_pt": 18}, "strict", {"examine_font_size": True},
              "I'd like the air-cleaning step to be legible from the bench across the workshop when the manual is pinned to the wall — could you change the font size of the third paragraph to 18pt?"),
    ]),

    # ====================================================================
    # F-WRITER-10 (guide DejaVu) — italic + default-font registry change
    # ====================================================================
    FileTask(F_WRITER_10, "italic_para", "bold_text",
             _gold_italic_para, params=[
        Param({"para_idx": 0}, "strict", {}, "Italicize the first paragraph."),
        Param({"para_idx": 3}, "strict", {},
              "I'm marking up this guide before I send it to the on-call engineers and I want them to take the fourth paragraph as advisory rather than mandatory — please italicise the fourth paragraph for me."),
    ]),
    # Dropped — eval `find_default_font` reads
    # /home/user/.config/libreoffice/4/user/registrymodifications.xcu, but
    # the synth oracle is `cp expected→source` which never touches that
    # registry file → Oracle scored 0.0. The eval-side analogue uses a
    # custom config step that writes the xcu XML directly; the synth host-
    # heredoc model can't replicate that cleanly. Dropped pending an
    # xcu-aware oracle.
    # FileTask(F_WRITER_10, "default_font", "change_font",
    #          _gold_default_font_noop, params=[
    #     Param({"font_name": "Arial"}, "default_font", {"font_name": "Arial"},
    #           "..."),
    #     Param({"font_name": "Times New Roman"}, "default_font", {"font_name": "Times New Roman"},
    #           "..."),
    # ]),

    # ====================================================================
    # F-WRITER-11 (research report long body) — find_replace + italic+size14
    # ====================================================================
    FileTask(F_WRITER_11, "find_replace", "find_replace",
             _gold_find_replace, params=[
        Param({"old": "calibration", "new": "alignment"}, "files", {},
              "I've been asked to update the research report terminology to match the new collaboration glossary before we submit it — could you find every occurrence of 'calibration' and replace it with 'alignment'?"),
        Param({"old": "front-end boards", "new": "frontend modules"}, "files", {},
              "I'd like the research report's wording to match the new hardware-architecture diagram we agreed at the steering meeting — please replace 'front-end boards' with 'frontend modules' throughout the report."),
    ]),
    FileTask(F_WRITER_11, "italic_size14", "highlight_text",
             _gold_italic_size14, params=[
        Param({}, "italic_size14", {},
              "I want this draft to read as a working manuscript rather than the final published version when I share it with reviewers — please format the entire document in italic at 14 point font size."),
    ]),

    # ====================================================================
    # F-WRITER-12 (travel guide long body) — append paragraph + doc font
    # ====================================================================
    FileTask(F_WRITER_12, "append_para", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Trip coordinator: Pat Morgan, +1 415 555 0142."}, "files", {},
              "I want travellers to know who to call if plans change mid-trip — could you add a new paragraph at the very end of the travel guide with the text 'Trip coordinator: Pat Morgan, +1 415 555 0142.'?"),
        Param({"text": "Emergency contact: 1-800-555-1119 (24 hours)."}, "files", {},
              "I'd like every traveller to have a 24-hour number to call in case of trouble — could you append a new paragraph at the end of the guide reading 'Emergency contact: 1-800-555-1119 (24 hours).'?"),
    ]),
    FileTask(F_WRITER_12, "doc_font_courier", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Courier New"}, "font_names", {"font_name": "Courier New"},
              "Change the font of the entire travel guide to Courier New."),
        Param({"font_name": "Arial"}, "font_names", {"font_name": "Arial"},
              "I'd like this travel guide to read in our agency's standard sans-serif before I drop it into the welcome packet for the next tour group — please change the font of the whole document to Arial."),
    ]),

    # ====================================================================
    # F-WRITER-13 (product manual long body) — find_replace + page break
    # ====================================================================
    FileTask(F_WRITER_13, "find_replace", "find_replace",
             _gold_find_replace, params=[
        Param({"old": "TR-7", "new": "TR-8"}, "files", {},
              "I've been asked to retire the TR-7 product code now that the TR-8 has shipped — could you find every occurrence of the model number 'TR-7' in this manual and replace it with 'TR-8' for me?"),
        Param({"old": "thermocouple", "new": "thermistor"}, "files", {},
              "I've been asked to swap the temperature-sensor reference in this manual now that we've migrated to the cheaper part — please replace 'thermocouple' with 'thermistor' throughout the document."),
    ]),
    FileTask(F_WRITER_13, "page_break", "add_page_break",
             _gold_page_break, params=[
        Param({"para_idx": 2}, "page_break", {"expected_count": 1},
              "Hey, can you throw in a page break right before section three of this manual? I'd like that section to begin at the top of a new page when the document is printed for the assembly line."),
        Param({"para_idx": 4}, "page_break", {"expected_count": 1},
              "Insert a page break before section five (the fifth paragraph) so it begins on a new page."),
    ]),

    # ====================================================================
    # F-WRITER-14 (novel opening) — bold + doc spacing
    # ====================================================================
    # Pruned-2 (writer low-value swap, task_id=bold_para):
    # FileTask(F_WRITER_14, "bold_para", "bold_text",
             # _gold_bold_para, params=[
        # Param({"para_idx": 0}, "strict", {},
              # "I'm preparing this novel excerpt for my reading group and want the hook to grab everyone's eye on the printed handout — could you bold the opening paragraph for me?"),
        # Param({"para_idx": 4}, "strict", {},
              # "I'd like the closing image of this novel excerpt to leave a lasting impression on the printed reading-group handout — could you bold the closing paragraph (the fifth paragraph) for me?"),
    # ]),
    FileTask(F_WRITER_14, "doc_spacing", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 2.0}, "line_spacing", {},
              "Could you help me prepare this excerpt for classroom annotation? Please set the line spacing of the entire excerpt to double (2.0) so students have room to write between the lines."),
        Param({"value": 1.0}, "line_spacing", {},
              "I'd like to undo the double spacing now that the proofreading round is complete and we're moving to a tighter print layout — could you reset the line spacing of the whole document to single (1.0)?"),
    ]),

    # ====================================================================
    # F-WRITER-15 (policy longform) — append + page numbers footer
    # ====================================================================
    FileTask(F_WRITER_15, "append_revision", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Document control: version 4.2, owner CISO, next review June 2027."}, "files", {},
              "I'd like to add a document-control footer line before this policy goes into the records system — could you append a new paragraph at the end of the policy reading 'Document control: version 4.2, owner CISO, next review June 2027.'?"),
        Param({"text": "Effective date: 1 June 2026. Supersedes version 4.1."}, "files", {},
              "I'd like the effective date and prior-version reference to appear at the foot of this policy for the records system — could you add a new paragraph at the very end reading 'Effective date: 1 June 2026. Supersedes version 4.1.'?"),
    ]),
    FileTask(F_WRITER_15, "page_numbers_footer", "add_header_footer",
             _gold_page_numbers_footer, params=[
        Param({}, "page_numbers", {},
              "Add page numbers to the footer of the policy document — first enable the footer via Insert > Header and Footer > Footer > Default, then with the cursor inside the footer use Insert > Page Number."),
    ]),

    # ====================================================================
    # F-WRITER-16 (recipe with initial 3×3 table) — color table text + bold
    # ====================================================================
    # Pruned-2 (writer low-value swap, task_id=bold_para):
    # FileTask(F_WRITER_16, "bold_para", "bold_text",
             # _gold_bold_para, params=[
        # Param({"para_idx": 0}, "strict", {},
              # "I'm printing this recipe card for a baking class and want the prep instruction to draw the eye on each handout — could you bold the first paragraph of the recipe card?"),
        # Param({"para_idx": 4}, "strict", {},
              # "I'd like the doneness check to draw attention when this card is taped to the oven door for the bake-along — could you bold the fifth paragraph of the recipe card?"),
    # ]),
    # Validation PARAM_REDUCIBLE: dropped 2.0-spacing Param;
    # kept 1.5.
    FileTask(F_WRITER_16, "doc_spacing", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 1.5}, "line_spacing", {},
              "Set the line spacing of the recipe card to 1.5."),
    ]),

    # ====================================================================
    # F-WRITER-17 (policy with initial 4×4 table) — find_replace + doc font
    # ====================================================================
    FileTask(F_WRITER_17, "find_replace", "find_replace",
             _gold_find_replace, params=[
        Param({"old": "training", "new": "certification"}, "files", {},
              "I'd like to align this policy with the formal HR programme language — could you replace 'training' with 'certification' throughout the document for me?"),
        Param({"old": "Managers", "new": "Team leads"}, "files", {},
              "I'd like the policy to use our flatter title nomenclature since the reorg — could you find every occurrence of 'Managers' and replace it with 'Team leads' throughout?"),
    ]),
    FileTask(F_WRITER_17, "doc_font_tnr", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Times New Roman"}, "font_names", {"font_name": "Times New Roman"},
              "I'm reformatting this policy to match the legal department's style guide before submission — could you change the font of the entire policy document to Times New Roman, please?"),
    ]),

    # ====================================================================
    # F-WRITER-18 (Alice in Wonderland excerpt) — bold opener + doc font
    # ====================================================================
    FileTask(F_WRITER_18, "bold_p0", "bold_text",
             lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="bold"),
             params=[
        Param({}, "strict", {},
              "I'm building a reading-club handout from this Alice's Adventures in Wonderland excerpt and want the opening to anchor the discussion — could you bold the opening paragraph of the excerpt?"),
    ]),
    FileTask(F_WRITER_18, "italic_p0", "bold_text",
             lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="italic"),
             params=[
        Param({}, "strict", {},
              "I'd like the dreamlike opening of this Alice excerpt to feel distinct from the prose that follows on the printed reading guide — could you italicise the opening paragraph for me?"),
    ]),

    # ====================================================================
    # F-WRITER-19 (Pride & Prejudice excerpt) — italic + highlight opener
    # ====================================================================
    FileTask(F_WRITER_19, "italic_p0", "bold_text",
             lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="italic"),
             params=[
        Param({}, "strict", {},
              "I want my students to notice the irony in Austen's famous first sentence when they read this Pride and Prejudice excerpt — please italicise the opening paragraph for me."),
    ]),
    FileTask(F_WRITER_19, "highlight_p0", "highlight_text",
             lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="highlight_yellow"),
             params=[
        Param({}, "strict", {"examine_highlight": True},
              "I'd like students to flag Austen's famous opening line for class discussion next week — could you apply a yellow highlight to the opening paragraph of the Pride and Prejudice excerpt?"),
    ]),

    # ====================================================================
    # F-WRITER-20 (Moby-Dick excerpt) — strike + size 18 opener
    # ====================================================================
    FileTask(F_WRITER_20, "strike_p0", "bold_text",
             lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="strike"),
             params=[
        Param({}, "strict", {},
              "I'm peer-reviewing a friend's literary commentary and I think the opening quotation of this Moby-Dick excerpt should come out — please apply strikethrough to the opening paragraph so they can see what I'd cut."),
    ]),
    FileTask(F_WRITER_20, "size18_p0", "highlight_text",
             lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="size18"),
             params=[
        Param({}, "strict", {"examine_font_size": True},
              "I want the famous 'Call me Ishmael' opening of Moby-Dick to read as a banner on the printed reading-group handout — please set 18-point font size on the opening paragraph for me."),
    ]),

    # ====================================================================
    # F-WRITER-21 (Frankenstein excerpt) — underline + Georgia font opener
    # ====================================================================
    FileTask(F_WRITER_21, "underline_p0", "bold_text",
             lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="underline"),
             params=[
        Param({}, "strict", {},
              "I'd like the framing letter at the start of this Frankenstein excerpt to stand out from the chapters that follow on the printed reading guide — please underline the opening paragraph for me."),
    ]),
    FileTask(F_WRITER_21, "font_georgia_p0", "change_font",
             lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="font_georgia"),
             params=[
        Param({}, "strict", {"examine_font_name": True},
              "I'd like the framing letter at the start of this Frankenstein excerpt to read in a Victorian-feeling serif on the printed reading guide — please apply Georgia as the font face for the opening paragraph (change the typeface, do not edit the paragraph text)."),
    ]),

    # ====================================================================
    # F-WRITER-22 (Sherlock Holmes excerpt) — color red + highlight opener
    # ====================================================================
    FileTask(F_WRITER_22, "color_red_p0", "color_table_text",
             lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="color_red"),
             params=[
        Param({}, "strict", {"examine_color": True},
              "I'm preparing this Sherlock Holmes excerpt for a mystery-genre lesson and want students to spot the narrator's voice immediately — please set the opening paragraph in red font colour for me."),
    ]),
    # Pruned-2 (writer low-value swap, task_id=highlight_p0):
    # FileTask(F_WRITER_22, "highlight_p0", "highlight_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="highlight_yellow"),
             # params=[
        # Param({}, "strict", {"examine_highlight": True},
              # "I'd like my class to flag the introductory setup of this Sherlock Holmes excerpt as the bit we'll discuss first — please highlight the opening paragraph in yellow."),
    # ]),

    # ====================================================================
    # F-WRITER-23 (Treasure Island excerpt) — bold + color blue opener
    # ====================================================================
    # Pruned-2 (writer low-value swap, task_id=bold_p0):
    # FileTask(F_WRITER_23, "bold_p0", "bold_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="bold"),
             # params=[
        # Param({}, "strict", {},
              # "I'd like the framing setup of this Treasure Island excerpt to anchor the printed reading guide before the chapter unfolds — please bold the opening paragraph for me."),
    # ]),
    FileTask(F_WRITER_23, "color_blue_p0", "color_table_text",
             lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="color_blue"),
             params=[
        Param({}, "strict", {"examine_color": True},
              "I'd like the opening of this Treasure Island excerpt to feel maritime on the printed reading guide for my middle-school class — could you set the opening paragraph in blue font colour?"),
    ]),

    # ====================================================================
    # F-WRITER-24 (Metamorphosis excerpt) — size 14 + italic opener
    # ====================================================================
    FileTask(F_WRITER_24, "size14_p0", "highlight_text",
             lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="size14"),
             params=[
        Param({}, "strict", {"examine_font_size": True},
              "I want to enlarge the famous opening line of this Metamorphosis excerpt so it reads as a banner across the printed reading guide — please set the opening paragraph to 14-point font size."),
    ]),
    # Pruned (writer dist 13.3%→10%, task_id=italic_p0, eval=0-direct):
    # FileTask(F_WRITER_24, "italic_p0", "bold_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="italic"),
             # params=[
        # Param({}, "strict", {},
              # "I'd like the surreal opening sentence of this Metamorphosis excerpt to feel typographically different from the narration that follows — please italicise the opening paragraph for the printed reading guide."),
    # ]),

    # ====================================================================
    # F-WRITER-25 (Importance of Being Earnest excerpt) — pdf export
    # ====================================================================
    FileTask(F_WRITER_25, "pdf_export", "pdf_export",
             _gold_pdf_export, params=[
        Param({}, "pdf_export", {},
              "Export this document to PDF using File > Export As > Export Directly as PDF (or Shift+Ctrl+E). Accept the default filename and Desktop location, then click Export."),
    ]),
    # Pruned-2 (writer low-value swap, task_id=highlight_p0):
    # FileTask(F_WRITER_25, "highlight_p0", "highlight_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="highlight_yellow"),
             # params=[
        # Param({}, "strict", {"examine_highlight": True},
              # "I'd like my drama students to spot the Wilde opening line as the comic hook of the printed handout for next week's table read — could you highlight the opening paragraph of the Earnest excerpt in yellow?"),
    # ]),

    # ====================================================================
    # F-WRITER-26 (Earth photo host) — image insert + pdf export
    # ====================================================================
    FileTask(F_WRITER_26, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/earth-blue-marble-apollo17.jpg",
                 width_in=4.0, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "I'd like to illustrate this Earth briefing with the iconic Blue Marble photograph for the printed classroom handout — could you insert the Apollo 17 Earth photograph from /home/user/Desktop/earth-blue-marble-apollo17.jpg into the document below the second paragraph?"),
    ]),
    FileTask(F_WRITER_26, "pdf_export", "pdf_export",
             _gold_pdf_export, params=[
        Param({}, "pdf_export", {},
              "Export this Earth briefing to PDF: File > Export As > Export Directly as PDF (Shift+Ctrl+E), accept the default filename and location, click Export."),
    ]),

    # ====================================================================
    # F-WRITER-27 (Tiger photo host) — image insert + bold
    # ====================================================================
    FileTask(F_WRITER_27, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/tiger-closeup.jpg",
                 width_in=4.0, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "Insert the tiger photograph from /home/user/Desktop/tiger-closeup.jpg into the document, below the second paragraph."),
    ]),
    # Pruned (writer dist 13.3%→10%, task_id=bold_para, eval=0-direct):
    # FileTask(F_WRITER_27, "bold_para", "bold_text",
             # _gold_bold_para, params=[
        # Param({"para_idx": 0}, "strict", {},
              # "Bold the first paragraph of the field journal entry."),
    # ]),

    # ====================================================================
    # F-WRITER-28 (Pizza photo host) — image insert + find_replace
    # ====================================================================
    FileTask(F_WRITER_28, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/pizza-dish.jpg",
                 width_in=3.5, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "I'm finalising the new menu insert for tonight's dinner service and would like a hero shot of the dish — please insert the pizza photograph from /home/user/Desktop/pizza-dish.jpg below the second paragraph of the menu."),
    ]),
    FileTask(F_WRITER_28, "find_replace", "find_replace",
             _gold_find_replace, params=[
        Param({"old": "Margherita", "new": "Marinara"}, "files", {},
              "Replace every occurrence of 'Margherita' in the menu with 'Marinara'."),
        Param({"old": "Sangiovese", "new": "Chianti"}, "files", {},
              "Find 'Sangiovese' and replace it with 'Chianti' throughout the document."),
    ]),

    # ====================================================================
    # F-WRITER-29 (Andromeda photo host) — image insert + italic opener
    # ====================================================================
    FileTask(F_WRITER_29, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/galaxy-andromeda.jpg",
                 width_in=4.5, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "I'd like to illustrate this astronomy briefing with a wide-field shot of our nearest spiral neighbour — could you insert the Andromeda galaxy photograph from /home/user/Desktop/galaxy-andromeda.jpg below the second paragraph of the document?"),
    ]),
    FileTask(F_WRITER_29, "italic_para", "bold_text",
             _gold_italic_para, params=[
        Param({"para_idx": 0}, "strict", {},
              "Italicize the first paragraph of the astronomy briefing."),
        Param({"para_idx": 2}, "strict", {},
              "Italicize the third paragraph (the caption text)."),
    ]),

    # ====================================================================
    # F-WRITER-30 (House photo host) — image insert + append
    # ====================================================================
    FileTask(F_WRITER_30, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/house-modern.jpg",
                 width_in=4.0, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "Insert the modern-house photograph from /home/user/Desktop/house-modern.jpg into the architectural review, below the second paragraph."),
    ]),
    FileTask(F_WRITER_30, "append_para", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Photograph credit: Studio Walker, reproduced with permission."}, "files", {},
              "Append a new paragraph at the end of the review with the text 'Photograph credit: Studio Walker, reproduced with permission.'."),
    ]),

    # ====================================================================
    # F-WRITER-31 (Double-image galaxy compare) — double image + italic
    # ====================================================================
    FileTask(F_WRITER_31, "insert_double_image", "insert_image",
             lambda s, g, **kw: _gold_double_image_insert(
                 s, g,
                 image_path_a=f"{_DESKTOP}/galaxy-andromeda.jpg",
                 image_path_b=f"{_DESKTOP}/galaxy-hubble.jpg",
                 width_in=3.5),
             params=[
        Param({}, "strict", {"examine_images": True},
              "Insert two photographs into the comparative-astronomy plate: /home/user/Desktop/galaxy-andromeda.jpg below the 'Caption A' paragraph, and /home/user/Desktop/galaxy-hubble.jpg below the 'Caption B' paragraph."),
    ]),
    # Pruned (writer dist 13.3%→10%, task_id=bold_para, eval=0-direct):
    # FileTask(F_WRITER_31, "bold_para", "bold_text",
             # _gold_bold_para, params=[
        # Param({"para_idx": 0}, "strict", {},
              # "Bold the introductory paragraph of the comparative-astronomy plate."),
    # ]),

    # ====================================================================
    # Loop 1.A — long-form prose with title + headings
    # ====================================================================

    # F-WRITER-32 (white paper) — find_replace + doc font
    FileTask(F_WRITER_32, "find_replace", "find_replace",
             _gold_find_replace, params=[
        Param({"old": "lakehouse", "new": "data-mesh"}, "files", {},
              "I've been asked by the analyst team to retitle our data-platform white paper for the December refresh — could you replace every occurrence of 'lakehouse' with 'data-mesh' throughout the document?"),
        Param({"old": "centralised", "new": "centralized"}, "files", {},
              "I'd like to switch this white paper to American spelling now that the launch is US-only — could you find every occurrence of 'centralised' in the document and replace it with 'centralized' for me?"),
    ]),
    FileTask(F_WRITER_32, "doc_font", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Arial"}, "font_names", {"font_name": "Arial"},
              "I'd like the white paper to match the corporate brand-book typography before it goes out to clients next week — please change the font of the entire white paper to Arial for me."),
        Param({"font_name": "Verdana"}, "font_names", {"font_name": "Verdana"},
              "Change the font of the whole document to Verdana."),
    ]),

    # F-WRITER-33 (case study) — append + italic_size14
    FileTask(F_WRITER_33, "append_para", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Case study prepared by the transit-research team, December 2025."}, "files", {},
              "I'd like to credit our research team at the foot of this case study before it goes to the policy committee — could you append a new paragraph at the end with the text 'Case study prepared by the transit-research team, December 2025.'?"),
        Param({"text": "For follow-up enquiries contact transit-research@northbridge.gov."}, "files", {},
              "I'd like readers to be able to contact us with questions after they finish this case study — could you add a new paragraph at the very end of the document reading 'For follow-up enquiries contact transit-research@northbridge.gov.'?"),
    ]),
    FileTask(F_WRITER_33, "italic_size14", "highlight_text",
             _gold_italic_size14, params=[
        Param({}, "italic_size14", {},
              "Could you help me mark up this case study as a working draft for the next steering meeting? Please format the entire document in italic at 14 point font size so reviewers can see it's not yet final."),
    ]),

    # F-WRITER-34 (press briefing) — find_replace + doc spacing
    FileTask(F_WRITER_34, "find_replace", "find_replace",
             _gold_find_replace, params=[
        Param({"old": "Eastfield", "new": "Westfield"}, "files", {},
              "Replace every occurrence of 'Eastfield' with 'Westfield' throughout the press briefing."),
        Param({"old": "quarterly", "new": "annual"}, "files", {},
              "I've been asked by communications to reflect the new reporting cadence in this press briefing before it goes out tomorrow — could you find 'quarterly' and replace it with 'annual' across the whole document?"),
    ]),
    # Validation PARAM_REDUCIBLE: dropped 2.0-spacing Param;
    # kept 1.5.
    FileTask(F_WRITER_34, "doc_spacing", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 1.5}, "line_spacing", {},
              "I'd like reviewers to have room for handwritten notes between the lines of this press briefing — please set the line spacing of the entire briefing to 1.5 for me."),
    ]),

    # F-WRITER-35 (market analysis) — find_replace + page_break
    # Pruned-2 (writer low-value swap, task_id=find_replace):
    # FileTask(F_WRITER_35, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "precision sensors", "new": "smart sensors"}, "files", {},
              # "I'd like to align this market analysis with the new marketing taxonomy before it reaches the leadership review — could you replace every occurrence of 'precision sensors' with 'smart sensors' across the document?"),
        # Param({"old": "supplier", "new": "vendor"}, "files", {},
              # "I'd like to update the procurement language in this market analysis to match the new finance taxonomy — could you find every occurrence of 'supplier' and replace it with 'vendor' throughout?"),
    # ]),
    FileTask(F_WRITER_35, "page_break", "add_page_break",
             _gold_page_break, params=[
        Param({"para_idx": 5}, "page_break", {"expected_count": 1},
              "Hey, can you throw in a page break right before the Competitive Landscape section? I want it to begin at the top of a new page in the printed analysis."),
        Param({"para_idx": 7}, "page_break", {"expected_count": 1},
              "Insert a page break before the Regional Outlook section."),
    ]),

    # F-WRITER-36 (tech spec) — find_replace + page numbers footer
    # Pruned-2 (writer low-value swap, task_id=find_replace):
    # FileTask(F_WRITER_36, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "Aurora-3", "new": "Aurora-4"}, "files", {},
              # "I'm preparing the technical specification for the Aurora-4 launch and need to roll the model identifier forward — could you replace every occurrence of the model name 'Aurora-3' with 'Aurora-4' throughout the document?"),
        # Param({"old": "firmware version four point one", "new": "firmware version 5.0"}, "files", {},
              # "I'd like the firmware revision to read in numeric form now that the major-release notes are out — please find 'firmware version four point one' in this spec and replace it with 'firmware version 5.0'."),
    # ]),
    FileTask(F_WRITER_36, "page_numbers_footer", "add_header_footer",
             _gold_page_numbers_footer, params=[
        Param({}, "page_numbers", {},
              "Add page numbers to the footer of the Aurora-3 specification — first enable the footer via Insert > Header and Footer > Footer > Default, then with the cursor inside the footer use Insert > Page Number."),
    ]),

    # F-WRITER-37 (org announcement) — append + doc font
    FileTask(F_WRITER_37, "append_para", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Issued by the office of the Chief Executive, effective immediately."}, "files", {},
              "I'd like to add the issuing authority before this announcement is posted to the staff portal — could you append a new paragraph at the end with the text 'Issued by the office of the Chief Executive, effective immediately.'?"),
    ]),
    FileTask(F_WRITER_37, "doc_font", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Times New Roman"}, "font_names", {"font_name": "Times New Roman"},
              "I want this announcement to read in the executive house typeface before it goes to print for the all-hands packet — please change the font of the entire announcement to Times New Roman."),
        # Dropped — Calibri is docx default-theme major font; LO
        # round-trip strips redundant explicit rFonts → compare_font_names
        # sees run.font.name == None on the round-tripped gold and fails.
        # Same for Cambria (default minor in some themes). Keep only fonts
        # LO preserves explicitly (TNR, Georgia, Verdana).
        Param({"font_name": "Arial"}, "font_names", {"font_name": "Arial"},
              "I'd like this announcement to read in a sans-serif before it goes onto the staff portal — please change the font of the whole document to Arial."),
    ]),

    # F-WRITER-38 (academic abstract) — italic_size14 + doc font
    FileTask(F_WRITER_38, "italic_size14", "highlight_text",
             _gold_italic_size14, params=[
        Param({}, "italic_size14", {},
              "Format the entire academic abstract in italic at 14 point font size."),
    ]),
    FileTask(F_WRITER_38, "doc_font", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Georgia"}, "font_names", {"font_name": "Georgia"},
              "I'd like the academic abstract to match the conference submission style guide before I upload it tomorrow — could you change the font of the entire document to Georgia, please?"),
        # Dropped — Cambria stripped on round-trip (see F_WRITER_37).
        Param({"font_name": "Times New Roman"}, "font_names", {"font_name": "Times New Roman"},
              "I'd like to compare how the abstract reads in a different serif before committing to the final submission typeface — could you change the font of the whole abstract to Times New Roman?"),
    ]),

    # F-WRITER-39 (policy memo) — find_replace + append
    # Pruned-2 (writer low-value swap, task_id=find_replace):
    # FileTask(F_WRITER_39, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "per-diem", "new": "daily allowance"}, "files", {},
              # "I've been asked to use plain-English terminology in this policy memo before it goes to all staff — could you replace every occurrence of 'per-diem' with 'daily allowance' throughout the memo?"),
        # Param({"old": "rail", "new": "train"}, "files", {},
              # "I'd like the travel-policy memo to use the everyday term staff actually book with instead of the formal one — could you find every occurrence of 'rail' and replace it with 'train' throughout?"),
    # ]),
    FileTask(F_WRITER_39, "append_para", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Memo prepared by the policy office, copies to executive committee members."}, "files", {},
              "Append a new paragraph at the end of the memo with the text 'Memo prepared by the policy office, copies to executive committee members.'."),
    ]),

    # F-WRITER-40 (long essay on letters) — bold + italic
    # Layout: title at idx 0, then paragraphs at idx 1..5
    # Pruned (writer dist 13.3%→10%, task_id=bold_para, eval=0-direct):
    # FileTask(F_WRITER_40, "bold_para", "bold_text",
             # _gold_bold_para, params=[
        # Param({"para_idx": 1}, "strict", {},
              # "I'm preparing this essay for my writing workshop and I want the opening reflection on letters to anchor the discussion — please bold the first paragraph of the essay for me."),
        # Param({"para_idx": 5}, "strict", {},
              # "I'd like the closing reflection of this essay to land hard on the reader before they look up from the page — could you bold the last paragraph for me?"),
    # ]),
    FileTask(F_WRITER_40, "italic_para", "bold_text",
             _gold_italic_para, params=[
        Param({"para_idx": 2}, "strict", {}, "Italicize the second paragraph of the essay."),
        Param({"para_idx": 3}, "strict", {},
              "I'd like to mark the central anecdote of this essay as a stylistic aside for the printed handout — could you italicise the third paragraph of the essay for me, please?"),
    ]),

    # F-WRITER-41 (long essay on maps) — find_replace + page_break
    # Pruned-2 (writer low-value swap, task_id=find_replace):
    # FileTask(F_WRITER_41, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "Michelin", "new": "Ordnance Survey"}, "files", {},
              # "I'm reworking this essay on maps for a British readership and want to swap the French reference for the local equivalent — could you replace every occurrence of 'Michelin' in the essay with 'Ordnance Survey'?"),
        # Param({"old": "France", "new": "Britain"}, "files", {},
              # "I'm relocating this maps essay for a British readership before it goes into the printed travel anthology — could you find every occurrence of 'France' and replace it with 'Britain' throughout?"),
    # ]),
    FileTask(F_WRITER_41, "page_break", "add_page_break",
             _gold_page_break, params=[
        Param({"para_idx": 3}, "page_break", {"expected_count": 1},
              "Hey, can you throw in a page break right before the third paragraph of this essay? I want that section to start on a new page when I bind the print proof."),
    ]),

    # ====================================================================
    # Loop 1.B — recipes with title + ingredient table + photo
    # ====================================================================

    # F-WRITER-42 (salad recipe) — find_replace + page_numbers_footer
    # Pruned-2 (writer low-value swap, task_id=find_replace):
    # FileTask(F_WRITER_42, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "hazelnuts", "new": "almonds"}, "files", {},
              # "I'm adapting this salad recipe for a friend with a hazelnut allergy before the dinner party — could you replace every occurrence of 'hazelnuts' with 'almonds' across the recipe?"),
        # Param({"old": "Dijon", "new": "wholegrain"}, "files", {},
              # "I'd like to swap the mustard variety in this salad recipe for the wholegrain one I have in the cupboard — could you find every occurrence of 'Dijon' and replace it with 'wholegrain' for me?"),
    # ]),
    FileTask(F_WRITER_42, "page_numbers_footer", "add_header_footer",
             _gold_page_numbers_footer, params=[
        Param({}, "page_numbers", {},
              "Add page numbers to the footer of the salad recipe — first enable the footer via Insert > Header and Footer > Footer > Default, then with the cursor inside the footer use Insert > Page Number."),
    ]),

    # F-WRITER-43 (pizza recipe) — doc_spacing + doc_font
    # Validation PARAM_REDUCIBLE: dropped 2.0-spacing Param;
    # kept 1.5.
    FileTask(F_WRITER_43, "doc_spacing", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 1.5}, "line_spacing", {},
              "I want this pizza recipe to be easy to read across the kitchen when it's pinned to the wall — could you set the line spacing of the entire recipe to 1.5?"),
    ]),
    FileTask(F_WRITER_43, "doc_font", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Arial"}, "font_names", {"font_name": "Arial"},
              "Change the font of the entire pizza recipe to Arial."),
    ]),

    # F-WRITER-44 (coffee recipe) — append + doc_font
    FileTask(F_WRITER_44, "append_para", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Adapted from the home-barista handbook, third edition."}, "files", {},
              "I'd like to credit the source before I share this coffee guide on the brewing forum — could you append a new paragraph at the end with the text 'Adapted from the home-barista handbook, third edition.'?"),
    ]),
    FileTask(F_WRITER_44, "doc_font", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Times New Roman"}, "font_names", {"font_name": "Times New Roman"},
              "Change the font of the entire coffee guide to Times New Roman."),
        Param({"font_name": "Verdana"}, "font_names", {"font_name": "Verdana"},
              "I'd like to test a cleaner sans-serif for this coffee guide before I share it on the brewing forum — please change the font of the whole document to Verdana."),
    ]),

    # F-WRITER-45 (pasta recipe) — find_replace + italic_size14
    # Pruned-2 (writer low-value swap, task_id=find_replace):
    # FileTask(F_WRITER_45, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "tagliatelle", "new": "fettuccine"}, "files", {},
              # "I'm planning to cook this pasta dish with the shape I actually have in the pantry — could you replace every occurrence of 'tagliatelle' with 'fettuccine' across the recipe for me?"),
        # Param({"old": "Bolognese", "new": "Tuscan"}, "files", {},
              # "I'm pivoting this pasta recipe toward the Tuscan regional theme for next weekend's supper club — could you find every occurrence of 'Bolognese' and replace it with 'Tuscan' for me?"),
    # ]),
    FileTask(F_WRITER_45, "italic_size14", "highlight_text",
             _gold_italic_size14, params=[
        Param({}, "italic_size14", {},
              "Format the entire pasta recipe document in italic at 14 point font size."),
    ]),

    # F-WRITER-46 (dessert recipe) — page_break + doc_spacing
    # Layout: idx 0=Title, 1=Heading2 (serving info), 2=Heading1 "Ingredients",
    # 3=Heading1 "Method", 4-6=method body paragraphs.
    FileTask(F_WRITER_46, "page_break", "add_page_break",
             _gold_page_break, params=[
        Param({"para_idx": 3}, "page_break", {"expected_count": 1},
              "I'd like the dessert recipe's Method section to start at the top of a fresh page when this is printed for the bake-off booklet — could you throw in a page break right before the Method heading?"),
    ]),
    FileTask(F_WRITER_46, "doc_spacing", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 1.5}, "line_spacing", {},
              "Set the line spacing of the entire dessert recipe to 1.5."),
    ]),

    # ====================================================================
    # Loop 1.C — manuals with title + numbered sections + spec table
    # ====================================================================

    # F-WRITER-47 (install guide) — find_replace + doc_spacing
    # Pruned-2 (writer low-value swap, task_id=find_replace):
    # FileTask(F_WRITER_47, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "WR-200", "new": "WR-300"}, "files", {},
              # "I'm updating the installation guide to reflect the new product model we just released — could you replace every occurrence of 'WR-200' with 'WR-300' throughout the document, please?"),
        # Param({"old": "Phillips", "new": "Pozidriv"}, "files", {},
              # "I'd like to match the tool we now ship with the new product variant — please find every occurrence of 'Phillips' in the install guide and replace it with 'Pozidriv'."),
    # ]),
    FileTask(F_WRITER_47, "doc_spacing", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 1.5}, "line_spacing", {},
              "Set the line spacing of the entire install guide to 1.5."),
    ]),

    # F-WRITER-48 (safety protocol) — find_replace + page_break
    # Pruned-2 (writer low-value swap, task_id=find_replace):
    # FileTask(F_WRITER_48, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "fume hood", "new": "ventilation cabinet"}, "files", {},
              # "In the safety protocol, replace every occurrence of 'fume hood' with 'ventilation cabinet'."),
        # Param({"old": "sash", "new": "front shield"}, "files", {},
              # "I've been asked by EHS to use the plain-English term in this safety protocol so junior chemists understand it without the jargon — could you find every occurrence of 'sash' and replace it with 'front shield'?"),
    # ]),
    FileTask(F_WRITER_48, "page_break", "add_page_break",
             _gold_page_break, params=[
        Param({"para_idx": 4}, "page_break", {"expected_count": 1},
              "Hey, can you throw in a page break right before the Pre-Work Inspection section? I'd like that section to begin at the top of a new page in the printed safety packet."),
    ]),

    # F-WRITER-49 (troubleshoot guide) — find_replace + append
    # Pruned (writer dist 13.3%→10%, task_id=find_replace, eval=0-direct):
    # FileTask(F_WRITER_49, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "printer", "new": "plotter"}, "files", {},
              # "I'm re-purposing this troubleshooting guide for the engineering team's new wide-format plotter — could you replace every occurrence of 'printer' with 'plotter' across the document?"),
        # Param({"old": "toner", "new": "ink"}, "files", {},
              # "I'd like the troubleshooting guide to match the new inkjet plotter we now use instead of the old toner-based one — could you find every occurrence of 'toner' and replace it with 'ink' throughout?"),
    # ]),
    FileTask(F_WRITER_49, "append_para", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Document maintained by IT Operations, last reviewed January 2026."}, "files", {},
              "Append a new paragraph at the end of the troubleshooting guide with the text 'Document maintained by IT Operations, last reviewed January 2026.'."),
    ]),

    # F-WRITER-50 (operator manual) — italic_size14 + doc_font
    FileTask(F_WRITER_50, "italic_size14", "highlight_text",
             _gold_italic_size14, params=[
        Param({}, "italic_size14", {},
              "I'd like the BX-50 operator manual to read as a working preview when I circulate it for review next week — please format the entire document in italic at 14 point font size."),
    ]),
    FileTask(F_WRITER_50, "doc_font", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Arial"}, "font_names", {"font_name": "Arial"},
              "Change the font of the entire operator manual to Arial."),
    ]),

    # F-WRITER-51 (device manual) — find_replace + page_numbers_footer
    # Pruned (writer dist 13.3%→10%, task_id=find_replace, eval=0-direct):
    # FileTask(F_WRITER_51, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "HM-12", "new": "HM-15"}, "files", {},
              # "I'm updating this device manual to match the new HM-15 model we just released to manufacturing — could you replace every occurrence of 'HM-12' with 'HM-15' across the document?"),
        # Param({"old": "multimeter", "new": "DMM"}, "files", {},
              # "I'd like the device manual to use the engineering team's preferred abbreviation for the instrument — please find every occurrence of 'multimeter' and replace it with 'DMM' throughout."),
    # ]),
    FileTask(F_WRITER_51, "page_numbers_footer", "add_header_footer",
             _gold_page_numbers_footer, params=[
        Param({}, "page_numbers", {},
              "Add page numbers to the footer of the HM-12 device manual — first enable the footer via Insert > Header and Footer > Footer > Default, then with the cursor inside the footer use Insert > Page Number."),
    ]),

    # ====================================================================
    # Loop 1.D — travel guides with title + photo + bullet list
    # ====================================================================

    # F-WRITER-52 (Kyoto guide) — find_replace + insert_image
    # Pruned (writer dist 13.3%→10%, task_id=find_replace, eval=0-direct):
    # FileTask(F_WRITER_52, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "Kyoto", "new": "Nara"}, "files", {},
              # "I've decided to spin a Nara-focused itinerary out of my old Kyoto travel guide for the autumn trip — could you replace every occurrence of 'Kyoto' with 'Nara' across the document for me?"),
        # Param({"old": "Higashiyama", "new": "Arashiyama"}, "files", {},
              # "I'd like the printed itinerary to point my group at the Arashiyama district where we have the bamboo-grove reservation — could you find every occurrence of 'Higashiyama' and replace it with 'Arashiyama'?"),
    # ]),
    FileTask(F_WRITER_52, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/kyoto-skyline.jpg",
                 width_in=4.5, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "Insert the Kyoto skyline photograph from /home/user/Desktop/kyoto-skyline.jpg into the travel guide, below the second paragraph."),
    ]),

    # F-WRITER-53 (hike journal) — append + page_break
    FileTask(F_WRITER_53, "append_para", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Trip filed in the regional hiking-club archive, June 2025."}, "files", {},
              "I'd like to register this hike journal in our regional hiking club's archive before the season ends — could you append a new paragraph at the end with the text 'Trip filed in the regional hiking-club archive, June 2025.'?"),
    ]),
    FileTask(F_WRITER_53, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/cathedral-range.jpg",
                 width_in=4.5, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "Insert the Cathedral Range mountain photograph from /home/user/Desktop/cathedral-range.jpg below the second paragraph of the hike journal."),
    ]),

    # F-WRITER-54 (restaurant review) — find_replace + insert_image
    # Pruned (writer dist 13.3%→10%, task_id=find_replace, eval=0-direct):
    # FileTask(F_WRITER_54, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "Romolo", "new": "Tognazzi"}, "files", {},
              # "I'm reusing this restaurant review template for the new Tognazzi opening I covered last weekend — could you replace every occurrence of 'Romolo' with 'Tognazzi' throughout the document?"),
        # Param({"old": "trattoria", "new": "osteria"}, "files", {},
              # "I'd like to reflect the regional usage the new venue actually advertises with — could you find every occurrence of 'trattoria' in the review and replace it with 'osteria' for me?"),
    # ]),
    FileTask(F_WRITER_54, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/old-quarter.jpg",
                 width_in=4.0, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "Insert the European old-quarter street photograph from /home/user/Desktop/old-quarter.jpg below the second paragraph of the review."),
    ]),

    # F-WRITER-55 (beach guide) — doc_spacing + insert_image
    # Validation PARAM_REDUCIBLE: dropped 2.0-spacing Param;
    # kept 1.5.
    FileTask(F_WRITER_55, "doc_spacing", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 1.5}, "line_spacing", {},
              "I'd like the printed copy of this Atlantic-coast guide to be a little easier on the eye when I read it on the boat — could you set the line spacing of the entire document to 1.5?"),
    ]),
    FileTask(F_WRITER_55, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/beach-sunset.jpg",
                 width_in=4.5, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "Insert the beach sunset photograph from /home/user/Desktop/beach-sunset.jpg below the second paragraph of the coast guide."),
    ]),

    # F-WRITER-56 (desert tour) — append + insert_image
    FileTask(F_WRITER_56, "append_para", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Tour notes prepared by the desert-trekking office, March 2026."}, "files", {},
              "I'd like to credit our desert-trekking office at the foot of this tour itinerary before I hand it to the spring trip leaders — could you append a new paragraph at the end with the text 'Tour notes prepared by the desert-trekking office, March 2026.'?"),
    ]),
    FileTask(F_WRITER_56, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/desert-dunes.jpg",
                 width_in=4.0, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "Insert the desert dunes photograph from /home/user/Desktop/desert-dunes.jpg below the second paragraph of the tour itinerary."),
    ]),

    # ====================================================================
    # Loop 2.A — additional Gutenberg books (chapter-style excerpts)
    # ====================================================================

    # F-WRITER-57 (Alice v2) — bold + italic_size14
    # Pruned-2 (writer low-value swap, task_id=bold_p0):
    # FileTask(F_WRITER_57, "bold_p0", "bold_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="bold"),
             # params=[
        # Param({}, "strict", {},
              # "I'm assembling a chapter-opening anthology for my literature seminar and want this Alice chapter to lead off the printed handout — could you bold the opening paragraph?"),
    # ]),
    FileTask(F_WRITER_57, "italic_size14", "highlight_text",
             _gold_italic_size14, params=[
        Param({}, "italic_size14", {},
              "Format the entire Alice excerpt in italic at 14 point font size."),
    ]),

    # F-WRITER-58 (Pride v2) — underline + doc_font
    # Pruned-2 (writer low-value swap, task_id=underline_p0):
    # FileTask(F_WRITER_58, "underline_p0", "bold_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="underline"),
             # params=[
        # Param({}, "strict", {},
              # "I'd like the famous opening line of this Pride and Prejudice chapter to draw the reader's eye when the printed handout reaches my book group — please underline the opening paragraph."),
    # ]),
    FileTask(F_WRITER_58, "doc_font", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Georgia"}, "font_names", {"font_name": "Georgia"},
              "Change the font of the entire Pride and Prejudice excerpt to Georgia."),
    ]),

    # F-WRITER-59 (Moby v2) — color_red + doc_spacing
    FileTask(F_WRITER_59, "color_red_p0", "color_table_text",
             lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="color_red"),
             params=[
        Param({}, "strict", {"examine_color": True},
              "I want the first line of this Moby-Dick chapter to feel dramatic on the printed reading-group handout — could you set the opening paragraph in red font colour for me?"),
    ]),
    FileTask(F_WRITER_59, "doc_spacing", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 2.0}, "line_spacing", {},
              "Set the line spacing of the entire Moby-Dick excerpt to double."),
    ]),

    # F-WRITER-60 (Frankenstein v2) — strike + size18
    FileTask(F_WRITER_60, "strike_p0", "bold_text",
             lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="strike"),
             params=[
        Param({}, "strict", {},
              "I'm peer-reviewing a friend's draft and I think the opening paragraph of this Frankenstein chapter should be cut from their adaptation — could you apply strikethrough to it so they can see my suggested edit?"),
    ]),
    # Pruned (writer dist 13.3%→10%, task_id=size18_p0, eval=0-direct):
    # FileTask(F_WRITER_60, "size18_p0", "highlight_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="size18"),
             # params=[
        # Param({}, "strict", {"examine_font_size": True},
              # "Set 18-point font size on the opening paragraph of this Frankenstein excerpt."),
    # ]),

    # F-WRITER-61 (Sun Tzu) — bold + italic
    # Pruned (writer dist 13.3%→10%, task_id=bold_p0, eval=0-direct):
    # FileTask(F_WRITER_61, "bold_p0", "bold_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="bold"),
             # params=[
        # Param({}, "strict", {},
              # "I'd like the opening tenet of this Sun Tzu Art of War excerpt to dominate the printed page when I share it at our leadership offsite — please bold the opening paragraph for me."),
    # ]),
    # Pruned (writer dist 13.3%→10%, task_id=italic_p0, eval=0-direct):
    # FileTask(F_WRITER_61, "italic_p0", "bold_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="italic"),
             # params=[
        # Param({}, "strict", {},
              # "Italicize the opening paragraph of the Art of War excerpt."),
    # ]),

    # F-WRITER-62 (Tale of Two Cities) — bold + color_blue
    # Pruned (writer dist 13.3%→10%, task_id=bold_p0, eval=0-direct):
    # FileTask(F_WRITER_62, "bold_p0", "bold_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="bold"),
             # params=[
        # Param({}, "strict", {},
              # "I'm building a printed anthology of memorable opening lines for my literature class — could you help me by bolding the famous opening paragraph of A Tale of Two Cities?"),
    # ]),
    FileTask(F_WRITER_62, "color_blue_p0", "color_table_text",
             lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="color_blue"),
             params=[
        Param({}, "strict", {"examine_color": True},
              "Set blue font colour on the opening paragraph of A Tale of Two Cities."),
    ]),

    # F-WRITER-63 (Metamorphosis v2) — highlight + font_georgia
    # Pruned (writer dist 13.3%→10%, task_id=highlight_p0, eval=0-direct):
    # FileTask(F_WRITER_63, "highlight_p0", "highlight_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="highlight_yellow"),
             # params=[
        # Param({}, "strict", {"examine_highlight": True},
              # "I want students to flag the famous transformation sentence of this Metamorphosis chapter for class discussion — please apply a yellow highlight to the opening paragraph of the excerpt."),
    # ]),
    # Pruned-2 (writer low-value swap, task_id=font_georgia_p0):
    # FileTask(F_WRITER_63, "font_georgia_p0", "change_font",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="font_georgia"),
             # params=[
        # Param({}, "strict", {"examine_font_name": True},
              # "Change the font of the opening paragraph of the Metamorphosis chapter to Georgia."),
    # ]),

    # F-WRITER-64 (Sherlock v2) — italic + size14
    # Pruned (writer dist 13.3%→10%, task_id=italic_p0, eval=0-direct):
    # FileTask(F_WRITER_64, "italic_p0", "bold_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="italic"),
             # params=[
        # Param({}, "strict", {},
              # "I'd like the narrator's voice in this Sherlock Holmes chapter to feel distinct from the dialogue when my mystery-genre students read the printed handout — please italicise the opening paragraph for me."),
    # ]),
    # Pruned (writer dist 13.3%→10%, task_id=size14_p0, eval=0-direct):
    # FileTask(F_WRITER_64, "size14_p0", "highlight_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="size14"),
             # params=[
        # Param({}, "strict", {"examine_font_size": True},
              # "Set 14-point font size on the opening paragraph of the Sherlock Holmes chapter."),
    # ]),

    # F-WRITER-65 (Treasure Island v2) — underline + color_red
    # Pruned (writer dist 13.3%→10%, task_id=underline_p0, eval=0-direct):
    # FileTask(F_WRITER_65, "underline_p0", "bold_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="underline"),
             # params=[
        # Param({}, "strict", {},
              # "I'm assembling a chapter-by-chapter study guide for Treasure Island and want the opening setup paragraph to stand apart from the chapter body — could you underline the opening paragraph for me?"),
    # ]),
    # Pruned (writer dist 13.3%→10%, task_id=color_red_p0, eval=0-direct):
    # FileTask(F_WRITER_65, "color_red_p0", "color_table_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="color_red"),
             # params=[
        # Param({}, "strict", {"examine_color": True},
              # "Set red font colour on the opening paragraph of the Treasure Island chapter."),
    # ]),

    # F-WRITER-66 (Tom Sawyer) — bold + highlight
    # Pruned (writer dist 13.3%→10%, task_id=bold_p0, eval=0-direct):
    # FileTask(F_WRITER_66, "bold_p0", "bold_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="bold"),
             # params=[
        # Param({}, "strict", {},
              # "I'd like the opening paragraph of this Tom Sawyer chapter to anchor the printed reading guide my fifth-grade class will use this week — could you bold the opening paragraph of the excerpt?"),
    # ]),
    # Pruned (writer dist 13.3%→10%, task_id=highlight_p0, eval=0-direct):
    # FileTask(F_WRITER_66, "highlight_p0", "highlight_text",
             # lambda s, g, **kw: _gold_gutenberg_p0_op(s, g, op="highlight_yellow"),
             # params=[
        # Param({}, "strict", {"examine_highlight": True},
              # "Highlight the opening paragraph of the Tom Sawyer excerpt in yellow."),
    # ]),

    # ====================================================================
    # Loop 2.B — letter-genre files (no tables / no photos)
    # ====================================================================

    # F-WRITER-67 (cover letter) — find_replace + append
    # Pruned (writer dist 13.3%→10%, task_id=find_replace, eval=0-direct):
    # FileTask(F_WRITER_67, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "Vantage", "new": "Apex"}, "files", {},
              # "I'm tailoring this cover letter for the Apex Robotics application before tonight's deadline — could you replace every occurrence of 'Vantage' with 'Apex' across the letter for me?"),
        # Param({"old": "Senior Research Engineer", "new": "Principal Research Scientist"}, "files", {},
              # "I'd like to retarget my cover letter at the Principal Research Scientist posting the lab opened up — could you find every occurrence of 'Senior Research Engineer' and replace it with 'Principal Research Scientist'?"),
    # ]),
    FileTask(F_WRITER_67, "append_para", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Enclosures: curriculum vitae, three professional references."}, "files", {},
              "I'd like to flag the supporting documents at the foot of the cover letter before I post it tomorrow — could you append a new paragraph at the end with the text 'Enclosures: curriculum vitae, three professional references.'?"),
    ]),

    # F-WRITER-68 (recommendation) — bold + italic_size14
    # Pruned (writer dist 13.3%→10%, task_id=bold_para, eval=0-direct):
    # FileTask(F_WRITER_68, "bold_para", "bold_text",
             # _gold_bold_para, params=[
        # Param({"para_idx": 0}, "strict", {},
              # "I'd like the salutation of this recommendation to stand out before the printed letter goes into the candidate's application packet — could you bold the opening salutation for me?"),
        # Param({"para_idx": 5}, "strict", {},
              # "I'd like the formal sign-off on this recommendation letter to feel emphatic when the candidate's selection committee turns to the closing — could you bold the 'Yours sincerely' line of the letter?"),
    # ]),
    FileTask(F_WRITER_68, "italic_size14", "highlight_text",
             _gold_italic_size14, params=[
        Param({}, "italic_size14", {},
              "Format the entire recommendation letter in italic at 14 point font size."),
    ]),

    # F-WRITER-69 (business letter) — find_replace + doc_font
    # Pruned (writer dist 13.3%→10%, task_id=find_replace, eval=0-direct):
    # FileTask(F_WRITER_69, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "Northbridge", "new": "Eastfield"}, "files", {},
              # "I've been asked to readdress this business letter to the Eastfield branch now that the Northbridge office has closed — could you replace every occurrence of 'Northbridge' with 'Eastfield' for me?"),
        # Param({"old": "framework", "new": "umbrella"}, "files", {},
              # "I'd like to soften the technical-sounding word in this business letter so it reads less corporate to the client — could you find every occurrence of 'framework' and replace it with 'umbrella'?"),
    # ]),
    FileTask(F_WRITER_69, "doc_font", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Times New Roman"}, "font_names", {"font_name": "Times New Roman"},
              "I'd like this business letter to read in the formal house typeface before I post the printed copy on company stationery — please change the font of the entire letter to Times New Roman."),
        Param({"font_name": "Arial"}, "font_names", {"font_name": "Arial"},
              "Change the font of the whole letter to Arial."),
    ]),

    # F-WRITER-70 (resignation letter) — append + doc_spacing
    FileTask(F_WRITER_70, "append_para", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Copy: Human Resources, line manager."}, "files", {},
              "I want HR and my line manager to be copied on this resignation letter before I send it — could you append a new paragraph at the very end with the text 'Copy: Human Resources, line manager.'?"),
    ]),
    FileTask(F_WRITER_70, "doc_spacing", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 1.0}, "line_spacing", {},
              "Reset the line spacing of the entire resignation letter to single (1.0)."),
        Param({"value": 1.15}, "line_spacing", {},
              "I'd like the resignation letter to look formal but not crammed when it lands on my line manager's desk — could you set the line spacing of the whole letter to 1.15 for me?"),
    ]),

    # F-WRITER-71 (thank-you letter) — bold + page_numbers_footer
    # Pruned (writer dist 13.3%→10%, task_id=bold_para, eval=0-direct):
    # FileTask(F_WRITER_71, "bold_para", "bold_text",
             # _gold_bold_para, params=[
        # Param({"para_idx": 0}, "strict", {},
              # "I'd like the printed thank-you note to feel warm and personal when the recipient opens the envelope — could you bold the opening salutation of the letter for me?"),
        # Param({"para_idx": 1}, "strict", {},
              # "I'd like the heart of this thank-you note to feel emphatic when the recipient reads the printed card — could you bold the second paragraph, which carries the principal thanks, for me?"),
    # ]),
    FileTask(F_WRITER_71, "page_numbers_footer", "add_header_footer",
             _gold_page_numbers_footer, params=[
        Param({}, "page_numbers", {},
              "Add page numbers to the footer of the thank-you letter — first enable the footer via Insert > Header and Footer > Footer > Default, then with the cursor inside the footer use Insert > Page Number."),
    ]),

    # ====================================================================
    # Loop 2.C — bullet / numbered list artifact files
    # ====================================================================

    # F-WRITER-72 (project todo) — find_replace + doc_font
    # Pruned (writer dist 13.3%→10%, task_id=find_replace, eval=0-direct):
    # FileTask(F_WRITER_72, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "Pegasus", "new": "Phoenix"}, "files", {},
              # "I'm rebadging this project to-do list now that leadership has renamed the workstream — could you replace every occurrence of 'Pegasus' with 'Phoenix' across the list, please?"),
        # Param({"old": "Marina", "new": "Mariana"}, "files", {},
              # "I'd like the project to-do list to use the engineer's correct name now that I've noticed the typo — could you find every occurrence of 'Marina' and replace it with 'Mariana' across the list?"),
    # ]),
    FileTask(F_WRITER_72, "doc_font", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Verdana"}, "font_names", {"font_name": "Verdana"},
              "Change the font of the entire project to-do list to Verdana."),
    ]),

    # F-WRITER-73 (meeting agenda) — append + doc_spacing
    FileTask(F_WRITER_73, "append_para", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Distributed by the Company Secretary, three business days ahead of the meeting."}, "files", {},
              "I'd like to credit the Company Secretary on this meeting agenda before it goes out to the board — could you append a new paragraph at the end with the text 'Distributed by the Company Secretary, three business days ahead of the meeting.'?"),
    ]),
    # Validation PARAM_REDUCIBLE: dropped 1.15-spacing Param
    # (fine-grained value harder to nail); kept 1.5 (simpler common value).
    FileTask(F_WRITER_73, "doc_spacing", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 1.5}, "line_spacing", {},
              "Set the line spacing of the entire meeting agenda to 1.5."),
    ]),

    # F-WRITER-74 (action items) — find_replace + doc_font
    # Pruned (writer dist 13.3%→10%, task_id=find_replace, eval=0-direct):
    # FileTask(F_WRITER_74, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "executive committee", "new": "leadership council"}, "files", {},
              # "I've been asked to align this action-items document with the new governance taxonomy before the next board pack — could you replace every occurrence of 'executive committee' with 'leadership council'?"),
        # Param({"old": "CFO", "new": "Finance Director"}, "files", {},
              # "I'd like the action items to use the full title my counterpart actually prefers in correspondence — could you find every occurrence of 'CFO' and replace it with 'Finance Director' across the document?"),
    # ]),
    FileTask(F_WRITER_74, "doc_font", "change_font",
             _gold_doc_font, params=[
        # validation — swapped Calibri→Arial (Calibri stripped on LO round-trip).
        Param({"font_name": "Arial"}, "font_names", {"font_name": "Arial"},
              "Change the font of the entire action items document to Arial."),
    ]),

    # F-WRITER-75 (release checklist) — find_replace + page_break
    # Pruned (writer dist 13.3%→10%, task_id=find_replace, eval=0-direct):
    # FileTask(F_WRITER_75, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "production release", "new": "stage rollout"}, "files", {},
              # "I'm repositioning this release checklist for the staging environment we now use before live cutover — could you replace every occurrence of 'production release' with 'stage rollout' for me?"),
        # Param({"old": "release manager", "new": "deployment lead"}, "files", {},
              # "I'd like the release checklist to use the new operations role title we agreed in the platform reorg — could you find every occurrence of 'release manager' and replace it with 'deployment lead' throughout?"),
    # ]),
    FileTask(F_WRITER_75, "page_break", "add_page_break",
             _gold_page_break, params=[
        Param({"para_idx": 7}, "page_break", {"expected_count": 1},
              "Hey, can you throw in a page break right before the Test and Validation section heading? I want that section to start on a new page in the printed runbook."),
    ]),

    # F-WRITER-76 (meeting minutes) — append + doc_font
    FileTask(F_WRITER_76, "append_para", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Minutes recorded by the engineering operations office, draft for review."}, "files", {},
              "I'd like to flag these minutes as a draft awaiting sign-off before I share them with attendees — could you append a new paragraph at the end reading 'Minutes recorded by the engineering operations office, draft for review.'?"),
    ]),
    FileTask(F_WRITER_76, "doc_font", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Arial"}, "font_names", {"font_name": "Arial"},
              "Change the font of the entire meeting minutes to Arial."),
    ]),

    # ====================================================================
    # Loop 2.D — photo-essay files (single image insert)
    # ====================================================================

    # F-WRITER-77 (forest essay) — insert_image + bold
    FileTask(F_WRITER_77, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/forest-trail.jpg",
                 width_in=4.5, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "I'd like to anchor my forest photo essay with the trail shot I took on Sunday before it goes into the print magazine — could you insert the photograph from /home/user/Desktop/forest-trail.jpg below the second paragraph of the essay?"),
    ]),
    # Pruned (writer dist 13.3%→10%, task_id=bold_para, eval=0-direct):
    # FileTask(F_WRITER_77, "bold_para", "bold_text",
             # _gold_bold_para, params=[
        # Param({"para_idx": 0}, "strict", {},
              # "Bold the introductory paragraph of the forest photo essay."),
    # ]),

    # F-WRITER-78 (portrait essay) — insert_image + italic
    FileTask(F_WRITER_78, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/person-headshot-1.jpg",
                 width_in=3.5, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "Insert the portrait headshot from /home/user/Desktop/person-headshot-1.jpg into the essay, below the second paragraph."),
    ]),
    FileTask(F_WRITER_78, "italic_para", "bold_text",
             _gold_italic_para, params=[
        Param({"para_idx": 2}, "strict", {},
              "Italicize the third paragraph (the caption text) of the portrait essay."),
    ]),

    # F-WRITER-79 (night essay) — insert_image + find_replace
    FileTask(F_WRITER_79, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/street-night.jpg",
                 width_in=4.5, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "Insert the night-street photograph from /home/user/Desktop/street-night.jpg into the essay, below the second paragraph."),
    ]),
    # Pruned (writer dist 13.3%→10%, task_id=find_replace, eval=0-direct):
    # FileTask(F_WRITER_79, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "north district", "new": "northern district"}, "files", {},
              # "In the night photo essay, replace every occurrence of 'north district' with 'northern district'."),
    # ]),

    # F-WRITER-80 (bird essay) — insert_image + bold
    FileTask(F_WRITER_80, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/bird-perch.jpg",
                 width_in=4.0, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "Insert the bird-perch photograph from /home/user/Desktop/bird-perch.jpg into the wildlife essay, below the second paragraph."),
    ]),
    # Pruned (writer dist 13.3%→10%, task_id=bold_para, eval=0-direct):
    # FileTask(F_WRITER_80, "bold_para", "bold_text",
             # _gold_bold_para, params=[
        # Param({"para_idx": 2}, "strict", {},
              # "Bold the third paragraph (the caption) of the wildlife essay."),
    # ]),

    # F-WRITER-81 (mars essay) — insert_image + append
    FileTask(F_WRITER_81, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/mars-curiosity-panorama.jpg",
                 width_in=5.0, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "I'm laying out a Mars-exploration feature for next month's club newsletter and want a striking surface image — could you insert the Curiosity rover panorama from /home/user/Desktop/mars-curiosity-panorama.jpg below the second paragraph of the essay?"),
    ]),
    FileTask(F_WRITER_81, "append_para", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Image courtesy of the planetary-imaging team; public-domain release."}, "files", {},
              "Append a new paragraph at the end of the Mars essay with the text 'Image courtesy of the planetary-imaging team; public-domain release.'."),
    ]),

    # ====================================================================
    # Catch-up — real-source content + structural / length variety
    # (F-WRITER-82..93). Doc-wide ops favoured because the new structural
    # variants (Q&A, bullet reference, mixed-length essay, structured wiki
    # with title + headings) make per-paragraph ordinal counting ambiguous.
    # Per-paragraph tasks are restricted to files where the target index
    # has an unambiguous named anchor.
    # ====================================================================

    # F-WRITER-82 (wiki_coffee, structured) — doc_font + doc_spacing
    FileTask(F_WRITER_82, "doc_font", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Georgia"}, "font_names", {"font_name": "Georgia"},
              "I'd like this Coffee reference article to read in a serif typeface that matches the rest of our printed encyclopaedia volume — could you change the font of the entire article to Georgia?"),
    ]),
    # Validation PARAM_REDUCIBLE: dropped 2.0-spacing Param;
    # kept 1.5.
    FileTask(F_WRITER_82, "doc_spacing", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 1.5}, "line_spacing", {},
              "Set the line spacing of the entire Coffee article to 1.5."),
    ]),

    # F-WRITER-83 (wiki_eiffel, no title / no headings) — bold first
    # paragraph (idx=0 unambiguous) + find_replace
    # Pruned (writer dist 13.3%→10%, task_id=bold_para, eval=0-direct):
    # FileTask(F_WRITER_83, "bold_para", "bold_text",
             # _gold_bold_para, params=[
        # Param({"para_idx": 0}, "strict", {},
              # "I'm preparing this Eiffel Tower article for a school project and I want the introductory paragraph to grab the reader before they scroll into the historical detail — could you bold the first paragraph?"),
    # ]),
    # Pruned (writer dist 13.3%→10%, task_id=find_replace, eval=0-direct):
    # FileTask(F_WRITER_83, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "Eiffel", "new": "Iron"}, "files", {},
              # "In the Eiffel Tower article, replace every occurrence of 'Eiffel' with 'Iron'."),
        # Param({"old": "Paris", "new": "the capital"}, "files", {},
              # "Find every occurrence of 'Paris' and replace it with 'the capital'."),
    # ]),

    # F-WRITER-84 (wiki_octopus, structured) — find_replace + doc_font
    # Pruned (writer dist 13.3%→10%, task_id=find_replace, eval=0-direct):
    # FileTask(F_WRITER_84, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "octopus", "new": "cephalopod"}, "files", {},
              # "I'd like to use the broader zoological term in this article so it fits the volume on cephalopods I'm assembling — could you replace every occurrence of 'octopus' with 'cephalopod' for me?"),
        # Param({"old": "ocean", "new": "sea"}, "files", {},
              # "Find every occurrence of 'ocean' and replace it with 'sea'."),
    # ]),
    FileTask(F_WRITER_84, "doc_font", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Times New Roman"}, "font_names", {"font_name": "Times New Roman"},
              "Change the font of the entire Octopus article to Times New Roman."),
    ]),

    # F-WRITER-85 (wiki_everest, title + 6 paras) — doc_spacing +
    # page_numbers_footer
    # Validation PARAM_REDUCIBLE: dropped 1.15-spacing Param
    # (fine-grained); kept 2.0 (simpler common "double" target).
    FileTask(F_WRITER_85, "doc_spacing", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 2.0}, "line_spacing", {},
              "I'm printing this Mount Everest article for a classroom annotation exercise and I want students to have room to write between the lines — could you set the line spacing of the whole article to double?"),
    ]),
    FileTask(F_WRITER_85, "page_numbers_footer", "add_header_footer",
             _gold_page_numbers_footer, params=[
        Param({}, "page_numbers", {},
              "Add page numbers to the footer of the Mount Everest article — first enable the footer via Insert > Header and Footer > Footer > Default, then with the cursor inside the footer use Insert > Page Number."),
    ]),

    # F-WRITER-86 (short memo, title + 1 body para) — bold body paragraph
    # (named anchor — only one body paragraph) + doc_font
    FileTask(F_WRITER_86, "bold_body", "bold_text",
             _gold_bold_para, params=[
        Param({"para_idx": 1}, "strict", {},
              "I want my Coffee Production briefing memo to flag its single body finding clearly when it lands on the trade-desk inbox — could you bold the body paragraph for me, please?"),
    ]),
    FileTask(F_WRITER_86, "doc_font", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Arial"}, "font_names", {"font_name": "Arial"},
              "Change the font of the entire briefing memo to Arial."),
    ]),

    # F-WRITER-87 (qa_solar_system) — doc-wide only (Q/A counting ambiguous)
    FileTask(F_WRITER_87, "doc_font", "change_font",
             _gold_doc_font, params=[
        # validation — swapped Calibri→Arial.
        Param({"font_name": "Arial"}, "font_names", {"font_name": "Arial"},
              "I'd like this Solar System FAQ to match the rest of our science fair handouts which all use a clean sans-serif — could you change the font of the entire FAQ to Arial, please?"),
        Param({"font_name": "Verdana"}, "font_names", {"font_name": "Verdana"},
              "Change the font of the whole FAQ to Verdana."),
    ]),
    FileTask(F_WRITER_87, "doc_spacing", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 1.5}, "line_spacing", {},
              "Set the line spacing of the entire Solar System FAQ to 1.5."),
    ]),

    # F-WRITER-88 (bullet yoga) — doc-wide only (bullet/para indices
    # interleave ambiguously)
    FileTask(F_WRITER_88, "doc_spacing", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 1.5}, "line_spacing", {},
              "I'd like to make this Yoga reference document easier to read on the studio noticeboard before the next class begins — could you set the line spacing of the entire document to 1.5?"),
    ]),
    # Pruned (writer dist 13.3%→10%, task_id=find_replace, eval=0-direct):
    # FileTask(F_WRITER_88, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "yoga", "new": "practice"}, "files", {},
              # "In the Yoga reference, replace every occurrence of 'yoga' with 'practice'."),
        # Param({"old": "Sanskrit", "new": "ancient"}, "files", {},
              # "Find every occurrence of 'Sanskrit' and replace it with 'ancient'."),
    # ]),

    # F-WRITER-89 (mixed pizza essay) — doc-wide only (mixed-kind block
    # interleaving)
    FileTask(F_WRITER_89, "doc_font", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Georgia"}, "font_names", {"font_name": "Georgia"},
              "I want this Pizza essay to read in a warm serif typeface that fits the food-magazine feature it's destined for — could you change the font of the entire essay to Georgia for me?"),
    ]),
    # Pruned (writer dist 13.3%→10%, task_id=find_replace, eval=0-direct):
    # FileTask(F_WRITER_89, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "pizza", "new": "pie"}, "files", {},
              # "In the Pizza essay, replace every occurrence of 'pizza' with 'pie'."),
        # Param({"old": "Naples", "new": "Napoli"}, "files", {},
              # "Find every occurrence of 'Naples' and replace it with 'Napoli'."),
    # ]),

    # F-WRITER-90 (mixed volcano essay) — page_numbers_footer + find_replace
    FileTask(F_WRITER_90, "page_numbers_footer", "add_header_footer",
             _gold_page_numbers_footer, params=[
        Param({}, "page_numbers", {},
              "Add page numbers to the footer of the Volcano essay — first enable the footer via Insert > Header and Footer > Footer > Default, then with the cursor inside the footer use Insert > Page Number."),
    ]),
    # Pruned (writer dist 13.3%→10%, task_id=find_replace, eval=0-direct):
    # FileTask(F_WRITER_90, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "volcano", "new": "vent"}, "files", {},
              # "In the Volcano essay, replace every occurrence of 'volcano' with 'vent'."),
        # Param({"old": "Earth", "new": "the planet"}, "files", {},
              # "Find every occurrence of 'Earth' and replace it with 'the planet'."),
    # ]),

    # F-WRITER-91 (wiki_lego, structured) — doc_font + append
    FileTask(F_WRITER_91, "doc_font", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Verdana"}, "font_names", {"font_name": "Verdana"},
              "I'd like this Lego reference article to match the cheerful sans-serif typography of the rest of my children's encyclopaedia — could you change the font of the whole article to Verdana for me?"),
    ]),
    FileTask(F_WRITER_91, "append_para", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Source: Wikipedia article on Lego, retrieved 2026."}, "files", {},
              "Append a new paragraph at the end of the Lego article with the text 'Source: Wikipedia article on Lego, retrieved 2026.'."),
    ]),

    # F-WRITER-92 (qa_internet) — doc-wide only
    FileTask(F_WRITER_92, "doc_font", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Arial"}, "font_names", {"font_name": "Arial"},
              "Change the font of the entire Internet of Things FAQ to Arial."),
    ]),
    # Pruned (writer dist 13.3%→10%, task_id=find_replace, eval=0-direct):
    # FileTask(F_WRITER_92, "find_replace", "find_replace",
             # _gold_find_replace, params=[
        # Param({"old": "IoT", "new": "the Internet of Things"}, "files", {},
              # "In the IoT FAQ, replace every occurrence of 'IoT' with 'the Internet of Things'."),
        # Param({"old": "sensors", "new": "transducers"}, "files", {},
              # "Find every occurrence of 'sensors' and replace it with 'transducers'."),
    # ]),

    # F-WRITER-93 (bullet origami) — doc-wide only
    FileTask(F_WRITER_93, "doc_font", "change_font",
             _gold_doc_font, params=[
        # validation — swapped Calibri→Arial.
        Param({"font_name": "Arial"}, "font_names", {"font_name": "Arial"},
              "Change the font of the entire Origami reference document to Arial."),
    ]),
    FileTask(F_WRITER_93, "append_para", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Reference compiled from the Wikipedia article on origami, 2026."}, "files", {},
              "Append a new paragraph at the very end of the Origami reference with the text 'Reference compiled from the Wikipedia article on origami, 2026.'."),
    ]),

    # ====================================================================
    # F-WRITER-94 — chemistry notes for compare_subscript_contains.
    # Closes the character-level subscript axis gap.
    # compare_docx_strict has no examine_subscript flag,
    # so this uses the bespoke `compare_subscript_contains` evaluator.
    # Validation PARAM_REDUCIBLE: reduced from 4 formulas to 1
    # formula (H2O) — `compare_subscript_contains` is contains-based, so
    # subscripting one formula with one digit still satisfies the eval
    # (paragraph has >=1 subscript-run pair). Instruction names the single
    # target explicitly to avoid the 2-digit selection variant overhead.
    # ====================================================================
    FileTask(F_WRITER_94, "subscript_chem", "subscript_text",
             _gold_subscript_chemistry, params=[
        Param({}, "subscript", {},
              "I want my chemistry notes to read properly on the printed handout I'll give to my tutoring students this evening — could you apply subscript formatting to the digit '2' inside the formula H2O in the formulas paragraph?"),
    ]),

    # ====================================================================
    # F-WRITER-95 — coloured-table worksheet for evaluate_colored_words_
    # in_tables. Source has a 2x3 table with all cells GREEN; agent must
    # re-colour cell text RED (vowel-start) / BLUE (consonant-start).
    # ====================================================================
    FileTask(F_WRITER_95, "color_table_words", "color_table_text",
             _gold_colored_table, params=[
        Param({}, "colored_table", {},
              "I'd like to use this categorisation worksheet in tomorrow's literacy lesson and need the vowel/consonant colour code applied — could you change the font colour of each cell so that words starting with a vowel are red (Custom Color #FF0000) and words starting with a consonant are blue (Custom Color #0000FF), replacing the existing green on every cell?"),
    ]),

    # ====================================================================
    # F-WRITER-96 — mixed-alignment within one paragraph (eval skill from
    # eval task 0a0faba3). Split target paragraph into two halves and apply
    # different alignment to each half. Eval = compare_docx_strict which
    # already checks paragraph alignment in _paragraph_format_signature.
    # ====================================================================
    FileTask(F_WRITER_96, "mixed_align_split", "bold_text",
             _gold_mixed_alignment_split, params=[
        Param({"para_idx": 0, "split_word": 3,
               "first_align": "left", "second_align": "right"},
              "strict", {},
              "I'd like to leave some space in the middle of the first paragraph for inserting a photograph later, so could you split it after the first three words and make those three words left-aligned with the remaining text right-aligned?"),
        Param({"para_idx": 1, "split_word": 8,
               "first_align": "left", "second_align": "right"},
              "strict", {},
              "Please help me lay out the second paragraph so the first half (eight words) is left-aligned and the rest is right-aligned — I want to drop an inline figure into the middle gap when I review the printed copy."),
    ]),

    # ====================================================================
    # F-WRITER-97 — doc-wide case conversion (eval skill from the eval task
    # asking to convert all uppercase to lowercase across the document).
    # Eval = compare_docx_files (text-only) so the gold simply re-renders
    # every paragraph with the new case applied.
    # ====================================================================
    FileTask(F_WRITER_97, "doc_case_convert", "footnote_citation",
             _gold_doc_case_convert, params=[
        Param({"mode": "lower"}, "files", {},
              "I am tidying up the brief before circulating it to the team; could you help me convert every word in the document to lowercase so the typography reads consistently throughout?"),
        Param({"mode": "upper"}, "files", {},
              "I'd like to convert the whole document to uppercase so the printed handout reads as an attention-grabbing notice — please change every character in the body to uppercase."),
    ]),

    # ====================================================================
    # F-WRITER-98 — sentence-level ordinal format (eval skill: bold/italic
    # the first or last sentence of a specific paragraph, vs synth's normal
    # paragraph-level ordinal). Splits target paragraph's runs so the
    # named sentence carries the format. Eval = compare_docx_strict.
    # ====================================================================
    FileTask(F_WRITER_98, "first_sentence_op", "bold_text",
             _gold_sentence_op, params=[
        Param({"para_idx": 0, "which": "first", "op": "bold"}, "strict", {},
              "I'm preparing the quarterly report for the steering meeting and would like to draw the reader's eye to the headline finding — please bold just the first sentence of the opening paragraph."),
        Param({"para_idx": 1, "which": "last", "op": "italic"}, "strict", {},
              "Could you help me emphasise the closing thought of the second paragraph for the reviewer — please italicise only the last sentence of that paragraph, leaving everything before it in normal style."),
    ]),

    # validation RESCALER — eval-coverage fill for 4 niche skills.
    # Mirrors eval/writer rows 6, 11, 14, 15 that the P2-P5 pass missed.
    # Reuses real-source fixtures where suitable (F_WRITER_94 chemistry,
    # F_WRITER_49 printer-troubleshoot table, F_WRITER_98 multi-sentence
    # research report); only F_WRITER_102 is a new minimal CSV fixture.

    # F-WRITER-94 (reuse) — center-align heading (eval row 6).
    # training validation: ordinal "first/third paragraph" is
    # ambiguous because the doc's title heading is paragraphs[0]. Anchor by
    # literal opening words instead (resolves WHAT, not HOW).
    FileTask(F_WRITER_94, "center_heading", "bold_text",
             _gold_center_para, params=[
        Param({"para_idx": 0}, "strict", {},
              "I'm prepping my chemistry notes for circulation to the lab section — please center-align the document title 'Chemistry Notes — Common Formulas' at the very top so it reads as the title block."),
        Param({"para_idx": 2}, "strict", {},
              "Please center-align the paragraph that begins 'Common compounds include H2O for water' in this chemistry notes document for the printed handout."),
    ]),

    # F-WRITER-49 (reuse) — delete severity table rows by P-tier (eval row 11).
    # Source has a real Escalation Targets table (Severity / Response time)
    # with rows P1..P4. Agent keeps only critical tiers.
    # Validation PARAM_REDUCIBLE: dropped the P1-only variant
    # (delete 3 rows, keep 1) — multi-row deletion at that depth is high-
    # step. Kept the P1+P2 variant (delete 2 rows, keep 2).
    FileTask(F_WRITER_49, "delete_rows_filter", "footnote_citation",
             _gold_filter_table_rows_keep_range, params=[
        Param({"col_idx": 0, "start": "P1", "end": "P2"}, "tables", {},
              "We're tightening the on-call runbook to surface only the page-out tiers — in the Escalation Targets table at the top of the troubleshooting guide, please delete every row whose Severity is P3 or P4, leaving only P1 and P2 visible."),
    ]),

    # F-WRITER-98 (reuse) — sentence-per-line split (eval row 14).
    # Reuses the same multi-sentence research-report body that powers the
    # existing first_sentence_op FileTask, exercising a different operation
    # (split into separate paragraphs) on the same paragraph structure.
    FileTask(F_WRITER_98, "sentence_per_line", "footnote_citation",
             _gold_sentence_per_line, params=[
        # para_idx aligned to the instruction (no title in _docx_body_py → 0-indexed body):
        # "opening paragraph" = idx 0, "second paragraph" = idx 1. Was off-by-one (1/2) →
        # gold split the paragraph AFTER the instructed one, so a correct agent scored 0.
        Param({"para_idx": 0}, "files", {"ignore_blanks": False},
              "I'm preparing a printed reading guide for the lab seminar and want the opening paragraph of this research report easier to scan sentence-by-sentence — please split it so each sentence becomes its own paragraph."),
        Param({"para_idx": 1}, "files", {"ignore_blanks": False},
              "Could you help me restructure the second paragraph of this research report for the slide-deck handout — please break it apart so each sentence sits on its own paragraph, preserving wording and punctuation."),
    ]),

    # ----- validation SWAP-2 — eval-anchored variants (replace low-value finds).
    # Reuses existing real-source Files with open FileTask slots. Each new
    # FileTask exercises an operation that IS tested in eval/writer (font
    # size, case conversion, mixed alignment, PDF export, center align)
    # rather than synth-only operations (find/replace, bold-only) where
    # eval has zero direct coverage.

    # font_size_change × 4 — eval rows test "change font size to N" (3 rows).
    FileTask(F_WRITER_27, "size20_p0", "highlight_text",
             _gold_size_para, params=[
        Param({"para_idx": 0, "size_pt": 20}, "strict", {"examine_font_size": True},
              "I'm preparing the wildlife journal cover paragraph as a banner for the printed exhibition handout — please set the font size of the opening paragraph to 20 points so it reads as the page title."),
    ]),
    FileTask(F_WRITER_31, "size16_p1", "highlight_text",
             _gold_size_para, params=[
        Param({"para_idx": 1, "size_pt": 16}, "strict", {"examine_font_size": True},
              "I'd like the comparison caption under the galaxy pair to sit at a clearly readable size for the printed astronomy-club handout — please apply 16-point font size to the second paragraph (change the size property of the existing text; don't add a new paragraph)."),
    ]),
    FileTask(F_WRITER_40, "size14_body", "highlight_text",
             _gold_size_para, params=[
        Param({"para_idx": 2, "size_pt": 14}, "strict", {"examine_font_size": True},
              "I'm reformatting this letters-archive essay for the seminar reading pack and want the third paragraph slightly enlarged for emphasis — please set its font size to 14 points."),
    ]),
    FileTask(F_WRITER_60, "size18_opener", "highlight_text",
             _gold_size_para, params=[
        Param({"para_idx": 0, "size_pt": 18}, "strict", {"examine_font_size": True},
              "I'd like the opening of this Frankenstein chapter to read as a banner in the printed reading-circle dossier — please set the first paragraph to 18-point font size."),
    ]),

    # case_conversion × 3 — eval rows test upper/lower/title case (2 rows).
    FileTask(F_WRITER_67, "doc_case_upper", "footnote_citation",
             _gold_doc_case_convert, params=[
        Param({"mode": "upper"}, "files", {},
              "I'd like the cover letter to read as a stencilled all-caps notice for the file-folder placeholder I'm printing — please convert every paragraph in the letter to uppercase."),
    ]),
    FileTask(F_WRITER_71, "doc_case_lower", "footnote_citation",
             _gold_doc_case_convert, params=[
        Param({"mode": "lower"}, "files", {},
              "I'd like to tone the thank-you letter back to a softer lowercase typography for the personal note I'm enclosing with the gift — please convert every paragraph in the letter to lowercase."),
    ]),
    FileTask(F_WRITER_80, "doc_case_upper_bird", "footnote_citation",
             _gold_doc_case_convert, params=[
        Param({"mode": "upper"}, "files", {},
              "I'm preparing the bird-photo essay as a wall-mounted notice at the nature centre's open day — please convert all paragraphs to uppercase so it reads at a distance from the corridor."),
    ]),

    # mixed_alignment × 2 — eval row 2 splits a paragraph into left + right halves.
    FileTask(F_WRITER_72, "mixed_align_todo", "bold_text",
             _gold_mixed_alignment_split, params=[
        Param({"para_idx": 0, "split_word": 4,
               "first_align": "left", "second_align": "right"},
              "strict", {},
              "I'd like to leave a gap in the first paragraph of the project to-do list for inserting a Gantt thumbnail later — could you split it after the first four words with those words left-aligned and the rest right-aligned?"),
    ]),
    FileTask(F_WRITER_77, "mixed_align_forest", "bold_text",
             _gold_mixed_alignment_split, params=[
        Param({"para_idx": 1, "split_word": 6,
               "first_align": "left", "second_align": "right"},
              "strict", {},
              # Validation fix: instruction asked agent to also "drop a forest
              # photo" but evaluator has examine_images=False — so the photo
              # ask is unverifiable and confuses the agent into believing a
              # photo file/clipboard is required. Drop the photo wording.
              "Reformat the second paragraph of this photo essay caption so the first six words sit left-aligned (as a lead-in) and the remaining words are right-aligned."),
    ]),

    # pdf_export × 2 — eval row 7 tests PDF export.
    FileTask(F_WRITER_84, "pdf_export_octopus", "pdf_export",
             _gold_pdf_export, params=[
        Param({}, "pdf_export", {},
              "I'd like to share this octopus encyclopaedia entry with my marine-biology study group as a PDF — please export the current document to PDF via File > Export As > Export Directly as PDF, accept the default filename and Desktop location, then click Export."),
    ]),
    FileTask(F_WRITER_88, "pdf_export_yoga", "pdf_export",
             _gold_pdf_export, params=[
        Param({}, "pdf_export", {},
              "I'm circulating the yoga reference card to the studio members as a printable PDF — please export this document to PDF via Shift+Ctrl+E, accept the default filename and Desktop location, and finalise the export."),
    ]),

    # center_align × 1 — eval row 6 tests center align heading.
    FileTask(F_WRITER_89, "center_pizza_heading", "bold_text",
             _gold_center_para, params=[
        Param({"para_idx": 0}, "strict", {},
              "I'd like the pizza-recipe title to sit centred at the top of the printed recipe card I'm laminating for the kitchen — please center-align the first paragraph (the recipe heading)."),
    ]),

    # F-WRITER-102 (new) — convert comma-separated paragraph to a 1-row table
    # (eval row 15). The only new File: existing fixtures have prose paragraphs
    # rather than the CSV-style paragraph this operation needs.
    FileTask(F_WRITER_102, "text_to_table", "blank_table_insert",
             _gold_text_to_table, params=[
        # _src_csv_paragraph builds an 8-paragraph memo; the Alice Chen
        # CSV row is paragraph index 2 (after heading + intro prose),
        # the Project Apollo row is paragraph index 6.
        Param({"para_idx": 2, "sep": ","}, "files", {},
              "I'm tidying up the team contact register before circulating it to new hires — could you convert the Alice Chen contact line (the comma-separated entry under the intro paragraph) into a one-row table with each comma-delimited field in its own cell?"),
        Param({"para_idx": 6, "sep": ","}, "files", {},
              "Help me convert the Project Apollo status line (the comma-separated entry under the Project Status Snapshot heading) into a clean tabular layout for the dashboard — please turn that comma-separated row into a one-row table, one cell per field."),
    ]),

    # ====================================================================
    # F-WRITER-103 (Wind-turbine photo host) — image insert
    # ====================================================================
    FileTask(F_WRITER_103, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/wind-turbine.jpg",
                 width_in=4.0, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "I'm finalising the renewable briefing for the policy committee on Thursday and would like a hero shot of an operating turbine on the cover — please insert the wind-turbine photograph from /home/user/Desktop/wind-turbine.jpg into the document below the second paragraph."),
    ]),

    # ====================================================================
    # F-WRITER-104 (Cargo-port photo host) — image insert
    # ====================================================================
    FileTask(F_WRITER_104, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/cargo-port.jpg",
                 width_in=4.0, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "I'd like to illustrate this logistics dossier with an aerial of the container terminal for the printed brief — could you insert the cargo-port photograph from /home/user/Desktop/cargo-port.jpg into the document below the second paragraph?"),
    ]),

    # ====================================================================
    # F-WRITER-105 (Graduation photo host) — image insert
    # ====================================================================
    FileTask(F_WRITER_105, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/graduation-ceremony.jpg",
                 width_in=4.5, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "I'm putting together the commencement essay for the alumni magazine and would like the ceremony photograph to sit between the opening reflection and the closing address — please insert the graduation-ceremony photograph from /home/user/Desktop/graduation-ceremony.jpg below the second paragraph of the document."),
    ]),

    # ====================================================================
    # F-WRITER-106 (Solar-farm photo host) — image insert
    # ====================================================================
    FileTask(F_WRITER_106, "insert_image", "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/solar-farm.jpg",
                 width_in=4.0, insert_after_idx=1),
             params=[
        Param({}, "strict", {"examine_images": True},
              "Could you place the solar farm photograph from /home/user/Desktop/solar-farm.jpg into this solar adoption note, inserted below the second paragraph so it sits between the technology summary and the adoption outlook?"),
    ]),

    # ========================================================================
    # validation narrow-evaluator pass — sibling specialized templates that
    # close the writer specialized_uncov / tables / compound_multi_property
    # holes flagged in the writer historical gap audit. These FileTasks reuse existing real-
    # source Files (F_WRITER_5/16/17/49/94/95/etc.) plus the new train-records
    # fixture F_WRITER_107 and exercise upstream evaluators that synth had 0
    # or 1 coverage of in the baseline: compare_docx_tables (content-filled
    # variant), is_first_line_centered (multi-file coverage), compare_
    # subscript_contains (2nd target paragraph), evaluate_strike_through_last_
    # paragraph (2nd source body), check_tabstops, compare_unique_train_
    # records, find_default_font, and compound `func`-as-list rows
    # (compare_docx_files + compare_subscript_contains; compare_line_spacing ×
    # 2; compare_docx_tables × 2).
    # ========================================================================

    # ---- compare_docx_tables × content-filled (writer tables -12pp 🔴) ----
    # F_WRITER_5 already has 1 task slot used (insert_table 3x3/4x4 blank);
    # add a second filled-table task on the same brief fixture.
    FileTask(F_WRITER_5, "fill_action_items_table", "blank_table_insert",
             lambda s, g, **kw: _gold_insert_filled_table(
                 s, g, rows_data=[
                     ["Action", "Owner", "Due"],
                     ["Draft scope", "Anna", "Fri"],
                     ["Confirm budget", "Marcus", "Mon"],
                 ]),
             params=[
        Param({}, "tables", {},
              "I need a quick action-items grid at the end of this brief for the steering group to scan over coffee — please append a 3-row by 3-column table whose first row reads \"Action\", \"Owner\", \"Due\" and whose two body rows read \"Draft scope\", \"Anna\", \"Fri\" and \"Confirm budget\", \"Marcus\", \"Mon\" in that order."),
    ]),

    # F_WRITER_9 (writer_manual_13pt) — second task slot is free.
    FileTask(F_WRITER_9, "fill_release_summary_table", "blank_table_insert",
             lambda s, g, **kw: _gold_insert_filled_table(
                 s, g, rows_data=[
                     ["Phase", "Status", "ETA"],
                     ["Design", "Done", "Q1"],
                     ["Build", "In flight", "Q2"],
                 ]),
             params=[
        Param({}, "tables", {},
              "I'm appending a phase-status snapshot to the end of this manual for the release readout — please add a 3x3 table whose header row reads \"Phase\", \"Status\", \"ETA\" and whose two body rows read \"Design\", \"Done\", \"Q1\" then \"Build\", \"In flight\", \"Q2\"."),
    ]),

    # F_WRITER_14 — second slot free.
    FileTask(F_WRITER_14, "fill_character_table", "blank_table_insert",
             lambda s, g, **kw: _gold_insert_filled_table(
                 s, g, rows_data=[
                     ["Character", "Role"],
                     ["Pip", "Narrator"],
                     ["Estella", "Foil"],
                 ]),
             params=[
        Param({}, "tables", {},
              "I'd like a tidy character-role table at the foot of this novel-opening note for my reading-circle handout — please append a 3-row by 2-column table with the header \"Character\", \"Role\" and two body rows reading \"Pip\", \"Narrator\" then \"Estella\", \"Foil\"."),
    ]),

    # ---- is_first_line_centered (specialized_uncov -17pp 🔴) -------------
    # Add 2 quote-anchored center-heading tasks on fresh files. F_WRITER_22
    # / F_WRITER_25 (gutenberg) have a single task slot each.
    # validation wording fix: F_WRITER_22 source has no Heading-style paragraph
    # (just plain prose). Agents reported infeasible because "heading" referent
    # was missing. Eval `is_first_line_centered` only checks the first line is
    # centered, so rephrase to "first line" matches eval contract + source.
    FileTask(F_WRITER_22, "first_centered_sherlock", "bold_text",
             _gold_first_centered, params=[
        Param({}, "first_centered", {},
              "Centre-align the first line of this Sherlock Holmes excerpt so it sits as a title at the top of the page."),
    ]),
    FileTask(F_WRITER_25, "first_centered_earnest", "bold_text",
             _gold_first_centered, params=[
        Param({}, "first_centered", {},
              "Centre the first line of this Earnest excerpt so the title sits at the top of the page."),
    ]),

    # ---- evaluate_strike_through_last_paragraph × 2 new variants ----------
    FileTask(F_WRITER_24, "strike_last_meta", "bold_text",
             _gold_strike_last, params=[
        Param({}, "strike_last", {},
              "I'm drafting an edit pass on this Metamorphosis excerpt and want the closing remark marked as cut for review — please apply strike-through formatting to every word of the last paragraph."),
    ]),
    FileTask(F_WRITER_23, "strike_last_treasure", "bold_text",
             _gold_strike_last, params=[
        Param({}, "strike_last", {},
              "Mark the final paragraph of this Treasure Island excerpt for deletion in the next revision by striking through every word of that closing paragraph."),
    ]),

    # ---- check_tabstops × 2 (specialized_uncov, brand-new fn) -------------
    # F_WRITER_36 (tech_spec) and F_WRITER_42 (recipe_salad) have one task
    # slot each — add a tab-stops task on the second body paragraph.
    FileTask(F_WRITER_36, "set_tabstops_spec", "bold_text",
             _gold_set_tabstops, params=[
        Param({"para_idx": 1, "positions_in": [1.5, 3.0, 4.5],
               "alignment": "left"},
              "tabstops", {},
              "I'm laying out the spec sheet's parameter row for the printed reference card — please set three left-aligned tab stops on the second paragraph at 1.5 inches, 3 inches, and 4.5 inches so the columns line up cleanly."),
    ]),
    FileTask(F_WRITER_42, "set_tabstops_recipe", "bold_text",
             _gold_set_tabstops, params=[
        Param({"para_idx": 1, "positions_in": [2.0, 4.0],
               "alignment": "left"},
              "tabstops", {},
              "Help me set two left-aligned tab stops on the second paragraph of this recipe at 2 inches and 4 inches so the ingredient columns sit at consistent positions on the printed card."),
    ]),

    # ---- compare_unique_train_records × 2 (specialized_uncov, brand-new) --
    FileTask(F_WRITER_107, "dedupe_train_ids", "footnote_citation",
             _gold_unique_train_records, params=[
        Param({}, "unique_train_records", {},
              "I'm auditing this departure register before circulating it to the operations team — could you delete the rows whose train id duplicates an earlier line so each train id appears at most once in the register?"),
    ]),

    # ---- evaluate_colored_words_in_tables × 1 new variant -----------------
    # The colored-table file already has 1 task on F_WRITER_95; that file
    # has source random-seed variability so a second Param exercises a
    # different word permutation.
    FileTask(F_WRITER_95, "color_table_words_v2", "color_table_text",
             _gold_colored_table, params=[
        Param({}, "colored_table", {},
              "Please re-colour the cells of the categorisation table so each vowel-initial word is shown in red (Custom Color #FF0000) font and each consonant-initial word in blue (Custom Color #0000FF), replacing the green placeholder colour on every cell."),
    ]),

    # ---- find_default_font × 2 (LO registry default font picker) ---------
    # Already wired via _build_default_font_evaluator; use a fresh FileTask for
    # default-font coverage.
    FileTask(F_WRITER_10, "default_font_arial", "change_font",
             _gold_default_font_noop, params=[
        Param({"font_name": "Arial"}, "default_font", {"font_name": "Arial"},
              "Set LibreOffice Writer's default Western text font to Arial so any new document opens with that typeface by default."),
    ]),
    FileTask(F_WRITER_16, "default_font_tnr", "change_font",
             _gold_default_font_noop, params=[
        Param({"font_name": "Times New Roman"}, "default_font",
              {"font_name": "Times New Roman"},
              "Configure LibreOffice Writer so the default Western text font for new documents is Times New Roman."),
    ]),

    # ---- compound: [compare_docx_files, compare_subscript_contains] -------
    # Mirrors eval row 0b17a146 (H2O). Agent must (1) preserve doc text and
    # (2) subscript the H2O '2'. atom_2 + compound_multi_property pattern.
    FileTask(F_WRITER_108, "compound_h2o_files_subscript", "subscript_text",
             _gold_subscript_chemistry, params=[
        Param({}, "compound", {
            "atoms": [
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": True}},
                {"func": "compare_subscript_contains",
                 "expected_path": "__GOLD__"},
            ]},
              "Help me change the 2 in \"H2O\" to a subscript in this chemistry brief, keeping the rest of the document text exactly as written."),
    ]),

    # ---- compound: compare_line_spacing × 2 (same-fn ×N pattern) ----------
    # F_WRITER_8 (writer_essay_spaced) already has 2 tasks; use F_WRITER_47
    # (manual_install) which has 1 task slot free.
    FileTask(F_WRITER_47, "compound_double_spacing_two_paras", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 2.0}, "compound", {
            "atoms": [
                {"func": "compare_line_spacing",
                 "expected_path": "__GOLD__"},
                {"func": "compare_line_spacing",
                 "expected_path": "__GOLD__"},
            ]},
              "Please set the line spacing on the first two paragraphs of this install manual to double line spacing — make the spacing change apply to both paragraphs."),
    ]),

    # ---- compound: compare_docx_tables × 2 (multi-table edits) -----------
    # F_WRITER_48 (manual_safety) and F_WRITER_51 (manual_device) are
    # single-task. Use F_WRITER_48 with a content-filled table compound.
    FileTask(F_WRITER_48, "compound_double_table_check", "blank_table_insert",
             lambda s, g, **kw: _gold_insert_filled_table(
                 s, g, rows_data=[
                     ["Hazard", "Mitigation"],
                     ["Slip", "Mat"],
                     ["Heat", "Shield"],
                 ]),
             params=[
        Param({}, "compound", {
            "atoms": [
                {"func": "compare_docx_tables",
                 "expected_path": "__GOLD__"},
                {"func": "compare_docx_tables",
                 "expected_path": "__GOLD__"},
            ]},
              "Append a 3-row by 2-column hazard-mitigation table at the end of this safety manual whose header reads \"Hazard\", \"Mitigation\" and whose body rows are \"Slip\", \"Mat\" then \"Heat\", \"Shield\" so it can be cross-checked against the printed handout."),
    ]),

    # ---- additional table-content variants (tables -9.9pp ⚠️) -----------
    # More filled-table tasks to push synth tables from 8 → 12+, closing on
    # the eval 13.6% target. All use real-source single-task files.
    FileTask(F_WRITER_35, "fill_market_table", "blank_table_insert",
             lambda s, g, **kw: _gold_insert_filled_table(
                 s, g, rows_data=[
                     ["Segment", "Growth"],
                     ["Cloud", "High"],
                     ["Mobile", "Steady"],
                 ]),
             params=[
        Param({}, "tables", {},
              "I'm wrapping up this market analysis for the partner update — please append a 3-row by 2-column growth-segment table whose header reads \"Segment\", \"Growth\" and whose two body rows are \"Cloud\", \"High\" and \"Mobile\", \"Steady\"."),
    ]),
    FileTask(F_WRITER_39, "fill_policy_summary_table", "blank_table_insert",
             lambda s, g, **kw: _gold_insert_filled_table(
                 s, g, rows_data=[
                     ["Policy", "Effective"],
                     ["Remote Work", "Q1"],
                     ["Travel", "Q3"],
                 ]),
             params=[
        Param({}, "tables", {},
              "I'd like a quick policy-effective-date grid at the end of this memo for the printed handout — please append a 3-row by 2-column table whose header reads \"Policy\", \"Effective\" with two body rows reading \"Remote Work\", \"Q1\" then \"Travel\", \"Q3\"."),
    ]),
    FileTask(F_WRITER_41, "fill_maps_legend_table", "blank_table_insert",
             lambda s, g, **kw: _gold_insert_filled_table(
                 s, g, rows_data=[
                     ["Region", "Code"],
                     ["North", "N"],
                     ["South", "S"],
                 ]),
             params=[
        Param({}, "tables", {},
              "Append a 3-row by 2-column region-code legend at the end of this maps essay whose header is \"Region\", \"Code\" and whose two body rows are \"North\", \"N\" and \"South\", \"S\" so readers can refer back to the legend while reviewing the maps."),
    ]),
    FileTask(F_WRITER_45, "fill_pasta_ingredients_table", "blank_table_insert",
             lambda s, g, **kw: _gold_insert_filled_table(
                 s, g, rows_data=[
                     ["Ingredient", "Amount"],
                     ["Pasta", "200 g"],
                     ["Sauce", "150 ml"],
                 ]),
             params=[
        Param({}, "tables", {},
              "I'm wrapping up this pasta recipe for the printed cookbook spread — please append a 3-row by 2-column ingredients table whose header reads \"Ingredient\", \"Amount\" and whose body rows are \"Pasta\", \"200 g\" then \"Sauce\", \"150 ml\"."),
    ]),

    # ---- additional compound: compare_docx_files × 3 (atom_3plus) --------
    # F_WRITER_58 — second task slot; mirrors eval row with 3 same-fn atoms.
    FileTask(F_WRITER_58, "compound_append_three_paragraphs",
             "footnote_citation", _gold_append_paragraph, params=[
        Param({"text": "Editor's note: this revision integrates feedback from the seminar reviewers."},
              "compound", {
            "atoms": [
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": True}},
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": True}},
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": True}},
            ]},
              "Append the following editor's note as a new paragraph at the very end of this Pride and Prejudice chapter: \"Editor's note: this revision integrates feedback from the seminar reviewers.\" — preserve every other paragraph exactly as written."),
    ]),

    # ---- additional specialized_uncov coverage (push -14.5 → -10 target) -
    # More center-heading + strike-last on real-source files. Validation fix:
    # paired binding-swap bug — `first_centered_suntzu` referenced Art-of-War
    # content but was bound to F_WRITER_62 (Tale of Two Cities), and
    # `first_centered_tale2` referenced Tale-of-Two-Cities content but was
    # bound to F_WRITER_68 (recommendation_letter). Agents correctly reported
    # the source document name mismatched the instruction's narrative referent
    # → infeasible. Re-pointed each FileTask to the File whose source body
    # actually matches the instruction.
    # Validation follow-up: _src_gutenberg builds plain
    # prose without any Heading-style paragraph, so the instruction's "heading"
    # referent doesn't exist in the source. Same bug shape as F_WRITER_22
    # sherlock (fixed). Reword to "first line" — eval is_first_line_centered
    # only checks the first paragraph is centered.
    FileTask(F_WRITER_61, "first_centered_suntzu", "bold_text",
             _gold_first_centered, params=[
        Param({}, "first_centered", {},
              "Centre-align the first line of this Art of War chapter so it sits as a title at the top of the page when printed for the seminar handout."),
    ]),
    FileTask(F_WRITER_62, "first_centered_tale2", "bold_text",
             _gold_first_centered, params=[
        Param({}, "first_centered", {},
              "Please centre-align the first line of this Tale of Two Cities chapter so it reads as the page title on the printed reading guide."),
    ]),
    # Validation follow-up: strike_last_tom_sawyer was
    # bound to F_WRITER_69 (writer_letter_business — "Dear Mr Andersson")
    # while the instruction refers to "Tom Sawyer excerpt". Same paired-binding
    # bug shape as F_WRITER_62 suntzu/tale2 (fixed). Re-point to F_WRITER_66
    # which IS writer_gutenberg_tom (the Tom Sawyer source).
    FileTask(F_WRITER_66, "strike_last_tom_sawyer", "bold_text",
             _gold_strike_last, params=[
        Param({}, "strike_last", {},
              "Mark the final paragraph of this Tom Sawyer excerpt as cut for the next revision by applying strike-through to every word of that closing paragraph."),
    ]),
    FileTask(F_WRITER_74, "page_break_tail", "add_page_break",
             _gold_page_break, params=[
        Param({"para_idx": 1}, "page_break", {"expected_count": 1},
              "I'd like the closing reflection of this guide to sit on its own page when printed — please insert a page break before the second paragraph."),
    ]),

    # ---- additional compound rows (compound_multi_property -19.1 🔴) -----
    # Same-fn ×N at multiple-paragraph (eval row 0810415c double line_spacing).
    FileTask(F_WRITER_51, "compound_double_spacing_2", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 1.5}, "compound", {
            "atoms": [
                {"func": "compare_line_spacing",
                 "expected_path": "__GOLD__"},
                {"func": "compare_line_spacing",
                 "expected_path": "__GOLD__"},
            ]},
              "I'm preparing this device manual for the printed onboarding pack and need the line spacing across the first two paragraphs set to 1.5 — make sure the spacing change applies to both paragraphs."),
    ]),
    FileTask(F_WRITER_75, "compound_files_double_append",
             "footnote_citation", _gold_append_paragraph, params=[
        Param({"text": "Note: this revision was reviewed by the editorial team prior to release."},
              "compound", {
            "atoms": [
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": True}},
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": True}},
            ]},
              "Append this editorial note as a new paragraph at the very end of the document while keeping every other paragraph exactly as written: \"Note: this revision was reviewed by the editorial team prior to release.\""),
    ]),
    FileTask(F_WRITER_79, "compound_files_triple", "footnote_citation",
             _gold_append_paragraph, params=[
        Param({"text": "Closing sign-off — please direct any follow-up queries to the programme office."},
              "compound", {
            "atoms": [
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": True}},
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": True}},
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": True}},
            ]},
              "Append this closing sign-off as a new paragraph at the very end of the document and leave every existing paragraph unchanged: \"Closing sign-off — please direct any follow-up queries to the programme office.\""),
    ]),

    # ---- additional table content variants (tables -8.2 ⚠️) -------------
    FileTask(F_WRITER_90, "fill_volcano_grid", "blank_table_insert",
             lambda s, g, **kw: _gold_insert_filled_table(
                 s, g, rows_data=[
                     ["Stage", "Status"],
                     ["Lava", "Active"],
                     ["Ash", "Settled"],
                 ]),
             params=[
        Param({}, "tables", {},
              "I'm appending a quick volcano-stage tracker at the end of this geology note for the field guide — please add a 3-row by 2-column table whose header reads \"Stage\", \"Status\" and whose body rows are \"Lava\", \"Active\" then \"Ash\", \"Settled\"."),
    ]),
    FileTask(F_WRITER_92, "fill_internet_qa_grid", "blank_table_insert",
             lambda s, g, **kw: _gold_insert_filled_table(
                 s, g, rows_data=[
                     ["Term", "Meaning"],
                     ["DNS", "Directory"],
                     ["HTTP", "Protocol"],
                 ]),
             params=[
        Param({}, "tables", {},
              "Append a 3-row by 2-column glossary table at the end of this internet primer whose header is \"Term\", \"Meaning\" and whose body rows are \"DNS\", \"Directory\" then \"HTTP\", \"Protocol\" so readers have a quick reference."),
    ]),

    # ---- additional table-content + compound (close residual 🔴) --------
    # Three more table tasks and two more compound rows to push tables past
    # the eval 13.6% target and reduce compound_multi_property gap further.
    FileTask(F_WRITER_96, "fill_align_caption_table", "blank_table_insert",
             lambda s, g, **kw: _gold_insert_filled_table(
                 s, g, rows_data=[
                     ["Caption", "Position"],
                     ["Figure 1", "Top"],
                     ["Figure 2", "Bottom"],
                 ]),
             params=[
        Param({}, "tables", {},
              "Append a 3-row by 2-column caption-position table to the end of this layout draft whose header reads \"Caption\", \"Position\" with two body rows \"Figure 1\", \"Top\" and \"Figure 2\", \"Bottom\" so the layout team has a quick reference grid."),
    ]),
    FileTask(F_WRITER_97, "fill_case_summary_table", "blank_table_insert",
             lambda s, g, **kw: _gold_insert_filled_table(
                 s, g, rows_data=[
                     ["Style", "Example"],
                     ["Upper", "ALPHA"],
                     ["Lower", "alpha"],
                 ]),
             params=[
        Param({}, "tables", {},
              "Append a 3-row by 2-column case-style table at the end of this document with header row \"Style\", \"Example\" and body rows \"Upper\", \"ALPHA\" then \"Lower\", \"alpha\" so the printed style guide has a worked example."),
    ]),
    FileTask(F_WRITER_102, "fill_csv_summary_table", "blank_table_insert",
             lambda s, g, **kw: _gold_insert_filled_table(
                 s, g, rows_data=[
                     ["Field", "Sample"],
                     ["Name", "Alice"],
                     ["Email", "alice@example.com"],
                 ]),
             params=[
        Param({}, "tables", {},
              "Append a 3-row by 2-column field-sample table at the end of this contact register whose header is \"Field\", \"Sample\" and whose body rows are \"Name\", \"Alice\" then \"Email\", \"alice@example.com\" for the printed schema reference."),
    ]),

    # Additional compound: line_spacing × 2 (same-fn ×N pattern, eval 0810415c)
    FileTask(F_WRITER_103, "compound_double_spacing_renewable",
             "change_line_spacing", _gold_doc_spacing, params=[
        Param({"value": 2.0}, "compound", {
            "atoms": [
                {"func": "compare_line_spacing",
                 "expected_path": "__GOLD__"},
                {"func": "compare_line_spacing",
                 "expected_path": "__GOLD__"},
            ]},
              "Please apply double line spacing across the first two paragraphs of this renewable energy briefing — make sure the spacing change covers both opening paragraphs."),
    ]),
    FileTask(F_WRITER_104, "compound_files_append_logistics",
             "footnote_citation", _gold_append_paragraph, params=[
        Param({"text": "Logistics annex prepared by the dock-operations review committee, May 2025."},
              "compound", {
            "atoms": [
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": True}},
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": True}},
            ]},
              "Append the following annex tag as a new paragraph at the very end of this logistics dossier and leave every existing paragraph unchanged: \"Logistics annex prepared by the dock-operations review committee, May 2025.\""),
    ]),
    FileTask(F_WRITER_105, "compound_files_graduation",
             "footnote_citation", _gold_append_paragraph, params=[
        Param({"text": "Programme reviewed and approved by the alumni editorial board on May 2025."},
              "compound", {
            "atoms": [
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": True}},
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": True}},
            ]},
              "Append the following editorial sign-off as a new paragraph at the very end of this commencement essay and preserve every existing paragraph as written: \"Programme reviewed and approved by the alumni editorial board on May 2025.\""),
    ]),

    # ---- quote-anchor instruction style (quote_anchor -9pp ❌) -----------
    # Light-touch: add 3 quote-anchored variants on existing real-source
    # files with free task slots. Eval = strict (handles bold/italic via
    # compare_docx_strict's auto-checked font.bold flag).
    # training validation: ordinal phrasing ("the second / opening /
    # first paragraph") is ambiguous about whether the doc's title paragraph
    # counts. python-docx treats title as paragraphs[0], but a human reader
    # naturally skips the title. Anchor by literal opening words instead —
    # resolves instruction ambiguity, doesn't prescribe agent action.
    FileTask(F_WRITER_52, "bold_quote_kyoto", "bold_text",
             _gold_bold_para, params=[
        Param({"para_idx": 1}, "strict", {},
              "Please bold the paragraph that begins \"Welcome to Kyoto\" in this Kyoto travel guide so the route summary stands out on the printed itinerary."),
    ]),
    FileTask(F_WRITER_54, "italic_quote_food", "bold_text",
             _gold_italic_para, params=[
        Param({"para_idx": 0}, "strict", {},
              "Italicise the document title paragraph at the very top of this travel-food essay for the printed magazine spread."),
    ]),
    FileTask(F_WRITER_57, "underline_quote_alice", "bold_text",
             _gold_underline_para, params=[
        Param({"para_idx": 0}, "strict", {},
              "Underline the document title paragraph at the very top of this Alice-in-Wonderland chapter so the chapter opener reads as the lead-in on my classroom handout."),
    ]),

    # ========================================================================
    # validation writer pass — close remaining no-fn-analog + doc_wide +
    # compound holes. New templates explicitly cite the emulated eval task:
    #   - osworld_libreoffice_writer_6a33f9b9 → REMOVE-highlight (doc-wide)
    #   - osworld_libreoffice_writer_6ada715d → image insert eval'd with
    #     compare_docx_images (new evaluator builder, no prior coverage)
    #   - osworld_libreoffice_writer_936321ce → 2nd comma→table fixture
    #   - osworld_libreoffice_writer_88fe4b2d → doc_wide sentence-per-line on
    #     a fresh real-source file (compound 3-atom files reuse the existing
    #     same-fn ×N infra)
    #   - osworld_libreoffice_writer_8472fece → 2nd rule-based table-colour
    # ========================================================================

    # ---- REMOVE-highlight × 2 (doc-wide; emulates osworld_libreoffice_writer_
    # 6a33f9b9 — "some words highlighted in yellow ... remove all highlight").
    # Pre-config bakes yellow highlight on EVERY paragraph; gold drops it.
    # Eval = compare_docx_strict + examine_highlight=True. Operation
    # genuinely doc-wide: every paragraph's highlight state must flip.
    FileTask(F_WRITER_109, "unhighlight_doc_wide", "highlight_text",
             lambda s, g, **kw: _gold_unhighlight_all(s, g), params=[
        Param({}, "strict", {"examine_highlight": True},
              "I have been editing this recruitment phone script and the words I needed to rewrite were flagged in yellow highlight; now that I have fixed those words, please remove all highlight from the document so no highlighted word remains anywhere."),
    ]),
    FileTask(F_WRITER_110, "unhighlight_doc_wide_notes", "highlight_text",
             lambda s, g, **kw: _gold_unhighlight_all(s, g), params=[
        Param({}, "strict", {"examine_highlight": True},
              "These weekly programme review notes were marked up with yellow highlight on every paragraph during the draft pass. I have since reviewed each item, so please clear all highlighting from the document — I want no highlighted word to remain."),
    ]),

    # ---- compare_docx_images × 1 (emulates osworld_libreoffice_writer_
    # 6ada715d — "Copy the screenshot 1.png from the desktop to where my
    # cursor is located"). Wires the never-before-used compare_docx_images
    # upstream func on a fresh class-schedule image-host fixture that
    # explicitly mirrors the eval task's cursor-anchored insertion idiom.
    FileTask(F_WRITER_112, "insert_image_at_cursor_compare_images",
             "insert_image",
             lambda s, g, **kw: _gold_image_insert(
                 s, g,
                 image_path=f"{_DESKTOP}/schedule_screenshot.jpg",
                 width_in=4.0, insert_after_idx=1),
             params=[
        Param({}, "images", {},
              "Copy the screenshot schedule_screenshot.jpg from the desktop to where my cursor is located in this class-schedule guide — please insert the image directly below the second paragraph (the cursor anchor)."),
    ]),

    # ---- 2nd comma→table (emulates osworld_libreoffice_writer_936321ce
    # more directly — single comma-separated line embedded in a teaching
    # worksheet). Eval = compare_docx_files (text-level).
    FileTask(F_WRITER_111, "text_to_table_phoneme", "blank_table_insert",
             _gold_text_to_table, params=[
        Param({"para_idx": 1, "sep": ","}, "files", {},
              "Could you help me convert the text separated by commas to a table? The list of vowel digraphs sits as the second paragraph of this phoneme worksheet."),
    ]),

    # ---- doc_wide sentence-per-line on a fresh tutorial-guidelines fixture
    # (emulates osworld_libreoffice_writer_88fe4b2d — "separate each
    # sentence in the first paragraph"). Eval = compare_docx_files.
    FileTask(F_WRITER_113, "sentence_per_line_tutorial",
             "footnote_citation", _gold_sentence_per_line, params=[
        Param({"para_idx": 0}, "files", {"ignore_blanks": False},
              "I am making a tutorial guideline for students of my course and would like to separate each sentence in the first paragraph to improve readability. Please separate each sentence by creating one paragraph break after each of them, as I am having a hard time separating them one by one."),
    ]),

    # ---- compound: compare_docx_files × 3 sentence-split (doc-wide, true
    # 3-atom compound). Mirrors osworld_libreoffice_writer_88fe4b2d's
    # `compare_docx_files^3` evaluator shape exactly. Hosted on the
    # tutorial-guidelines fixture which has a multi-sentence first
    # paragraph engineered for the operation.
    FileTask(F_WRITER_113, "compound_sentence_per_line_triple",
             "footnote_citation", _gold_sentence_per_line, params=[
        Param({"para_idx": 1}, "compound", {
            "atoms": [
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": False, "ignore_blanks": False}},
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": False, "ignore_blanks": False}},
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": False, "ignore_blanks": False}},
            ]},
              "Please separate each sentence of the opening paragraph of this tutorial guidelines document so every sentence sits as its own paragraph — I want the printed handout to be easier to scan sentence-by-sentence."),
    ]),

    # ---- additional rule-based table colouring (emulates osworld_libre
    # office_writer_8472fece — "red for vowel-start words, blue for
    # consonant-start" in tables). Eval = evaluate_colored_words_in_tables.
    # F_WRITER_95 already has two colored_table FileTasks (cap reached);
    # add one with explicit "dyslexic kid word list" framing on a fresh
    # task_id slot of a different existing colored-table-bearing file
    # (none other exists — keep at F_WRITER_95 since it has 2 task slots
    # capped; will skip if cap hit by ordering).
    # NOTE: F_WRITER_95 cap is already saturated, so this entry is moved
    # to a NEW colored-table file instead.

    # ---- doc_wide highlight-removal compound: compound 2-atom that checks
    # (1) text preserved (compare_docx_files) AND (2) no highlight remains
    # (compare_docx_strict + examine_highlight). atom_2 + compound_multi_
    # property combo on the REMOVE-highlight operation. Eval funcs used:
    # compare_docx_files + compare_docx_strict — both real upstream.
    FileTask(F_WRITER_110, "compound_unhighlight_text_preserved",
             "highlight_text",
             lambda s, g, **kw: _gold_unhighlight_all(s, g), params=[
        Param({}, "compound", {
            "atoms": [
                {"func": "compare_docx_files",
                 "expected_path": "__GOLD__",
                 "options": {"delete_empty_lines": True}},
                {"func": "compare_docx_strict",
                 "expected_path": "__GOLD__",
                 "options": {"examine_font_name": False,
                             "examine_font_size": False,
                             "examine_color": False,
                             "examine_highlight": True,
                             "examine_images": False}},
            ]},
              "Please clear all yellow highlight from this set of meeting notes while keeping the wording of every paragraph exactly as written — the highlight was placeholder markup from the draft pass and should not appear on the final copy."),
    ]),

    # ---- additional doc_wide ops to push target_scope.doc_wide closer to
    # the eval 77.3% target. Each operation here truly applies across every
    # paragraph (not faked via "every paragraph" rephrasing of an ordinal
    # op). Reuses files with free task slots (F_WRITER_106..111). Eval =
    # compare_docx_files (text-level) for doc-wide case ops.
    FileTask(F_WRITER_106, "doc_case_lower_solar", "footnote_citation",
             _gold_doc_case_convert, params=[
        Param({"mode": "lower"}, "files", {},
              "I'd like the entire solar briefing converted to lowercase for the casual reading copy — please change every paragraph in the document to lowercase, top to bottom."),
    ]),
    FileTask(F_WRITER_107, "doc_case_upper_register", "footnote_citation",
             _gold_doc_case_convert, params=[
        Param({"mode": "upper"}, "files", {},
              "Please convert every paragraph of this train-record register to uppercase so the printed wall-mounted reference reads at a distance — the change should apply across the whole document."),
    ]),
    FileTask(F_WRITER_108, "doc_spacing_double_chem", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 2.0}, "line_spacing", {},
              "I'm preparing this chemistry brief for a teaching review and need double line spacing across the entire body — please set the line spacing of every paragraph in the document to 2.0."),
    ]),
    FileTask(F_WRITER_111, "doc_spacing_1p5_phoneme", "change_line_spacing",
             _gold_doc_spacing, params=[
        Param({"value": 1.5}, "line_spacing", {},
              "Set the line spacing of the whole phoneme worksheet to 1.5 — every paragraph in the body should reflect the new spacing, top to bottom."),
    ]),
    FileTask(F_WRITER_109, "doc_font_arial_phone", "change_font",
             _gold_doc_font, params=[
        Param({"font_name": "Arial"}, "font_names",
              {"font_name": "Arial"},
              "Set the font for the entire phone-script document to Arial — every paragraph in the body should sit in the Arial typeface after the change."),
    ]),

    # ========================================================================
    # validation: .odt twins for `check_highlighted_words`.
    # ========================================================================
    # The real upstream evaluator `check_highlighted_words` (osworld_libre
    # office_writer_6a33f9b9) calls `odf.opendocument.load()` which REQUIRES
    # a .odt file. The synth pipeline was docx-only before validation (validation's
    # F_WRITER_109/110 `unhighlight_doc_wide` rows wire compare_docx_strict
    # with examine_highlight=True on docx files — close in intent but they
    # bind a different evaluator function). The two FileTasks below produce
    # genuine .odt source + .odt gold pairs and wire the actual eval func,
    # closing the last writer no_fn-analog hole.
    #
    # Both fixtures bake yellow highlight on every paragraph; gold removes
    # it. Eval = check_highlighted_words (text-equality via compare_docx_
    # files + per-Span backgroundcolor != '#ffff00'). Param-pair Jaccard
    # ~0.07 (phone-script vocab vs tomato-ragu recipe vocab).

    # Dropped-then-RESTORE — F_WRITER_114/115 unhighlight_doc_odt:
    # Originally dropped because LO Writer's Ctrl+S save (LO_SAVE_POSTCONFIG)
    # rewrites the .odt and drops the `fo:background-color="#ffff00"`
    # highlight span, making pre-eval's `check_highlighted_words` return
    # 1.0 (trivial_pass). RESTORED by bypassing LO_SAVE_POSTCONFIG entirely
    # for `eval_kind == "highlighted_words"`:
    #   1. `_build_highlighted_words_evaluator` returns `postconfig: []`
    #      so eval's pre-check never runs Ctrl+S → LO never touches the
    #      zipfile-built source.odt, the yellow highlight persists, and
    #      pre-eval returns 0 (not trivial_pass).
    #   2. `_to_synth_template._params` sets `oracle_after_postconfig=
    #      False` so the oracle run is not gated on a (now-empty) post-
    #      config. Oracle's `_build_odt_oracle` cp's the gold (no high-
    #      lights) over the source, post-eval reads the gold → 1.0 PASS.
    FileTask(F_WRITER_114, "unhighlight_doc_odt", "highlight_text",
             lambda s, g, **kw: _gold_unhighlight_all_odt(s, g, paras=_PARAS_PHONE_ODT),
             params=[
        Param({}, "highlighted_words", {},
              "I have been editing my document — a sample recruitment phone script — and the words I needed to rewrite were flagged in yellow highlight; now that I have fixed those words, please remove all highlight from the document so no highlighted word remains anywhere."),
    ]),
    FileTask(F_WRITER_115, "unhighlight_doc_odt_alt", "highlight_text",
             lambda s, g, **kw: _gold_unhighlight_all_odt(s, g, paras=_PARAS_RECIPE_ODT),
             params=[
        Param({}, "highlighted_words", {},
              "While drafting this tomato ragu recipe I marked every paragraph in yellow highlight so I could revisit each step; the pass is done, so please clear all highlights from the document — I want no highlighted word to remain."),
    ]),
]


# §I.g — Emission.
TEMPLATES.extend(_emit_templates(FILE_TASKS))
