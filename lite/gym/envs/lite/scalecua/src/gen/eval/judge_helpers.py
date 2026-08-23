"""Shared read-helpers, patched into the judge overlays at import.

These are read-side fixes for the upstream ScaleCUA `judge_functions`, which
proxy-check the *on-disk artifact* (static `cell.fill`, a formula string, a
direct-run font, a raw string) instead of the achieved end-state, so genuine
agent successes score reward=0 (false negatives that poison the RL gradient).

Rather than hand-edit the pulled `.cache/.../judge_functions/*.py` (clobbered on
every re-pull), we follow the same **harden-at-import** pattern as
`normalize_postconfig` / `_rewrite_urls`: `patch_judge_functions()` injects this
helper library into every materialized getter/metric module during
`dataset.import_all()` (before catalog validation), so the regenerated /
patched checkers resolve the helpers by name and the fix survives every
`.cache` pull.

The helper functions are defined here as ordinary Python so they can be unit
tested directly (hand-built CF xlsx, formula cell, `w:shd` docx). The exact same
source text is what gets injected (built via `inspect.getsource`), so what the
tests exercise is what runs inside the overlays.

Run / test:
    uv run pytest tests/gym/envs/lite/scalecua/test_scalecua_a1_judge_helpers.py
"""

from __future__ import annotations

import inspect
import re
from pathlib import Path

# --- Calc / openpyxl -----------------------------------------------------------

def scalecua_xlsx_conditional_formats(path, sheet=None):
    """Surface conditional-formatting rules the static-fill proxy misses.

    Returns a list of {"range", "type", "operator", "formula", "fill"} for every
    CF rule on the sheet. A checker that only reads ``cell.fill`` sees a blank
    cell when a CF rule paints it (e.g. "colour weekends red"); this exposes the
    rule + its dxf fill so the CF can be asserted directly.
    """
    import openpyxl

    wb = openpyxl.load_workbook(path, data_only=False)
    ws = wb[sheet] if sheet else wb.active
    out = []
    # openpyxl exposes conditional_formatting as a mapping range -> [rules].
    for cf_range in ws.conditional_formatting:
        try:
            rules = ws.conditional_formatting[cf_range]
        except Exception:
            rules = getattr(cf_range, "rules", []) or []
        rng = getattr(cf_range, "sqref", cf_range)
        for rule in rules:
            out.append({
                "range": str(rng),
                "type": getattr(rule, "type", None),
                "operator": getattr(rule, "operator", None),
                "formula": list(getattr(rule, "formula", []) or []),
                "text": getattr(rule, "text", None),
                "fill": _scalecua_dxf_fill_hex(getattr(rule, "dxf", None)),
            })
    wb.close()
    return out


def _scalecua_dxf_fill_hex(dxf):
    if dxf is None:
        return None
    fill = getattr(dxf, "fill", None)
    if fill is None:
        return None
    for attr in ("bgColor", "fgColor", "start_color", "end_color"):
        color = getattr(fill, attr, None)
        rgb = getattr(color, "rgb", None)
        if isinstance(rgb, str) and rgb and rgb != "00000000":
            return rgb.upper()
    return None


def scalecua_xlsx_effective_fill(path, sheet, coord, value=None):
    """The *effective* fill of a cell: a matching CF rule wins over static fill.

    ``value`` (the cell's computed value) is used to evaluate ``cellIs`` /
    ``containsText`` rules; when it is None the computed value is loaded. Rules
    that cannot be evaluated statically (arbitrary ``expression`` formulas) are
    skipped here — use ``scalecua_xlsx_conditional_formats`` to assert those.
    """
    import openpyxl
    from openpyxl.utils import coordinate_to_tuple
    from openpyxl.utils.cell import range_boundaries

    if value is None:
        wb_v = openpyxl.load_workbook(path, data_only=True)
        ws_v = wb_v[sheet] if sheet else wb_v.active
        value = ws_v[coord].value
        wb_v.close()

    wb = openpyxl.load_workbook(path, data_only=False)
    ws = wb[sheet] if sheet else wb.active
    row, col = coordinate_to_tuple(coord)
    static = _scalecua_cell_fill_hex(ws[coord])
    effective = static
    # Excel renders the fill of the matching CF rule with the LOWEST priority
    # NUMBER (highest precedence). A naive `sorted(..., priority asc)` + last-wins
    # loop would keep the LAST-iterated = highest-number = LOWEST-precedence rule's
    # fill -> inverted color when >=2 overlapping rules match (confirmed by a
    # 2-rule openpyxl repro). Priority is a SHEET-GLOBAL ordering, so track the min
    # across ALL matching rules (every cf_range) — this also honors cross-range
    # precedence a per-block sort never would. Single-rule cells (the norm) and
    # no-match cells are unchanged -> no new FP.
    best_priority = None
    for cf_range in ws.conditional_formatting:
        try:
            rules = ws.conditional_formatting[cf_range]
        except Exception:
            rules = getattr(cf_range, "rules", []) or []
        sqref = str(getattr(cf_range, "sqref", cf_range))
        if not _scalecua_coord_in_sqref(row, col, sqref, range_boundaries):
            continue
        for rule in rules:
            if _scalecua_cf_rule_matches(rule, value):
                fill = _scalecua_dxf_fill_hex(getattr(rule, "dxf", None))
                if fill:
                    prio = getattr(rule, "priority", None)
                    prio = prio if prio is not None else 10**9
                    if best_priority is None or prio < best_priority:
                        best_priority = prio
                        effective = fill
    wb.close()
    return effective


def _scalecua_cell_fill_hex(cell):
    fill = getattr(cell, "fill", None)
    color = getattr(fill, "fgColor", None) or getattr(fill, "start_color", None)
    rgb = getattr(color, "rgb", None)
    if isinstance(rgb, str) and rgb and rgb != "00000000":
        return rgb.upper()
    return None


def _scalecua_coord_in_sqref(row, col, sqref, range_boundaries):
    for part in str(sqref).replace(",", " ").split():
        if not part:
            continue
        try:
            min_c, min_r, max_c, max_r = range_boundaries(part)
        except Exception:
            continue
        max_c = max_c or col
        max_r = max_r or row
        if min_r <= row <= max_r and min_c <= col <= max_c:
            return True
    return False


def _scalecua_cf_rule_matches(rule, value):
    rtype = getattr(rule, "type", None)
    formula = list(getattr(rule, "formula", []) or [])
    if rtype == "cellIs":
        op = getattr(rule, "operator", None)
        nums = []
        for f in formula:
            try:
                nums.append(float(str(f).strip().lstrip("=")))
            except (TypeError, ValueError):
                nums.append(str(f))
        try:
            v = float(value)
        except (TypeError, ValueError):
            v = value
        if not nums:
            return False
        a = nums[0]
        b = nums[1] if len(nums) > 1 else None
        try:
            if op in ("equal", "equals"):
                return v == a
            if op == "notEqual":
                return v != a
            if op == "greaterThan":
                return v > a
            if op == "greaterThanOrEqual":
                return v >= a
            if op == "lessThan":
                return v < a
            if op == "lessThanOrEqual":
                return v <= a
            if op == "between" and b is not None:
                return a <= v <= b
            if op == "notBetween" and b is not None:
                return not (a <= v <= b)
        except TypeError:
            return False
    if rtype in ("containsText", "notContainsText"):
        text = getattr(rule, "text", None)
        has = bool(text) and text in str(value or "")
        return has if rtype == "containsText" else not has
    return False


def scalecua_xlsx_cell_formula_and_value(path, sheet, coord):
    """Return both the stored formula and the computed value of a cell.

    Fixes the formula-vs-value proxy: a checker expecting ``=AVERAGE(...)`` fails
    when the agent typed the equivalent literal (and vice-versa). Callers can
    then accept either the formula or the value.
    """
    import openpyxl

    wb_f = openpyxl.load_workbook(path, data_only=False)
    ws_f = wb_f[sheet] if sheet else wb_f.active
    formula = ws_f[coord].value
    wb_f.close()
    wb_v = openpyxl.load_workbook(path, data_only=True)
    ws_v = wb_v[sheet] if sheet else wb_v.active
    value = ws_v[coord].value
    wb_v.close()
    return {"formula": formula, "value": value}


def scalecua_formula_equivalent(a, b):
    """Whitespace/case/``$``/leading-``=`` insensitive formula comparison."""
    def norm(x):
        s = str(x if x is not None else "").strip()
        if s.startswith("="):
            s = s[1:]
        return s.replace("$", "").replace(" ", "").upper()
    return norm(a) == norm(b)


def scalecua_xlsx_freeze_panes(path, sheet=None):
    """The freeze-panes anchor (e.g. ``"A2"``) or None — recognises frozen rows/cols."""
    import openpyxl

    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb[sheet] if sheet else wb.active
    fp = ws.freeze_panes
    wb.close()
    return str(fp) if fp else None


def scalecua_xlsx_cell_date(path, sheet, coord):
    """Return {"value", "number_format", "is_date"} for date-format checks."""
    import openpyxl
    from openpyxl.utils.datetime import from_excel  # noqa: F401  (import parity)

    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb[sheet] if sheet else wb.active
    cell = ws[coord]
    out = {
        "value": cell.value,
        "number_format": cell.number_format,
        "is_date": cell.is_date,
    }
    wb.close()
    return out


def scalecua_date_equal(a, b):
    """Compare a date-string against a datetime/date (day resolution)."""
    from datetime import date, datetime

    def to_date(x):
        if isinstance(x, datetime):
            return x.date()
        if isinstance(x, date):
            return x
        s = str(x).strip()
        for fmt in (
            "%Y-%m-%d", "%m/%d/%Y", "%d/%m/%Y", "%Y/%m/%d",
            "%d-%m-%Y", "%m-%d-%Y", "%Y-%m-%d %H:%M:%S", "%d.%m.%Y",
        ):
            try:
                return datetime.strptime(s, fmt).date()
            except ValueError:
                continue
        return None
    da, db = to_date(a), to_date(b)
    return da is not None and da == db


# --- Writer / python-docx --------------------------------------------------

_SCALECUA_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def scalecua_docx_paragraph_shading(path, para_index):
    """The ``w:shd`` fill of a paragraph (or its runs) — highlight via shading.

    The upstream checker reads only the run ``highlight`` property and misses a
    highlight applied as paragraph/run shading (``<w:shd w:fill=.../>``).
    Returns an upper-case hex fill or None.
    """
    import docx

    doc = docx.Document(path)
    paras = doc.paragraphs
    if para_index < 0 or para_index >= len(paras):
        return None
    para = paras[para_index]
    fill = _scalecua_shd_fill(para._p.find(f"{_SCALECUA_W_NS}pPr"))
    if fill:
        return fill
    for run in para.runs:
        fill = _scalecua_shd_fill(run._r.find(f"{_SCALECUA_W_NS}rPr"))
        if fill:
            return fill
    return None


def _scalecua_shd_fill(pr_element):
    if pr_element is None:
        return None
    shd = pr_element.find(f"{_SCALECUA_W_NS}shd")
    if shd is None:
        return None
    fill = shd.get(f"{_SCALECUA_W_NS}fill")
    if fill and fill.lower() not in ("auto", "ffffff", "000000", ""):
        return fill.upper()
    return None


def scalecua_docx_effective_run_format(path, para_index, run_index=0):
    """Resolve run formatting through the style hierarchy.

    Bold/italic/font/size applied at the *style* level (not the direct run) are
    invisible to a checker that reads ``run.bold`` only. Walk run -> paragraph
    style -> based-on styles for the effective value.
    """
    import docx

    doc = docx.Document(path)
    paras = doc.paragraphs
    if para_index < 0 or para_index >= len(paras):
        return {}
    para = paras[para_index]
    runs = para.runs
    run = runs[run_index] if 0 <= run_index < len(runs) else None

    def resolve(attr):
        if run is not None:
            direct = getattr(run.font, attr, None)
            if direct is not None:
                return direct
        style = para.style
        seen = set()
        while style is not None and id(style) not in seen:
            seen.add(id(style))
            val = getattr(getattr(style, "font", None), attr, None)
            if val is not None:
                return val
            style = getattr(style, "base_style", None)
        return None

    return {
        "bold": resolve("bold"),
        "italic": resolve("italic"),
        "underline": resolve("underline"),
        "name": resolve("name"),
        "size": resolve("size"),
    }


def scalecua_bake_effective_run_bold(doc):
    """Bake style-inherited BOLD onto each run's direct formatting, in place, so ANY
    getter that reads ``run.bold`` / ``run.font.bold`` / ``r.font.bold`` RAW sees the
    effective value. python-docx returns None when bold comes from the paragraph/
    character STYLE rather than a direct run override; walk run.font -> paragraph
    style -> based-on chain for the effective value. Only fills runs whose DIRECT
    bold is None (never overrides an explicit direct value) -> already-correct reads
    are unchanged (no new false positives). This GENERALIZES the effective-bold fix
    to every ``get_docx_para_bold__*`` getter form (bold-ratio / all()-comprehension /
    run.bold / r.font.bold) via one uniform inject after the doc open, replacing the
    original per-hash source .replace() that covered only 1 of 4 sibling getters."""
    for para in getattr(doc, "paragraphs", []):
        for run in para.runs:
            if run.font.bold is not None:
                continue
            style = para.style
            seen = set()
            while style is not None and id(style) not in seen:
                seen.add(id(style))
                val = getattr(getattr(style, "font", None), "bold", None)
                if val is not None:
                    run.font.bold = val
                    break
                style = getattr(style, "base_style", None)


def scalecua_docx_body_paragraph_index(path, target_text=None, skip_headings=True):
    """Map a *content* paragraph to its real index, skipping title/headings.

    Upstream checks address "the Nth paragraph" but count from paragraph 0,
    which is often the Title/Heading; the agent counts body paragraphs. Returns
    the real index of ``target_text`` (or the list of body indices when
    ``target_text`` is None).
    """
    import docx

    doc = docx.Document(path)
    body = []
    for i, para in enumerate(doc.paragraphs):
        style_name = (getattr(para.style, "name", "") or "").lower()
        if skip_headings and (
            style_name.startswith("heading") or style_name == "title"
        ):
            continue
        if not para.text.strip():
            continue
        body.append(i)
        if target_text is not None and para.text.strip() == str(target_text).strip():
            return i
    return None if target_text is not None else body


# --- String normalize -----------------------------------------------------------

def scalecua_normalize_compare_string(s):
    """Normalize before comparing: strip newline / leading ``$`` / case / en-dash."""
    if s is None:
        return ""
    text = str(s).replace("\r", "").strip()
    if text.startswith("$"):
        text = text[1:]
    text = text.replace("–", "-").replace("—", "-")  # en / em dash
    text = " ".join(text.split())
    return text.casefold()


# --- Unpacked chrome extension path ---------------------------------------------

def scalecua_find_unpacked_extension_manifest(root, max_depth=6):
    """Depth-tolerant search for an unpacked extension's ``manifest.json``.

    The upstream check probes a fixed depth; agents unpack extensions at varying
    depths. Returns the shallowest matching manifest path, or None.
    """
    import os

    root = os.fspath(root)
    best = None
    best_depth = None
    for dirpath, _dirs, files in os.walk(root):
        depth = dirpath[len(root):].count(os.sep)
        if depth > max_depth:
            continue
        if "manifest.json" in files:
            if best_depth is None or depth < best_depth:
                best = os.path.join(dirpath, "manifest.json")
                best_depth = depth
    return best


# --- VS Code JSON: default / scope / argv comments ------------------------------

def scalecua_parse_jsonc(text):
    """Parse comment-laden / trailing-comma JSON (VS Code ``argv.json`` etc.)."""
    import json
    import re

    if isinstance(text, (bytes, bytearray)):
        text = text.decode("utf-8", errors="replace")
    text = str(text)
    # Strip block and line comments outside strings (best-effort, VS Code style).
    text = re.sub(r"/\*.*?\*/", "", text, flags=re.DOTALL)
    text = re.sub(r"(^|[^:])//[^\n\r]*", lambda m: m.group(1), text)
    # Drop trailing commas.
    text = re.sub(r",(\s*[}\]])", r"\1", text)
    try:
        return json.loads(text)
    except Exception:
        return {}


def scalecua_vscode_setting_effective(user, workspace, key, default=None):
    """Effective VS Code setting: workspace overrides user overrides default.

    Fixes the false negative where the correct value equals the VS Code default
    and is therefore never written to ``settings.json``.
    """
    for scope in (workspace, user):
        if isinstance(scope, dict) and key in scope:
            return scope[key]
    return default


def scalecua_keybinding_matches(entries, key, command):
    """Membership test for a keybinding ignoring the UI-injected ``when`` clause.

    Treats ``editorFocus`` and ``editorTextFocus`` as equivalent and ignores
    ``when`` entirely, so an agent-created binding still matches.
    """
    def norm_key(k):
        # Normalize modifier order *within* each chord; preserve chord sequence
        # order (``ctrl+k ctrl+c`` != ``ctrl+c ctrl+k``).
        chords = (k or "").lower().split()
        return " ".join("+".join(sorted(c.split("+"))) for c in chords)

    want_key = norm_key(key)
    for entry in entries or []:
        if not isinstance(entry, dict):
            continue
        if entry.get("command") != command:
            continue
        if norm_key(entry.get("key")) == want_key:
            return True
    return False


# --- import-pipeline injection (harden-at-import) -----------------------------

HELPER_MARKER = "# __SCALECUA_A1_READ_HELPERS__ (issue #154 G1-7/G1-8)"

_HELPER_FUNCS = [
    scalecua_xlsx_conditional_formats,
    _scalecua_dxf_fill_hex,
    scalecua_xlsx_effective_fill,
    _scalecua_cell_fill_hex,
    _scalecua_coord_in_sqref,
    _scalecua_cf_rule_matches,
    scalecua_xlsx_cell_formula_and_value,
    scalecua_formula_equivalent,
    scalecua_xlsx_freeze_panes,
    scalecua_xlsx_cell_date,
    scalecua_date_equal,
    scalecua_docx_paragraph_shading,
    _scalecua_shd_fill,
    scalecua_docx_effective_run_format,
    scalecua_bake_effective_run_bold,
    scalecua_docx_body_paragraph_index,
    scalecua_normalize_compare_string,
    scalecua_find_unpacked_extension_manifest,
    scalecua_parse_jsonc,
    scalecua_vscode_setting_effective,
    scalecua_keybinding_matches,
]


def helper_source() -> str:
    """The self-contained source block injected into each overlay module."""
    parts = [
        HELPER_MARKER,
        "# Injected by lite.scalecua import pipeline; re-applied on every "
        ".cache pull.",
        f'_SCALECUA_W_NS = "{_SCALECUA_W_NS}"',
    ]
    for fn in _HELPER_FUNCS:
        parts.append(inspect.getsource(fn))
    return "\n\n".join(parts) + "\n"


def _overlay_module_files(judges_root: Path):
    for split in ("train", "rl"):
        base = judges_root / split
        if not base.is_dir():
            continue
        for name in ("getters.py", "metrics.py"):
            path = base / name
            if path.is_file():
                yield path
        for sub in ("verigen_getters", "verigen_metrics"):
            subdir = base / sub
            if not subdir.is_dir():
                continue
            for path in sorted(subdir.glob("*.py")):
                if path.name == "__init__.py":
                    continue
                yield path


_CF_COLOR_GETTER_IDS = (
    "03a99dd86c6d30983d2467bc1177489c_qw35sft2_df3a9ab2",
    "1a5d60ee39c11d25031d40117b38947f_qw35sft2_ceaed29a",
    "cc3c6531bd3eb1e9b2189f82dea73399_qw35sft2_64440f9c",
)

_VSCODE_KEYBINDING_IDS = (
    "ed42af8e24d428ca40b99f135c68274f_qw35sft2",
    "70885a60e67c4def20a9d0e3284c1793_qw35sft2",
)


def _rewrite_judge_helpers(text: str) -> tuple[str, bool]:
    original = text
    for identifier in _CF_COLOR_GETTER_IDS:
        fn = f"get_xlsx_cell_colors__{identifier}"
        current = _function_source(text, fn)
        if not current or "scalecua_xlsx_effective_fill" in current:
            continue
        rewritten = current.replace(
            """            cell = ws[cell_ref]
            fill = cell.fill
            if fill and fill.fgColor and (fill.fgColor.type == 'rgb'):
                color = fill.fgColor.rgb
            else:
                color = '00000000'
            result[cell_ref] = color
""",
            """            cell = ws[cell_ref]
            color = scalecua_xlsx_effective_fill(
                tmp_path, ws.title, cell_ref, cell.value
            ) or '00000000'
            result[cell_ref] = color
""",
        )
        if rewritten != current:
            text = _replace_function_source(text, fn, rewritten)[0]
    for identifier in _VSCODE_KEYBINDING_IDS:
        getter_prefix = f"get_vscode_keybindings__{identifier}"
        metric_prefix = f"check_vscode_keybinding__{identifier}"
        text = _rewrite_vscode_keybinding_getter(text, getter_prefix)
        text = _rewrite_vscode_keybinding_metric(text, metric_prefix)
    text = _rewrite_docx_para_bold_getters(text)
    text = _rewrite_sum_formula_metrics(text)
    return text, text != original


def _function_source(text: str, function_name: str) -> str:
    start = text.find(f"def {function_name}(")
    if start < 0:
        return ""
    next_start = text.find("\ndef ", start + 1)
    return text[start:] if next_start < 0 else text[start:next_start]


def _replace_function_source(text: str, function_name: str, new_source: str) -> tuple[str, bool]:
    start = text.find(f"def {function_name}(")
    if start < 0:
        return text, False
    next_start = text.find("\ndef ", start + 1)
    if next_start < 0:
        return text[:start] + new_source.rstrip() + "\n", True
    return text[:start] + new_source.rstrip() + "\n" + text[next_start + 1:], True


def _find_function_by_prefix(text: str, prefix: str) -> str | None:
    marker = f"def {prefix}"
    start = text.find(marker)
    if start < 0:
        return None
    name_end = text.find("(", start)
    if name_end < 0:
        return None
    return text[start + 4:name_end]


def _rewrite_vscode_keybinding_getter(text: str, prefix: str) -> str:
    name = _find_function_by_prefix(text, prefix)
    if not name:
        return text
    current = _function_source(text, name)
    if "scalecua_parse_jsonc" in current:
        return text
    new_source = f'''def {name}(env, config: dict):
    """Read VS Code keybindings.json and return parsed entries."""
    keybindings_path = '/home/user/.config/Code/User/keybindings.json'
    file_bytes = env.controller.get_file(keybindings_path)
    if not file_bytes:
        return {{'error': 'keybindings.json not found', 'entries': []}}
    try:
        entries = scalecua_parse_jsonc(file_bytes.decode('utf-8'))
        return {{'entries': entries if isinstance(entries, list) else []}}
    except Exception as e:
        return {{'error': str(e), 'entries': []}}
'''
    return _replace_function_source(text, name, new_source)[0]


def _rewrite_vscode_keybinding_metric(text: str, prefix: str) -> str:
    name = _find_function_by_prefix(text, prefix)
    if not name:
        return text
    current = _function_source(text, name)
    if "scalecua_keybinding_matches" in current:
        return text
    new_source = f'''def {name}(result, expected, **options):
    """Check if VS Code keybindings.json contains the expected keybinding entry."""
    if result.get('error') or not result.get('entries'):
        return 0.0
    entries = result.get('entries', [])
    expected_key = expected.get('key', '')
    expected_command = expected.get('command', '')
    return 1.0 if scalecua_keybinding_matches(entries, expected_key, expected_command) else 0.0
'''
    return _replace_function_source(text, name, new_source)[0]


# PREDICATES, not ID allow-lists: match by getter/metric
# NAME-PREFIX so every present + future sibling of the class is covered, not one
# hand-listed hash. Each rewrite is self-limiting — the .replace()/inject no-ops on
# any function whose source doesn't carry the pattern, and the idempotency guard
# skips already-rewritten shards.
_DOCX_PARA_BOLD_GETTER_RE = re.compile(r"get_docx_para_bold__[0-9a-f]+")
_SUM_FORMULA_METRIC_RE = re.compile(r"check_sum_formula__[0-9a-f]+")


def _rewrite_docx_para_bold_getters(text: str) -> str:
    # Every get_docx_para_bold__* getter reads bold RAW (run.font.bold / run.bold /
    # r.font.bold across >=4 syntactic forms) and so misses STYLE-inherited bold ->
    # false-negative. They all open the doc with `doc = Document(tmp_path)`, so inject
    # one scalecua_bake_effective_run_bold(doc) right after that uniform open: every
    # read form then sees the baked effective value. Covers all 4 known hashes
    # (7a1ffa51/1defcaa1/7fd194e2/e827cac2) + any future form via the name prefix.
    for fn in sorted(set(_DOCX_PARA_BOLD_GETTER_RE.findall(text))):
        current = _function_source(text, fn)
        if not current or "scalecua_bake_effective_run_bold" in current:
            continue
        rewritten = current.replace(
            "        doc = Document(tmp_path)\n",
            "        doc = Document(tmp_path)\n"
            "        scalecua_bake_effective_run_bold(doc)\n",
            1,
        )
        if rewritten != current:
            text = _replace_function_source(text, fn, rewritten)[0]
    return text


def _rewrite_sum_formula_metrics(text: str) -> str:
    # Every check_sum_formula__* metric keys the formula credit on exact-string /
    # literal 'C2'&'C8' substring match, which $-absolute refs (=SUM($C$2:$C$8))
    # break. Route both branches through scalecua_formula_equivalent ($/ws/case-
    # insensitive). Prefix-scan (currently 1 hash: 9956b78e) so a future sibling is
    # covered automatically; .replace() no-ops on any form lacking the pattern.
    for fn in sorted(set(_SUM_FORMULA_METRIC_RE.findall(text))):
        current = _function_source(text, fn)
        if not current or "scalecua_formula_equivalent" in current:
            continue
        rewritten = current.replace(
            "        if any((f.upper() == formula_upper for f in valid_formulas)):\n"
            "            score += 0.5\n"
            "        elif 'SUM' in formula_upper and 'C2' in formula_upper and "
            "('C8' in formula_upper):\n"
            "            score += 0.5\n",
            "        if any(scalecua_formula_equivalent(formula, f) for f in valid_formulas):\n"
            "            score += 0.5\n"
            "        elif 'SUM' in formula_upper and scalecua_formula_equivalent("
            "formula.replace('$', ''), '=SUM(C2:C8)'):\n"
            "            score += 0.5\n",
        )
        if rewritten != current:
            text = _replace_function_source(text, fn, rewritten)[0]
    return text


def patch_judge_functions(judges_root: Path | str) -> dict:
    """Inject the A-1 read-helper library into every pulled judge overlay module.

    Idempotent (guarded by ``HELPER_MARKER``). Runs on the *staging* tree during
    ``dataset.import_all`` before catalog validation, so the fix survives every
    re-pull of the upstream ``judge_functions`` cache and is not a manual
    ``.cache`` edit.
    """
    judges_root = Path(judges_root)
    source_block = "\n\n" + helper_source()
    report = {"patched": [], "already_patched": [], "rewritten": []}
    for path in _overlay_module_files(judges_root):
        text = path.read_text(encoding="utf-8")
        text, rewritten = _rewrite_judge_helpers(text)
        if HELPER_MARKER in text:
            if rewritten:
                path.write_text(text, encoding="utf-8")
                report["rewritten"].append(str(path))
            report["already_patched"].append(str(path))
            continue
        if not text.endswith("\n"):
            text += "\n"
        path.write_text(text + source_block, encoding="utf-8")
        report["patched"].append(str(path))
        if rewritten:
            report["rewritten"].append(str(path))
    return report
