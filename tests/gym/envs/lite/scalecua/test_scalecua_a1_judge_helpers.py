"""ScaleCUA A1 judge helper patch and runtime helper tests."""

from __future__ import annotations

import ast
import importlib.util
import textwrap
from types import SimpleNamespace

from lite.gym.envs.lite.scalecua.src.gen.eval import judge_helpers as scalecua_judge_helpers

# ---------------------------------------------------------------------------
# #154 G1-7 / G1-8 — A-1 read-helpers + import-pipeline patch (judge_helpers)
# ---------------------------------------------------------------------------


def test_scalecua_a1_helper_source_is_valid_and_self_contained():
    src = scalecua_judge_helpers.helper_source()
    ns: dict = {}
    exec(compile(src, "<a1-helpers>", "exec"), ns)
    assert "scalecua_xlsx_effective_fill" in ns
    assert "scalecua_docx_paragraph_shading" in ns
    assert "scalecua_normalize_compare_string" in ns


def test_scalecua_a1_patch_is_injected_and_idempotent(tmp_path):
    for split in ("train", "rl"):
        (tmp_path / split / "verigen_getters").mkdir(parents=True)
        (tmp_path / split / "getters.py").write_text("from .verigen_getters import *\n")
        (tmp_path / split / "metrics.py").write_text("X = 1\n")
        (tmp_path / split / "verigen_getters" / "__init__.py").write_text("from .calc import *\n")
        (tmp_path / split / "verigen_getters" / "calc.py").write_text(
            "def get_x(env, c):\n    return 1\n"
        )

    first = scalecua_judge_helpers.patch_judge_functions(tmp_path)
    second = scalecua_judge_helpers.patch_judge_functions(tmp_path)
    # first pass patches getters.py + metrics.py + verigen_getters/calc.py per split
    assert len(first["patched"]) == 6 and not first["already_patched"] and not first["rewritten"]
    # re-pull is not clobbered: second pass is a no-op (marker guard)
    assert not second["patched"] and len(second["already_patched"]) == 6 and not second["rewritten"]

    patched = (tmp_path / "rl" / "verigen_getters" / "calc.py").read_text()
    ast.parse(patched)  # still valid python
    assert patched.count(scalecua_judge_helpers.HELPER_MARKER) == 1
    assert "def get_x" in patched  # original checker preserved
    assert "def scalecua_xlsx_effective_fill" in patched  # helper resolvable in-module


def test_scalecua_a1_patch_survives_clean_repull_and_helpers_execute(tmp_path):
    import docx
    import openpyxl
    from docx.oxml.ns import qn
    from openpyxl.formatting.rule import CellIsRule
    from openpyxl.styles import PatternFill

    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = 10
    ws["A2"] = 90
    ws["A3"] = "=AVERAGE(A1:A2)"
    ws.freeze_panes = "B2"
    red = PatternFill(start_color="FFFF0000", end_color="FFFF0000", fill_type="solid")
    ws.conditional_formatting.add(
        "A1:A2", CellIsRule(operator="greaterThan", formula=["50"], fill=red)
    )
    xlsx_path = tmp_path / "a1.xlsx"
    wb.save(xlsx_path)

    doc = docx.Document()
    run = doc.add_paragraph().add_run("highlighted")
    rpr = run._r.get_or_add_rPr()
    shd = rpr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): "FFFF00"})
    rpr.append(shd)
    docx_path = tmp_path / "a1.docx"
    doc.save(docx_path)

    raw_overlay = textwrap.dedent(
        """
        def probe_a1_helpers(xlsx_path, docx_path):
            formula = scalecua_xlsx_cell_formula_and_value(xlsx_path, None, "A3")["formula"]
            return {
                "fill": scalecua_xlsx_effective_fill(xlsx_path, None, "A2"),
                "formula_ok": scalecua_formula_equivalent(formula, "=AVERAGE(A1:A2)"),
                "freeze": scalecua_xlsx_freeze_panes(xlsx_path),
                "shading": scalecua_docx_paragraph_shading(docx_path, 0),
            }
        """
    ).lstrip()

    for pull_index in (1, 2):
        judges_root = tmp_path / f"clean_pull_{pull_index}"
        module_path = judges_root / "rl" / "metrics.py"
        module_path.parent.mkdir(parents=True)
        module_path.write_text(raw_overlay, encoding="utf-8")
        assert scalecua_judge_helpers.HELPER_MARKER not in module_path.read_text(encoding="utf-8")

        report = scalecua_judge_helpers.patch_judge_functions(judges_root)
        assert report == {"patched": [str(module_path)], "already_patched": [], "rewritten": []}

        spec = importlib.util.spec_from_file_location(f"a1_repull_{pull_index}", module_path)
        assert spec and spec.loader
        mod = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(mod)

        assert mod.probe_a1_helpers(str(xlsx_path), str(docx_path)) == {
            "fill": "FFFF0000",
            "formula_ok": True,
            "freeze": "B2",
            "shading": "FFFF00",
        }


def test_scalecua_a1_patch_rewrites_real_cf_color_getter(tmp_path):
    import importlib.util

    import openpyxl
    from openpyxl.formatting.rule import CellIsRule
    from openpyxl.styles import PatternFill

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Calendar"
    ws["B6"] = 7
    red = PatternFill(start_color="FFFF0000", end_color="FFFF0000", fill_type="solid")
    ws.conditional_formatting.add(
        "B6:B6", CellIsRule(operator="greaterThan", formula=["5"], fill=red)
    )
    xlsx_path = tmp_path / "cf.xlsx"
    wb.save(xlsx_path)

    judges_root = tmp_path / "judges"
    module_path = judges_root / "rl" / "verigen_getters" / "calc.py"
    module_path.parent.mkdir(parents=True)
    module_path.write_text(
        textwrap.dedent(
            """
            def get_xlsx_cell_colors__03a99dd86c6d30983d2467bc1177489c_qw35sft2_df3a9ab2(
                env, config: dict
            ):
                import tempfile, os, openpyxl
                path = config.get('path', '/home/user/Calendar_Highlight_Weekend_Days.xlsx')
                cells_to_check = config.get('cells', [])
                file_bytes = env.controller.get_file(path)
                if not file_bytes:
                    return {'error': 'File not found'}
                with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
                    tmp.write(file_bytes)
                    tmp_path = tmp.name
                try:
                    wb = openpyxl.load_workbook(tmp_path, data_only=True)
                    ws = wb.worksheets[0]
                    result = {}
                    for cell_ref in cells_to_check:
                        cell = ws[cell_ref]
                        fill = cell.fill
                        if fill and fill.fgColor and (fill.fgColor.type == 'rgb'):
                            color = fill.fgColor.rgb
                        else:
                            color = '00000000'
                        result[cell_ref] = color
                    return result
                finally:
                    os.unlink(tmp_path)
            """
        ).lstrip(),
        encoding="utf-8",
    )

    report = scalecua_judge_helpers.patch_judge_functions(judges_root)
    assert report["rewritten"] == [str(module_path)]
    patched = module_path.read_text(encoding="utf-8")
    assert "scalecua_xlsx_effective_fill" in patched

    spec = importlib.util.spec_from_file_location("a1_cf_rewrite", module_path)
    assert spec and spec.loader
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    env = SimpleNamespace(controller=SimpleNamespace(get_file=lambda _path: xlsx_path.read_bytes()))

    assert mod.get_xlsx_cell_colors__03a99dd86c6d30983d2467bc1177489c_qw35sft2_df3a9ab2(
        env, {"cells": ["B6"]}
    ) == {"B6": "FFFF0000"}


def test_scalecua_a1_cf_getter_flips_full_metric_verdict_0_to_1(tmp_path):
    """#154 A-1 recovery evidence: the landed patch flips the real CF weekend-red
    checker (calc_8b1ce5f2 traj_verify_6) from reward 0 -> 1 on a CF-painted
    artifact, stays backward-compatible on static fill, and keeps a negative
    control at 0. Asserts the end-to-end getter+metric verdict, not just the
    getter's raw output. The arbiter is the PULLED xlsx + its key mutation (the
    conditional-formatting rule), per #155.
    """
    import importlib.util

    import openpyxl
    from openpyxl.formatting.rule import CellIsRule
    from openpyxl.styles import PatternFill

    getter_name = "get_xlsx_cell_colors__03a99dd86c6d30983d2467bc1177489c_qw35sft2_df3a9ab2"
    red = PatternFill(start_color="FFFF0000", end_color="FFFF0000", fill_type="solid")

    def make_xlsx(name, *, cf=False, static=False):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Calendar"
        for cell in ("B6", "B7"):
            ws[cell] = 7  # weekend marker value (>5)
            if static:
                ws[cell].fill = red
        if cf:
            # KEY MUTATION: a conditional-format rule paints the cells red; the
            # cells carry NO static fill, so a static cell.fill read misses it.
            ws.conditional_formatting.add(
                "B1:B30", CellIsRule(operator="greaterThan", formula=["5"], fill=red)
            )
        path = tmp_path / name
        wb.save(path)
        return path

    cf_only = make_xlsx("cf_only.xlsx", cf=True)
    static_fill = make_xlsx("static.xlsx", static=True)
    unformatted = make_xlsx("none.xlsx")

    # The real upstream (OLD, CF-blind) getter body for this exact checker id.
    old_getter_src = (
        textwrap.dedent(
            """
        def GETTER(env, config: dict):
            import tempfile, os, openpyxl
            file_bytes = env.controller.get_file(config.get('path', '/x.xlsx'))
            if not file_bytes:
                return {'error': 'File not found'}
            with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            try:
                wb = openpyxl.load_workbook(tmp_path, data_only=True)
                ws = wb.worksheets[0]
                result = {}
                for cell_ref in config.get('cells', []):
                    cell = ws[cell_ref]
                    fill = cell.fill
                    if fill and fill.fgColor and (fill.fgColor.type == 'rgb'):
                        color = fill.fgColor.rgb
                    else:
                        color = '00000000'
                    result[cell_ref] = color
                return result
            finally:
                os.unlink(tmp_path)
        """
        )
        .replace("GETTER", getter_name)
        .lstrip()
    )

    # The real upstream metric for this checker id (partial credit).
    metric_src = textwrap.dedent(
        """
        def check_cells_red(result, expected, **options):
            if not isinstance(result, dict) or result.get('error'):
                return 0.0
            target_cells = expected.get('target_cells', [])
            expected_color = expected.get('expected_color', 'FFFF0000').upper()
            if not target_cells:
                return 0.0
            correct = sum(1 for c in target_cells
                          if result.get(c, '').upper() == expected_color)
            return correct / len(target_cells)
        """
    ).lstrip()

    expected = {"target_cells": ["B6", "B7"], "expected_color": "FFFF0000"}
    cfg = {"path": "/x.xlsx", "cells": ["B6", "B7"], "sheet": 0}

    def verdicts(module, tag):
        getter = getattr(module, getter_name)
        check = module.check_cells_red
        out = {}
        for label, art in (("cf", cf_only), ("static", static_fill), ("none", unformatted)):
            env = SimpleNamespace(
                controller=SimpleNamespace(get_file=lambda _p, a=art: a.read_bytes())
            )
            out[label] = check(getter(env, cfg), expected)
        return out

    def load_module(source, stem):
        path = tmp_path / stem / "rl" / "verigen_getters" / "calc.py"
        path.parent.mkdir(parents=True)
        path.write_text(source, encoding="utf-8")
        return path

    # OLD: raw upstream getter body, executed WITHOUT the patch.
    old_path = load_module(old_getter_src + "\n" + metric_src, "old")
    old_spec = importlib.util.spec_from_file_location("a1_cf_old", old_path)
    old_mod = importlib.util.module_from_spec(old_spec)
    old_spec.loader.exec_module(old_mod)
    old = verdicts(old_mod, "old")

    # PATCHED: run the REAL patch_judge_functions over the same overlay; it
    # rewrites the getter body to use scalecua_xlsx_effective_fill + injects the
    # helper library.
    patched_path = load_module(old_getter_src + "\n" + metric_src, "patched")
    scalecua_judge_helpers.patch_judge_functions(patched_path.parents[2])
    assert "scalecua_xlsx_effective_fill" in patched_path.read_text(encoding="utf-8")
    patched_spec = importlib.util.spec_from_file_location("a1_cf_patched", patched_path)
    patched_mod = importlib.util.module_from_spec(patched_spec)
    patched_spec.loader.exec_module(patched_mod)
    patched = verdicts(patched_mod, "patched")

    # The recovery: CF-painted artifact flips 0 -> 1.
    assert old["cf"] == 0.0 and patched["cf"] == 1.0
    # Backward compatible: a real static fill still scores 1 either way.
    assert old["static"] == 1.0 and patched["static"] == 1.0
    # Negative control: an unformatted artifact stays 0 after the patch.
    assert old["none"] == 0.0 and patched["none"] == 0.0


def test_scalecua_a1_patch_rewrites_real_vscode_keybinding_pair(tmp_path):
    import importlib.util

    judges_root = tmp_path / "judges"
    getter_path = judges_root / "rl" / "verigen_getters" / "vscode.py"
    metric_path = judges_root / "rl" / "verigen_metrics" / "vscode.py"
    getter_path.parent.mkdir(parents=True)
    metric_path.parent.mkdir(parents=True)
    getter_path.write_text(
        textwrap.dedent(
            r"""
            import json, re

            def get_vscode_keybindings__ed42af8e24d428ca40b99f135c68274f_qw35sft2_44312464(
                env, config: dict
            ):
                keybindings_path = '/home/user/.config/Code/User/keybindings.json'
                file_bytes = env.controller.get_file(keybindings_path)
                if not file_bytes:
                    return {'error': 'keybindings.json not found', 'entries': []}
                try:
                    content = file_bytes.decode('utf-8')
                    content = re.sub('//[^\n]*', '', content)
                    content = re.sub('/\*.*?\*/', '', content, flags=re.DOTALL)
                    entries = json.loads(content.strip())
                    return {'entries': entries if isinstance(entries, list) else []}
                except Exception as e:
                    return {'error': str(e), 'entries': []}
            """
        ).lstrip(),
        encoding="utf-8",
    )
    metric_path.write_text(
        textwrap.dedent(
            """
            def check_vscode_keybinding__ed42af8e24d428ca40b99f135c68274f_qw35sft2_bb493910(
                result, expected, **options
            ):
                if result.get('error') or not result.get('entries'):
                    return 0.0
                entries = result.get('entries', [])
                expected_key = expected.get('key', '').lower().replace(' ', '')
                expected_command = expected.get('command', '')
                for entry in entries:
                    if not isinstance(entry, dict):
                        continue
                    entry_key = entry.get('key', '').lower().replace(' ', '')
                    entry_command = entry.get('command', '')
                    if entry_command.startswith('-'):
                        continue
                    if entry_key == expected_key and entry_command == expected_command:
                        return 1.0
                return 0.0
            """
        ).lstrip(),
        encoding="utf-8",
    )

    report = scalecua_judge_helpers.patch_judge_functions(judges_root)
    assert set(report["rewritten"]) == {str(getter_path), str(metric_path)}

    getter_spec = importlib.util.spec_from_file_location("a1_vscode_getter", getter_path)
    metric_spec = importlib.util.spec_from_file_location("a1_vscode_metric", metric_path)
    assert getter_spec and getter_spec.loader and metric_spec and metric_spec.loader
    getter_mod = importlib.util.module_from_spec(getter_spec)
    metric_mod = importlib.util.module_from_spec(metric_spec)
    getter_spec.loader.exec_module(getter_mod)
    metric_spec.loader.exec_module(metric_mod)

    keybindings = b"""[
      // UI-created binding
      {"key": "shift+ctrl+j", "command": "workbench.action.togglePanel",},
    ]"""
    env = SimpleNamespace(controller=SimpleNamespace(get_file=lambda _path: keybindings))
    result = getter_mod.get_vscode_keybindings__ed42af8e24d428ca40b99f135c68274f_qw35sft2_44312464(
        env, {}
    )
    expected = {"key": "ctrl+shift+j", "command": "workbench.action.togglePanel"}

    assert (
        metric_mod.check_vscode_keybinding__ed42af8e24d428ca40b99f135c68274f_qw35sft2_bb493910(
            result, expected
        )
        == 1.0
    )
    assert (
        metric_mod.check_vscode_keybinding__ed42af8e24d428ca40b99f135c68274f_qw35sft2_bb493910(
            result, {"key": "ctrl+shift+j", "command": "wrong.command"}
        )
        == 0.0
    )


def test_scalecua_a1_conditional_formatting_fill(tmp_path):
    import openpyxl
    from openpyxl.formatting.rule import CellIsRule
    from openpyxl.styles import PatternFill

    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = 10
    ws["A2"] = 90
    red = PatternFill(start_color="FFFF0000", end_color="FFFF0000", fill_type="solid")
    ws.conditional_formatting.add(
        "A1:A2", CellIsRule(operator="greaterThan", formula=["50"], fill=red)
    )
    path = tmp_path / "cf.xlsx"
    wb.save(path)

    rules = scalecua_judge_helpers.scalecua_xlsx_conditional_formats(path)
    assert rules and rules[0]["fill"] == "FFFF0000"
    # end-state the static-fill proxy misses: A2 satisfies the rule -> red
    assert scalecua_judge_helpers.scalecua_xlsx_effective_fill(path, None, "A2") == "FFFF0000"
    assert scalecua_judge_helpers.scalecua_xlsx_effective_fill(path, None, "A1") is None


def test_scalecua_a1_formula_value_and_freeze_and_date(tmp_path):
    from datetime import datetime

    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = 2
    ws["A2"] = 4
    ws["A3"] = "=AVERAGE(A1:A2)"
    ws["B1"] = datetime(2024, 3, 15)
    ws["B1"].number_format = "YYYY-MM-DD"
    ws.freeze_panes = "B2"
    path = tmp_path / "f.xlsx"
    wb.save(path)

    fv = scalecua_judge_helpers.scalecua_xlsx_cell_formula_and_value(path, None, "A3")
    assert fv["formula"] == "=AVERAGE(A1:A2)"
    assert scalecua_judge_helpers.scalecua_formula_equivalent("=AVERAGE(A1:A2)", "average(a1:a2)")
    assert scalecua_judge_helpers.scalecua_formula_equivalent("=$A$1+$B$1", "A1+B1")
    assert scalecua_judge_helpers.scalecua_xlsx_freeze_panes(path) == "B2"
    date_info = scalecua_judge_helpers.scalecua_xlsx_cell_date(path, None, "B1")
    assert date_info["is_date"]
    assert scalecua_judge_helpers.scalecua_date_equal(date_info["value"], "2024-03-15")
    assert scalecua_judge_helpers.scalecua_date_equal("03/15/2024", datetime(2024, 3, 15))


def test_scalecua_a1_docx_shd_highlight_and_paragraph_index(tmp_path):
    import docx
    from docx.oxml.ns import qn

    doc = docx.Document()
    doc.add_heading("Title Here", level=0)
    para = doc.add_paragraph()
    run = para.add_run("highlighted text")
    rpr = run._r.get_or_add_rPr()
    shd = rpr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): "FFFF00"})
    rpr.append(shd)
    doc.add_paragraph("plain body para")
    path = tmp_path / "h.docx"
    doc.save(path)

    # w:shd highlight recognised where run.highlight would read None
    assert scalecua_judge_helpers.scalecua_docx_paragraph_shading(path, 1) == "FFFF00"
    assert scalecua_judge_helpers.scalecua_docx_paragraph_shading(path, 2) is None
    # body-paragraph index skips the title/heading
    body = scalecua_judge_helpers.scalecua_docx_body_paragraph_index(path)
    assert 0 not in body and 1 in body and 2 in body
    assert (
        scalecua_judge_helpers.scalecua_docx_body_paragraph_index(
            path, target_text="plain body para"
        )
        == 2
    )


def test_scalecua_a1_docx_effective_run_format(tmp_path):
    import docx

    doc = docx.Document()
    para = doc.add_paragraph()
    run = para.add_run("bold text")
    run.font.bold = True
    path = tmp_path / "b.docx"
    doc.save(path)

    fmt = scalecua_judge_helpers.scalecua_docx_effective_run_format(path, 0, 0)
    assert fmt["bold"] is True


def test_scalecua_bake_effective_run_bold_generalizes_whole_class(tmp_path):
    # De-overfit guard (RC-FN-6/7): the docx style-bold fix must cover EVERY
    # get_docx_para_bold__* getter via a name-prefix predicate, not a hash allow-list.
    import docx
    from docx.enum.style import WD_STYLE_TYPE

    jh = scalecua_judge_helpers

    # (1) mechanism: STYLE-inherited bold (run.font.bold is None) resolved to True;
    # a plain run and an explicit direct-False run are left untouched (no false pos).
    d = docx.Document()
    st = d.styles.add_style("BoldBody", WD_STYLE_TYPE.PARAGRAPH)
    st.font.bold = True
    p = d.add_paragraph("x")
    p.style = st
    assert p.runs[0].font.bold is None  # the FN condition
    jh.scalecua_bake_effective_run_bold(d)
    assert p.runs[0].font.bold is True  # resolved

    d2 = docx.Document()
    p2 = d2.add_paragraph("y")
    jh.scalecua_bake_effective_run_bold(d2)
    assert p2.runs[0].font.bold in (None, False)  # plain: no false positive

    d3 = docx.Document()
    p3 = d3.add_paragraph()
    p3.style = st
    r3 = p3.add_run("z")
    r3.font.bold = False  # explicit direct value under a bold style
    jh.scalecua_bake_effective_run_bold(d3)
    assert r3.font.bold is False  # direct value preserved (never overridden)

    # (2) CLASS COVERAGE: two DIFFERENT hashes in DIFFERENT read-forms must BOTH get
    # the bake injected after their uniform `doc = Document(tmp_path)` open.
    text = (
        "def get_docx_para_bold__aaaa1111(config, tmp_path):\n"
        "    from docx import Document\n"
        "    if True:\n"
        "        doc = Document(tmp_path)\n"
        "        para = doc.paragraphs[0]\n"
        "        return all(run.font.bold for run in para.runs)\n"
        "\n\n"
        "def get_docx_para_bold__bbbb2222(config, tmp_path):\n"
        "    from docx import Document\n"
        "    if True:\n"
        "        doc = Document(tmp_path)\n"
        "        para = doc.paragraphs[0]\n"
        "        return [1 for run in para.runs if not run.bold]\n"
    )
    out, changed = jh._rewrite_judge_helpers(text)
    assert changed
    # BOTH getters baked (2), not just one — the overfit this test guards against.
    assert out.count("scalecua_bake_effective_run_bold(doc)") == 2


def test_scalecua_a1_string_ext_and_vscode_read_bugs(tmp_path):
    # RC-FN-15 string normalize
    norm = scalecua_judge_helpers.scalecua_normalize_compare_string
    assert norm(" $Total–Value\n") == "total-value"
    assert norm("ABC") == norm("abc")

    # RC-FN-11 depth-tolerant unpacked-extension manifest search
    (tmp_path / "ext" / "a" / "b" / "c").mkdir(parents=True)
    (tmp_path / "ext" / "a" / "b" / "c" / "manifest.json").write_text("{}")
    (tmp_path / "ext" / "top").mkdir()
    (tmp_path / "ext" / "top" / "manifest.json").write_text("{}")
    found = scalecua_judge_helpers.scalecua_find_unpacked_extension_manifest(tmp_path / "ext")
    assert found.endswith("top/manifest.json")

    # RC-FN-9 vs_code: comment-laden argv.json, default/scope, keybinding when-clause
    parsed = scalecua_judge_helpers.scalecua_parse_jsonc('{ // c\n "a":1, /* b */ "c":2, }')
    assert parsed == {"a": 1, "c": 2}
    assert scalecua_judge_helpers.scalecua_vscode_setting_effective({"x": 1}, {"x": 2}, "x") == 2
    assert (
        scalecua_judge_helpers.scalecua_vscode_setting_effective({}, {}, "editor.tabSize", 4) == 4
    )
    kb = [{"key": "ctrl+shift+p", "command": "foo", "when": "editorTextFocus"}]
    assert scalecua_judge_helpers.scalecua_keybinding_matches(kb, "shift+ctrl+p", "foo")
    assert not scalecua_judge_helpers.scalecua_keybinding_matches(kb, "ctrl+p", "foo")
